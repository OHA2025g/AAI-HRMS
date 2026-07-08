"""Resume text extraction for career trajectory analysis."""

from __future__ import annotations

import io
from typing import Optional

from docx import Document


def _pdf_text_pdfplumber(raw: bytes) -> str:
    try:
        import pdfplumber

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _pdf_text_pypdf(raw: bytes) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _docx_text(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def extract_text_from_bytes(raw: bytes, filename: str = "") -> str:
    """Extract plain text from PDF, DOCX, or TXT bytes."""
    name = (filename or "").lower()
    if name.endswith(".txt") or not name:
        try:
            return raw.decode("utf-8", errors="ignore").strip()
        except Exception:
            return ""
    if name.endswith(".docx") or name.endswith(".doc"):
        return _docx_text(raw)
    if name.endswith(".pdf"):
        text = _pdf_text_pdfplumber(raw)
        if len(text.strip()) < 80:
            text = _pdf_text_pypdf(raw)
        return text.strip()
    return raw.decode("utf-8", errors="ignore").strip()
