from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks, Request, Body, Header, Query
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from starlette.responses import Response
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, ConfigDict, EmailStr, Field, model_validator
from typing import List, Optional, Dict, Any, Union, Literal, Set, Tuple
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import json
import re
import io
import csv
import random
import httpx
from PyPDF2 import PdfReader
from docx import Document
import time
import contextvars
import hashlib
import asyncio
from collections import defaultdict

from talent_acquisition.connector_fetch import fetch_connector_candidates
from talent_acquisition.ingestion_queue import INGESTION_JOBS_COLLECTION, run_unified_ingestion
from talent_acquisition.candidate_source import (
    all_talent_pool_mongo_filter,
    display_channel_mongo_filter,
    is_linkedin_sourced_candidate,
    is_talent_pool_candidate,
)
from talent_acquisition.candidate_fit_filter import candidate_ids_matching_fit_range
from talent_acquisition.match_candidate_pool import (
    applied_ids_excluding_fit_seeds,
    gather_job_match_candidates,
    load_persisted_fit_score,
    merge_fit_with_seed_persisted,
)
from talent_acquisition.match_ordering import (
    DEFAULT_TOTAL_MATCH_LIMIT,
    count_match_buckets,
    order_job_match_results,
)
from talent_acquisition.hiring_analytics_events import log_find_matches_event
from talent_acquisition.hiring_dashboard import build_hiring_dashboard_pack
from talent_acquisition.hiring_dashboard_access import enforce_hiring_dashboard_scope
from talent_acquisition.hiring_dashboard_perf import log_slow_hiring_pack_query
from talent_acquisition.hiring_dashboard_schemas import (
    HiringAlertDismissRequest,
    HiringAlertDismissalsResponse,
    HiringDashboardConfigResponse,
    HiringDashboardConfigUpdate,
    HiringDashboardPack,
    HiringDashboardTrends,
    HiringSnapshotHealth,
)
from talent_acquisition.hiring_dashboard_config import (
    config_to_json,
    get_hiring_dashboard_config,
    upsert_hiring_dashboard_config,
)
from talent_acquisition.hiring_alert_dismissals import (
    dismiss_alert as dismiss_hiring_alert,
    list_dismissed_alert_ids,
    restore_alert as restore_hiring_alert,
)
from talent_acquisition.hiring_pack_cache import (
    get_cached_hiring_pack,
    invalidate_hiring_pack_cache,
    set_cached_hiring_pack,
)
from talent_acquisition.hiring_snapshots import (
    get_hiring_dashboard_trends,
    get_hiring_snapshot_health,
    seed_hiring_snapshots_if_sparse,
    write_hiring_dashboard_snapshot,
)
from talent_acquisition.assessments_routes import create_assessments_router
from m2_employee_lifecycle.state_machine import (
    approval_rule_for_event,
    target_status_for_event,
    validate_direct_status_transition,
    validate_lifecycle_event_for_status,
)
from m2_employee_lifecycle.document_sla import default_sla_due, is_past_iso

from m3_workforce_intel.baseline_model import BaselineParams, evaluate_on_history, fit_per_skill_baseline, predict_demand
from m3_workforce_intel.constants import (
    COL_DRIFT_EVENTS,
    COL_ETL_RUNS,
    COL_EVAL_RUNS,
    COL_HIST_FEATURES,
    COL_MODELS,
    COL_MODEL_STATE,
    COL_MONITORING_STATE,
    MODEL_STATE_DOC_ID,
    MONITORING_STATE_DOC_ID,
)
from m3_workforce_intel.features import extract_workforce_intel_feature_rows
from m3_workforce_intel.hist_store import load_demand_series_by_skill
from m3_workforce_intel.monitoring import evaluate_active_model_vs_current, retrain_trigger_evaluation
from m3_workforce_intel.pipeline import etl_backfill_demo, etl_snapshot

import m3_workforce_intel.prom_metrics  # noqa: F401 — register M3 Prometheus gauges

from m4_resource_optimization.constants import COL_ALLOCATION_SCENARIOS, COL_ALLOCATION_SETTINGS, SETTINGS_DOC_ID
from m4_resource_optimization.io import (
    apply_assignments_to_project_allocations,
    ensure_default_settings,
    get_merged_settings,
    run_allocation_solve,
)
from m4_resource_optimization.solver import compare_solve_results

from m5_training.constants import (
    COL_ASSIGNMENTS,
    COL_CERTIFICATIONS,
    COL_LEARNING_PATH_TEMPLATES,
    COL_LMS_COURSES,
    COL_LMS_SYNC_RUNS,
    DEFAULT_LMS_PROVIDER,
)
from m5_training.lms_sync import run_lms_catalog_sync
from m5_training.recommendation_rules import build_employee_recommendation_payloads
from m5_training.service import load_path_templates_map, manager_team_training_summary, scan_certification_expiry

from m6_engagement.audit_log import log_engagement_privacy_event
from m6_engagement.constants import (
    COL_PRIVACY_AUDIT,
    COL_SURVEY_SCHEDULES,
    COL_SURVEY_TEMPLATES,
)
from m6_engagement.privacy import anonymity_min_threshold, redacted_dashboard_payload, should_redact_survey_aggregates
from m6_engagement.schedules import next_run_after, parse_iso_dt
from m6_engagement.sentiment import compute_sentiment as m6_compute_sentiment
from m6_engagement.topics import aggregate_topic_counts, classify_topic, confidence_tier, weekly_rating_trends

from m7_automation.constants import (
    COL_HR_COPILOT_AUDIT,
    COL_MANUAL_WORKFLOW_BASELINES,
    COL_WORKFLOW_RULES,
    COL_WORKFLOW_RUNS,
    COPILOT_ENGINE_VERSION,
    WORKFLOW_ENGINE_VERSION,
)
from m7_automation.copilot_hf import resolve_copilot_intent_async
from m7_automation.copilot_intent import extract_employee_code_hint, help_text
from m7_automation.retry import run_with_retries
from m7_automation.savings import baseline_map, compute_savings_totals
from m7_automation.workflow_flow import execute_flow_graph
from m7_automation.workflow_triggers import should_execute_trigger
from m7_automation.workflow_webhook import execute_http_webhook

from m8_retention.constants import (
    ATTRITION_MODEL_VERSION,
    COL_ATTRITION_MODEL_STATE,
    COL_ATTRITION_SCORES_LATEST,
    COL_RETENTION_INTERVENTIONS,
    COL_RETENTION_PLAYBOOKS,
    COL_RETENTION_SEGMENT_SETTINGS,
)
from m8_retention import service as m8_retention_service

from m9_analytics.constants import COL_M9_KPI_DEFINITIONS, COL_M9_LEADERSHIP_SNAPSHOTS

from m10_allocation_section.constants import (
    COL_ACTIVITY_LOGS,
    COL_AI_INSIGHTS,
    COL_ALERTS,
    COL_BENCH_MATCHES,
    COL_CHANGES,
    COL_CONFLICTS,
    COL_DOCUMENTS,
    COL_FORECAST_SNAPSHOTS,
    COL_NOTES,
    COL_RELEASES,
    COL_ROLL_EVENTS,
    COL_STAFFING_REQUEST_HISTORY,
    COL_STAFFING_REQUESTS,
    COL_WORKFLOW_APPROVALS,
)
from m10_allocation_section.routes import create_allocation_section_router
from m11_resource_section.constants import (
    COL_ACTIVITY as RS_COL_ACTIVITY,
    COL_AI_INSIGHTS as RS_COL_AI_INSIGHTS,
    COL_APPROVALS as RS_COL_APPROVALS,
    COL_ATTENDANCE_IMPACT as RS_COL_ATTENDANCE_IMPACT,
    COL_AVAILABILITY as RS_COL_AVAILABILITY,
    COL_BENCH_RECORDS as RS_COL_BENCH_RECORDS,
    COL_CAREER as RS_COL_CAREER,
    COL_CERTIFICATIONS as RS_COL_CERTIFICATIONS,
    COL_CLASSIFICATIONS as RS_COL_CLASSIFICATIONS,
    COL_COMPLIANCE as RS_COL_COMPLIANCE,
    COL_COST_PROFILES as RS_COL_COST_PROFILES,
    COL_DEMAND_MATCHES as RS_COL_DEMAND_MATCHES,
    COL_DOCUMENTS as RS_COL_RESOURCE_DOCUMENTS,
    COL_NOTES as RS_COL_RESOURCE_NOTES,
    COL_FORECASTS as RS_COL_FORECASTS,
    COL_LEARNING as RS_COL_LEARNING,
    COL_MOBILITY as RS_COL_MOBILITY,
    COL_NOTES as RS_COL_RESOURCE_NOTES,
    COL_PROFILES as RS_COL_PROFILES,
    COL_READINESS as RS_COL_READINESS,
    COL_SKILL_RECORDS as RS_COL_SKILL_RECORDS,
    COL_UTIL_SNAPSHOTS as RS_COL_UTIL_SNAPSHOTS,
)
from m11_resource_section.routes import create_resource_section_router
from m12_training_development.routes import create_training_development_router
from m12_training_development.constants import (
    COL_APPROVAL_REQUESTS as TD_COL_APPROVAL_REQUESTS,
    COL_ASSESSMENT_RESULTS as TD_COL_ASSESSMENT_RESULTS,
    COL_ASSESSMENTS as TD_COL_ASSESSMENTS,
    COL_CATALOG_ITEMS as TD_COL_CATALOG_ITEMS,
    COL_EXTENDED_RECORDS as TD_COL_EXTENDED_RECORDS,
    COL_TRAINING_ATTENDANCE as TD_COL_TRAINING_ATTENDANCE,
    COL_TRAINING_BATCHES as TD_COL_TRAINING_BATCHES,
    COL_TRAINING_ENROLLMENTS as TD_COL_TRAINING_ENROLLMENTS,
    COL_TRAINING_PROGRAMS as TD_COL_TRAINING_PROGRAMS,
    COL_TRAINING_SESSIONS as TD_COL_TRAINING_SESSIONS,
)
from m13_high_skill_talent_retention.routes import create_high_skill_retention_router
from career_trajectory.async_jobs import recover_stale_analyze_jobs
from career_trajectory.auto_analyze import trigger_auto_analyze_if_eligible
from career_trajectory.routes import create_career_trajectory_router
from candidate_fit_phase2.routes import create_phase2_fit_router
from m13_high_skill_talent_retention.constants import (
    COL_AI_FLIGHT_RISK as HSR_COL_AI_FLIGHT_RISK,
    COL_AI_RECOMMENDATIONS as HSR_COL_AI_RECOMMENDATIONS,
    COL_ATTRITION_PREDICTIONS as HSR_COL_ATTRITION_PREDICTIONS,
    COL_CRITICAL_TALENT_PROFILES as HSR_COL_CRITICAL_TALENT_PROFILES,
    COL_ENGAGEMENT_ACTION_PLANS as HSR_COL_ENGAGEMENT_ACTION_PLANS,
    COL_EXIT_RISK_TRIGGERS as HSR_COL_EXIT_RISK_TRIGGERS,
    COL_RISK_ASSESSMENTS as HSR_COL_RISK_ASSESSMENTS,
    COL_SEARCH_LOGS as HSR_COL_SEARCH_LOGS,
    COL_STABILITY_FORECASTS as HSR_COL_STABILITY_FORECASTS,
    COL_STAY_INTERVIEWS as HSR_COL_STAY_INTERVIEWS,
    COL_TALENT_CRITICALITY_TAGS as HSR_COL_TALENT_CRITICALITY_TAGS,
    COL_TALENT_SEGMENTS as HSR_COL_TALENT_SEGMENTS,
    COL_RETENTION_CASES as HSR_COL_RETENTION_CASES,
)
from m14_employee_lifecycle_management.routes import create_employee_lifecycle_management_router
from m14_employee_lifecycle_management.constants import (
    COL_PREBOARDING as ELM_COL_PREBOARDING,
    COL_ONBOARDING as ELM_COL_ONBOARDING,
    COL_PROBATION as ELM_COL_PROBATION,
    COL_CONFIRMATION as ELM_COL_CONFIRMATION,
    COL_EMPLOYEE_DOCUMENTS as ELM_COL_EMPLOYEE_DOCUMENTS,
    COL_BGV as ELM_COL_BGV,
    COL_POLICY_CONSENTS as ELM_COL_POLICY_CONSENTS,
    COL_ACCESS_PROVISIONING as ELM_COL_ACCESS_PROVISIONING,
    COL_PAYROLL_LINKAGE as ELM_COL_PAYROLL_LINKAGE,
    COL_APPROVAL_REQUESTS as ELM_COL_APPROVAL_REQUESTS,
    COL_RETENTION_SIGNALS as ELM_COL_RETENTION_SIGNALS,
    COL_RESIGNATION as ELM_COL_RESIGNATION,
    COL_NOTICE as ELM_COL_NOTICE,
    COL_EXIT_INTERVIEW as ELM_COL_EXIT_INTERVIEW,
    COL_CLEARANCE as ELM_COL_CLEARANCE,
    COL_FORECASTS as ELM_COL_FORECASTS,
    COL_AI_INSIGHTS as ELM_COL_AI_INSIGHTS,
    COL_LIFECYCLE_NOTES as ELM_COL_LIFECYCLE_NOTES,
    COL_ACTIVITY_LOGS as ELM_COL_ACTIVITY_LOGS,
)
from m15_workforce_intelligence.routes import create_workforce_intelligence_router
from m15_workforce_intelligence.constants import (
    COL_SNAPSHOT_RECORDS as WFI_COL_SNAPSHOT_RECORDS,
    COL_HEADCOUNT_RECORDS as WFI_COL_HEADCOUNT_RECORDS,
    COL_DEMOGRAPHIC_SNAPSHOTS as WFI_COL_DEMOGRAPHIC_SNAPSHOTS,
    COL_SKILL_VISIBILITY_RECORDS as WFI_COL_SKILL_VISIBILITY_RECORDS,
    COL_UTILIZATION_SNAPSHOTS as WFI_COL_UTILIZATION_SNAPSHOTS,
    COL_ENGAGEMENT_VISIBILITY_RECORDS as WFI_COL_ENGAGEMENT_VISIBILITY_RECORDS,
    COL_PERFORMANCE_VISIBILITY_RECORDS as WFI_COL_PERFORMANCE_VISIBILITY_RECORDS,
    COL_COMPLIANCE_VISIBILITY_RECORDS as WFI_COL_COMPLIANCE_VISIBILITY_RECORDS,
    COL_COST_VISIBILITY_RECORDS as WFI_COL_COST_VISIBILITY_RECORDS,
    COL_WORKFORCE_PLANS as WFI_COL_WORKFORCE_PLANS,
    COL_DEMAND_SUPPLY_RECORDS as WFI_COL_DEMAND_SUPPLY_RECORDS,
    COL_SCENARIO_MODELS as WFI_COL_SCENARIO_MODELS,
    COL_MANAGER_EFFECTIVENESS_RECORDS as WFI_COL_MANAGER_EFFECTIVENESS_RECORDS,
    COL_FORECASTS as WFI_COL_FORECASTS,
    COL_ATTRITION_PREDICTIONS as WFI_COL_ATTRITION_PREDICTIONS,
    COL_BURNOUT_PREDICTIONS as WFI_COL_BURNOUT_PREDICTIONS,
    COL_SKILL_RISK_PREDICTIONS as WFI_COL_SKILL_RISK_PREDICTIONS,
    COL_COST_RISK_PREDICTIONS as WFI_COL_COST_RISK_PREDICTIONS,
    COL_COMPLIANCE_RISK_PREDICTIONS as WFI_COL_COMPLIANCE_RISK_PREDICTIONS,
    COL_AI_RECOMMENDATIONS as WFI_COL_AI_RECOMMENDATIONS,
    COL_COPILOT_QUERIES as WFI_COL_COPILOT_QUERIES,
    COL_STRATEGIC_RISK_SNAPSHOTS as WFI_COL_STRATEGIC_RISK_SNAPSHOTS,
    COL_STRATEGIC_OPPORTUNITY_SNAPSHOTS as WFI_COL_STRATEGIC_OPPORTUNITY_SNAPSHOTS,
    COL_EXECUTIVE_SUMMARY_SNAPSHOTS as WFI_COL_EXECUTIVE_SUMMARY_SNAPSHOTS,
    COL_ACTIVITY_LOGS as WFI_COL_ACTIVITY_LOGS,
)
from m16_cost_optimization_automation.constants import ALL_INDEXED_COLLECTIONS
from m16_cost_optimization_automation.routes import create_cost_optimization_automation_router
from m17_employee_satisfaction_engagement.routes import create_employee_satisfaction_engagement_router
from m17_employee_satisfaction_engagement import service as m17_ese_service
from m9_analytics.export_packs import (
    create_full_leadership_pack_zip,
    create_monthly_cron_snapshots,
    create_monthly_snapshot_and_deliver,
    deliver_snapshot_webhook,
    format_snapshot_csv,
    format_snapshot_pdf,
    get_snapshot_doc,
    list_snapshots,
)
from m9_analytics.talent_kpis import compute_talent_acquisition_metrics
from m9_analytics.freshness import compute_source_freshness
from m9_analytics.compare import compare_snapshots
from m9_analytics.bundle import get_executive_dashboard_bundle
from m9_analytics.constants import COL_M9_KPI_THRESHOLDS
from m9_analytics.service import drill_filter_options, get_drill_dashboard_cached, get_kpi_pack, load_merged_kpi_definitions
from m9_analytics.definition_config import delete_definition_override
from m9_analytics.threshold_config import delete_threshold_override, list_threshold_overrides, upsert_threshold_override
from m9_analytics.predictive import get_executive_predictive_views
from m9_analytics.trends import get_kpi_trends
from m9_analytics.strategic_aggregate import build_strategic_dashboard_data

from m10_events.constants import COL_M10_EVENTS, COL_M10_HANDLER_AUDIT, COL_M10_IDEMPOTENCY
from m10_events.consumer import spawn_consumer_task
from m10_events.replay import replay_events
from m10_events.schemas import M10ReplayRequest
from m10_events.topics import (
    TOPIC_EMPLOYEE_LIFECYCLE_EVENT_CREATED,
    TOPIC_WORKFLOW_RUN_COMPLETED,
    TOPIC_WORKFLOW_RUN_FAILED,
)

from prometheus_client import CONTENT_TYPE_LATEST, Counter, Histogram, generate_latest

# Correlation ID for structured logs (M0-3).
request_id_ctx: contextvars.ContextVar[str] = contextvars.ContextVar("request_id", default="-")

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Optional: HashiCorp Vault KV or AWS Secrets Manager (SECRET_STORE) — see secrets_loader.py
from secrets_loader import apply_secret_store

apply_secret_store()

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Config
JWT_SECRET = os.environ.get('JWT_SECRET', 'default_secret')
JWT_ALGORITHM = os.environ.get('JWT_ALGORITHM', 'HS256')
JWT_EXPIRY_HOURS = int(os.environ.get('JWT_EXPIRY_HOURS', 24))

# Create the main app
app = FastAPI(title="AAI-HRMS API", version="1.0.0")
api_router = APIRouter(prefix="/api")
security = HTTPBearer()


class _RequestIdLoggingFilter(logging.Filter):
    """Injects request_id from context for structured / text logs (M0-3)."""

    def filter(self, record: logging.LogRecord) -> bool:
        record.request_id = request_id_ctx.get("-")
        return True


def _configure_logging() -> None:
    """
    Text logs by default; set LOG_FORMAT=json for one-line JSON (log shipping friendly).
    See memory/runbooks/structured-logging-standard.md.
    """
    log_format = os.environ.get("LOG_FORMAT", "text").strip().lower()
    service = os.environ.get("SERVICE_NAME", "aai-hrms-api").strip() or "aai-hrms-api"
    root = logging.getLogger()
    root.handlers.clear()
    handler = logging.StreamHandler()
    handler.addFilter(_RequestIdLoggingFilter())

    if log_format in ("json", "structured"):
        class JsonFormatter(logging.Formatter):
            def format(self, record: logging.LogRecord) -> str:
                payload: Dict[str, Any] = {
                    "ts": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
                    "service": service,
                    "request_id": getattr(record, "request_id", "-"),
                    "level": record.levelname,
                    "logger": record.name,
                    "msg": record.getMessage(),
                }
                if record.exc_info:
                    payload["exc"] = self.formatException(record.exc_info)
                return json.dumps(payload, ensure_ascii=False)

        handler.setFormatter(JsonFormatter())
    else:
        handler.setFormatter(
            logging.Formatter(
                f"%(asctime)s - {service} - %(name)s - %(levelname)s - [%(request_id)s] - %(message)s"
            )
        )

    root.addHandler(handler)
    root.setLevel(logging.INFO)


_configure_logging()
logger = logging.getLogger(__name__)


async def _m10_publish_safe(**kwargs: Any) -> None:
    """Best-effort M10 outbox publish — never raises to callers."""
    try:
        from m10_events.producer import publish_event

        await publish_event(db, **kwargs)
    except Exception:
        logger.exception("m10_events publish failed")


# Basic in-memory API metrics for M0 observability baseline (admin JSON: GET /api/metrics).
API_METRICS: Dict[str, Any] = {
    "started_at": datetime.now(timezone.utc).isoformat(),
    "total_requests": 0,
    "total_errors": 0,
    "by_path": defaultdict(lambda: {"count": 0, "errors": 0, "total_ms": 0.0}),
}


def _http_status_class(code: int) -> str:
    if 200 <= code < 300:
        return "2xx"
    if 300 <= code < 400:
        return "3xx"
    if 400 <= code < 500:
        return "4xx"
    if 500 <= code < 600:
        return "5xx"
    return "other"


# Prometheus metrics (low-cardinality labels; scrape GET /metrics)
HTTP_REQUESTS_TOTAL = Counter(
    "aai_http_requests_total",
    "Total HTTP requests",
    ["method", "status_class"],
)
HTTP_DURATION_SECONDS = Histogram(
    "aai_http_request_duration_seconds",
    "HTTP request duration in seconds",
    ["method"],
    buckets=(0.01, 0.025, 0.05, 0.1, 0.25, 0.5, 1.0, 2.5, 5.0, 10.0, 30.0),
)
HIRING_PACK_REQUESTS = Counter(
    "aai_hiring_pack_requests_total",
    "Smart Hiring Dashboard hiring-pack requests",
    ["cache_hit"],
)
HIRING_PACK_DURATION_SECONDS = Histogram(
    "aai_hiring_pack_duration_seconds",
    "hiring-pack aggregation duration in seconds",
    buckets=(0.05, 0.1, 0.25, 0.5, 0.8, 1.0, 2.0, 5.0, 10.0),
)


@app.middleware("http")
async def observability_middleware(request: Request, call_next):
    """
    Request correlation ID + in-memory admin metrics + Prometheus histogram/counter.
    Skips metric recording for GET /metrics (scrape endpoint) to avoid noise.
    """
    rid_header = os.environ.get("REQUEST_ID_HEADER", "X-Request-ID").strip() or "X-Request-ID"
    incoming = request.headers.get(rid_header)
    rid = incoming.strip() if incoming else str(uuid.uuid4())
    request.state.request_id = rid
    ctx_token = request_id_ctx.set(rid)
    path = request.url.path
    skip_metrics = path == "/metrics"
    start = time.perf_counter()
    status_code = 500
    try:
        response = await call_next(request)
        status_code = response.status_code
        if not skip_metrics:
            elapsed = time.perf_counter() - start
            method = request.method.upper()
            sc = _http_status_class(status_code)
            HTTP_REQUESTS_TOTAL.labels(method=method, status_class=sc).inc()
            HTTP_DURATION_SECONDS.labels(method=method).observe(elapsed)
            elapsed_ms = elapsed * 1000.0
            API_METRICS["total_requests"] += 1
            row = API_METRICS["by_path"][path]
            row["count"] += 1
            row["total_ms"] += elapsed_ms
            if status_code >= 500:
                API_METRICS["total_errors"] += 1
                row["errors"] += 1
        response.headers[rid_header] = rid
        return response
    finally:
        request_id_ctx.reset(ctx_token)


@app.get("/metrics", include_in_schema=False)
def prometheus_metrics(request: Request):
    """
    Prometheus text exposition (M0-3). Scrape this path from Prometheus.
    Set PROMETHEUS_SCRAPE_TOKEN and send Authorization: Bearer <token> to restrict access.
    """
    expected = os.environ.get("PROMETHEUS_SCRAPE_TOKEN", "").strip()
    if expected:
        auth = request.headers.get("Authorization", "")
        if auth != f"Bearer {expected}":
            raise HTTPException(status_code=401, detail="Invalid Prometheus scrape authorization")
    return Response(content=generate_latest(), media_type=CONTENT_TYPE_LATEST)

# ========================
# ATS / AGGREGATOR MODELS
# ========================

class CandidateMatch(BaseModel):
    candidate_id: str
    source: str
    score: float
    title_score: float
    skill_score: float
    must_have_ok: bool
    description_score: float

class JobMatchesResponse(BaseModel):
    job_id: str
    generated_at: str
    matches: List[CandidateMatch]

# Demo candidates
class DemoCandidatesRequest(BaseModel):
    count: int = 50

class DemoCandidatesResponse(BaseModel):
    job_id: str
    created: int

# ========================
# ADMIN / CONNECTOR CONFIG
# ========================

CONNECTOR_CONFIG_COLLECTION = "connector_configs"
CANDIDATE_DEDUP_AUDIT_COLLECTION = "candidate_dedup_audit"
LIFECYCLE_AUDIT_COLLECTION = "employee_lifecycle_audit_logs"
COMPLIANCE_DOCS_COLLECTION = "employee_compliance_documents"

class ConnectorConfigUpdate(BaseModel):
    enabled: Optional[bool] = None

    # Company DB candidates connector (external Mongo)
    mongo_url: Optional[str] = None
    db_name: Optional[str] = None
    collection_name: Optional[str] = None

    # External platform connectors (placeholders for official APIs)
    client_id: Optional[str] = None
    client_secret: Optional[str] = None
    base_url: Optional[str] = None
    scopes: Optional[Union[List[str], str]] = None
    # M1-1 OAuth productionization
    oauth_token_url: Optional[str] = None
    redirect_uri: Optional[str] = None
    refresh_token: Optional[str] = None
    access_token: Optional[str] = None
    token_expires_at: Optional[str] = None
    # M1-2 paging / resilience / throttling
    page_size: Optional[int] = None
    max_retries: Optional[int] = None
    min_interval_ms: Optional[int] = None
    search_path: Optional[Union[str, List[str]]] = None

class AdminUserRoleUpdate(BaseModel):
    # Roles supported by Phase-1 permissions.
    role: Literal["admin", "recruiter", "hr_admin", "hr_viewer"]


class CandidateMergeRequest(BaseModel):
    """Admin: merge two duplicate candidate profiles (M1-3 override control)."""

    keep_candidate_id: str
    merge_candidate_id: str


def _normalize_scopes(scopes: Optional[Union[List[str], str]]) -> Optional[List[str]]:
    if scopes is None:
        return None
    if isinstance(scopes, list):
        return [s for s in scopes if isinstance(s, str) and s.strip()]
    if isinstance(scopes, str):
        # Accept comma-separated or space-separated scopes
        parts = re.split(r"[,\n\r\t ]+", scopes)
        return [p for p in parts if p.strip()]
    return None

async def get_connector_config(name: str) -> Dict[str, Any]:
    return await db[CONNECTOR_CONFIG_COLLECTION].find_one({"name": name}, {"_id": 0}) or {}


async def _append_lifecycle_audit(
    *,
    employee_code: str,
    action: str,
    from_status: Optional[str],
    to_status: Optional[str],
    event_type: Optional[str],
    event_id: Optional[str],
    actor_id: str,
    notes: Optional[str] = None,
) -> None:
    doc = {
        "id": str(uuid.uuid4()),
        "employee_code": employee_code,
        "action": action,
        "from_status": from_status,
        "to_status": to_status,
        "event_type": event_type,
        "event_id": event_id,
        "actor_id": actor_id,
        "notes": notes,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    try:
        await db[LIFECYCLE_AUDIT_COLLECTION].insert_one(doc)
    except Exception as e:
        logger.warning("lifecycle audit insert failed: %s", e)


async def _hr_escalation_recipient_ids() -> List[str]:
    rows = await db.users.find({"role": {"$in": ["admin", "hr_admin"]}}, {"_id": 0, "id": 1}).to_list(500)
    return [r["id"] for r in rows if r.get("id")]


def _require_admin(current_user: dict):
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin only")
    return current_user


def _require_allocation_approver(current_user: dict):
    """M4-3: scenario approval — admin or hr_admin."""
    role = (current_user.get("role") or "").lower()
    if role not in ("admin", "hr_admin"):
        raise HTTPException(status_code=403, detail="Approver role required (admin or hr_admin)")
    return current_user


def _require_engagement_raw_privileged(current_user: dict):
    """M6-3: raw survey responses (PII) — admin or hr_admin only."""
    role = (current_user.get("role") or "").lower()
    if role not in ("admin", "hr_admin"):
        raise HTTPException(
            status_code=403,
            detail="Raw engagement responses require admin or hr_admin (privacy policy).",
        )
    return current_user


PHASE1_PERMISSIONS = {
    # Existing roles remain compatible.
    "admin": {
        "employees_read",
        "employees_write",
        "skills_read",
        "skills_write",
        "kpi_read",
        "lifecycle_write",
        "engagement_read",
        "engagement_write",
        "engagement_analytics",
        "engagement_executive",
        "engagement_ai",
        "engagement_privacy_raw",
    },
    # Recruiters can manage lifecycle/approvals without full employee master write.
    "recruiter": {
        "employees_read",
        "skills_read",
        "kpi_read",
        "lifecycle_write",
        "engagement_read",
    },
    # New optional granular roles
    "hr_admin": {
        "employees_read",
        "employees_write",
        "skills_read",
        "skills_write",
        "kpi_read",
        "lifecycle_write",
        "engagement_read",
        "engagement_write",
        "engagement_analytics",
        "engagement_executive",
        "engagement_ai",
        "engagement_privacy_raw",
    },
    "hr_viewer": {
        "employees_read",
        "skills_read",
        "kpi_read",
        "engagement_read",
        "engagement_analytics",
    },
}

def _require_phase1_access(current_user: dict, permission: str):
    role = (current_user.get("role") or "").lower()
    allowed = PHASE1_PERMISSIONS.get(role, set())
    if permission not in allowed:
        raise HTTPException(status_code=403, detail=f"Missing permission: {permission}")
    return current_user


def _user_has_phase1_permission(current_user: dict, permission: str) -> bool:
    role = (current_user.get("role") or "").lower()
    return permission in PHASE1_PERMISSIONS.get(role, set())


def _require_engagement_executive(current_user: dict):
    """CHRO / leadership views — admin, hr_admin only by default."""
    _require_phase1_access(current_user, "engagement_executive")
    return current_user


def _require_engagement_ai(current_user: dict):
    """AI engagement copilot and WFI-sourced AI recs — admin, hr_admin."""
    _require_phase1_access(current_user, "engagement_ai")
    return current_user


# ========================
# PYDANTIC MODELS
# ========================

# Auth Models
class UserCreate(BaseModel):
    # Use string instead of EmailStr so local/dev domains (e.g. *.local) are allowed.
    email: str
    password: str
    full_name: str
    role: str = "recruiter"

class UserLogin(BaseModel):
    # Use string instead of EmailStr so local/dev domains (e.g. *.local) are allowed.
    email: str
    password: str

class UserResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    email: str
    full_name: str
    role: str
    created_at: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse


def _normalize_auth_email(email: str) -> str:
    return str(email or "").strip().lower()


def _user_email_match_query(email: str) -> Dict[str, Any]:
    """Case-insensitive exact match on email (handles legacy rows with mixed casing)."""
    em = _normalize_auth_email(email)
    if not em:
        return {"email": "__no_match__"}
    return {"email": {"$regex": f"^{re.escape(em)}$", "$options": "i"}}


def _user_doc_to_response(user: Dict[str, Any]) -> UserResponse:
    """Safe UserResponse from Mongo user doc (datetime created_at, missing full_name, etc.)."""
    uid = user.get("id")
    if not uid:
        raise HTTPException(status_code=500, detail="Invalid user record")
    ca = user.get("created_at")
    if isinstance(ca, datetime):
        ca = ca.isoformat()
    elif ca is not None and not isinstance(ca, str):
        ca = str(ca)
    if not ca:
        ca = datetime.now(timezone.utc).isoformat()
    email = str(user.get("email") or "").strip()
    fn = user.get("full_name")
    full_name = (str(fn).strip() if fn is not None else "") or (email.split("@")[0] if "@" in email else "") or "User"
    return UserResponse(
        id=str(uid),
        email=email,
        full_name=full_name,
        role=str(user.get("role") or "recruiter"),
        created_at=ca,
    )


# Job Models
class SkillInput(BaseModel):
    skill_name: str
    skill_type: str = "GOOD_TO_HAVE"  # MUST_HAVE or GOOD_TO_HAVE
    weight: float = 1.0

class JobCreate(BaseModel):
    title: str
    description: str
    skills_needed: List[str]
    must_have_skills: List[str]
    location: Optional[str] = None
    work_mode: Optional[str] = "hybrid"  # remote, hybrid, onsite
    seniority: Optional[str] = None
    business_pillar: Optional[str] = None
    business_department: Optional[str] = None
    business_sub_department: Optional[str] = None
    project_id: Optional[str] = None

class JobResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    title: str
    normalized_title: Optional[str] = None
    description: str
    seniority: Optional[str] = None
    domain: Optional[str] = None
    business_pillar: Optional[str] = None
    business_department: Optional[str] = None
    business_sub_department: Optional[str] = None
    project_id: Optional[str] = None
    location: Optional[str] = None
    work_mode: Optional[str] = None
    status: str
    skills: List[Dict[str, Any]] = []
    activities: List[Dict[str, Any]] = []
    scoring_rubric: Optional[Dict[str, Any]] = None
    pin_rank: Optional[int] = None  # higher = listed first (e.g. Excel JD imports)
    seed_marker: Optional[str] = None
    import_source_file: Optional[str] = None
    import_stable_id: Optional[str] = None
    created_by: str
    created_at: str
    candidate_count: int = 0

class JobUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    location: Optional[str] = None
    work_mode: Optional[str] = None
    seniority: Optional[str] = None
    business_pillar: Optional[str] = None
    business_department: Optional[str] = None
    business_sub_department: Optional[str] = None
    project_id: Optional[str] = None
    skills: Optional[List[Dict[str, Any]]] = None

# Candidate Models
class ExperienceInput(BaseModel):
    company: str
    title: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    description: Optional[str] = None

class CandidateCreate(BaseModel):
    full_name: str
    email: Optional[EmailStr] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    headline: Optional[str] = None
    total_experience_years: Optional[float] = None
    skills: List[str] = []
    source: str = "DIRECT_UPLOAD"
    experience: List[ExperienceInput] = []
    resume_text: Optional[str] = None

class CandidateResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    headline: Optional[str] = None
    total_experience_years: Optional[float] = None
    skills: List[Dict[str, Any]] = []
    source: str
    experience: List[Dict[str, Any]] = []
    pin_rank: Optional[int] = None  # higher = listed first (e.g. Excel imports)
    seed_marker: Optional[str] = None
    import_source_file: Optional[str] = None
    created_at: str

class CandidateUpdate(BaseModel):
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    headline: Optional[str] = None
    total_experience_years: Optional[float] = None
    # Accept list[str] or list[{skill_name, proficiency}]
    skills: Optional[List[Any]] = None

class CandidatesPagedResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    items: List[CandidateResponse]
    total: int
    page: int
    page_size: int
    total_pages: int

# Application/Pipeline Models
class ApplicationCreate(BaseModel):
    job_id: str
    candidate_id: str
    stage: str = "SOURCED"

class ApplicationUpdate(BaseModel):
    stage: str
    reason: Optional[str] = None

class OfferStatusUpdate(BaseModel):
    offer_status: Literal["SENT", "NEGOTIATION", "ACCEPTED", "DECLINED"]

class ApplicationStageHistoryItem(BaseModel):
    from_stage: str | None = None
    to_stage: str
    changed_at: str
    days_in_stage: float | None = None
    reason: str | None = None
    offer_status: str | None = None

class ApplicationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    job_id: str
    candidate_id: str
    stage: str
    status: str
    fit_score: Optional[Dict[str, Any]] = None
    candidate: Optional[Dict[str, Any]] = None
    job: Optional[Dict[str, Any]] = None
    created_at: str
    updated_at: str

# Referral Models
class ReferralCreate(BaseModel):
    job_id: str
    candidate_name: str
    candidate_email: Optional[EmailStr] = None
    candidate_phone: Optional[str] = None
    resume_text: Optional[str] = None
    note: Optional[str] = None

class ReferralResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    job_id: str
    candidate_id: str
    referred_by: str
    note: Optional[str] = None
    status: str
    candidate: Optional[Dict[str, Any]] = None
    created_at: str
    fit_score: Optional[Dict[str, Any]] = None

# Assessment Models
class AssessmentCreate(BaseModel):
    assessment_type: str = "CORE_SKILL"  # SCREENING, CORE_SKILL, WORK_SIMULATION, BEHAVIORAL
    title: str
    duration_minutes: int = 60

class QuestionInput(BaseModel):
    question_type: str  # MCQ, SHORT_ANSWER, CODING, SQL, CASE_STUDY
    question_text: str
    options: Optional[List[str]] = None
    answer_key: Optional[str] = None
    max_marks: int = 10
    difficulty: str = "MEDIUM"

class AssessmentResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    job_id: str
    assessment_type: str
    title: str
    duration_minutes: int
    total_marks: int
    questions: List[Dict[str, Any]] = []
    rubric: Optional[Dict[str, Any]] = None
    created_at: str

# Fit Score Models
class FitScoreResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    job_id: str
    candidate_id: str
    title_score: float
    skill_match_pct: float
    activity_match_pct: float
    experience_score: float
    final_score: float
    must_have_ok: bool
    explanation: Optional[Dict[str, Any]] = None
    computed_at: str

# Dashboard Stats
class DashboardStats(BaseModel):
    total_jobs: int
    open_jobs: int
    total_candidates: int
    total_applications: int
    applications_by_stage: Dict[str, int]
    recent_activities: List[Dict[str, Any]]

# Phase-1: Employee Master + Workforce + Executive KPIs
class EmployeeCreate(BaseModel):
    employee_code: str
    full_name: str
    email: Optional[EmailStr] = None
    department: str
    role_title: str
    manager_id: Optional[str] = None
    location: Optional[str] = None
    status: Literal["ONBOARDING", "ACTIVE", "INACTIVE", "EXITED"] = "ACTIVE"
    skills: List[str] = []
    join_date: Optional[str] = None
    # M8 retention segmentation / compensation proxy (optional)
    compensation_band: Optional[Literal["LOW", "MID", "HIGH", "LEAD"]] = None
    last_promotion_at: Optional[str] = None
    high_performer: Optional[bool] = None
    critical_role: Optional[bool] = None
    # HRIS / external comp positioning (optional; feeds M8 compensation feature)
    comp_market_percentile: Optional[float] = Field(default=None, ge=0.0, le=100.0)
    hris_last_sync_at: Optional[str] = None
    hris_comp_source: Optional[str] = Field(default=None, max_length=120)

class EmployeeUpdate(BaseModel):
    full_name: Optional[str] = None
    email: Optional[EmailStr] = None
    department: Optional[str] = None
    role_title: Optional[str] = None
    manager_id: Optional[str] = None
    location: Optional[str] = None
    status: Optional[Literal["ONBOARDING", "ACTIVE", "INACTIVE", "EXITED"]] = None
    skills: Optional[List[str]] = None
    join_date: Optional[str] = None
    compensation_band: Optional[Literal["LOW", "MID", "HIGH", "LEAD"]] = None
    last_promotion_at: Optional[str] = None
    high_performer: Optional[bool] = None
    critical_role: Optional[bool] = None
    comp_market_percentile: Optional[float] = None
    hris_last_sync_at: Optional[str] = None
    hris_comp_source: Optional[str] = None

class EmployeeResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    employee_code: str
    full_name: str
    email: Optional[str] = None
    department: str
    role_title: str
    manager_id: Optional[str] = None
    location: Optional[str] = None
    status: str
    skills: List[str] = []
    join_date: Optional[str] = None
    compensation_band: Optional[str] = None
    last_promotion_at: Optional[str] = None
    high_performer: Optional[bool] = None
    critical_role: Optional[bool] = None
    comp_market_percentile: Optional[float] = None
    hris_last_sync_at: Optional[str] = None
    hris_comp_source: Optional[str] = None
    created_at: str
    updated_at: Optional[str] = None


def _employee_doc_to_response(doc: Dict[str, Any]) -> EmployeeResponse:
    """
    Build EmployeeResponse from a Mongo document. Coerces legacy / CSV-import shapes so
    list/get endpoints don't return 500 (e.g. skills=null, numeric timestamps).
    """
    d = dict(doc)
    sk = d.get("skills")
    if sk is None:
        d["skills"] = []
    elif isinstance(sk, list):
        d["skills"] = [str(x).strip() for x in sk if x is not None and str(x).strip()]
    elif isinstance(sk, str) and sk.strip():
        d["skills"] = [s.strip() for s in re.split(r"[,;|]", sk) if s.strip()]
    else:
        d["skills"] = []

    cmp_raw = d.get("comp_market_percentile")
    if cmp_raw is not None and not isinstance(cmp_raw, (int, float)):
        try:
            d["comp_market_percentile"] = float(str(cmp_raw).strip())
        except (TypeError, ValueError):
            d["comp_market_percentile"] = None

    created = d.get("created_at")
    if not created:
        d["created_at"] = d.get("updated_at") or datetime.now(timezone.utc).isoformat()
    elif not isinstance(created, str):
        d["created_at"] = str(created)

    updated = d.get("updated_at")
    if updated is not None and not isinstance(updated, str):
        d["updated_at"] = str(updated)

    for fk in ("join_date", "last_promotion_at", "hris_last_sync_at"):
        v = d.get(fk)
        if v is not None and not isinstance(v, str):
            d[fk] = str(v)

    return EmployeeResponse(**d)


class SkillInventoryCreate(BaseModel):
    skill_name: str
    demand_count: int = 0
    supply_count: int = 0
    category: Optional[str] = None
    priority: Literal["HIGH", "MEDIUM", "LOW"] = "MEDIUM"
    notes: Optional[str] = None

class SkillInventoryUpdate(BaseModel):
    demand_count: Optional[int] = None
    supply_count: Optional[int] = None
    category: Optional[str] = None
    priority: Optional[Literal["HIGH", "MEDIUM", "LOW"]] = None
    notes: Optional[str] = None

class SkillInventoryResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    skill_name: str
    demand_count: int
    supply_count: int
    gap: int
    category: Optional[str] = None
    priority: str
    notes: Optional[str] = None
    updated_at: str

class ExecutiveKpiResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")

    employee_count: int
    active_employee_count: int
    attrition_count: int
    attrition_rate_pct: float
    avg_skills_per_employee: float
    top_skill_gaps: List[Dict[str, Any]]
    hiring_demand_total: int
    workforce_supply_total: int
    skill_coverage_pct: float
    talent_acquisition: Dict[str, Any] = Field(default_factory=dict)


# ========================
# Phase-2: Employee Lifecycle Management
# ========================

EmployeeLifecycleEventType = Literal[
    "ONBOARDED",
    "ACTIVATED",
    "ROLE_CHANGED",
    "DOCUMENT_ADDED",
    "EXITED",
]


class EmployeeLifecycleEventCreate(BaseModel):
    employee_code: str
    event_type: EmployeeLifecycleEventType
    # Optional date provided by HR; defaults to `created_at` on the server.
    effective_date: Optional[str] = None
    # Free-form details (e.g., manager change, document name, notes, metadata).
    details: Optional[Dict[str, Any]] = None


class EmployeeLifecycleEventUpdate(BaseModel):
    event_type: Optional[EmployeeLifecycleEventType] = None
    effective_date: Optional[str] = None
    details: Optional[Dict[str, Any]] = None


LifecycleProcessingStatus = Literal["PENDING", "PROCESSED", "FAILED", "REJECTED"]


class EmployeeLifecycleEventResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    employee_code: str
    event_type: str
    effective_date: Optional[str] = None
    details: Dict[str, Any] = {}
    created_at: str
    updated_at: Optional[str] = None
    processing_status: LifecycleProcessingStatus = "PENDING"
    attempts: int = 0
    processed_at: Optional[str] = None
    processing_error: Optional[str] = None
    requires_approval: bool = False
    approval_status: Optional[Literal["PENDING", "APPROVED", "REJECTED"]] = None
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None
    rejection_reason: Optional[str] = None
    escalated_at: Optional[str] = None


class EmployeeLifecycleEventsPagedResponse(BaseModel):
    items: List[EmployeeLifecycleEventResponse]
    page: int
    page_size: int
    total: int
    total_pages: int
    sort_by: str
    sort_dir: Literal["asc", "desc"]


class EmployeeLifecycleDashboardResponse(BaseModel):
    total_events: int
    last_30_days_events: int
    counts_by_event_type: Dict[str, int]
    onboarded_total: int
    activated_total: int
    role_changed_total: int
    document_added_total: int
    exited_total: int


class LifecycleEventRejectRequest(BaseModel):
    reason: Optional[str] = None


class ComplianceDocumentCreate(BaseModel):
    employee_code: str
    document_type: str
    title: str
    storage_uri: Optional[str] = None
    expires_at: Optional[str] = None
    content_base64: Optional[str] = None


class ComplianceDocumentVerifyRequest(BaseModel):
    notes: Optional[str] = None


class ComplianceDocumentResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    employee_code: str
    document_type: str
    title: str
    storage_uri: Optional[str] = None
    status: str
    uploaded_by: str
    uploaded_at: str
    verified_at: Optional[str] = None
    verified_by: Optional[str] = None
    expires_at: Optional[str] = None
    sla_due_at: Optional[str] = None
    reminder_sent_at: Optional[str] = None
    sla_breached_at: Optional[str] = None
    updated_at: Optional[str] = None


# ========================
# Phase-3: Workforce Intelligence (Demand-Supply)
# ========================

class WorkforceIntelligenceSkillForecast(BaseModel):
    model_config = ConfigDict(extra="ignore")
    skill_name: str
    priority: str
    demand_current: int
    supply_count: int
    demand_forecast: int
    forecast_gap: int
    forecast_gap_pct: float


class WorkforceIntelligenceResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    horizon_months: int
    generated_at: str
    skills_total: int
    demand_current_total: int
    workforce_supply_total: int
    demand_forecast_total: int
    forecast_gap_total: int
    top_forecast_gaps: List[WorkforceIntelligenceSkillForecast]


class WorkforceIntelModelForecastRow(BaseModel):
    model_config = ConfigDict(extra="ignore")
    skill_name: str
    skill_name_lc: str
    priority: str
    demand_current: int
    supply_count: int
    demand_forecast_model: int
    forecast_gap: int
    forecast_gap_pct: float
    horizon_steps: int


class WorkforceIntelModelForecastResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    version_id: str
    horizon_months: int
    generated_at: str
    demand_source: str
    supply_source: str
    skills_total: int
    demand_forecast_total: int
    forecast_gap_total: int
    top_forecast_gaps: List[WorkforceIntelModelForecastRow]


class M3EtlBackfillRequest(BaseModel):
    days: int = 30
    seed: int = 42


class M3TrainModelRequest(BaseModel):
    activate: bool = False
    max_snapshots: int = 200


# Phase-3 M4: Resource vs Project Optimization (MVP metrics)
class ResourceOptimizationSkillMetrics(BaseModel):
    model_config = ConfigDict(extra="ignore")
    skill_name: str
    priority: str
    demand_count: int
    supply_count: int
    utilization_pct: float
    allocation_status: Literal["OVER_ALLOCATED", "UNDER_ALLOCATED", "BALANCED"]
    bench_count: int
    shortage_count: int


class ResourceOptimizationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    generated_at: str
    skills_total: int
    total_bench: int
    total_shortage: int
    total_demand: int
    total_supply: int
    under_allocated: List[ResourceOptimizationSkillMetrics]
    over_allocated: List[ResourceOptimizationSkillMetrics]


class AllocationOptimizationSettingsResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    max_projects_per_employee: int
    max_seats_per_employee_per_project: int
    shortage_penalty_hard: float
    shortage_penalty_soft: float
    utilization_weight: float
    target_utilization_pct: float


class AllocationOptimizationSettingsUpdate(BaseModel):
    max_projects_per_employee: Optional[int] = None
    max_seats_per_employee_per_project: Optional[int] = None
    shortage_penalty_hard: Optional[float] = None
    shortage_penalty_soft: Optional[float] = None
    utilization_weight: Optional[float] = None
    target_utilization_pct: Optional[float] = None


class AllocationSimulateRequest(BaseModel):
    demand_overrides: List[Dict[str, Any]] = Field(default_factory=list)
    constraint_overrides: Dict[str, Any] = Field(default_factory=dict)


class AllocationScenarioCreate(BaseModel):
    """If result is omitted, server re-runs simulate using overrides."""

    name: str
    description: Optional[str] = None
    demand_overrides: List[Dict[str, Any]] = Field(default_factory=list)
    constraint_overrides: Dict[str, Any] = Field(default_factory=dict)
    result: Optional[Dict[str, Any]] = None


class AllocationScenarioRejectRequest(BaseModel):
    reason: Optional[str] = None


class AllocationScenarioApplyRequest(BaseModel):
    dry_run: bool = False


# ========================
# Phase-3 M4 extension: Project skill demands (project-resource mapping MVP)
# ========================

class ProjectRequiredSkill(BaseModel):
    skill_name: str
    level: Optional[str] = None
    required_count: int = 0

class ProjectMilestone(BaseModel):
    name: str
    due_date: Optional[str] = None
    status: Optional[str] = "PENDING"

class ProjectDocument(BaseModel):
    doc_name: str
    url: str

class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = None
    client_name: Optional[str] = None
    business_unit: Optional[str] = None
    project_type: Optional[Literal["INTERNAL", "EXTERNAL"]] = None
    status: Optional[str] = "ACTIVE"  # ACTIVE, PAUSED, CLOSED
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    project_manager_id: Optional[str] = None
    budget: Optional[float] = None
    billing_type: Optional[Literal["FIXED", "TIME_MATERIAL"]] = None
    required_skills: List[ProjectRequiredSkill] = Field(default_factory=list)
    milestones: List[ProjectMilestone] = Field(default_factory=list)
    documents: List[ProjectDocument] = Field(default_factory=list)

class ProjectUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    client_name: Optional[str] = None
    business_unit: Optional[str] = None
    project_type: Optional[Literal["INTERNAL", "EXTERNAL"]] = None
    status: Optional[str] = None  # ACTIVE, PAUSED, CLOSED
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    project_manager_id: Optional[str] = None
    budget: Optional[float] = None
    billing_type: Optional[Literal["FIXED", "TIME_MATERIAL"]] = None
    required_skills: Optional[List[ProjectRequiredSkill]] = None
    milestones: Optional[List[ProjectMilestone]] = None
    documents: Optional[List[ProjectDocument]] = None

class ProjectResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: Optional[str] = None
    client_name: Optional[str] = None
    business_unit: Optional[str] = None
    project_type: Optional[str] = None
    status: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    project_manager_id: Optional[str] = None
    budget: Optional[float] = None
    billing_type: Optional[str] = None
    required_skills: List[Dict[str, Any]] = []
    milestones: List[Dict[str, Any]] = []
    documents: List[Dict[str, Any]] = []
    created_at: str
    updated_at: Optional[str] = None

class ProjectSkillDemandRow(BaseModel):
    skill_name: str
    demand_count: int = 0
    priority: Optional[str] = None  # optional override; if omitted uses workforce_skills priority/default
    # M4-1: capacity constraint bounds (default to demand_count when omitted on write paths)
    demand_min: Optional[int] = None
    demand_max: Optional[int] = None
    constraint_type: Optional[Literal["HARD", "SOFT"]] = None  # HARD = must-fill min; SOFT = stretch max with soft penalty

class ProjectSkillDemandBulkRequest(BaseModel):
    rows: List[ProjectSkillDemandRow]
    mode: Literal["skip", "upsert"] = "upsert"

class ProjectSkillDemandResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    project_id: str
    skill_name: str
    demand_count: int
    priority: str
    updated_at: str
    demand_min: int = 0
    demand_max: int = 0
    constraint_type: str = "HARD"


# ========================
# Phase-3 M6: Project vs Resource allocation (individual allocations)
# ========================

class AllocationCreate(BaseModel):
    project_id: str
    employee_id: str
    role: Optional[str] = None
    allocation_percentage: int = 100  # 0..100
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    billable: bool = True
    allocation_type: Optional[Literal["FULL", "PARTIAL", "SHADOW"]] = None
    status: Optional[Literal["ACTIVE", "PENDING", "CLOSED"]] = "PENDING"
    cost_rate: Optional[float] = None
    billing_rate: Optional[float] = None

class AllocationUpdate(BaseModel):
    role: Optional[str] = None
    allocation_percentage: Optional[int] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    billable: Optional[bool] = None
    allocation_type: Optional[Literal["FULL", "PARTIAL", "SHADOW"]] = None
    status: Optional[Literal["ACTIVE", "PENDING", "CLOSED"]] = None
    approval_status: Optional[Literal["PENDING", "APPROVED", "REJECTED"]] = None
    cost_rate: Optional[float] = None
    billing_rate: Optional[float] = None

class AllocationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    project_id: str
    employee_id: str
    role: Optional[str] = None
    allocation_percentage: int
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    billable: bool = True
    allocation_type: Optional[str] = None
    status: str
    approval_status: str
    cost_rate: Optional[float] = None
    billing_rate: Optional[float] = None
    created_at: str
    updated_at: Optional[str] = None
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None

class AllocationApproveRequest(BaseModel):
    action: Literal["approve", "reject"]
    reason: Optional[str] = None


# ========================
# Project Section (Enterprise Project Management) — Phase-3 M7
# ========================

PROJECT_LIFECYCLE_STATES = [
    "draft",
    "proposed",
    "under_review",
    "approved",
    "active",
    "on_hold",
    "completed",
    "closed",
    "cancelled",
]

PROJECT_TYPES = ["internal", "external", "r&d", "support"]
PROJECT_PRIORITIES = ["low", "medium", "high", "critical"]
PROJECT_HEALTH = ["green", "amber", "red"]
APPROVAL_TYPES = ["project_approval", "budget_approval", "lifecycle_transition", "change_request", "closure_approval"]


def _require_project_view(current_user: dict):
    return _require_phase1_access(current_user, "kpi_read")


def _require_project_edit(current_user: dict):
    return _require_phase1_access(current_user, "skills_write")


def _require_project_approve(current_user: dict):
    # Same approver gate used across other approval flows
    role = (current_user.get("role") or "").lower()
    if role not in ("admin", "hr_admin"):
        raise HTTPException(status_code=403, detail="Approver role required (admin or hr_admin)")
    return current_user


class ProjectMasterCreate(BaseModel):
    project_name: str
    project_code: Optional[str] = None  # if omitted auto-generated
    client_name: Optional[str] = None
    project_type: Optional[str] = "external"
    business_unit: Optional[str] = None
    department: Optional[str] = None
    cost_center: Optional[str] = None
    project_manager_id: Optional[str] = None
    delivery_manager_id: Optional[str] = None
    account_manager_id: Optional[str] = None
    project_owner_id: Optional[str] = None
    project_priority: Optional[str] = "medium"
    project_category: Optional[str] = None
    project_status: Optional[str] = "draft"
    project_health: Optional[str] = "green"
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    actual_end_date: Optional[str] = None
    billing_type: Optional[str] = None
    currency: Optional[str] = "INR"
    project_budget: Optional[float] = None
    expected_revenue: Optional[float] = None
    location: Optional[str] = None
    geography: Optional[str] = None
    work_mode: Optional[str] = None
    description: Optional[str] = None
    objectives: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    remarks: Optional[str] = None


class ProjectMasterUpdate(BaseModel):
    project_name: Optional[str] = None
    project_code: Optional[str] = None
    client_name: Optional[str] = None
    project_type: Optional[str] = None
    business_unit: Optional[str] = None
    department: Optional[str] = None
    cost_center: Optional[str] = None
    project_manager_id: Optional[str] = None
    delivery_manager_id: Optional[str] = None
    account_manager_id: Optional[str] = None
    project_owner_id: Optional[str] = None
    project_priority: Optional[str] = None
    project_category: Optional[str] = None
    project_status: Optional[str] = None
    project_health: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    actual_end_date: Optional[str] = None
    billing_type: Optional[str] = None
    currency: Optional[str] = None
    project_budget: Optional[float] = None
    expected_revenue: Optional[float] = None
    location: Optional[str] = None
    geography: Optional[str] = None
    work_mode: Optional[str] = None
    description: Optional[str] = None
    objectives: Optional[str] = None
    tags: Optional[List[str]] = None
    remarks: Optional[str] = None
    is_archived: Optional[bool] = None


class ProjectMasterResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    project_id: str
    project_name: str
    project_code: str
    client_name: Optional[str] = None
    project_type: str
    business_unit: Optional[str] = None
    department: Optional[str] = None
    cost_center: Optional[str] = None
    project_manager_id: Optional[str] = None
    delivery_manager_id: Optional[str] = None
    account_manager_id: Optional[str] = None
    project_owner_id: Optional[str] = None
    project_priority: str
    project_category: Optional[str] = None
    project_status: str
    project_health: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    planned_duration_days: Optional[int] = None
    actual_end_date: Optional[str] = None
    billing_type: Optional[str] = None
    currency: str = "INR"
    project_budget: Optional[float] = None
    expected_revenue: Optional[float] = None
    location: Optional[str] = None
    geography: Optional[str] = None
    work_mode: Optional[str] = None
    description: Optional[str] = None
    objectives: Optional[str] = None
    tags: List[str] = []
    remarks: Optional[str] = None
    is_archived: bool = False
    created_at: str
    updated_at: Optional[str] = None
    created_by: Optional[str] = None
    updated_by: Optional[str] = None


class ProjectLifecycleTransitionRequest(BaseModel):
    to_state: str
    reason: Optional[str] = None


class ProjectApprovalActionRequest(BaseModel):
    action: Literal["approve", "reject"]
    reason: Optional[str] = None


# Phase-3 M5: Employee Training & Skill Development (MVP: Recommendations)
class TrainingPathStep(BaseModel):
    step_title: str
    description: str


class TrainingSkillRecommendation(BaseModel):
    skill_name: str
    priority: str
    reason: str
    path_steps: List[TrainingPathStep]


class EmployeeTrainingRecommendation(BaseModel):
    employee_code: str
    full_name: str
    recommended_skills: List[TrainingSkillRecommendation]


class TrainingRecommendationsResponse(BaseModel):
    generated_at: str
    page: int
    page_size: int
    total: int
    total_pages: int
    max_skills_per_employee: int
    recommendations: List[EmployeeTrainingRecommendation]


class LearningPathStepTemplate(BaseModel):
    step_title: str
    description: str = ""


class LearningPathTemplateUpsert(BaseModel):
    skill_name: str
    steps: List[LearningPathStepTemplate]


class TrainingAssignmentCreate(BaseModel):
    employee_code: str
    skill_name: str


class TrainingProgressUpdate(BaseModel):
    progress_pct: Optional[float] = None
    status: Optional[Literal["ASSIGNED", "IN_PROGRESS", "COMPLETED"]] = None


class TrainingCertificationCreate(BaseModel):
    employee_code: str
    title: str
    issued_at: str
    expires_at: Optional[str] = None


class LmsSyncRequest(BaseModel):
    provider: Optional[str] = None


# ========================
# Phase-4 M6: Employee Satisfaction & Engagement (Pulse MVP)
# ========================

class PulseSurveyCreate(BaseModel):
    title: str
    question: str
    rating_min: int = 1
    rating_max: int = 5
    active: bool = True
    target_all: bool = True
    target_departments: List[str] = Field(default_factory=list)


class PulseSurveyResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    title: str
    question: str
    rating_min: int
    rating_max: int
    active: bool
    created_at: str
    target_all: bool = True
    target_departments: List[str] = Field(default_factory=list)


class PulseSurveyResponseCreate(BaseModel):
    survey_id: str
    employee_code: str
    rating: int
    response_text: Optional[str] = None


class PulseSurveyResponseResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    survey_id: str
    employee_code: str
    rating: int
    response_text: Optional[str] = None
    sentiment_label: str
    sentiment_score: float
    topic_primary: Optional[str] = None
    sentiment_pipeline_version: Optional[str] = None
    created_at: str
    updated_at: Optional[str] = None


class PulseSurveyDashboardResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    total_responses: int
    avg_rating: float
    last_30_days_responses: int
    sentiment_counts: Dict[str, int]
    topic_counts: Dict[str, int] = Field(default_factory=dict)
    weekly_trend: List[Dict[str, Any]] = Field(default_factory=list)
    display_confidence: str = "MEDIUM"
    confidence_rationale: str = ""
    anonymity_note: Optional[str] = None


class EngagementSurveyTemplateCreate(BaseModel):
    name: str
    description: Optional[str] = None
    default_title: str
    default_question: str
    rating_min: int = 1
    rating_max: int = 5
    target_all: bool = True
    target_departments: List[str] = Field(default_factory=list)


class EngagementSurveyTemplateUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    default_title: Optional[str] = None
    default_question: Optional[str] = None
    rating_min: Optional[int] = None
    rating_max: Optional[int] = None
    target_all: Optional[bool] = None
    target_departments: Optional[List[str]] = None


class EngagementSurveyFromTemplateCreate(BaseModel):
    template_id: str
    title_override: Optional[str] = None
    question_override: Optional[str] = None


class EngagementScheduleCreate(BaseModel):
    template_id: str
    cadence: Literal["WEEKLY", "MONTHLY", "QUARTERLY"] = "MONTHLY"
    enabled: bool = True
    next_run_at: Optional[str] = None


class PulseSurveyResponsesPagedResponse(BaseModel):
    items: List[PulseSurveyResponseResponse]
    page: int
    page_size: int
    total: int
    total_pages: int
    sort_by: str
    sort_dir: Literal["asc", "desc"]


# Phase-4 M8: High-Skill Talent Retention (MVP)
class RetentionCriticalSkill(BaseModel):
    skill_name: str
    priority: str
    demand_count: int
    supply_count: int
    shortage_count: int
    shortage_ratio: float
    risk_score: float


class RetentionRiskEmployeeSkill(BaseModel):
    skill_name: str
    risk_score: float


class RetentionRiskEmployee(BaseModel):
    employee_code: str
    full_name: str
    critical_skills_matched: int
    risk_score: float
    risk_label: str
    skills: List[RetentionRiskEmployeeSkill]


class AttritionV1FactorBrief(BaseModel):
    model_config = ConfigDict(extra="ignore")
    feature: str
    contribution: float
    direction: str


class AttritionV1EmployeeBrief(BaseModel):
    model_config = ConfigDict(extra="ignore")
    employee_id: str
    employee_code: str
    full_name: str
    attrition_risk: float
    confidence: float
    risk_band: str
    segments: List[str]
    top_factors: List[AttritionV1FactorBrief]


class AttritionV1DashboardSummary(BaseModel):
    model_config = ConfigDict(extra="ignore")
    model_version: str
    last_computed_at: Optional[str] = None
    scored_employee_count: int
    avg_attrition_risk: float
    top_at_risk: List[AttritionV1EmployeeBrief]


class RetentionDashboardResponse(BaseModel):
    generated_at: str
    critical_skills: List[RetentionCriticalSkill]
    total_high_skill_employees: int
    avg_risk_score: float
    top_risk_employees: List[RetentionRiskEmployee]
    # M8-1: model v1 summary (empty until POST .../score-run)
    attrition_v1: Optional[AttritionV1DashboardSummary] = None


class M8AttritionTrainLabel(BaseModel):
    employee_id: str
    churned: bool


class M8AttritionTrainRequest(BaseModel):
    labels: List[M8AttritionTrainLabel] = Field(default_factory=list)
    use_gradient_boosting: bool = False
    interaction_features: bool = False


class M8ModelRuntimePatch(BaseModel):
    ensemble_mode: Optional[Literal["linear", "gb", "avg"]] = None
    interaction_features_enabled: Optional[bool] = None


class M8RetentionPlaybookCreate(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    description: str = ""
    category: str = "GENERAL"
    suggested_duration_days: int = Field(default=30, ge=1, le=365)


class M8RetentionInterventionCreate(BaseModel):
    employee_id: str
    playbook_id: str
    notes: str = ""


class M8RetentionTimelineEvent(BaseModel):
    event_type: str = "NOTE"
    note: str = ""


class M8RetentionOutcomeUpdate(BaseModel):
    outcome: Literal["RETAINED", "EXITED", "UNKNOWN", "CLOSED"]
    note: str = ""


class M8SegmentSettingsUpdate(BaseModel):
    high_risk_score_min: Optional[float] = Field(default=None, ge=0.0, le=1.0)
    medium_risk_score_min: Optional[float] = Field(default=None, ge=0.0, le=1.0)
    require_critical_role_for_segment: Optional[bool] = None

class EmployeeBulkImportRequest(BaseModel):
    rows: List[EmployeeCreate]
    mode: Literal["skip", "upsert"] = "skip"
    dry_run: bool = False

class SkillBulkImportRow(BaseModel):
    skill_name: str
    demand_count: int = 0
    supply_count: int = 0
    category: Optional[str] = None
    priority: Literal["HIGH", "MEDIUM", "LOW"] = "MEDIUM"
    notes: Optional[str] = None

class SkillBulkImportRequest(BaseModel):
    rows: List[SkillBulkImportRow]
    mode: Literal["skip", "upsert"] = "skip"
    dry_run: bool = False

# Project Skill Demand CSV/MVP bulk import request
class ProjectSkillDemandBulkImportRow(BaseModel):
    skill_name: str
    demand_count: int = 0
    priority: Optional[str] = None  # optional override; defaults to MEDIUM
    demand_min: Optional[int] = None
    demand_max: Optional[int] = None
    constraint_type: Optional[Literal["HARD", "SOFT"]] = None

class ProjectSkillDemandBulkImportRequest(BaseModel):
    rows: List[ProjectSkillDemandBulkImportRow]
    mode: Literal["skip", "upsert"] = "skip"
    dry_run: bool = False

# Project Skill Allocation CSV/MVP bulk import request (Phase-3 M4 supply governance)
class ProjectSkillAllocationBulkImportRow(BaseModel):
    skill_name: str
    allocated_count: int = 0

class ProjectSkillAllocationBulkImportRequest(BaseModel):
    rows: List[ProjectSkillAllocationBulkImportRow]
    mode: Literal["skip", "upsert"] = "skip"
    dry_run: bool = False

# Minimal response row for listing allocations
class ProjectSkillAllocationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    project_id: str
    skill_name: str
    allocated_count: int
    updated_at: str

# Interview Models
class InterviewCreate(BaseModel):
    application_id: str
    round: int = 1
    mode: str = "VIRTUAL"  # VIRTUAL, ONSITE, PHONE
    scheduled_start: str
    scheduled_end: str
    meeting_link: Optional[str] = None
    interviewers: List[str] = []
    notes: Optional[str] = None

class InterviewUpdate(BaseModel):
    scheduled_start: Optional[str] = None
    scheduled_end: Optional[str] = None
    meeting_link: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None

class InterviewFeedbackCreate(BaseModel):
    decision: str  # STRONG_YES, YES, MAYBE, NO, STRONG_NO
    score: Optional[float] = None
    strengths: Optional[str] = None
    concerns: Optional[str] = None
    notes: Optional[str] = None

class InterviewResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    application_id: str
    round: int
    mode: str
    scheduled_start: str
    scheduled_end: str
    meeting_link: Optional[str] = None
    status: str
    interviewers: List[str] = []
    notes: Optional[str] = None
    feedback: List[Dict[str, Any]] = []
    candidate: Optional[Dict[str, Any]] = None
    job: Optional[Dict[str, Any]] = None
    created_at: str

# ========================
# Interview Proposal Models (M1 HR Approval Scheduling)
# ========================
InterviewProposalStatus = Literal["PENDING", "APPROVED", "REJECTED"]

class InterviewProposedSlot(BaseModel):
    slot_index: int
    scheduled_start: str
    scheduled_end: str

class InterviewProposalApproveRequest(BaseModel):
    slot_index: int = 0
    interviewers: List[str] = []
    meeting_link: Optional[str] = None
    notes: Optional[str] = None

class InterviewProposalRejectRequest(BaseModel):
    reason: Optional[str] = "Rejected by HR"

class InterviewProposalResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    job_id: str
    candidate_id: str
    round: int = 1
    mode: str = "VIRTUAL"
    status: InterviewProposalStatus = "PENDING"
    proposed_slots: List[InterviewProposedSlot] = []
    application_id: Optional[str] = None
    interview_id: Optional[str] = None
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None
    rejected_reason: Optional[str] = None
    created_at: str
    updated_at: Optional[str] = None
    candidate: Optional[Dict[str, Any]] = None
    job: Optional[Dict[str, Any]] = None

# Notification Models
class NotificationCreate(BaseModel):
    recipient_id: str
    type: str  # STAGE_CHANGE, INTERVIEW_SCHEDULED, REFERRAL_UPDATE, ASSESSMENT_INVITE
    title: str
    message: str
    metadata: Optional[Dict[str, Any]] = None

class NotificationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    recipient_id: str
    type: str
    title: str
    message: str
    read: bool
    metadata: Optional[Dict[str, Any]] = None
    created_at: str

# Extended Candidate Response with full profile
class CandidateProfileResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    headline: Optional[str] = None
    total_experience_years: Optional[float] = None
    skills: List[Dict[str, Any]] = []
    source: str
    experience: List[Dict[str, Any]] = []
    education: List[Dict[str, Any]] = []
    resume_url: Optional[str] = None
    resume_text: Optional[str] = None
    applications: List[Dict[str, Any]] = []
    interviews: List[Dict[str, Any]] = []
    created_at: str

# ========================
# AUTH UTILITIES
# ========================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRY_HOURS),
        "iat": datetime.now(timezone.utc)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if user is None:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ========================
# AI AGENT UTILITIES
# ========================

async def _hf_chat(system_message: str, user_text: str) -> str:
    """
    Hugging Face Inference API (OpenAI-compatible) chat completion.
    Env:
      - HF_TOKEN (required)
      - HF_MODEL (optional, default: meta-llama/Llama-3.1-8B-Instruct)
      - HF_BASE_URL (optional, default: https://api-inference.huggingface.co)
    """
    hf_token = os.environ.get("HF_TOKEN")
    if not hf_token:
        raise ValueError("Set HF_TOKEN")
    model = os.environ.get("HF_MODEL") or "meta-llama/Llama-3.1-8B-Instruct"
    base = (os.environ.get("HF_BASE_URL") or "https://api-inference.huggingface.co").rstrip("/")

    import httpx

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": user_text},
        ],
        "temperature": 0.2,
    }

    async with httpx.AsyncClient(timeout=httpx.Timeout(60.0)) as client:
        r = await client.post(
            f"{base}/v1/chat/completions",
            headers={"Authorization": f"Bearer {hf_token}"},
            json=payload,
        )
        r.raise_for_status()
        data = r.json()
        return ((data.get("choices") or [{}])[0].get("message") or {}).get("content", "").strip()

async def _llm_chat(system_message: str, user_text: str) -> str:
    """Call LLM (HF, Emergent, OpenAI). Returns response text."""
    hf_token = os.environ.get("HF_TOKEN")
    api_key = os.environ.get('EMERGENT_LLM_KEY') or os.environ.get('OPENAI_API_KEY')
    if not (hf_token or api_key):
        raise ValueError("Set HF_TOKEN or EMERGENT_LLM_KEY or OPENAI_API_KEY")

    # Prefer Hugging Face if configured
    if hf_token:
        try:
            return await _hf_chat(system_message, user_text)
        except Exception as e:
            logger.error(f"Hugging Face LLM call failed, falling back: {e}")

    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        chat = LlmChat(
            api_key=api_key,
            session_id=f"session_{uuid.uuid4()}",
            system_message=system_message,
        ).with_model("openai", "gpt-5.2")
        user_message = UserMessage(text=user_text)
        return await chat.send_message(user_message)
    except ImportError:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=api_key)
        r = await client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_text},
            ],
        )
        return (r.choices[0].message.content or "").strip()
    except Exception as e:
        logger.error(f"LLM call failed: {e}")
        raise

async def analyze_jd_with_ai(title: str, description: str, skills_needed: List[str], must_have_skills: List[str]) -> Dict[str, Any]:
    """Use LLM to analyze and structure JD"""
    try:
        system_message = """You are an HR domain expert. Extract structured information from job descriptions.
Always respond with valid JSON only. No markdown formatting, just raw JSON."""
        prompt = f"""Analyze this Job Description and extract structured data:

Job Title: {title}
Skills Needed: {', '.join(skills_needed)}
Must-Have Skills: {', '.join(must_have_skills)}
Description: {description}

Return JSON with:
{{
  "normalized_title": "standardized job title",
  "seniority": "Intern|Junior|Mid|Senior|Lead|null",
  "domain": "industry domain (BFSI, Healthcare, SaaS, etc.)",
  "skills": [
    {{"skill_name": "string", "skill_type": "MUST_HAVE|GOOD_TO_HAVE", "weight": 1.0}}
  ],
  "activities": [
    {{"activity_text": "specific responsibility", "category": "Technical|Communication|Leadership", "priority": "HIGH|MEDIUM|LOW"}}
  ],
  "scoring_rubric": {{
    "min_skill_match_pct": 70,
    "min_activity_match_pct": 60,
    "weights": {{"title": 0.2, "skill": 0.4, "activity": 0.3, "experience": 0.1}}
  }}
}}"""
        response = await _llm_chat(system_message, prompt)
        try:
            clean_response = response.strip()
            if clean_response.startswith("```"):
                clean_response = re.sub(r'^```(?:json)?\n?', '', clean_response)
                clean_response = re.sub(r'\n?```$', '', clean_response)
            return json.loads(clean_response)
        except json.JSONDecodeError:
            logger.error(f"Failed to parse AI response: {response}")
            return generate_default_jd_analysis(title, skills_needed, must_have_skills)
    except Exception as e:
        logger.error(f"AI JD analysis failed: {e}")
        return generate_default_jd_analysis(title, skills_needed, must_have_skills)

def _norm(text: Optional[str]) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip().lower()

def _tokenize(text: str) -> set:
    tokens = re.findall(r"[a-z0-9\+\.#]+", _norm(text))
    return set(t for t in tokens if len(t) >= 2)

def _jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    inter = len(a.intersection(b))
    union = len(a.union(b))
    return inter / union if union else 0.0

def _job_skill_sets(job: Dict[str, Any]) -> tuple[set, set]:
    skills_all = set()
    must_have = set()
    for s in job.get("skills", []) or []:
        name = _norm(s.get("skill_name"))
        if not name:
            continue
        skills_all.add(name)
        if s.get("skill_type") == "MUST_HAVE":
            must_have.add(name)
    return skills_all, must_have

def _candidate_skill_set(candidate: Dict[str, Any]) -> set:
    out = set()
    for s in candidate.get("skills", []) or []:
        name = _norm(s.get("skill_name"))
        if name:
            out.add(name)
    return out

def compute_match_score(job: Dict[str, Any], candidate: Dict[str, Any]) -> Dict[str, Any]:
    """
    Deterministic matcher for the in-house aggregator.
    External sources should normalize into the same candidate shape before scoring.
    """
    job_title = _norm(job.get("normalized_title") or job.get("title"))
    cand_title = _norm(candidate.get("headline") or "")

    job_skills, must_have = _job_skill_sets(job)
    cand_skills = _candidate_skill_set(candidate)

    title_score = _jaccard(_tokenize(job_title), _tokenize(cand_title)) * 100.0
    skill_score = _jaccard(job_skills, cand_skills) * 100.0 if job_skills else 0.0
    must_have_ok = must_have.issubset(cand_skills) if must_have else True

    # Description similarity: compare JD text to candidate resume_text (if present)
    jd_text = _norm(job.get("description"))
    resume_text = _norm(candidate.get("resume_text") or "")
    description_score = _jaccard(_tokenize(jd_text), _tokenize(resume_text)) * 100.0 if jd_text and resume_text else 0.0

    weights = (job.get("scoring_rubric") or {}).get("weights") or {"title": 0.2, "skill": 0.5, "activity": 0.0, "experience": 0.0}
    score = (
        title_score * float(weights.get("title", 0.2)) +
        skill_score * float(weights.get("skill", 0.5)) +
        description_score * 0.3
    )
    if not must_have_ok:
        score *= 0.25  # hard penalty for missing must-haves

    return {
        "title_score": round(title_score, 2),
        "skill_score": round(skill_score, 2),
        "must_have_ok": must_have_ok,
        "description_score": round(description_score, 2),
        "score": round(score, 2),
    }

def _norm_email(email: Optional[str]) -> Optional[str]:
    return email.strip().lower() if isinstance(email, str) and email.strip() else None

def _norm_full_name(full_name: Optional[str]) -> str:
    return _norm(full_name or "")


def _norm_phone_digits(phone: Optional[str]) -> Optional[str]:
    if not phone or not isinstance(phone, str):
        return None
    digits = re.sub(r"\D", "", phone.strip())
    if len(digits) >= 10:
        return digits[-10:]
    return digits or None


def _resume_content_hash(text: Optional[str]) -> Optional[str]:
    if not text or not isinstance(text, str):
        return None
    norm = re.sub(r"\s+", " ", text.strip().lower())
    if len(norm) < 40:
        return None
    return hashlib.sha256(norm.encode("utf-8")).hexdigest()


async def _append_dedup_audit(
    candidate_id: str,
    action: str,
    keys_matched: List[str],
    incoming_source: Optional[str],
    extra: Optional[Dict[str, Any]] = None,
) -> None:
    doc = {
        "id": str(uuid.uuid4()),
        "candidate_id": candidate_id,
        "action": action,
        "keys_matched": keys_matched,
        "incoming_source": incoming_source,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "detail": extra or {},
    }
    try:
        await db[CANDIDATE_DEDUP_AUDIT_COLLECTION].insert_one(doc)
    except Exception as e:
        logger.warning("dedup audit log failed: %s", e)


def _unique_sources(*sources: Optional[str]) -> List[str]:
    out: List[str] = []
    for s in sources:
        if isinstance(s, str) and s.strip():
            out.append(s.strip())
    # stable-unique
    return list(dict.fromkeys(out))

def _merge_candidate_docs(existing: Dict[str, Any], incoming: Dict[str, Any]) -> Dict[str, Any]:
    # Merge skills by skill_name (case-insensitive).
    existing_skills = existing.get("skills") or []
    incoming_skills = incoming.get("skills") or []

    skill_map: Dict[str, Dict[str, Any]] = {}
    for s in existing_skills:
        if isinstance(s, dict) and s.get("skill_name"):
            key = _norm(str(s.get("skill_name")))
            skill_map[key] = {**s, "skill_name": str(s.get("skill_name")).strip()}

    for s in incoming_skills:
        if isinstance(s, dict) and s.get("skill_name"):
            key = _norm(str(s.get("skill_name")))
            # Keep incoming proficiency if present, otherwise preserve existing.
            prev = skill_map.get(key)
            if prev and prev.get("proficiency") is None and s.get("proficiency") is not None:
                skill_map[key] = {**prev, **s}
            else:
                skill_map[key] = skill_map.get(key) or s

    merged_skills = list(skill_map.values())

    # Choose "best" text fields.
    def pick_longer(a: Optional[str], b: Optional[str]) -> Optional[str]:
        a = a or ""
        b = b or ""
        return b if len(b) > len(a) else (a if a else None)

    existing_resume = existing.get("resume_text") or None
    incoming_resume = incoming.get("resume_text") or None

    # Prefer non-empty scalar fields from incoming when existing is empty.
    def pick_non_empty(existing_val: Any, incoming_val: Any) -> Any:
        if incoming_val is None:
            return existing_val
        if isinstance(incoming_val, str) and not incoming_val.strip():
            return existing_val
        if existing_val is None:
            return incoming_val
        # keep existing unless it's empty-ish
        if isinstance(existing_val, str) and not existing_val.strip():
            return incoming_val
        return existing_val

    existing_sources = existing.get("sources") or []
    if not isinstance(existing_sources, list):
        existing_sources = []
    incoming_source = incoming.get("source") or ""
    merged_sources = _unique_sources(*existing_sources, incoming_source)

    merged = {
        **existing,
        # Canonical identity keys (never overwrite `id`).
        "full_name": pick_non_empty(existing.get("full_name"), incoming.get("full_name")),
        "email": pick_non_empty(existing.get("email"), incoming.get("email")),
        "phone": pick_non_empty(existing.get("phone"), incoming.get("phone")),
        "location": pick_non_empty(existing.get("location"), incoming.get("location")),
        "headline": pick_non_empty(existing.get("headline"), incoming.get("headline")),
        "total_experience_years": existing.get("total_experience_years") or incoming.get("total_experience_years"),
        "resume_text": pick_longer(existing_resume, incoming_resume),
        "skills": merged_skills,
        "experience": existing.get("experience") or incoming.get("experience") or [],
        "source": "|".join(merged_sources[:3]) if merged_sources else existing.get("source"),
        "sources": merged_sources,
        "email_lc": _norm_email(incoming.get("email")) or existing.get("email_lc"),
        "full_name_lc": _norm_full_name(incoming.get("full_name")) or existing.get("full_name_lc"),
        "phone_lc": _norm_phone_digits(incoming.get("phone")) or existing.get("phone_lc"),
        "resume_content_hash": _resume_content_hash(incoming.get("resume_text")) or existing.get("resume_content_hash"),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    return merged

async def upsert_candidate_dedup(incoming_candidate: Dict[str, Any]) -> Dict[str, Any]:
    """
    Canonicalize candidate profile across sources with dedup by email, phone, resume hash, then name.
    Merge strategy is deterministic and meant for ingestion-time consolidation.
    """
    # Ensure minimum fields exist.
    incoming_candidate = incoming_candidate.copy()
    incoming_candidate["full_name"] = (incoming_candidate.get("full_name") or "").strip()

    email_lc = _norm_email(incoming_candidate.get("email"))
    full_name_lc = _norm_full_name(incoming_candidate.get("full_name"))
    phone_lc = _norm_phone_digits(incoming_candidate.get("phone"))
    resume_hash = _resume_content_hash(incoming_candidate.get("resume_text"))

    # Generate stable candidate id when possible.
    if not incoming_candidate.get("id"):
        if email_lc:
            incoming_candidate["id"] = str(uuid.uuid5(uuid.NAMESPACE_DNS, email_lc))
        elif resume_hash:
            incoming_candidate["id"] = str(uuid.uuid5(uuid.NAMESPACE_URL, f"resume:{resume_hash}"))
        elif phone_lc:
            incoming_candidate["id"] = str(uuid.uuid5(uuid.NAMESPACE_URL, f"phone:{phone_lc}"))
        else:
            incoming_candidate["id"] = str(uuid.uuid4())

    existing: Optional[Dict[str, Any]] = None
    keys_matched: List[str] = []

    if email_lc:
        existing = await db.candidates.find_one({"email_lc": email_lc}, {"_id": 0})
        if not existing:
            existing = await db.candidates.find_one(
                {"email": {"$regex": f"^{re.escape(email_lc)}$", "$options": "i"}}, {"_id": 0}
            )
        if existing:
            keys_matched.append("email")

    if not existing and phone_lc:
        existing = await db.candidates.find_one({"phone_lc": phone_lc}, {"_id": 0})
        if existing:
            keys_matched.append("phone")

    if not existing and resume_hash:
        existing = await db.candidates.find_one({"resume_content_hash": resume_hash}, {"_id": 0})
        if existing:
            keys_matched.append("resume_hash")

    if not existing and full_name_lc:
        existing = await db.candidates.find_one({"full_name_lc": full_name_lc}, {"_id": 0})
        if not existing:
            existing = await db.candidates.find_one(
                {"full_name": {"$regex": f"^{re.escape(full_name_lc)}$", "$options": "i"}}, {"_id": 0}
            )
        if existing:
            keys_matched.append("name")

    if existing:
        merged = _merge_candidate_docs(existing, incoming_candidate)
        await db.candidates.update_one({"id": existing["id"]}, {"$set": merged})
        await _append_dedup_audit(
            existing["id"],
            "MERGE",
            keys_matched,
            str(incoming_candidate.get("source") or ""),
            {"incoming_id": incoming_candidate.get("id")},
        )
        await trigger_auto_analyze_if_eligible(
            db,
            candidate_id=merged["id"],
            resume_text=merged.get("resume_text"),
            created_by="system",
        )
        return merged

    now = datetime.now(timezone.utc).isoformat()
    candidate_doc = {
        "id": incoming_candidate["id"],
        "full_name": incoming_candidate.get("full_name") or "",
        "email": incoming_candidate.get("email"),
        "phone": incoming_candidate.get("phone"),
        "location": incoming_candidate.get("location"),
        "headline": incoming_candidate.get("headline"),
        "total_experience_years": incoming_candidate.get("total_experience_years"),
        "skills": incoming_candidate.get("skills") or [],
        "source": incoming_candidate.get("source") or "EXTERNAL",
        "sources": _unique_sources(incoming_candidate.get("source")),
        "experience": incoming_candidate.get("experience") or [],
        "resume_text": incoming_candidate.get("resume_text"),
        "created_at": incoming_candidate.get("created_at") or now,
        "updated_at": now,
        "email_lc": email_lc,
        "full_name_lc": full_name_lc,
        "phone_lc": phone_lc,
        "resume_content_hash": resume_hash,
    }
    await db.candidates.insert_one(candidate_doc)
    await _append_dedup_audit(
        candidate_doc["id"],
        "CREATE",
        [],
        str(incoming_candidate.get("source") or ""),
        {},
    )
    await trigger_auto_analyze_if_eligible(
        db,
        candidate_id=candidate_doc["id"],
        resume_text=candidate_doc.get("resume_text"),
        created_by="system",
    )
    return candidate_doc

async def company_db_search_candidates(job: Dict[str, Any], limit: int = 500) -> List[Dict[str, Any]]:
    """
    Our in-house ATS/aggregator connector: "company database candidates".
    - If admin config is enabled, we pull from the configured external Mongo.
    - Otherwise, we pull from this app's Mongo `candidates`.

    This returns candidates in the app's canonical schema.
    """
    cfg = await get_connector_config("COMPANY_DB_CANDIDATES")
    if cfg.get("enabled") and cfg.get("mongo_url") and cfg.get("db_name"):
        ext_client = AsyncIOMotorClient(cfg["mongo_url"])
        ext_coll_name = cfg.get("collection_name") or "candidates"
        ext_coll = ext_client[cfg["db_name"]][ext_coll_name]

        docs = await ext_coll.find({}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)
        normalized: List[Dict[str, Any]] = []

        for doc in docs:
            # Best-effort normalization into our canonical candidate schema.
            candidate_id = doc.get("id")
            if not candidate_id and doc.get("email"):
                candidate_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, doc["email"]))
            if not candidate_id:
                continue

            raw_skills = doc.get("skills") or []
            skills = []
            if isinstance(raw_skills, list):
                for s in raw_skills:
                    if isinstance(s, str):
                        skills.append({"skill_name": s, "proficiency": None})
                    elif isinstance(s, dict):
                        if s.get("skill_name"):
                            skills.append({"skill_name": s.get("skill_name"), "proficiency": s.get("proficiency")})

            experience = doc.get("experience") or []
            candidate_doc = {
                "id": candidate_id,
                "full_name": doc.get("full_name") or doc.get("name") or "",
                "email": doc.get("email"),
                "phone": doc.get("phone"),
                "location": doc.get("location"),
                "headline": doc.get("headline"),
                "total_experience_years": doc.get("total_experience_years"),
                "skills": skills,
                "source": doc.get("source") or "EXTERNAL_COMPANY_DB",
                "experience": experience,
                "resume_text": doc.get("resume_text"),
                "created_at": doc.get("created_at") or datetime.now(timezone.utc).isoformat(),
            }

            # Upsert into this app's candidates collection so that pipeline operations work.
            await db.candidates.update_one(
                {"id": candidate_id},
                {"$set": candidate_doc},
                upsert=True
            )
            normalized.append(candidate_doc)

        return normalized

    # Default: pull from this app's Mongo `candidates`.
    return await db.candidates.find({}, {"_id": 0}).sort([("pin_rank", -1), ("created_at", -1)]).limit(limit).to_list(limit)

def _normalize_connector_skills(raw_skills: Any) -> List[Dict[str, Any]]:
    """
    Normalize connector skill payloads to:
      [{ "skill_name": <str>, "proficiency": <any> }, ...]
    """
    out: List[Dict[str, Any]] = []
    if not isinstance(raw_skills, list):
        return out
    for s in raw_skills:
        if isinstance(s, str) and s.strip():
            out.append({"skill_name": s.strip(), "proficiency": None})
        elif isinstance(s, dict) and s.get("skill_name"):
            out.append({"skill_name": str(s.get("skill_name")).strip(), "proficiency": s.get("proficiency")})
    return out

def _normalize_connector_candidate(raw: Dict[str, Any], source: str) -> Dict[str, Any]:
    email = raw.get("email")
    normalized: Dict[str, Any] = {
        "id": raw.get("id"),
        "full_name": raw.get("full_name") or raw.get("name") or raw.get("candidate_name") or "",
        "email": email,
        "phone": raw.get("phone"),
        "location": raw.get("location"),
        "headline": raw.get("headline") or raw.get("current_title") or raw.get("job_title"),
        "total_experience_years": raw.get("total_experience_years") or raw.get("experience_years"),
        "skills": _normalize_connector_skills(raw.get("skills") or raw.get("skill_set") or []),
        "experience": raw.get("experience") or raw.get("work_experience") or [],
        "resume_text": raw.get("resume_text") or raw.get("resume") or raw.get("summary"),
        "source": source,
        "created_at": raw.get("created_at"),
    }
    # Clean up empties.
    if isinstance(normalized.get("full_name"), str):
        normalized["full_name"] = normalized["full_name"].strip()
    return normalized

async def connector_ingest_source(connector_name: str, job: Dict[str, Any], limit: int) -> List[Dict[str, Any]]:
    """
    Fetch from a single connector (Mongo-first or HTTP+paging) and upsert with dedup.
    """
    cfg = await get_connector_config(connector_name)
    if not cfg.get("enabled"):
        return []

    candidates_raw = await fetch_connector_candidates(
        connector_name,
        cfg,
        job,
        connector_name,
        limit,
        db,
        CONNECTOR_CONFIG_COLLECTION,
    )

    upserted: List[Dict[str, Any]] = []
    for c in candidates_raw:
        try:
            # HTTP fetch already normalizes; allow legacy dicts through normalizer
            if not c.get("source"):
                c = _normalize_connector_candidate(c, connector_name)
            upserted.append(await upsert_candidate_dedup(c))
        except Exception as e:
            logger.error("%s: dedup upsert failed: %s", connector_name, e)
    return upserted


async def linkedIn_search_candidates(job: Dict[str, Any], limit: int = 200) -> List[Dict[str, Any]]:
    return await connector_ingest_source("LINKEDIN", job, limit)


async def naukri_search_candidates(job: Dict[str, Any], limit: int = 200) -> List[Dict[str, Any]]:
    return await connector_ingest_source("NAUKRI", job, limit)


async def monster_search_candidates(job: Dict[str, Any], limit: int = 200) -> List[Dict[str, Any]]:
    return await connector_ingest_source("MONSTER", job, limit)


async def ingest_candidates_for_job(job: Dict[str, Any], total_limit: int = 500) -> None:
    """
    Ingest candidates from enabled external connectors into this app's `candidates` collection
    using canonical dedup logic and a unified ingestion job record (M1-2).
    """
    order = ["LINKEDIN", "NAUKRI", "MONSTER"]
    enabled: List[str] = []
    for name in order:
        c = await get_connector_config(name)
        if c.get("enabled"):
            enabled.append(name)
    if not enabled:
        return

    async def runner(src: str, j: Dict[str, Any], lim: int) -> List[Dict[str, Any]]:
        return await connector_ingest_source(src, j, lim)

    await run_unified_ingestion(db, j=job, total_limit=total_limit, sources=enabled, runner=runner, created_by=None)

@api_router.get("/admin/connector-configs")
async def admin_get_connector_configs(current_user: dict = Depends(get_current_user)):
    """
    Admin-only: return connector configuration (secrets omitted).
    """
    current_user = _require_admin(current_user)
    names = [
        "COMPANY_DB_CANDIDATES",
        "LINKEDIN",
        "NAUKRI",
        "MONSTER",
    ]

    out: Dict[str, Any] = {}
    for name in names:
        cfg = await get_connector_config(name)
        cfg_out = {
            k: v
            for k, v in cfg.items()
            if k not in ("client_secret", "refresh_token", "access_token")
        }
        cfg_out["client_secret_set"] = bool(cfg.get("client_secret"))
        cfg_out["refresh_token_set"] = bool(cfg.get("refresh_token"))
        cfg_out["access_token_set"] = bool(cfg.get("access_token"))
        out[name] = cfg_out
    return out

@api_router.put("/admin/connector-configs/{name}")
async def admin_put_connector_config(name: str, update: ConnectorConfigUpdate, current_user: dict = Depends(get_current_user)):
    """
    Admin-only: save connector configuration.
    Note: returned payload never includes `client_secret`.
    """
    current_user = _require_admin(current_user)
    allowed = {"COMPANY_DB_CANDIDATES", "LINKEDIN", "NAUKRI", "MONSTER"}
    if name not in allowed:
        raise HTTPException(status_code=404, detail="Unknown connector")

    update_doc = update.model_dump(exclude_none=True)
    # Normalize scopes into a list for consistent storage.
    update_doc["scopes"] = _normalize_scopes(update_doc.get("scopes")) if "scopes" in update_doc else None
    if "scopes" in update_doc and update_doc["scopes"] is None:
        update_doc.pop("scopes", None)

    # Treat empty secret as "do not change".
    if "client_secret" in update_doc and isinstance(update_doc["client_secret"], str) and not update_doc["client_secret"].strip():
        update_doc.pop("client_secret", None)
    if "refresh_token" in update_doc and isinstance(update_doc["refresh_token"], str) and not update_doc["refresh_token"].strip():
        update_doc.pop("refresh_token", None)
    if "access_token" in update_doc and isinstance(update_doc["access_token"], str) and not update_doc["access_token"].strip():
        update_doc.pop("access_token", None)

    update_doc["name"] = name
    await db[CONNECTOR_CONFIG_COLLECTION].update_one(
        {"name": name},
        {"$set": update_doc},
        upsert=True
    )

    return {"message": "Connector config saved", "name": name}


@api_router.get("/admin/connectors/health")
async def admin_connectors_health(current_user: dict = Depends(get_current_user)):
    """
    M1-1: connector health / last fetch status (admin).
    """
    _require_admin(current_user)
    names = ["COMPANY_DB_CANDIDATES", "LINKEDIN", "NAUKRI", "MONSTER"]
    out: Dict[str, Any] = {}
    for n in names:
        cfg = await get_connector_config(n)
        out[n] = {
            "enabled": bool(cfg.get("enabled")),
            "health_ok": cfg.get("health_ok"),
            "health_checked_at": cfg.get("health_checked_at"),
            "health_detail": (cfg.get("health_detail") or "")[:500],
            "request_count_total": cfg.get("request_count_total", 0),
        }
    return out


@api_router.post("/admin/candidates/merge")
async def admin_merge_candidates(
    body: CandidateMergeRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    M1-3: merge duplicate candidates — keeps `keep_candidate_id`, removes the other, re-points applications.
    """
    _require_admin(current_user)
    if body.keep_candidate_id == body.merge_candidate_id:
        raise HTTPException(status_code=400, detail="Same candidate id")

    keep = await db.candidates.find_one({"id": body.keep_candidate_id}, {"_id": 0})
    merge = await db.candidates.find_one({"id": body.merge_candidate_id}, {"_id": 0})
    if not keep or not merge:
        raise HTTPException(status_code=404, detail="Candidate not found")

    merged = _merge_candidate_docs(keep, merge)
    await db.candidates.update_one({"id": keep["id"]}, {"$set": merged})
    await db.applications.update_many(
        {"candidate_id": merge["id"]},
        {"$set": {"candidate_id": keep["id"], "updated_at": datetime.now(timezone.utc).isoformat()}},
    )
    await db.interview_proposals.update_many(
        {"candidate_id": merge["id"]},
        {"$set": {"candidate_id": keep["id"], "updated_at": datetime.now(timezone.utc).isoformat()}},
    )
    await db.interviews.update_many(
        {"candidate_id": merge["id"]},
        {"$set": {"candidate_id": keep["id"]}},
    )
    await db.candidates.delete_one({"id": merge["id"]})
    await _append_dedup_audit(
        keep["id"],
        "ADMIN_MERGE",
        [],
        "manual",
        {"merged_candidate_id": merge["id"], "by": current_user.get("id")},
    )
    return {"message": "Candidates merged", "kept_id": keep["id"], "removed_id": merge["id"]}


@api_router.get("/admin/users")
async def admin_list_users(
    q: Optional[str] = None,
    role: Optional[Literal["admin", "recruiter", "hr_admin", "hr_viewer"]] = None,
    page: int = 1,
    page_size: int = 50,
    current_user: dict = Depends(get_current_user),
):
    """
    Admin-only: list application users (excluding password).
    """
    _require_admin(current_user)
    page = max(1, page)
    page_size = min(max(1, page_size), 200)

    query: Dict[str, Any] = {}
    if role:
        query["role"] = role
    if q:
        query["$or"] = [
            {"email": {"$regex": q, "$options": "i"}},
            {"full_name": {"$regex": q, "$options": "i"}},
            {"id": {"$regex": q, "$options": "i"}},
        ]

    total = await db.users.count_documents(query)
    users = (
        await db.users.find(
            query,
            {"_id": 0, "id": 1, "email": 1, "full_name": 1, "role": 1, "created_at": 1},
        )
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )

    return {
        "items": [_user_doc_to_response(u).model_dump() for u in users],
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
    }


@api_router.put("/admin/users/{user_id}/role")
async def admin_update_user_role(
    user_id: str,
    update: AdminUserRoleUpdate,
    current_user: dict = Depends(get_current_user),
):
    """
    Admin-only: update a user's role to control Phase-1 permissions.
    """
    _require_admin(current_user)

    existing = await db.users.find_one({"id": user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    await db.users.update_one({"id": user_id}, {"$set": {"role": update.role}})

    updated = await db.users.find_one({"id": user_id}, {"_id": 0})
    return _user_doc_to_response(updated)

async def generate_and_store_job_matches(job_id: str, top_k: int = 50) -> Dict[str, Any]:
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise ValueError("Job not found")

    # Aggregate across sources (for now: company DB only; external connectors are stubs)
    candidates: List[Dict[str, Any]] = []
    try:
        candidates.extend(await company_db_search_candidates(job, limit=1000))
    except Exception as e:
        logger.error(f"Company DB connector failed: {e}")

    scored = []
    for c in candidates:
        try:
            s = compute_match_score(job, c)
            scored.append({
                "candidate_id": c["id"],
                "source": c.get("source", "COMPANY_DB"),
                **s,
            })
        except Exception as e:
            logger.error(f"Failed scoring candidate {c.get('id')}: {e}")

    scored.sort(key=lambda x: x.get("score", 0.0), reverse=True)
    top = scored[:top_k]

    doc = {
        "job_id": job_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "matches": top,
    }
    await db.job_candidate_matches.update_one(
        {"job_id": job_id},
        {"$set": doc},
        upsert=True
    )
    return doc

def generate_default_jd_analysis(title: str, skills_needed: List[str], must_have_skills: List[str]) -> Dict[str, Any]:
    """Fallback JD analysis without AI"""
    skills = []
    for skill in must_have_skills:
        skills.append({"skill_name": skill, "skill_type": "MUST_HAVE", "weight": 1.5})
    for skill in skills_needed:
        if skill not in must_have_skills:
            skills.append({"skill_name": skill, "skill_type": "GOOD_TO_HAVE", "weight": 1.0})
    
    return {
        "normalized_title": title,
        "seniority": None,
        "domain": None,
        "skills": skills,
        "activities": [],
        "scoring_rubric": {
            "min_skill_match_pct": 70,
            "min_activity_match_pct": 60,
            "weights": {"title": 0.2, "skill": 0.4, "activity": 0.3, "experience": 0.1}
        }
    }

async def compute_fit_score(job: Dict, candidate: Dict) -> Dict[str, Any]:
    """Compute fit score between candidate and job"""
    try:
        system_message = """You are an HR matching expert. Compare candidates to job requirements and compute fit scores.
Always respond with valid JSON only. No markdown formatting."""
        job_skills = job.get("skills", [])
        candidate_skills = candidate.get("skills", [])
        prompt = f"""Compare this candidate to the job requirements:

JOB:
- Title: {job.get('title')}
- Required Skills: {json.dumps(job_skills)}
- Activities: {json.dumps(job.get('activities', []))}

CANDIDATE:
- Headline: {candidate.get('headline', 'N/A')}
- Skills: {json.dumps(candidate_skills)}
- Experience: {json.dumps(candidate.get('experience', []))}
- Total Experience: {candidate.get('total_experience_years', 0)} years

Compute matching scores and return JSON:
{{
  "title_score": 0-100,
  "skill_match_pct": 0-100,
  "activity_match_pct": 0-100,
  "experience_score": 0-100,
  "must_have_ok": true/false,
  "explanation": {{
    "matched_skills": ["skill1", "skill2"],
    "missing_must_have": ["skill"],
    "matched_activities": ["activity"],
    "strengths": ["strength1"],
    "concerns": ["concern1"]
  }}
}}"""
        response = await _llm_chat(system_message, prompt)
        try:
            clean_response = response.strip()
            if clean_response.startswith("```"):
                clean_response = re.sub(r'^```(?:json)?\n?', '', clean_response)
                clean_response = re.sub(r'\n?```$', '', clean_response)
            result = json.loads(clean_response)
            weights = job.get("scoring_rubric", {}).get("weights", {"title": 0.2, "skill": 0.4, "activity": 0.3, "experience": 0.1})
            final_score = (
                result.get("title_score", 0) * weights.get("title", 0.2) +
                result.get("skill_match_pct", 0) * weights.get("skill", 0.4) +
                result.get("activity_match_pct", 0) * weights.get("activity", 0.3) +
                result.get("experience_score", 0) * weights.get("experience", 0.1)
            )
            if not result.get("must_have_ok", True):
                final_score *= 0.25
            result["final_score"] = round(final_score, 2)
            result["score_source"] = "llm"
            w = weights
            result["score_factors"] = {
                "title_weighted": round(float(result.get("title_score", 0)) * float(w.get("title", 0.2)), 3),
                "skill_weighted": round(float(result.get("skill_match_pct", 0)) * float(w.get("skill", 0.4)), 3),
                "activity_weighted": round(float(result.get("activity_match_pct", 0)) * float(w.get("activity", 0.3)), 3),
                "experience_weighted": round(float(result.get("experience_score", 0)) * float(w.get("experience", 0.1)), 3),
            }
            return result
        except json.JSONDecodeError:
            return compute_basic_fit_score(job, candidate)
    except Exception as e:
        logger.error(f"AI fit scoring failed: {e}")
        return compute_basic_fit_score(job, candidate)

def _experience_fit_score(candidate: Dict[str, Any], job: Dict[str, Any]) -> float:
    """Map years of experience to 0–100 using job minimum when available."""
    raw = candidate.get("total_experience_years")
    if raw is None:
        return 50.0
    try:
        years = float(raw)
    except (TypeError, ValueError):
        return 50.0
    req = job.get("min_experience_years") or job.get("years_of_experience_min")
    if req is not None:
        try:
            required = max(float(req), 0.5)
            return round(min(100.0, max(0.0, (years / required) * 100.0)), 2)
        except (TypeError, ValueError):
            pass
    return round(min(100.0, max(0.0, years * 10.0)), 2)


def compute_basic_fit_score(job: Dict, candidate: Dict) -> Dict[str, Any]:
    """Deterministic fit score without LLM — per-candidate title, resume, and experience."""
    det = compute_match_score(job, candidate)
    job_skills, must_have = _job_skill_sets(job)
    candidate_skills = _candidate_skill_set(candidate)

    matched_skills = job_skills.intersection(candidate_skills)
    skill_match_pct = (len(matched_skills) / len(job_skills) * 100) if job_skills else float(det.get("skill_score", 0))
    must_have_ok = bool(det.get("must_have_ok", True))

    weights = (job.get("scoring_rubric") or {}).get("weights") or {"title": 0.2, "skill": 0.4, "activity": 0.3, "experience": 0.1}
    title_score = float(det.get("title_score", 0))

    resume_text = (candidate.get("resume_text") or "").strip()
    if resume_text:
        activity_match_pct = float(det.get("description_score", 0))
    else:
        job_acts = {_norm(a) for a in (job.get("activities") or []) if _norm(a)}
        cand_blob = _norm(
            " ".join(
                [candidate.get("headline") or "", " ".join(candidate_skills), candidate.get("summary") or ""]
            )
        )
        if job_acts:
            hits = sum(1 for act in job_acts if act in cand_blob)
            activity_match_pct = (hits / len(job_acts)) * 100.0
        else:
            activity_match_pct = title_score

    experience_score = _experience_fit_score(candidate, job)

    final = (
        title_score * float(weights.get("title", 0.2))
        + skill_match_pct * float(weights.get("skill", 0.4))
        + activity_match_pct * float(weights.get("activity", 0.3))
        + experience_score * float(weights.get("experience", 0.1))
    )
    if not must_have_ok:
        final *= 0.25
    return {
        "title_score": round(title_score, 2),
        "skill_match_pct": round(skill_match_pct, 2),
        "activity_match_pct": round(activity_match_pct, 2),
        "experience_score": experience_score,
        "final_score": round(final, 2),
        "must_have_ok": must_have_ok,
        "score_source": "basic",
        "score_factors": {
            "title_weighted": round(title_score * float(weights.get("title", 0.2)), 3),
            "skill_weighted": round(skill_match_pct * float(weights.get("skill", 0.4)), 3),
            "activity_weighted": round(activity_match_pct * float(weights.get("activity", 0.3)), 3),
            "experience_weighted": round(experience_score * float(weights.get("experience", 0.1)), 3),
        },
        "explanation": {
            "matched_skills": sorted(matched_skills),
            "missing_must_have": sorted(must_have - candidate_skills),
            "matched_activities": [],
            "strengths": [],
            "concerns": [],
        },
    }

async def generate_assessment_with_ai(job: Dict, assessment_type: str) -> Dict[str, Any]:
    """Generate assessment questions from JD using AI"""
    try:
        system_message = """You are an assessment designer. Create relevant test questions based on job requirements.
Always respond with valid JSON only. No markdown formatting."""
        prompt = f"""Create an assessment for this job:

Job Title: {job.get('title')}
Required Skills: {json.dumps(job.get('skills', []))}
Responsibilities: {json.dumps(job.get('activities', []))}
Assessment Type: {assessment_type}

Generate 10 questions appropriate for {assessment_type}. Return JSON:
{{
  "questions": [
    {{
      "question_type": "MCQ|SHORT_ANSWER|CODING|SQL|CASE_STUDY",
      "question_text": "question",
      "options": ["A", "B", "C", "D"] (for MCQ only),
      "answer_key": "correct answer or rubric",
      "max_marks": 10,
      "difficulty": "EASY|MEDIUM|HARD",
      "skill_tested": "skill name"
    }}
  ],
  "rubric": {{
    "pass_threshold": 70,
    "grading_guide": "description"
  }}
}}"""
        response = await _llm_chat(system_message, prompt)
        try:
            clean_response = response.strip()
            if clean_response.startswith("```"):
                clean_response = re.sub(r'^```(?:json)?\n?', '', clean_response)
                clean_response = re.sub(r'\n?```$', '', clean_response)
            return json.loads(clean_response)
        except json.JSONDecodeError:
            return generate_default_assessment(job, assessment_type)
    except Exception as e:
        logger.error(f"AI assessment generation failed: {e}")
        return generate_default_assessment(job, assessment_type)

def generate_default_assessment(job: Dict, assessment_type: str) -> Dict[str, Any]:
    """Default assessment without AI"""
    skills = job.get("skills", [])
    questions = []
    
    for i, skill in enumerate(skills[:5]):
        questions.append({
            "question_type": "MCQ",
            "question_text": f"What is your proficiency level in {skill.get('skill_name', 'this skill')}?",
            "options": ["Beginner", "Intermediate", "Advanced", "Expert"],
            "answer_key": "Self-assessment",
            "max_marks": 10,
            "difficulty": "EASY",
            "skill_tested": skill.get("skill_name", "")
        })
    
    return {
        "questions": questions,
        "rubric": {
            "pass_threshold": 70,
            "grading_guide": "Score based on accuracy and depth of answers"
        }
    }

# ========================
# RESUME PARSING UTILITIES
# ========================

def parse_pdf_resume(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        return ""

def parse_docx_resume(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""

def _extract_resume_data_basic(resume_text: str) -> Dict[str, Any]:
    """
    Heuristic fallback when no LLM key is configured.
    Goal: populate at least email/phone/skills/experience_years from OCR/parsed text.
    """
    text = (resume_text or "").strip()
    out: Dict[str, Any] = {
        "full_name": None,
        "email": None,
        "phone": None,
        "location": None,
        "headline": None,
        "total_experience_years": None,
        "skills": [],
        "experience": [],
        "education": [],
    }
    if not text:
        return out

    m = re.search(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", text, re.IGNORECASE)
    if m:
        out["email"] = m.group(1)

    # Very permissive phone grab (keeps last 10-15 digits).
    pm = re.search(r"(\+?\d[\d\s().-]{8,}\d)", text)
    if pm:
        out["phone"] = pm.group(1).strip()

    ym = re.search(r"(\d+(?:\.\d+)?)\s*\+?\s*years", text, re.IGNORECASE)
    if ym:
        try:
            out["total_experience_years"] = float(ym.group(1))
        except Exception:
            out["total_experience_years"] = None

    # Try to infer a name from the first non-empty line if it looks like a person name.
    first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
    if first_line and len(first_line) <= 60 and not re.search(r"@|http|www\.", first_line, re.IGNORECASE):
        if re.fullmatch(r"[A-Za-z][A-Za-z .'-]{2,}", first_line):
            out["full_name"] = first_line.strip()

    # Skills section extraction (common patterns: "Skills:", "Technical Skills", etc.)
    lines = [ln.strip() for ln in text.splitlines()]
    skills_blob = ""
    for i, ln in enumerate(lines[:200]):
        low = ln.lower()
        if low.startswith("skills") or "technical skills" in low or low.startswith("key skills"):
            # consume current line after ":" plus next 1-3 lines if present
            after = ln.split(":", 1)[1].strip() if ":" in ln else ""
            nxt = " ".join([lines[j] for j in range(i + 1, min(i + 4, len(lines)))])
            skills_blob = " ".join([after, nxt]).strip()
            break
    if skills_blob:
        raw = re.split(r"[,•|/]\s*|\s{2,}", skills_blob)
        cleaned = []
        for s in raw:
            s = re.sub(r"[\[\]().;:]+", " ", s).strip()
            if not s:
                continue
            if len(s) > 40:
                continue
            cleaned.append(s)
        # de-dupe preserving order
        seen = set()
        uniq = []
        for s in cleaned:
            key = s.lower()
            if key in seen:
                continue
            seen.add(key)
            uniq.append(s)
        out["skills"] = uniq[:50]

    return out

async def extract_resume_data_with_ai(resume_text: str) -> Dict[str, Any]:
    """Use AI to extract structured data from resume text"""
    try:
        system_message = """You are a resume parsing expert. Extract structured information from resume text.
Always respond with valid JSON only. No markdown formatting."""
        prompt = f"""Parse this resume and extract structured data:

{resume_text[:4000]}

Return JSON with:
{{
  "full_name": "string",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "headline": "current title or summary in 10 words",
  "total_experience_years": number or null,
  "skills": ["skill1", "skill2", ...],
  "experience": [
    {{"company": "string", "title": "string", "start_date": "YYYY-MM or null", "end_date": "YYYY-MM or null", "description": "brief summary"}}
  ],
  "education": [
    {{"institution": "string", "degree": "string", "field": "string or null", "year": "YYYY or null"}}
  ]
}}"""
        response = await _llm_chat(system_message, prompt)
        try:
            clean_response = response.strip()
            if clean_response.startswith("```"):
                clean_response = re.sub(r'^```(?:json)?\n?', '', clean_response)
                clean_response = re.sub(r'\n?```$', '', clean_response)
            return json.loads(clean_response)
        except json.JSONDecodeError:
            logger.error("Failed to parse AI response for resume")
            return _extract_resume_data_basic(resume_text)
    except Exception as e:
        logger.error(f"AI resume parsing failed: {e}")
        return _extract_resume_data_basic(resume_text)

# ========================
# NOTIFICATION UTILITIES
# ========================

async def create_notification(
    recipient_id: str,
    notification_type: str,
    title: str,
    message: str,
    metadata: Dict[str, Any] = None
):
    """Create and store a notification"""
    notification_doc = {
        "id": str(uuid.uuid4()),
        "recipient_id": recipient_id,
        "type": notification_type,
        "title": title,
        "message": message,
        "read": False,
        "metadata": metadata or {},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.notifications.insert_one(notification_doc)
    return notification_doc

async def notify_stage_change(
    application_id: str,
    old_stage: str,
    new_stage: str,
    candidate_id: str,
    job_title: str,
    changed_by_name: str
):
    """Send notification for stage change"""
    # Get all recruiters to notify
    recruiters = await db.users.find({"role": {"$in": ["recruiter", "admin"]}}, {"_id": 0}).to_list(100)
    
    candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    candidate_name = candidate.get("full_name", "Unknown") if candidate else "Unknown"
    
    for recruiter in recruiters:
        await create_notification(
            recipient_id=recruiter["id"],
            notification_type="STAGE_CHANGE",
            title=f"Pipeline Update: {candidate_name}",
            message=f"{candidate_name} moved from {old_stage.replace('_', ' ')} to {new_stage.replace('_', ' ')} for {job_title}",
            metadata={
                "application_id": application_id,
                "candidate_id": candidate_id,
                "old_stage": old_stage,
                "new_stage": new_stage,
                "changed_by": changed_by_name
            }
        )

def _interview_notification_copy(kind: str, candidate_name: str, job_title: str, scheduled_time: str) -> tuple[str, str]:
    """M1-5: centralized templates (extend via env or DB later)."""
    if kind == "REMINDER":
        return (
            f"Reminder: interview with {candidate_name}",
            f"Your interview for {job_title} is scheduled at {scheduled_time}. Please confirm your calendar.",
        )
    return (
        f"Interview Scheduled: {candidate_name}",
        f"You have an interview scheduled with {candidate_name} for {job_title} at {scheduled_time}.",
    )


async def sync_interview_calendar_event(interview_doc: Dict[str, Any]) -> Dict[str, Any]:
    """
    M1-5: optional outbound calendar sync. Set CALENDAR_WEBHOOK_URL to POST JSON payload to your bridge
    (Google Calendar, MS Graph, etc.). If unset, returns SKIPPED.
    """
    webhook = os.environ.get("CALENDAR_WEBHOOK_URL", "").strip()
    if not webhook:
        return {"status": "SKIPPED", "reason": "CALENDAR_WEBHOOK_URL not set"}
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            r = await client.post(webhook, json=interview_doc)
        ok = 200 <= r.status_code < 300
        return {"status": "SYNCED" if ok else "FAILED", "http_status": r.status_code, "body": (r.text or "")[:300]}
    except Exception as e:
        return {"status": "FAILED", "error": str(e)}


async def notify_interview_scheduled(
    interview_id: str,
    candidate_name: str,
    job_title: str,
    scheduled_time: str,
    interviewers: List[str]
):
    """Send notification for scheduled interview"""
    title, message = _interview_notification_copy("SCHEDULED", candidate_name, job_title, scheduled_time)
    for interviewer_id in interviewers:
        await create_notification(
            recipient_id=interviewer_id,
            notification_type="INTERVIEW_SCHEDULED",
            title=title,
            message=message,
            metadata={
                "interview_id": interview_id,
                "candidate_name": candidate_name,
                "job_title": job_title,
                "scheduled_time": scheduled_time,
                "template": "SCHEDULED",
            }
        )


async def notify_interview_reminder(
    interview_id: str,
    candidate_name: str,
    job_title: str,
    scheduled_time: str,
    interviewers: List[str],
):
    title, message = _interview_notification_copy("REMINDER", candidate_name, job_title, scheduled_time)
    for interviewer_id in interviewers:
        await create_notification(
            recipient_id=interviewer_id,
            notification_type="INTERVIEW_REMINDER",
            title=title,
            message=message,
            metadata={
                "interview_id": interview_id,
                "candidate_name": candidate_name,
                "job_title": job_title,
                "scheduled_time": scheduled_time,
                "template": "REMINDER",
            },
        )

async def notify_referral_status(
    referral_id: str,
    referred_by_id: str,
    candidate_name: str,
    status: str,
    job_title: str
):
    """Send notification for referral status update"""
    await create_notification(
        recipient_id=referred_by_id,
        notification_type="REFERRAL_UPDATE",
        title=f"Referral Update: {candidate_name}",
        message=f"Your referral {candidate_name} for {job_title} is now {status}",
        metadata={
            "referral_id": referral_id,
            "candidate_name": candidate_name,
            "status": status,
            "job_title": job_title
        }
    )

# ========================
# AUTH ROUTES
# ========================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    email_norm = _normalize_auth_email(str(user_data.email))
    if not email_norm:
        raise HTTPException(status_code=400, detail="Email is required")

    # Check if user exists (case-insensitive)
    existing = await db.users.find_one(_user_email_match_query(email_norm), {"_id": 0})
    if existing:
        # If an older/seeded record exists without a password, allow "claiming" it by setting password.
        existing_password = existing.get("password")
        if not existing_password:
            await db.users.update_one(
                {"id": existing["id"]},
                {"$set": {
                    "email": email_norm,
                    "password": hash_password(user_data.password),
                    "full_name": user_data.full_name,
                    "role": user_data.role,
                }}
            )
            user = await db.users.find_one({"id": existing["id"]}, {"_id": 0})
            token = create_token(user["id"], email_norm)
            return TokenResponse(access_token=token, user=_user_doc_to_response(user))

        raise HTTPException(status_code=400, detail="Email already registered")

    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": email_norm,
        "password": hash_password(user_data.password),
        "full_name": user_data.full_name,
        "role": user_data.role,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.users.insert_one(user_doc)

    token = create_token(user_id, email_norm)
    return TokenResponse(access_token=token, user=_user_doc_to_response(user_doc))

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    email_norm = _normalize_auth_email(str(credentials.email))
    user = await db.users.find_one(_user_email_match_query(email_norm), {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    stored_pw = user.get("password")
    if not stored_pw or not verify_password(credentials.password, stored_pw):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    token = create_token(user["id"], email_norm)
    return TokenResponse(access_token=token, user=_user_doc_to_response(user))

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return _user_doc_to_response(current_user)

# ========================
# JOB ROUTES
# ========================

@api_router.post("/jobs", response_model=JobResponse)
async def create_job(job_data: JobCreate, background_tasks: BackgroundTasks, current_user: dict = Depends(get_current_user)):
    job_id = str(uuid.uuid4())
    
    # Analyze JD with AI
    analysis = await analyze_jd_with_ai(
        job_data.title,
        job_data.description,
        job_data.skills_needed,
        job_data.must_have_skills
    )
    
    job_doc = {
        "id": job_id,
        "title": job_data.title,
        "normalized_title": analysis.get("normalized_title", job_data.title),
        "description": job_data.description,
        "seniority": analysis.get("seniority") or job_data.seniority,
        "domain": analysis.get("domain"),
        "business_pillar": job_data.business_pillar,
        "business_department": job_data.business_department,
        "business_sub_department": job_data.business_sub_department,
        "project_id": job_data.project_id,
        "location": job_data.location,
        "work_mode": job_data.work_mode,
        "status": "OPEN",
        "skills": analysis.get("skills", []),
        "activities": analysis.get("activities", []),
        "scoring_rubric": analysis.get("scoring_rubric"),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.jobs.insert_one(job_doc)
    invalidate_hiring_pack_cache(reason="job_created")

    # Our in-house aggregator: auto-generate Top 50 matches (company DB + connectors).
    background_tasks.add_task(generate_and_store_job_matches, job_id, 50)
    
    return JobResponse(**job_doc, candidate_count=0)

@api_router.get("/jobs", response_model=List[JobResponse])
async def list_jobs(status: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {}
    if status:
        query["status"] = status
    
    jobs = await db.jobs.find(query, {"_id": 0}).sort([("pin_rank", -1), ("created_at", -1)]).to_list(1000)
    
    # Get candidate counts
    for job in jobs:
        count = await db.applications.count_documents({"job_id": job["id"]})
        job["candidate_count"] = count
    
    return [JobResponse(**job) for job in jobs]

@api_router.get("/jobs/{job_id}", response_model=JobResponse)
async def get_job(job_id: str, current_user: dict = Depends(get_current_user)):
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    count = await db.applications.count_documents({"job_id": job_id})
    job["candidate_count"] = count
    return JobResponse(**job)

@api_router.get("/jobs/{job_id}/matches", response_model=JobMatchesResponse)
async def get_job_matches(job_id: str, limit: int = 50, refresh: bool = False, current_user: dict = Depends(get_current_user)):
    """
    Returns Top matches computed by our internal ATS/aggregator.
    - **refresh=true** recomputes now (useful after adding candidates via company DB connector).
    """
    if refresh:
        doc = await generate_and_store_job_matches(job_id, min(max(limit, 1), 200))
        return JobMatchesResponse(**doc)

    doc = await db.job_candidate_matches.find_one({"job_id": job_id}, {"_id": 0})
    if not doc:
        # Compute once on demand if background job hasn't run yet.
        doc = await generate_and_store_job_matches(job_id, min(max(limit, 1), 200))
    else:
        doc["matches"] = (doc.get("matches") or [])[: min(max(limit, 1), 200)]
    return JobMatchesResponse(**doc)

@api_router.put("/jobs/{job_id}", response_model=JobResponse)
async def update_job(job_id: str, job_data: JobUpdate, current_user: dict = Depends(get_current_user)):
    job = await db.jobs.find_one({"id": job_id})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    update_data = {k: v for k, v in job_data.model_dump().items() if v is not None}
    if update_data:
        await db.jobs.update_one({"id": job_id}, {"$set": update_data})
        invalidate_hiring_pack_cache(reason="job_updated")
    
    updated_job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    count = await db.applications.count_documents({"job_id": job_id})
    updated_job["candidate_count"] = count
    return JobResponse(**updated_job)

@api_router.delete("/jobs/{job_id}")
async def delete_job(job_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.jobs.delete_one({"id": job_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Job not found")
    invalidate_hiring_pack_cache(reason="job_deleted")
    return {"message": "Job deleted"}

# ========================
# CANDIDATE ROUTES
# ========================

@api_router.post("/candidates", response_model=CandidateResponse)
async def create_candidate(candidate_data: CandidateCreate, current_user: dict = Depends(get_current_user)):
    # Check for duplicate
    if candidate_data.email:
        existing = await db.candidates.find_one({"email": candidate_data.email})
        if existing:
            raise HTTPException(status_code=400, detail="Candidate with this email already exists")
    
    candidate_id = str(uuid.uuid4())
    
    # Process skills
    skills = [{"skill_name": s, "proficiency": None} for s in candidate_data.skills]
    
    # Process experience
    experience = [exp.model_dump() for exp in candidate_data.experience]
    
    candidate_doc = {
        "id": candidate_id,
        "full_name": candidate_data.full_name,
        "email": candidate_data.email,
        "phone": candidate_data.phone,
        "location": candidate_data.location,
        "headline": candidate_data.headline,
        "total_experience_years": candidate_data.total_experience_years,
        "skills": skills,
        "source": candidate_data.source,
        "experience": experience,
        "resume_text": candidate_data.resume_text,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.candidates.insert_one(candidate_doc)

    await trigger_auto_analyze_if_eligible(
        db,
        candidate_id=candidate_id,
        resume_text=candidate_doc.get("resume_text"),
        created_by=current_user.get("id") or "",
    )

    return CandidateResponse(**candidate_doc)

@api_router.get("/candidates", response_model=List[CandidateResponse])
async def list_candidates(
    source: Optional[str] = None,
    skill: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if source:
        query["source"] = source
    if skill:
        query["skills.skill_name"] = {"$regex": skill, "$options": "i"}
    
    candidates = await db.candidates.find(query, {"_id": 0}).sort([("pin_rank", -1), ("created_at", -1)]).to_list(1000)
    return [CandidateResponse(**c) for c in candidates]

# Must be registered BEFORE /candidates/{candidate_id} or \"paged\" is captured as a candidate_id (404).
@api_router.get("/candidates/paged", response_model=CandidatesPagedResponse)
async def list_candidates_paged(
    page: int = 1,
    page_size: int = 50,
    q: Optional[str] = None,
    source: Optional[str] = None,
    display_channel: Optional[str] = None,
    skill: Optional[str] = None,
    fit_min: Optional[float] = None,
    fit_max: Optional[float] = None,
    current_user: dict = Depends(get_current_user),
):
    page = max(1, int(page or 1))
    page_size = int(page_size or 50)
    page_size = min(max(page_size, 1), 200)

    query: Dict[str, Any] = {}
    source_norm = str(source or "").strip().upper()
    if source_norm == "TALENT_POOL":
        tp_filter = all_talent_pool_mongo_filter()
        query = tp_filter
    elif source:
        query["source"] = source
    if display_channel:
        ch_filter = display_channel_mongo_filter(display_channel)
        if ch_filter:
            query = {"$and": [query, ch_filter]} if query else ch_filter
    if skill:
        query["skills.skill_name"] = {"$regex": skill, "$options": "i"}
    if q and str(q).strip():
        qq = str(q).strip()
        text_or = {
            "$or": [
                {"full_name": {"$regex": qq, "$options": "i"}},
                {"email": {"$regex": qq, "$options": "i"}},
                {"headline": {"$regex": qq, "$options": "i"}},
            ]
        }
        query = {"$and": [query, text_or]} if query else text_or

    fit_ids = await candidate_ids_matching_fit_range(db, fit_min, fit_max)
    if fit_ids is not None:
        if not fit_ids:
            return CandidatesPagedResponse(
                items=[],
                total=0,
                page=1,
                page_size=int(page_size),
                total_pages=1,
            )
        fit_clause = {"id": {"$in": fit_ids}}
        query = {"$and": [query, fit_clause]} if query else fit_clause

    total = await db.candidates.count_documents(query)
    total_pages = max(1, (total + page_size - 1) // page_size) if total else 1
    page = min(page, total_pages)

    cursor = (
        db.candidates.find(query, {"_id": 0})
        .sort([("pin_rank", -1), ("created_at", -1)])
        .skip((page - 1) * page_size)
        .limit(page_size)
    )
    rows = await cursor.to_list(page_size)
    items = [CandidateResponse(**c) for c in rows]
    return CandidatesPagedResponse(
        items=items,
        total=int(total),
        page=int(page),
        page_size=int(page_size),
        total_pages=int(total_pages),
    )

@api_router.get("/candidates/{candidate_id}", response_model=CandidateResponse)
async def get_candidate(candidate_id: str, current_user: dict = Depends(get_current_user)):
    candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    return CandidateResponse(**candidate)

@api_router.get("/candidates/{candidate_id}/profile", response_model=CandidateProfileResponse)
async def get_candidate_profile(candidate_id: str, current_user: dict = Depends(get_current_user)):
    """Get full candidate profile with applications and interviews"""
    candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")

    # Avoid passing duplicate kwargs when candidate already includes these keys.
    # (Some ingestion / seed paths persist `education` / `applications` / `interviews` on the candidate doc.)
    candidate = dict(candidate)
    education = candidate.get("education", [])
    candidate.pop("applications", None)
    candidate.pop("interviews", None)
    # We pass `education` explicitly below with a default; remove from base dict to prevent double values.
    candidate.pop("education", None)
    
    # Get all applications for this candidate
    applications = await db.applications.find({"candidate_id": candidate_id}, {"_id": 0}).to_list(100)
    enriched_apps = []
    for app in applications:
        job = await db.jobs.find_one({"id": app["job_id"]}, {"_id": 0})
        fit_score = await db.fit_scores.find_one({"id": app.get("fit_score_id")}, {"_id": 0})
        enriched_apps.append({
            **app,
            "job": {"id": job["id"], "title": job["title"], "status": job.get("status")} if job else None,
            "fit_score": fit_score
        })
    
    # Get all interviews for this candidate
    interviews = await db.interviews.find({"candidate_id": candidate_id}, {"_id": 0}).sort("scheduled_start", -1).to_list(50)
    
    return CandidateProfileResponse(
        **candidate,
        education=education if isinstance(education, list) else [],
        applications=enriched_apps,
        interviews=interviews
    )

@api_router.post("/candidates/upload-resume")
async def upload_resume(
    file: UploadFile = File(...),
    source: str = Form("DIRECT_UPLOAD"),
    current_user: dict = Depends(get_current_user)
):
    """Upload and parse a resume file (PDF or DOCX)"""
    # Validate file type
    filename = file.filename.lower()
    if not (filename.endswith('.pdf') or filename.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    # Read file content
    content = await file.read()
    
    # Check file size (max 10MB)
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds 10MB limit")
    
    # Parse resume based on file type
    if filename.endswith('.pdf'):
        resume_text = parse_pdf_resume(content)
    else:
        resume_text = parse_docx_resume(content)
    
    if not resume_text:
        raise HTTPException(status_code=400, detail="Could not extract text from file")
    
    # Use AI to extract structured data
    extracted_data = await extract_resume_data_with_ai(resume_text)
    
    # Check for existing candidate by email
    candidate_id = str(uuid.uuid4())
    email = extracted_data.get("email")
    
    if email:
        existing = await db.candidates.find_one({"email": email})
        if existing:
            # Update existing candidate
            candidate_id = existing["id"]
            update_doc = {
                "resume_text": resume_text,
                "resume_filename": file.filename,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }
            if extracted_data.get("skills"):
                update_doc["skills"] = [{"skill_name": s, "proficiency": None} for s in extracted_data["skills"]]
            if extracted_data.get("experience"):
                update_doc["experience"] = extracted_data["experience"]
            if extracted_data.get("education"):
                update_doc["education"] = extracted_data["education"]
            
            await db.candidates.update_one({"id": candidate_id}, {"$set": update_doc})

            await trigger_auto_analyze_if_eligible(
                db,
                candidate_id=candidate_id,
                resume_text=resume_text,
                created_by=current_user.get("id") or "",
            )

            candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
            return {
                "message": "Existing candidate updated with new resume",
                "candidate_id": candidate_id,
                "candidate": candidate,
                "extracted_data": extracted_data,
                "is_new": False
            }
    
    # Create new candidate
    skills = [{"skill_name": s, "proficiency": None} for s in extracted_data.get("skills", [])]
    
    candidate_doc = {
        "id": candidate_id,
        "full_name": extracted_data.get("full_name", "Unknown"),
        "email": email,
        "phone": extracted_data.get("phone"),
        "location": extracted_data.get("location"),
        "headline": extracted_data.get("headline"),
        "total_experience_years": extracted_data.get("total_experience_years"),
        "skills": skills,
        "source": source,
        "experience": extracted_data.get("experience", []),
        "education": extracted_data.get("education", []),
        "resume_text": resume_text,
        "resume_filename": file.filename,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.candidates.insert_one(candidate_doc)

    await trigger_auto_analyze_if_eligible(
        db,
        candidate_id=candidate_id,
        resume_text=resume_text,
        created_by=current_user.get("id") or "",
    )

    return {
        "message": "Resume parsed and candidate created successfully",
        "candidate_id": candidate_id,
        "candidate": candidate_doc,
        "extracted_data": extracted_data,
        "is_new": True
    }

@api_router.put("/candidates/{candidate_id}")
async def update_candidate(
    candidate_id: str,
    payload: CandidateUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update candidate information"""
    candidate = await db.candidates.find_one({"id": candidate_id})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    update_doc = {"updated_at": datetime.now(timezone.utc).isoformat()}
    
    if payload.full_name is not None:
        update_doc["full_name"] = str(payload.full_name).strip()
    if payload.email is not None:
        update_doc["email"] = str(payload.email).strip()
    if payload.phone is not None:
        update_doc["phone"] = str(payload.phone).strip()
    if payload.location is not None:
        update_doc["location"] = str(payload.location).strip()
    if payload.headline is not None:
        update_doc["headline"] = str(payload.headline).strip()
    if payload.total_experience_years is not None:
        update_doc["total_experience_years"] = float(payload.total_experience_years)
    if payload.skills is not None:
        normalized: List[Dict[str, Any]] = []
        for s in payload.skills:
            if isinstance(s, str) and s.strip():
                normalized.append({"skill_name": s.strip(), "proficiency": None})
            elif isinstance(s, dict) and s.get("skill_name"):
                normalized.append(
                    {"skill_name": str(s.get("skill_name")).strip(), "proficiency": s.get("proficiency")}
                )
        update_doc["skills"] = normalized
    
    await db.candidates.update_one({"id": candidate_id}, {"$set": update_doc})
    
    updated = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    return CandidateResponse(**updated)

# ========================
# APPLICATION/PIPELINE ROUTES
# ========================

@api_router.post("/applications", response_model=ApplicationResponse)
async def create_application(app_data: ApplicationCreate, current_user: dict = Depends(get_current_user)):
    # Check job and candidate exist
    job = await db.jobs.find_one({"id": app_data.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    candidate = await db.candidates.find_one({"id": app_data.candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    # Check for existing application
    existing = await db.applications.find_one({
        "job_id": app_data.job_id,
        "candidate_id": app_data.candidate_id
    })
    if existing:
        raise HTTPException(status_code=400, detail="Application already exists")
    
    # Compute fit score
    fit_result = await compute_fit_score(job, candidate)
    
    app_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    # Store fit score
    fit_score_doc = {
        "id": str(uuid.uuid4()),
        "job_id": app_data.job_id,
        "candidate_id": app_data.candidate_id,
        **fit_result,
        "computed_at": now
    }
    await db.fit_scores.insert_one(fit_score_doc)
    
    # Create application
    app_doc = {
        "id": app_id,
        "job_id": app_data.job_id,
        "candidate_id": app_data.candidate_id,
        "stage": app_data.stage,
        "status": "ACTIVE",
        "fit_score_id": fit_score_doc["id"],
        "created_at": now,
        "updated_at": now
    }
    await db.applications.insert_one(app_doc)
    
    # Log stage history
    await db.application_stage_history.insert_one({
        "id": str(uuid.uuid4()),
        "application_id": app_id,
        "from_stage": None,
        "to_stage": app_data.stage,
        "changed_by": current_user["id"],
        "changed_at": now
    })

    invalidate_hiring_pack_cache(reason="application_created")
    
    return ApplicationResponse(
        **app_doc,
        fit_score=fit_result,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]}
    )

@api_router.get("/applications", response_model=List[ApplicationResponse])
async def list_applications(
    job_id: Optional[str] = None,
    stage: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if job_id:
        query["job_id"] = job_id
    if stage:
        query["stage"] = stage
    
    applications = await db.applications.find(query, {"_id": 0}).sort("updated_at", -1).to_list(500)
    
    result = []
    for app in applications:
        candidate = await db.candidates.find_one({"id": app["candidate_id"]}, {"_id": 0})
        job = await db.jobs.find_one({"id": app["job_id"]}, {"_id": 0})
        fit_score = await db.fit_scores.find_one({"id": app.get("fit_score_id")}, {"_id": 0})
        
        result.append(ApplicationResponse(
            **app,
            fit_score=fit_score,
            candidate=candidate,
            job={"id": job["id"], "title": job["title"]} if job else None
        ))
    
    return result

@api_router.put("/applications/{app_id}/stage", response_model=ApplicationResponse)
async def update_application_stage(
    app_id: str,
    stage_data: ApplicationUpdate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    app = await db.applications.find_one({"id": app_id})
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    
    old_stage = app["stage"]
    now = datetime.now(timezone.utc).isoformat()
    
    await db.applications.update_one(
        {"id": app_id},
        {"$set": {"stage": stage_data.stage, "updated_at": now}}
    )
    if stage_data.stage == "OFFER" and not app.get("offer_status"):
        await db.applications.update_one(
            {"id": app_id},
            {"$set": {"offer_status": "SENT"}},
        )
    
    # Log stage change
    history_doc = {
        "id": str(uuid.uuid4()),
        "application_id": app_id,
        "from_stage": old_stage,
        "to_stage": stage_data.stage,
        "reason": stage_data.reason,
        "changed_by": current_user["id"],
        "changed_at": now,
    }
    if stage_data.stage == "OFFER":
        history_doc["offer_status"] = app.get("offer_status") or "SENT"
    await db.application_stage_history.insert_one(history_doc)

    invalidate_hiring_pack_cache(reason="application_stage_updated")
    
    updated_app = await db.applications.find_one({"id": app_id}, {"_id": 0})
    candidate = await db.candidates.find_one({"id": updated_app["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": updated_app["job_id"]}, {"_id": 0})
    fit_score = await db.fit_scores.find_one({"id": updated_app.get("fit_score_id")}, {"_id": 0})
    
    # Send stage change notification
    if job and candidate:
        background_tasks.add_task(
            notify_stage_change,
            app_id,
            old_stage,
            stage_data.stage,
            candidate["id"],
            job["title"],
            current_user["full_name"]
        )
    
    return ApplicationResponse(
        **updated_app,
        fit_score=fit_score,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]} if job else None
    )


def _parse_stage_dt(value: Any) -> Optional[datetime]:
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z", "+00:00"))
    except ValueError:
        return None


@api_router.get("/applications/{app_id}/stage-history", response_model=List[ApplicationStageHistoryItem])
async def get_application_stage_history(
    app_id: str,
    current_user: dict = Depends(get_current_user),
):
    app = await db.applications.find_one({"id": app_id}, {"_id": 0, "id": 1, "updated_at": 1})
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")

    rows = await db.application_stage_history.find(
        {"application_id": app_id},
        {"_id": 0, "from_stage": 1, "to_stage": 1, "changed_at": 1, "reason": 1, "offer_status": 1},
    ).sort("changed_at", 1).to_list(500)

    out: List[ApplicationStageHistoryItem] = []
    for i, row in enumerate(rows):
        entered = _parse_stage_dt(row.get("changed_at"))
        days = None
        if entered:
            if i + 1 < len(rows):
                exited = _parse_stage_dt(rows[i + 1].get("changed_at"))
            else:
                exited = _parse_stage_dt(app.get("updated_at"))
            if exited and exited > entered:
                days = round((exited - entered).total_seconds() / 86400.0, 1)
        out.append(
            ApplicationStageHistoryItem(
                from_stage=row.get("from_stage"),
                to_stage=str(row.get("to_stage") or ""),
                changed_at=str(row.get("changed_at") or ""),
                days_in_stage=days,
                reason=row.get("reason"),
                offer_status=row.get("offer_status"),
            )
        )
    return out


@api_router.patch("/applications/{app_id}/offer-status", response_model=ApplicationResponse)
async def update_application_offer_status(
    app_id: str,
    body: OfferStatusUpdate,
    current_user: dict = Depends(get_current_user),
):
    app = await db.applications.find_one({"id": app_id})
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    if app.get("stage") != "OFFER":
        raise HTTPException(status_code=400, detail="Offer status can only be updated for applications in OFFER stage")

    old_status = app.get("offer_status")
    now = datetime.now(timezone.utc).isoformat()
    await db.applications.update_one(
        {"id": app_id},
        {"$set": {"offer_status": body.offer_status, "updated_at": now}},
    )
    await db.application_stage_history.insert_one({
        "id": str(uuid.uuid4()),
        "application_id": app_id,
        "from_stage": "OFFER",
        "to_stage": "OFFER",
        "offer_status": body.offer_status,
        "reason": f"Offer status: {old_status or 'unset'} → {body.offer_status}",
        "changed_by": current_user["id"],
        "changed_at": now,
    })
    invalidate_hiring_pack_cache(reason="offer_status_updated")

    updated_app = await db.applications.find_one({"id": app_id}, {"_id": 0})
    candidate = await db.candidates.find_one({"id": updated_app["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": updated_app["job_id"]}, {"_id": 0})
    fit_score = await db.fit_scores.find_one({"id": updated_app.get("fit_score_id")}, {"_id": 0})
    return ApplicationResponse(
        **updated_app,
        fit_score=fit_score,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]} if job else None
    )

@api_router.get("/pipeline/{job_id}")
async def get_pipeline(job_id: str, current_user: dict = Depends(get_current_user)):
    """Get applications grouped by stage for a job"""
    stages = ["SOURCED", "SCREENING", "ASSESSMENT_SENT", "ASSESSMENT_CLEARED", 
              "INTERVIEW_1", "INTERVIEW_2", "INTERVIEW_3", "HR_ROUND", 
              "OFFER", "OFFER_ACCEPTED", "JOINED", "REJECTED", "DROPPED"]
    
    pipeline = {}
    for stage in stages:
        apps = await db.applications.find(
            {"job_id": job_id, "stage": stage},
            {"_id": 0}
        ).to_list(100)
        
        enriched = []
        for app in apps:
            candidate = await db.candidates.find_one({"id": app["candidate_id"]}, {"_id": 0})
            fit_score = await db.fit_scores.find_one({"id": app.get("fit_score_id")}, {"_id": 0})
            enriched.append({
                **app,
                "candidate": candidate,
                "fit_score": fit_score
            })
        
        pipeline[stage] = enriched
    
    return pipeline

# ========================
# REFERRAL ROUTES
# ========================


def _apply_resume_extraction_to_candidate_doc(candidate_doc: Dict[str, Any], extracted: Dict[str, Any]) -> None:
    """Merge AI resume extraction into candidate_doc in place (form fields take precedence when present)."""
    if not extracted:
        return
    if not (candidate_doc.get("full_name") or "").strip() and extracted.get("full_name"):
        candidate_doc["full_name"] = extracted["full_name"]
    if not candidate_doc.get("email") and extracted.get("email"):
        candidate_doc["email"] = extracted["email"]
    if not candidate_doc.get("phone") and extracted.get("phone"):
        candidate_doc["phone"] = extracted["phone"]
    skills_raw = extracted.get("skills") or []
    if isinstance(skills_raw, list) and skills_raw:
        candidate_doc["skills"] = [
            {"skill_name": str(s), "proficiency": None}
            for s in skills_raw
            if s and str(s).strip()
        ]
    if extracted.get("experience"):
        candidate_doc["experience"] = extracted["experience"]
    if extracted.get("education"):
        candidate_doc["education"] = extracted["education"]
    if extracted.get("headline"):
        candidate_doc["headline"] = extracted["headline"]
    if extracted.get("location"):
        candidate_doc["location"] = extracted["location"]
    if extracted.get("total_experience_years") is not None:
        candidate_doc["total_experience_years"] = extracted["total_experience_years"]


async def _referral_create_application_with_fit(
    job: Dict[str, Any],
    job_id: str,
    candidate_id: str,
    changed_by: str,
) -> Dict[str, Any]:
    """
    Create SOURCED application + persisted fit score (job skills + JD vs candidate skills + resume_text),
    same contract as standard applications.
    """
    candidate = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=500, detail="Candidate record missing after referral insert")

    fit_result = compute_basic_fit_score(job, candidate)
    det = compute_match_score(job, candidate)
    fit_result["ranking_explainability"] = {
        "score_source": fit_result.get("score_source", "basic"),
        "weights_applied": (job.get("scoring_rubric") or {}).get("weights") or {},
        "deterministic_match": det,
        "score_factors": fit_result.get("score_factors"),
        "narrative": fit_result.get("explanation"),
    }

    now = datetime.now(timezone.utc).isoformat()
    fs_id = str(uuid.uuid4())
    fit_score_doc = {
        "id": fs_id,
        "job_id": job_id,
        "candidate_id": candidate_id,
        **{k: v for k, v in fit_result.items()},
        "computed_at": now,
    }
    await db.fit_scores.insert_one(fit_score_doc)

    app_id = str(uuid.uuid4())
    app_doc = {
        "id": app_id,
        "job_id": job_id,
        "candidate_id": candidate_id,
        "stage": "SOURCED",
        "status": "ACTIVE",
        "fit_score_id": fs_id,
        "created_at": now,
        "updated_at": now,
    }
    await db.applications.insert_one(app_doc)

    await db.application_stage_history.insert_one({
        "id": str(uuid.uuid4()),
        "application_id": app_id,
        "from_stage": None,
        "to_stage": "SOURCED",
        "changed_by": changed_by,
        "changed_at": now,
    })

    invalidate_hiring_pack_cache(reason="application_sourced")

    return {
        "final_score": fit_result.get("final_score"),
        "must_have_ok": fit_result.get("must_have_ok"),
        "score_source": fit_result.get("score_source"),
        "skill_match_pct": fit_result.get("skill_match_pct"),
        "activity_match_pct": fit_result.get("activity_match_pct"),
        "ranking_explainability": fit_result.get("ranking_explainability"),
    }


@api_router.post("/referrals", response_model=ReferralResponse)
async def create_referral(referral_data: ReferralCreate, current_user: dict = Depends(get_current_user)):
    # Check job exists
    job = await db.jobs.find_one({"id": referral_data.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Create candidate from referral
    candidate_id = str(uuid.uuid4())
    resume_text = (referral_data.resume_text or "").strip() or None
    candidate_doc = {
        "id": candidate_id,
        "full_name": referral_data.candidate_name,
        "email": referral_data.candidate_email,
        "phone": referral_data.candidate_phone,
        "source": "REFERRAL",
        "resume_text": resume_text,
        "skills": [],
        "experience": [],
        "created_at": datetime.now(timezone.utc).isoformat()
    }

    # Structured extract from pasted resume / summary (length gate avoids tiny blurbs)
    if resume_text and len(resume_text) >= 80:
        extracted = await extract_resume_data_with_ai(resume_text)
        _apply_resume_extraction_to_candidate_doc(candidate_doc, extracted)

    rh = _resume_content_hash(candidate_doc.get("resume_text"))
    if rh:
        candidate_doc["resume_content_hash"] = rh

    dedupe_email = candidate_doc.get("email") or referral_data.candidate_email
    if dedupe_email:
        existing = await db.referrals.find_one({
            "job_id": referral_data.job_id,
            "referred_by": current_user["id"],
            "candidate_email": dedupe_email,
        })
        if existing:
            raise HTTPException(status_code=400, detail="Candidate already referred for this job")

    await db.candidates.insert_one(candidate_doc)

    await trigger_auto_analyze_if_eligible(
        db,
        candidate_id=candidate_id,
        resume_text=candidate_doc.get("resume_text"),
        created_by=current_user.get("id") or "",
        job_id=referral_data.job_id,
    )

    # Create referral
    referral_id = str(uuid.uuid4())
    referral_doc = {
        "id": referral_id,
        "job_id": referral_data.job_id,
        "candidate_id": candidate_id,
        "candidate_email": dedupe_email,
        "referred_by": current_user["id"],
        "referred_by_name": current_user["full_name"],
        "note": referral_data.note,
        "status": "PENDING",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.referrals.insert_one(referral_doc)

    fit_summary = await _referral_create_application_with_fit(
        job, referral_data.job_id, candidate_id, current_user["id"]
    )

    stored = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})

    return ReferralResponse(
        id=referral_doc["id"],
        job_id=referral_doc["job_id"],
        candidate_id=referral_doc["candidate_id"],
        referred_by=referral_doc["referred_by"],
        note=referral_doc["note"],
        status=referral_doc["status"],
        candidate=stored,
        created_at=referral_doc["created_at"],
        fit_score=fit_summary,
    )


@api_router.post("/referrals/with-resume", response_model=ReferralResponse)
async def create_referral_with_resume(
    job_id: str = Form(...),
    candidate_name: str = Form(...),
    candidate_email: Optional[str] = Form(None),
    candidate_phone: Optional[str] = Form(None),
    note: Optional[str] = Form(None),
    resume_text_extra: Optional[str] = Form(None),
    resume_file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user),
):
    """
    Referral with optional PDF/DOCX resume attachment.
    Parses CV text, runs AI extraction for skills/experience, persists candidate,
    then creates pipeline application with job vs candidate fit score.
    """
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    resume_text = ""
    resume_filename = None
    if resume_file and resume_file.filename:
        resume_filename = resume_file.filename
        fn = resume_filename.lower()
        if not (fn.endswith(".pdf") or fn.endswith(".docx")):
            raise HTTPException(status_code=400, detail="Resume must be PDF or DOCX")
        content = await resume_file.read()
        if len(content) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Resume file exceeds 10MB limit")
        if fn.endswith(".pdf"):
            resume_text = parse_pdf_resume(content) or ""
        else:
            resume_text = parse_docx_resume(content) or ""
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from resume file")

    extra = (resume_text_extra or "").strip()
    if extra:
        resume_text = f"{resume_text}\n\n{extra}".strip() if resume_text else extra

    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="Attach a resume file or paste resume text")

    extracted = await extract_resume_data_with_ai(resume_text)

    candidate_id = str(uuid.uuid4())
    candidate_doc = {
        "id": candidate_id,
        "full_name": (candidate_name or "").strip() or (extracted.get("full_name") if extracted else None) or "Referral Candidate",
        "email": candidate_email.strip() if candidate_email else (extracted.get("email") if extracted else None),
        "phone": candidate_phone.strip() if candidate_phone else (extracted.get("phone") if extracted else None),
        "source": "REFERRAL",
        "resume_text": resume_text,
        "skills": [],
        "experience": [],
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    if resume_filename:
        candidate_doc["resume_filename"] = resume_filename

    _apply_resume_extraction_to_candidate_doc(candidate_doc, extracted)
    # Form name always wins when provided
    if (candidate_name or "").strip():
        candidate_doc["full_name"] = candidate_name.strip()

    rh = _resume_content_hash(resume_text)
    if rh:
        candidate_doc["resume_content_hash"] = rh

    dedupe_email = candidate_doc.get("email")
    if dedupe_email:
        existing = await db.referrals.find_one({
            "job_id": job_id,
            "referred_by": current_user["id"],
            "candidate_email": dedupe_email,
        })
        if existing:
            raise HTTPException(status_code=400, detail="Candidate already referred for this job")

    await db.candidates.insert_one(candidate_doc)

    await trigger_auto_analyze_if_eligible(
        db,
        candidate_id=candidate_id,
        resume_text=resume_text,
        created_by=current_user.get("id") or "",
        job_id=job_id,
    )

    referral_id = str(uuid.uuid4())
    referral_doc = {
        "id": referral_id,
        "job_id": job_id,
        "candidate_id": candidate_id,
        "candidate_email": dedupe_email,
        "referred_by": current_user["id"],
        "referred_by_name": current_user["full_name"],
        "note": note,
        "status": "PENDING",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.referrals.insert_one(referral_doc)

    fit_summary = await _referral_create_application_with_fit(job, job_id, candidate_id, current_user["id"])
    stored = await db.candidates.find_one({"id": candidate_id}, {"_id": 0})

    return ReferralResponse(
        id=referral_doc["id"],
        job_id=referral_doc["job_id"],
        candidate_id=referral_doc["candidate_id"],
        referred_by=referral_doc["referred_by"],
        note=referral_doc["note"],
        status=referral_doc["status"],
        candidate=stored,
        created_at=referral_doc["created_at"],
        fit_score=fit_summary,
    )

@api_router.get("/referrals", response_model=List[ReferralResponse])
async def list_referrals(current_user: dict = Depends(get_current_user)):
    query = {"referred_by": current_user["id"]}
    referrals = await db.referrals.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    result = []
    for ref in referrals:
        candidate = await db.candidates.find_one({"id": ref["candidate_id"]}, {"_id": 0})
        result.append(ReferralResponse(**ref, candidate=candidate))
    
    return result

@api_router.get("/referrals/all", response_model=List[ReferralResponse])
async def list_all_referrals(current_user: dict = Depends(get_current_user)):
    """Admin view - all referrals"""
    referrals = await db.referrals.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    
    result = []
    for ref in referrals:
        candidate = await db.candidates.find_one({"id": ref["candidate_id"]}, {"_id": 0})
        result.append(ReferralResponse(**ref, candidate=candidate))
    
    return result

# ========================
# ASSESSMENT ROUTES — see create_assessments_router() include below
# ========================

# ========================
# MATCH/SCORING ROUTES
# ========================

@api_router.post("/match/{job_id}")
async def match_candidates_to_job(job_id: str, current_user: dict = Depends(get_current_user)):
    """Find matching candidates for a job and compute scores.

    Order of operations:
    1) Search local Talent Pool (`candidates`) first (fast, avoids connector latency).
    2) If we still don't have enough candidates to rank, ingest from external connectors (best-effort).
    3) Compute deterministic/basic scores (no LLM fan-out) and return top matches.
    """
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    existing_apps = await db.applications.find({"job_id": job_id}).to_list(1000)
    applied_ids = applied_ids_excluding_fit_seeds(existing_apps)

    job_skill_names = [
        str(s.get("skill_name")).strip()
        for s in (job.get("skills") or [])
        if isinstance(s, dict) and s.get("skill_name")
    ]
    skill_or = []
    for sk in job_skill_names[:30]:
        skill_or.append({"skills.skill_name": {"$regex": re.escape(sk), "$options": "i"}})

    candidates = await gather_job_match_candidates(
        db, job, job_id, applied_ids, skill_or, per_bucket=200, max_total=600
    )

    if sum(1 for c in candidates if is_linkedin_sourced_candidate(c)) < 10:
        try:
            await ingest_candidates_for_job(job, total_limit=200)
        except Exception as e:
            logger.error(f"Connector ingestion failed for job {job_id}: {e}")
        extra = await gather_job_match_candidates(
            db, job, job_id, applied_ids, skill_or, per_bucket=80, max_total=600
        )
        seen_ids = {c["id"] for c in candidates if c.get("id")}
        for c in extra:
            if c.get("id") and c["id"] not in seen_ids:
                candidates.append(c)
                seen_ids.add(c["id"])

    deduped: List[Dict[str, Any]] = []
    seen_candidate_ids: Set[str] = set()
    for c in candidates[:500]:
        cid = c.get("id")
        if not cid or cid in seen_candidate_ids:
            continue
        seen_candidate_ids.add(cid)
        deduped.append(c)
    candidates = deduped

    # Bulk match must stay fast and timeout-safe: N×LLM calls (see compute_fit_score) routinely
    # exceed reverse-proxy limits (e.g. nginx default 60s) and are poor UX for "Play Demo".
    results = []
    for candidate in candidates:
        fit_result = compute_basic_fit_score(job, candidate)
        persisted = await load_persisted_fit_score(db, job_id, candidate.get("id") or "")
        fit_result = merge_fit_with_seed_persisted(job_id, candidate, fit_result, persisted)
        det = compute_match_score(job, candidate)
        fit_result["ranking_explainability"] = {
            "score_source": fit_result.get("score_source", "basic"),
            "weights_applied": (job.get("scoring_rubric") or {}).get("weights") or {},
            "deterministic_match": det,
            "score_factors": fit_result.get("score_factors"),
            "narrative": fit_result.get("explanation"),
        }
        results.append({
            "candidate": candidate,
            "fit_score": fit_result
        })
    
    # Grid order (3 cols): Excel | Talent pool | AI-generated fit 90%+, repeating for 50 matches.
    top50 = order_job_match_results(
        results, job_id=job_id, total_limit=DEFAULT_TOTAL_MATCH_LIMIT
    )
    bucket_counts = count_match_buckets(top50, job_id=job_id)

    # Create HR interview proposals for the top ranked matches (no auto-booking).
    try:
        await generate_interview_proposals_for_top_matches(
            job_id=job_id,
            top_matches=top50,
            created_by=current_user["id"],
            top_proposals=20,
            mode="VIRTUAL",
        )
    except Exception as e:
        logger.error(f"Failed to generate interview proposals for job {job_id}: {e}")

    try:
        await log_find_matches_event(
            db,
            job_id=job_id,
            user_id=current_user["id"],
            match_count=len(top50),
        )
    except Exception as e:
        logger.error(f"Failed to log find_matches analytics event for job {job_id}: {e}")

    return {
        "job_id": job_id,
        "matches": top50,
        "excel_count": bucket_counts["excel_count"],
        "talent_pool_count": bucket_counts["talent_pool_count"],
        "ai_high_match_count": bucket_counts["ai_high_match_count"],
        "source_order": ["excel", "talent_pool", "ai_high_match"],
    }

@api_router.post("/jobs/{job_id}/demo-candidates", response_model=DemoCandidatesResponse)
async def generate_demo_candidates(
    job_id: str,
    req: DemoCandidatesRequest = DemoCandidatesRequest(),
    current_user: dict = Depends(get_current_user)
):
    """
    Demo utility:
    - generates dummy candidates for the given job
    - stores them in Mongo `candidates`
    - does NOT create applications/pipeline stages (matches are computed by the existing /match route)
    """
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    count = int(req.count or 50)
    if count < 1:
        count = 1
    if count > 200:
        count = 200

    job_skills = job.get("skills", []) or []
    must_have = [s.get("skill_name") for s in job_skills if s.get("skill_type") == "MUST_HAVE" and s.get("skill_name")]
    good_to_have = [s.get("skill_name") for s in job_skills if s.get("skill_type") != "MUST_HAVE" and s.get("skill_name")]

    # Fallback: if no parsed skills exist, derive skills from description (very naive)
    if not must_have and not good_to_have:
        good_to_have = _tokenize(job.get("description") or "") or []

    first_names = ["Aarav", "Diya", "Vihaan", "Anaya", "Riya", "Kabir", "Ishaan", "Saanvi", "Arjun", "Meera", "Kunal", "Naina"]
    last_names = ["Sharma", "Patel", "Singh", "Gupta", "Verma", "Khan", "Das", "Iyer", "Nair", "Chopra", "Rao", "Bose"]
    headline_suffix = ["Engineer", "Specialist", "Developer", "Analyst", "Architect", "Manager", "Consultant", "Lead", "Associate"]

    created = 0

    for i in range(count):
        full_name = f"{first_names[i % len(first_names)]} {last_names[(i + i//2) % len(last_names)]}"
        email = f"demo.{job_id[:8]}.{i}@example.com".lower()
        candidate_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, email))

        # Make top subset match must-haves; others miss some must-haves.
        # This produces a realistic ranked list for demo.
        must_subset = []
        if must_have:
            if i < max(5, count // 2):
                must_subset = list(set(must_have))
            else:
                k = max(0, len(must_have) - random.randint(1, max(1, len(must_have))))
                must_subset = random.sample(list(set(must_have)), k=k) if k > 0 else []

        good_subset = []
        if good_to_have:
            k_good = min(len(good_to_have), 3 + (i % 5))
            good_subset = random.sample(list(set(good_to_have)), k=k_good) if k_good > 0 else []

        skill_names = list(set(must_subset + good_subset))
        skills = [{"skill_name": s, "proficiency": None} for s in skill_names]

        headline = f"{job.get('normalized_title') or job.get('title') or 'Candidate'} {headline_suffix[i % len(headline_suffix)]}"

        resume_text = f"{job.get('description') or ''}\n\nSkills: {', '.join(skill_names)}\nCandidate Headline: {headline}".strip()

        candidate_doc = {
            "id": candidate_id,
            "full_name": full_name,
            "email": email,
            "phone": None,
            "location": job.get("location"),
            "headline": headline,
            "total_experience_years": None,
            "skills": skills,
            "source": "DEMO",
            "experience": [],
            "resume_text": resume_text,
            "created_at": datetime.now(timezone.utc).isoformat()
        }

        await db.candidates.update_one(
            {"id": candidate_id},
            {"$set": candidate_doc},
            upsert=True
        )
        created += 1

    return {"job_id": job_id, "created": created}

@api_router.get("/fit-scores/{job_id}/{candidate_id}", response_model=FitScoreResponse)
async def get_fit_score(job_id: str, candidate_id: str, current_user: dict = Depends(get_current_user)):
    fit_score = await db.fit_scores.find_one(
        {"job_id": job_id, "candidate_id": candidate_id},
        {"_id": 0}
    )
    if not fit_score:
        raise HTTPException(status_code=404, detail="Fit score not found")
    return FitScoreResponse(**fit_score)

# ========================
# DASHBOARD ROUTES
# ========================

async def _attach_hiring_trends(pack: HiringDashboardPack, *, months: int) -> HiringDashboardPack:
    trends_data = await get_hiring_dashboard_trends(db, months=months)
    return pack.model_copy(update={"trends": HiringDashboardTrends(**trends_data)})


@api_router.get("/dashboard/stats", response_model=DashboardStats)
async def get_dashboard_stats(current_user: dict = Depends(get_current_user)):
    total_jobs = await db.jobs.count_documents({})
    open_jobs = await db.jobs.count_documents({"status": "OPEN"})
    total_candidates = await db.candidates.count_documents({})
    total_applications = await db.applications.count_documents({})
    
    # Applications by stage
    stages = ["SOURCED", "SCREENING", "ASSESSMENT_SENT", "ASSESSMENT_CLEARED",
              "INTERVIEW_1", "HR_ROUND", "OFFER", "JOINED", "REJECTED"]
    applications_by_stage = {}
    for stage in stages:
        count = await db.applications.count_documents({"stage": stage})
        applications_by_stage[stage] = count
    
    # Recent activities (last 10 applications)
    recent_apps = await db.applications.find({}, {"_id": 0}).sort("updated_at", -1).to_list(10)
    recent_activities = []
    for app in recent_apps:
        candidate = await db.candidates.find_one({"id": app["candidate_id"]}, {"_id": 0})
        job = await db.jobs.find_one({"id": app["job_id"]}, {"_id": 0})
        recent_activities.append({
            "type": "application",
            "candidate_name": candidate.get("full_name") if candidate else "Unknown",
            "job_title": job.get("title") if job else "Unknown",
            "stage": app["stage"],
            "timestamp": app["updated_at"]
        })
    
    return DashboardStats(
        total_jobs=total_jobs,
        open_jobs=open_jobs,
        total_candidates=total_candidates,
        total_applications=total_applications,
        applications_by_stage=applications_by_stage,
        recent_activities=recent_activities
    )


@api_router.get("/dashboard/hiring-pack", response_model=HiringDashboardPack)
async def get_hiring_dashboard_pack(
    window_days: int = 30,
    department: Optional[str] = None,
    scope: str = "all",
    job_id: Optional[str] = None,
    owner_id: Optional[str] = None,
    include_trends: bool = True,
    trends_months: int = 6,
    current_user: dict = Depends(get_current_user),
):
    """Smart Hiring Dashboard — windowed KPIs, funnel, source mix, alerts, and optional embedded trends."""
    scope, department, owner_id, job_id = await enforce_hiring_dashboard_scope(
        db,
        current_user=current_user,
        scope=scope,
        department=department,
        owner_id=owner_id,
        job_id=job_id,
    )
    cache_key = (
        f"{window_days}|{department or ''}|{scope}|{job_id or ''}|{owner_id or ''}|{current_user.get('id') or ''}"
    )
    cached = get_cached_hiring_pack(cache_key)
    if cached is not None:
        HIRING_PACK_REQUESTS.labels(cache_hit="true").inc()
        pack = cached.model_copy(update={"data_freshness": "cached"}) if hasattr(cached, "model_copy") else cached
    else:
        pack_start = time.perf_counter()
        pack = await build_hiring_dashboard_pack(
            db,
            window_days=window_days,
            department=department,
            scope=scope,
            user_id=current_user.get("id"),
            job_id=job_id,
            owner_id=owner_id,
        )
        duration = time.perf_counter() - pack_start
        HIRING_PACK_DURATION_SECONDS.observe(duration)
        HIRING_PACK_REQUESTS.labels(cache_hit="false").inc()
        log_slow_hiring_pack_query(
            duration,
            window_days=window_days,
            scope=scope,
            department=department,
            job_id=job_id,
            owner_id=owner_id,
        )
        set_cached_hiring_pack(cache_key, pack)

    if include_trends:
        pack = await _attach_hiring_trends(pack, months=trends_months)
    return pack


@api_router.get("/dashboard/trends", response_model=HiringDashboardTrends)
async def get_hiring_dashboard_trends_route(
    months: int = 6,
    current_user: dict = Depends(get_current_user),
):
    data = await get_hiring_dashboard_trends(db, months=months)
    return HiringDashboardTrends(**data)


@api_router.get("/dashboard/trends/health", response_model=HiringSnapshotHealth)
async def get_hiring_snapshot_health_route(
    current_user: dict = Depends(get_current_user),
):
    """Snapshot cron maturity and staleness for Smart Hiring trends."""
    data = await get_hiring_snapshot_health(db)
    return HiringSnapshotHealth(**data)


@api_router.get("/dashboard/hiring-alerts/dismissals", response_model=HiringAlertDismissalsResponse)
async def get_hiring_alert_dismissals(current_user: dict = Depends(get_current_user)):
    """List alert IDs dismissed by the current user (server-side persistence)."""
    user_id = current_user.get("id") or ""
    dismissed = await list_dismissed_alert_ids(db, user_id)
    return HiringAlertDismissalsResponse(dismissed=dismissed)


@api_router.post("/dashboard/hiring-alerts/dismissals")
async def post_hiring_alert_dismissal(
    body: HiringAlertDismissRequest,
    current_user: dict = Depends(get_current_user),
):
    user_id = current_user.get("id") or ""
    if not body.alert_id.strip():
        raise HTTPException(status_code=400, detail="alert_id is required")
    await dismiss_hiring_alert(db, user_id, body.alert_id.strip())
    return {"ok": True, "alert_id": body.alert_id.strip()}


@api_router.delete("/dashboard/hiring-alerts/dismissals/{alert_id}")
async def delete_hiring_alert_dismissal(alert_id: str, current_user: dict = Depends(get_current_user)):
    user_id = current_user.get("id") or ""
    await restore_hiring_alert(db, user_id, alert_id)
    return {"ok": True, "alert_id": alert_id}


@api_router.get("/admin/hiring-dashboard/config", response_model=HiringDashboardConfigResponse)
async def get_admin_hiring_dashboard_config(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    config = await get_hiring_dashboard_config(db)
    return HiringDashboardConfigResponse(**config_to_json(config))


@api_router.put("/admin/hiring-dashboard/config", response_model=HiringDashboardConfigResponse)
async def put_admin_hiring_dashboard_config(
    body: HiringDashboardConfigUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    existing = await get_hiring_dashboard_config(db)
    payload = config_to_json(existing)
    for key, value in body.model_dump(exclude_unset=True).items():
        if value is not None:
            payload[key] = value
    config = await upsert_hiring_dashboard_config(db, payload)
    invalidate_hiring_pack_cache(reason="hiring_dashboard_config_updated")
    return HiringDashboardConfigResponse(**config_to_json(config))


@api_router.post("/admin/hiring-dashboard/snapshot")
async def admin_hiring_dashboard_snapshot(current_user: dict = Depends(get_current_user)):
    """Ops/cron: persist daily hiring dashboard snapshot for trend charts."""
    _require_admin(current_user)
    pack = await build_hiring_dashboard_pack(db, window_days=30)
    doc = await write_hiring_dashboard_snapshot(db, pack.model_dump())
    seeded = await seed_hiring_snapshots_if_sparse(db)
    return {"ok": True, "snapshot": doc, "seeded_snapshots": seeded}


@api_router.post("/admin/hiring-dashboard/snapshot-cron")
async def hiring_dashboard_snapshot_cron(request: Request):
    """Scheduled snapshot (no JWT). Requires HIRING_SNAPSHOT_TOKEN and X-Hiring-Snapshot-Token."""
    expected = (os.environ.get("HIRING_SNAPSHOT_TOKEN") or "").strip()
    got = (request.headers.get("X-Hiring-Snapshot-Token") or "").strip()
    if not expected or got != expected:
        raise HTTPException(status_code=401, detail="Invalid or missing X-Hiring-Snapshot-Token")
    pack = await build_hiring_dashboard_pack(db, window_days=30)
    doc = await write_hiring_dashboard_snapshot(db, pack.model_dump())
    seeded = await seed_hiring_snapshots_if_sparse(db)
    return {"ok": True, "snapshot": doc, "seeded_snapshots": seeded}

# ========================
# PHASE-1 WORKFORCE MODULES
# ========================

@api_router.post("/employees", response_model=EmployeeResponse)
async def create_employee(payload: EmployeeCreate, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_write")
    existing = await db.employees.find_one({"employee_code": payload.employee_code}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Employee code already exists")

    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": str(uuid.uuid4()),
        "employee_code": payload.employee_code.strip(),
        "full_name": payload.full_name.strip(),
        "email": payload.email,
        "department": payload.department.strip(),
        "role_title": payload.role_title.strip(),
        "manager_id": payload.manager_id,
        "location": payload.location,
        "status": (payload.status or "ACTIVE").upper(),
        "skills": [s.strip() for s in (payload.skills or []) if isinstance(s, str) and s.strip()],
        "join_date": payload.join_date,
        "created_at": now,
        "updated_at": now,
    }
    if payload.compensation_band is not None:
        doc["compensation_band"] = str(payload.compensation_band).upper()
    if payload.last_promotion_at is not None:
        doc["last_promotion_at"] = payload.last_promotion_at
    if payload.high_performer is not None:
        doc["high_performer"] = bool(payload.high_performer)
    if payload.critical_role is not None:
        doc["critical_role"] = bool(payload.critical_role)
    if payload.comp_market_percentile is not None:
        doc["comp_market_percentile"] = float(payload.comp_market_percentile)
    if payload.hris_last_sync_at is not None:
        doc["hris_last_sync_at"] = payload.hris_last_sync_at
    if payload.hris_comp_source is not None:
        doc["hris_comp_source"] = str(payload.hris_comp_source).strip()[:120]
    await db.employees.insert_one(doc)
    return _employee_doc_to_response(doc)

@api_router.get("/employees", response_model=List[EmployeeResponse])
async def list_employees(
    status: Optional[str] = None,
    department: Optional[str] = None,
    q: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_read")
    query: Dict[str, Any] = {}
    if status:
        query["status"] = status.upper()
    if department:
        query["department"] = {"$regex": f"^{re.escape(department)}$", "$options": "i"}
    if q:
        query["$or"] = [
            {"full_name": {"$regex": q, "$options": "i"}},
            {"employee_code": {"$regex": q, "$options": "i"}},
            {"email": {"$regex": q, "$options": "i"}},
            {"role_title": {"$regex": q, "$options": "i"}},
        ]
    rows = await db.employees.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    return [_employee_doc_to_response(r) for r in rows]


# ========================
# Resource aliases (Resources == Employees)
# ========================

@api_router.post("/resources/create", response_model=EmployeeResponse)
async def create_resource(payload: EmployeeCreate, current_user: dict = Depends(get_current_user)):
    return await create_employee(payload, current_user)


@api_router.get("/resources/search")
async def search_resources(
    page: int = 1,
    page_size: int = 50,
    status: Optional[str] = None,
    department: Optional[str] = None,
    q: Optional[str] = None,
    skill: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    return await list_employees_paged(
        status=status,
        department=department,
        q=q,
        skill=skill,
        page=page,
        page_size=page_size,
        sort_by="created_at",
        sort_dir="desc",
        current_user=current_user,
    )


@api_router.get("/resources/{employee_id}", response_model=EmployeeResponse)
async def get_resource(employee_id: str, current_user: dict = Depends(get_current_user)):
    return await get_employee(employee_id, current_user)


@api_router.put("/resources/update", response_model=EmployeeResponse)
async def update_resource(
    employee_id: str,
    payload: EmployeeUpdate,
    current_user: dict = Depends(get_current_user),
):
    return await update_employee(employee_id, payload, current_user)


# Must be registered BEFORE /employees/{employee_id} or "paged" is captured as an employee_id (404).
@api_router.get("/employees/paged")
async def list_employees_paged(
    status: Optional[str] = None,
    department: Optional[str] = None,
    q: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "created_at",
    sort_dir: Literal["asc", "desc"] = "desc",
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_read")
    page = max(1, page)
    page_size = min(max(1, page_size), 200)
    allowed_sort = {"created_at", "updated_at", "full_name", "employee_code", "department", "status"}
    if sort_by not in allowed_sort:
        sort_by = "created_at"
    direction = 1 if sort_dir == "asc" else -1

    query: Dict[str, Any] = {}
    if status:
        query["status"] = status.upper()
    if department:
        query["department"] = {"$regex": f"^{re.escape(department)}$", "$options": "i"}
    if q:
        query["$or"] = [
            {"full_name": {"$regex": q, "$options": "i"}},
            {"employee_code": {"$regex": q, "$options": "i"}},
            {"email": {"$regex": q, "$options": "i"}},
            {"role_title": {"$regex": q, "$options": "i"}},
        ]

    total = await db.employees.count_documents(query)
    rows = await db.employees.find(query, {"_id": 0}).sort(sort_by, direction).skip((page - 1) * page_size).limit(page_size).to_list(page_size)
    return {
        "items": [_employee_doc_to_response(r).model_dump() for r in rows],
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
        "sort_by": sort_by,
        "sort_dir": sort_dir,
    }


@api_router.put("/employees/{employee_id}", response_model=EmployeeResponse)
async def update_employee(employee_id: str, payload: EmployeeUpdate, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_write")
    existing = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Employee not found")

    update_doc = payload.model_dump(exclude_none=True)
    if "skills" in update_doc:
        update_doc["skills"] = [s.strip() for s in (update_doc.get("skills") or []) if isinstance(s, str) and s.strip()]
    prev_status = (existing.get("status") or "ACTIVE").upper()
    if "status" in update_doc and isinstance(update_doc["status"], str):
        update_doc["status"] = update_doc["status"].upper()
        err = validate_direct_status_transition(prev_status, update_doc["status"])
        if err:
            raise HTTPException(status_code=400, detail=err)
    update_doc["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.employees.update_one({"id": employee_id}, {"$set": update_doc})
    if "status" in update_doc and update_doc.get("status") != prev_status:
        await _append_lifecycle_audit(
            employee_code=existing.get("employee_code") or "",
            action="STATUS_DIRECT_UPDATE",
            from_status=prev_status,
            to_status=update_doc.get("status"),
            event_type=None,
            event_id=None,
            actor_id=current_user["id"],
        )
    updated = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    return _employee_doc_to_response(updated)

@api_router.delete("/employees/{employee_id}")
async def delete_employee(employee_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_write")
    result = await db.employees.delete_one({"id": employee_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Employee not found")
    return {"message": "Employee deleted", "id": employee_id}


@api_router.get("/employees/{employee_id}", response_model=EmployeeResponse)
async def get_employee(employee_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_read")
    row = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    if not row:
        raise HTTPException(status_code=404, detail="Employee not found")
    return _employee_doc_to_response(row)


@api_router.get("/employees/{employee_id}/direct-reports", response_model=List[EmployeeResponse])
async def list_direct_reports(employee_id: str, current_user: dict = Depends(get_current_user)):
    """M2-2: employees whose manager_id points at this employee id."""
    _require_phase1_access(current_user, "employees_read")
    boss = await db.employees.find_one({"id": employee_id}, {"_id": 0, "id": 1})
    if not boss:
        raise HTTPException(status_code=404, detail="Employee not found")
    rows = await db.employees.find({"manager_id": employee_id}, {"_id": 0}).sort("full_name", 1).to_list(500)
    return [_employee_doc_to_response(r) for r in rows]


@api_router.get("/employees/{employee_id}/management-chain")
async def get_management_chain(employee_id: str, current_user: dict = Depends(get_current_user)):
    """M2-2: walk manager_id up to root (cycle-safe)."""
    _require_phase1_access(current_user, "employees_read")
    emp = await db.employees.find_one({"id": employee_id}, {"_id": 0})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")
    chain: List[Dict[str, Any]] = []
    seen: Set[str] = set()
    cur_id: Optional[str] = employee_id
    while cur_id and cur_id not in seen:
        seen.add(cur_id)
        doc = await db.employees.find_one({"id": cur_id}, {"_id": 0})
        if not doc:
            break
        chain.append(
            {
                "id": doc.get("id"),
                "employee_code": doc.get("employee_code"),
                "full_name": doc.get("full_name"),
                "role_title": doc.get("role_title"),
                "department": doc.get("department"),
            }
        )
        mid = doc.get("manager_id")
        cur_id = str(mid).strip() if mid else None
    return {"employee_id": employee_id, "chain": chain}


@api_router.get("/org/hierarchy")
async def get_org_hierarchy(
    root_id: Optional[str] = None,
    max_depth: int = 8,
    current_user: dict = Depends(get_current_user),
):
    """
    M2-2: nested tree of direct reports. If root_id omitted, builds multiple trees for employees
    with no manager or missing manager record.
    """
    _require_phase1_access(current_user, "employees_read")
    max_depth = min(max(1, max_depth), 20)
    all_emp = await db.employees.find({}, {"_id": 0}).to_list(10000)
    by_id = {e["id"]: e for e in all_emp if e.get("id")}
    children: Dict[str, List[str]] = defaultdict(list)
    for e in all_emp:
        eid = e.get("id")
        mid = e.get("manager_id")
        if eid and mid and str(mid) in by_id:
            children[str(mid)].append(str(eid))

    def subtree(eid: str, depth: int) -> Optional[Dict[str, Any]]:
        if depth > max_depth:
            return None
        node = by_id.get(eid)
        if not node:
            return None
        kids: List[Dict[str, Any]] = []
        for cid in children.get(eid, []):
            st = subtree(cid, depth + 1)
            if st:
                kids.append(st)
        return {
            "id": node.get("id"),
            "employee_code": node.get("employee_code"),
            "full_name": node.get("full_name"),
            "role_title": node.get("role_title"),
            "department": node.get("department"),
            "reports": kids,
        }

    if root_id:
        if root_id not in by_id:
            raise HTTPException(status_code=404, detail="root employee not found")
        return {"roots": [subtree(root_id, 1)]}

    roots_ids = []
    for e in all_emp:
        eid = e.get("id")
        mid = e.get("manager_id")
        if not eid:
            continue
        if not mid or str(mid) not in by_id:
            roots_ids.append(str(eid))
    return {"roots": [subtree(rid, 1) for rid in roots_ids if subtree(rid, 1)]}


@api_router.post("/workforce/skills", response_model=SkillInventoryResponse)
async def create_skill_inventory(payload: SkillInventoryCreate, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_write")
    key = payload.skill_name.strip().lower()
    if not key:
        raise HTTPException(status_code=400, detail="skill_name is required")
    existing = await db.workforce_skills.find_one({"skill_name_lc": key}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Skill already exists")

    now = datetime.now(timezone.utc).isoformat()
    demand = max(0, int(payload.demand_count))
    supply = max(0, int(payload.supply_count))
    doc = {
        "skill_name": payload.skill_name.strip(),
        "skill_name_lc": key,
        "demand_count": demand,
        "supply_count": supply,
        "gap": max(0, demand - supply),
        "category": payload.category,
        "priority": (payload.priority or "MEDIUM").upper(),
        "notes": payload.notes,
        "updated_at": now,
    }
    await db.workforce_skills.insert_one(doc)
    doc.pop("skill_name_lc", None)
    return SkillInventoryResponse(**doc)

@api_router.get("/workforce/skills", response_model=List[SkillInventoryResponse])
async def list_skill_inventory(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_read")
    skills = await db.workforce_skills.find({}, {"_id": 0, "skill_name_lc": 0}).sort("gap", -1).to_list(500)

    # Auto-derive supply from employee profiles where possible, and merge with manual supply.
    employees = await db.employees.find({}, {"_id": 0, "skills": 1}).to_list(2000)
    derived_supply: Dict[str, int] = {}
    for e in employees:
        uniq = set(s.strip().lower() for s in (e.get("skills") or []) if isinstance(s, str) and s.strip())
        for s in uniq:
            derived_supply[s] = derived_supply.get(s, 0) + 1

    out = []
    for row in skills:
        skill_key = (row.get("skill_name") or "").strip().lower()
        auto_supply = derived_supply.get(skill_key, 0)
        manual_supply = max(0, int(row.get("supply_count") or 0))
        merged_supply = max(manual_supply, auto_supply)
        demand = max(0, int(row.get("demand_count") or 0))
        row["supply_count"] = merged_supply
        row["gap"] = max(0, demand - merged_supply)
        out.append(SkillInventoryResponse(**row))

    out.sort(key=lambda x: x.gap, reverse=True)
    return out

@api_router.put("/workforce/skills/{skill_name}", response_model=SkillInventoryResponse)
async def update_skill_inventory(skill_name: str, payload: SkillInventoryUpdate, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_write")
    key = skill_name.strip().lower()
    existing = await db.workforce_skills.find_one({"skill_name_lc": key}, {"_id": 0, "skill_name_lc": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Skill not found")

    update_doc = payload.model_dump(exclude_none=True)
    if "priority" in update_doc and isinstance(update_doc["priority"], str):
        update_doc["priority"] = update_doc["priority"].upper()

    demand = max(0, int(update_doc.get("demand_count", existing.get("demand_count", 0))))
    supply = max(0, int(update_doc.get("supply_count", existing.get("supply_count", 0))))
    update_doc["demand_count"] = demand
    update_doc["supply_count"] = supply
    update_doc["gap"] = max(0, demand - supply)
    update_doc["updated_at"] = datetime.now(timezone.utc).isoformat()

    await db.workforce_skills.update_one({"skill_name_lc": key}, {"$set": update_doc})
    updated = await db.workforce_skills.find_one({"skill_name_lc": key}, {"_id": 0, "skill_name_lc": 0})
    return SkillInventoryResponse(**updated)

@api_router.delete("/workforce/skills/{skill_name}")
async def delete_skill_inventory(skill_name: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_write")
    key = skill_name.strip().lower()
    result = await db.workforce_skills.delete_one({"skill_name_lc": key})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Skill not found")
    return {"message": "Skill deleted", "skill_name": skill_name}

@api_router.get("/workforce/skills/paged")
async def list_skill_inventory_paged(
    q: Optional[str] = None,
    priority: Optional[Literal["HIGH", "MEDIUM", "LOW"]] = None,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "gap",
    sort_dir: Literal["asc", "desc"] = "desc",
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_read")
    page = max(1, page)
    page_size = min(max(1, page_size), 200)
    allowed_sort = {"skill_name", "demand_count", "supply_count", "gap", "priority", "updated_at"}
    if sort_by not in allowed_sort:
        sort_by = "gap"
    direction = 1 if sort_dir == "asc" else -1

    query: Dict[str, Any] = {}
    if q:
        query["skill_name"] = {"$regex": q, "$options": "i"}
    if priority:
        query["priority"] = priority

    total = await db.workforce_skills.count_documents(query)
    rows = await db.workforce_skills.find(query, {"_id": 0, "skill_name_lc": 0}).sort(sort_by, direction).skip((page - 1) * page_size).limit(page_size).to_list(page_size)
    items = [SkillInventoryResponse(**r).model_dump() for r in rows]
    return {
        "items": items,
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
        "sort_by": sort_by,
        "sort_dir": sort_dir,
    }

async def _write_import_audit(module: str, mode: str, dry_run: bool, summary: Dict[str, Any], details: List[Dict[str, Any]], actor_id: str):
    doc = {
        "id": str(uuid.uuid4()),
        "module": module,
        "mode": mode,
        "dry_run": dry_run,
        "summary": summary,
        "details": details[:1000],
        "created_by": actor_id,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.import_audit_logs.insert_one(doc)
    return doc

@api_router.post("/employees/bulk-import")
async def bulk_import_employees(payload: EmployeeBulkImportRequest, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_write")
    results: List[Dict[str, Any]] = []
    created = updated = failed = 0

    existing = await db.employees.find({}, {"_id": 0, "id": 1, "employee_code": 1}).to_list(100000)
    by_code = {str(r.get("employee_code", "")).strip(): r for r in existing if r.get("employee_code")}

    for idx, row in enumerate(payload.rows, start=1):
        code = (row.employee_code or "").strip()
        full_name = (row.full_name or "").strip()
        department = (row.department or "").strip()
        role_title = (row.role_title or "").strip()
        if not (code and full_name and department and role_title):
            failed += 1
            results.append({
                "row_number": idx,
                "employee_code": code,
                "full_name": full_name,
                "action": "FAILED",
                "reason": "Missing required fields (employee_code, full_name, department, role_title)",
            })
            continue

        row_doc = {
            "employee_code": code,
            "full_name": full_name,
            "email": row.email,
            "department": department,
            "role_title": role_title,
            "manager_id": row.manager_id,
            "location": row.location,
            "status": row.status,
            "skills": [s.strip() for s in (row.skills or []) if s.strip()],
            "join_date": row.join_date,
        }
        extra = row.model_dump(exclude_none=True)
        for k in (
            "compensation_band",
            "last_promotion_at",
            "high_performer",
            "critical_role",
            "comp_market_percentile",
            "hris_last_sync_at",
            "hris_comp_source",
        ):
            if k not in extra:
                continue
            val = extra[k]
            if k == "compensation_band" and val is not None:
                val = str(val).upper()
            row_doc[k] = val

        ex = by_code.get(code)
        if ex:
            if payload.mode == "upsert":
                updated += 1
                results.append({
                    "row_number": idx,
                    "employee_code": code,
                    "full_name": full_name,
                    "action": "UPDATE",
                    "reason": "",
                })
                if not payload.dry_run:
                    row_doc["updated_at"] = datetime.now(timezone.utc).isoformat()
                    await db.employees.update_one({"id": ex["id"]}, {"$set": row_doc})
            else:
                failed += 1
                results.append({
                    "row_number": idx,
                    "employee_code": code,
                    "full_name": full_name,
                    "action": "FAILED",
                    "reason": "Duplicate employee_code",
                })
            continue

        created += 1
        results.append({
            "row_number": idx,
            "employee_code": code,
            "full_name": full_name,
            "action": "CREATE",
            "reason": "",
        })
        if not payload.dry_run:
            now = datetime.now(timezone.utc).isoformat()
            doc = {
                "id": str(uuid.uuid4()),
                **row_doc,
                "created_at": now,
                "updated_at": now,
            }
            await db.employees.insert_one(doc)
            by_code[code] = {"id": doc["id"], "employee_code": code}

    summary = {"created": created, "updated": updated, "failed": failed, "total": len(payload.rows)}
    audit = await _write_import_audit("employees", payload.mode, payload.dry_run, summary, results, current_user["id"])
    return {"summary": summary, "rows": results, "audit_id": audit["id"]}

@api_router.post("/workforce/skills/bulk-import")
async def bulk_import_skills(payload: SkillBulkImportRequest, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_write")
    results: List[Dict[str, Any]] = []
    created = updated = failed = 0

    existing = await db.workforce_skills.find({}, {"_id": 0, "skill_name": 1, "skill_name_lc": 1}).to_list(100000)
    by_key = {str(r.get("skill_name_lc", "")).strip(): r for r in existing if r.get("skill_name_lc")}

    for idx, row in enumerate(payload.rows, start=1):
        name = (row.skill_name or '').strip()
        key = name.lower()
        if not name:
            failed += 1
            results.append({"row_number": idx, "skill_name": name, "action": "FAILED", "reason": "skill_name is required"})
            continue

        demand = max(0, int(row.demand_count))
        supply = max(0, int(row.supply_count))
        row_doc = {
            "skill_name": name,
            "skill_name_lc": key,
            "demand_count": demand,
            "supply_count": supply,
            "gap": max(0, demand - supply),
            "category": row.category,
            "priority": row.priority,
            "notes": row.notes,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }

        ex = by_key.get(key)
        if ex:
            if payload.mode == "upsert":
                updated += 1
                results.append({"row_number": idx, "skill_name": name, "action": "UPDATE", "reason": ""})
                if not payload.dry_run:
                    await db.workforce_skills.update_one({"skill_name_lc": key}, {"$set": row_doc})
            else:
                failed += 1
                results.append({"row_number": idx, "skill_name": name, "action": "FAILED", "reason": "Duplicate skill_name"})
            continue

        created += 1
        results.append({"row_number": idx, "skill_name": name, "action": "CREATE", "reason": ""})
        if not payload.dry_run:
            await db.workforce_skills.insert_one(row_doc)
            by_key[key] = {"skill_name_lc": key}

    summary = {"created": created, "updated": updated, "failed": failed, "total": len(payload.rows)}
    audit = await _write_import_audit("workforce_skills", payload.mode, payload.dry_run, summary, results, current_user["id"])
    return {"summary": summary, "rows": results, "audit_id": audit["id"]}

@api_router.get("/imports/audit")
async def list_import_audits(
    module: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    page = max(1, page)
    page_size = min(max(1, page_size), 100)
    query: Dict[str, Any] = {}
    if module:
        query["module"] = module
    total = await db.import_audit_logs.count_documents(query)
    rows = await db.import_audit_logs.find(query, {"_id": 0}).sort("created_at", -1).skip((page - 1) * page_size).limit(page_size).to_list(page_size)
    return {
        "items": rows,
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
    }

@api_router.get("/executive/kpis", response_model=ExecutiveKpiResponse)
async def get_executive_kpis(
    window_days: int = 30,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    window_days = max(1, min(int(window_days or 30), 365))
    employee_count = await db.employees.count_documents({})
    active_employee_count = await db.employees.count_documents({"status": {"$in": ["ACTIVE", "ONBOARDING"]}})
    attrition_count = await db.employees.count_documents({"status": "EXITED"})
    attrition_rate_pct = round((attrition_count / employee_count) * 100, 2) if employee_count else 0.0

    employees = await db.employees.find({}, {"_id": 0, "skills": 1}).to_list(5000)
    total_skill_entries = 0
    for e in employees:
        total_skill_entries += len([s for s in (e.get("skills") or []) if isinstance(s, str) and s.strip()])
    avg_skills_per_employee = round(total_skill_entries / employee_count, 2) if employee_count else 0.0

    skill_rows = await list_skill_inventory(current_user)
    top_skill_gaps = [
        {
            "skill_name": s.skill_name,
            "demand_count": s.demand_count,
            "supply_count": s.supply_count,
            "gap": s.gap,
            "priority": s.priority,
        }
        for s in skill_rows[:10]
    ]

    hiring_demand_total = sum(s["demand_count"] for s in top_skill_gaps) + sum(
        s.demand_count for s in skill_rows[10:]
    )
    workforce_supply_total = sum(s["supply_count"] for s in top_skill_gaps) + sum(
        s.supply_count for s in skill_rows[10:]
    )
    skill_coverage_pct = round((workforce_supply_total / hiring_demand_total) * 100, 2) if hiring_demand_total else 100.0

    talent_acq = await compute_talent_acquisition_metrics(db, window_days=window_days)

    return ExecutiveKpiResponse(
        employee_count=employee_count,
        active_employee_count=active_employee_count,
        attrition_count=attrition_count,
        attrition_rate_pct=attrition_rate_pct,
        avg_skills_per_employee=avg_skills_per_employee,
        top_skill_gaps=top_skill_gaps,
        hiring_demand_total=hiring_demand_total,
        workforce_supply_total=workforce_supply_total,
        skill_coverage_pct=skill_coverage_pct,
        talent_acquisition=talent_acq,
    )

# ========================
# Phase-2: Employee Lifecycle Management
# ========================

@api_router.post("/employee-lifecycle/events", response_model=EmployeeLifecycleEventResponse)
async def create_employee_lifecycle_event(
    payload: EmployeeLifecycleEventCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "lifecycle_write")

    employee_code = (payload.employee_code or "").strip()
    existing_emp = await db.employees.find_one({"employee_code": employee_code}, {"_id": 0})
    if not existing_emp:
        raise HTTPException(status_code=404, detail="Employee not found")

    cur_status = (existing_emp.get("status") or "ACTIVE").upper()
    ev_err = validate_lifecycle_event_for_status(payload.event_type, cur_status)
    if ev_err:
        raise HTTPException(status_code=400, detail=ev_err)

    rule = approval_rule_for_event(payload.event_type)
    requires_approval = rule is not None

    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": str(uuid.uuid4()),
        "employee_code": employee_code,
        "event_type": payload.event_type,
        "effective_date": payload.effective_date,
        "details": payload.details or {},
        "created_at": now,
        "updated_at": None,
        "processing_status": "PENDING",
        "attempts": 0,
        "processed_at": None,
        "processing_error": None,
        "requires_approval": requires_approval,
        "approval_status": "PENDING" if requires_approval else None,
        "approved_by": None,
        "approved_at": None,
        "rejection_reason": None,
        "escalated_at": None,
    }
    await db.employee_lifecycle_events.insert_one(doc)

    await _m10_publish_safe(
        topic=TOPIC_EMPLOYEE_LIFECYCLE_EVENT_CREATED,
        payload={
            "lifecycle_event_id": doc["id"],
            "employee_code": employee_code,
            "event_type": payload.event_type,
            "requires_approval": requires_approval,
        },
        idempotency_key=f"lifecycle_event:{doc['id']}",
        correlation_id=doc["id"],
    )

    await _append_lifecycle_audit(
        employee_code=employee_code,
        action="EVENT_CREATED",
        from_status=cur_status,
        to_status=target_status_for_event(payload.event_type),
        event_type=payload.event_type,
        event_id=doc["id"],
        actor_id=current_user["id"],
        notes="awaiting_approval" if requires_approval else "auto_process",
    )

    if not requires_approval:
        background_tasks.add_task(process_employee_lifecycle_event, doc["id"])
    return EmployeeLifecycleEventResponse(**doc)


@api_router.get("/employee-lifecycle/events", response_model=EmployeeLifecycleEventsPagedResponse)
async def list_employee_lifecycle_events(
    employee_code: Optional[str] = None,
    event_type: Optional[EmployeeLifecycleEventType] = None,
    q: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "created_at",
    sort_dir: Literal["asc", "desc"] = "desc",
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_read")

    page = max(1, page)
    page_size = min(max(1, page_size), 200)
    allowed_sort = {"created_at", "updated_at", "effective_date", "employee_code", "event_type"}
    if sort_by not in allowed_sort:
        sort_by = "created_at"
    direction = 1 if sort_dir == "asc" else -1

    query: Dict[str, Any] = {}
    if employee_code:
        query["employee_code"] = (employee_code or "").strip()
    if event_type:
        query["event_type"] = event_type
    if q:
        query["$or"] = [
            {"employee_code": {"$regex": q, "$options": "i"}},
            {"event_type": {"$regex": q, "$options": "i"}},
        ]

    total = await db.employee_lifecycle_events.count_documents(query)
    rows = (
        await db.employee_lifecycle_events.find(query, {"_id": 0})
        .sort(sort_by, direction)
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )

    return EmployeeLifecycleEventsPagedResponse(
        items=[EmployeeLifecycleEventResponse(**r) for r in rows],
        page=page,
        page_size=page_size,
        total=total,
        total_pages=max(1, (total + page_size - 1) // page_size),
        sort_by=sort_by,
        sort_dir=sort_dir,
    )


@api_router.put("/employee-lifecycle/events/{event_id}", response_model=EmployeeLifecycleEventResponse)
async def update_employee_lifecycle_event(
    event_id: str,
    payload: EmployeeLifecycleEventUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "lifecycle_write")
    existing = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Lifecycle event not found")

    update_doc = payload.model_dump(exclude_none=True)
    if update_doc.get("details") is None:
        update_doc.pop("details", None)

    now = datetime.now(timezone.utc).isoformat()
    update_doc["updated_at"] = now
    await db.employee_lifecycle_events.update_one({"id": event_id}, {"$set": update_doc})

    updated = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    return EmployeeLifecycleEventResponse(**updated)


@api_router.delete("/employee-lifecycle/events/{event_id}")
async def delete_employee_lifecycle_event(
    event_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "lifecycle_write")
    result = await db.employee_lifecycle_events.delete_one({"id": event_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Lifecycle event not found")
    return {"message": "Lifecycle event deleted", "event_id": event_id}


@api_router.post("/employee-lifecycle/events/{event_id}/approve", response_model=EmployeeLifecycleEventResponse)
async def approve_lifecycle_event(
    event_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    """M2-2: approval matrix — allowed roles per event type."""
    _require_phase1_access(current_user, "lifecycle_write")
    ev = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    if not ev:
        raise HTTPException(status_code=404, detail="Lifecycle event not found")
    if not ev.get("requires_approval"):
        raise HTTPException(status_code=400, detail="Event does not require approval")
    if ev.get("approval_status") != "PENDING":
        raise HTTPException(status_code=400, detail="Event is not pending approval")

    rule = approval_rule_for_event(ev.get("event_type") or "")
    if not rule:
        raise HTTPException(status_code=400, detail="No approval rule for this event")
    allowed_roles, _ = rule
    role = (current_user.get("role") or "").lower()
    if role not in allowed_roles:
        raise HTTPException(status_code=403, detail="Not authorized to approve this event")

    now = datetime.now(timezone.utc).isoformat()
    await db.employee_lifecycle_events.update_one(
        {"id": event_id},
        {"$set": {"approval_status": "APPROVED", "approved_by": current_user["id"], "approved_at": now, "updated_at": now}},
    )
    await _append_lifecycle_audit(
        employee_code=ev.get("employee_code") or "",
        action="EVENT_APPROVED",
        from_status=None,
        to_status=None,
        event_type=ev.get("event_type"),
        event_id=event_id,
        actor_id=current_user["id"],
    )
    background_tasks.add_task(process_employee_lifecycle_event, event_id)
    updated = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    return EmployeeLifecycleEventResponse(**updated)


@api_router.post("/employee-lifecycle/events/{event_id}/reject", response_model=EmployeeLifecycleEventResponse)
async def reject_lifecycle_event(
    event_id: str,
    payload: LifecycleEventRejectRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "lifecycle_write")
    ev = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    if not ev:
        raise HTTPException(status_code=404, detail="Lifecycle event not found")
    if not ev.get("requires_approval") or ev.get("approval_status") != "PENDING":
        raise HTTPException(status_code=400, detail="Event is not pending approval")
    rule = approval_rule_for_event(ev.get("event_type") or "")
    if not rule:
        raise HTTPException(status_code=400, detail="No approval rule for this event")
    allowed_roles, _ = rule
    role = (current_user.get("role") or "").lower()
    if role not in allowed_roles:
        raise HTTPException(status_code=403, detail="Not authorized to reject this event")

    now = datetime.now(timezone.utc).isoformat()
    reason = (payload.reason or "Rejected").strip()
    await db.employee_lifecycle_events.update_one(
        {"id": event_id},
        {
            "$set": {
                "approval_status": "REJECTED",
                "rejection_reason": reason,
                "processing_status": "REJECTED",
                "processed_at": now,
                "processing_error": reason,
                "updated_at": now,
            }
        },
    )
    await _append_lifecycle_audit(
        employee_code=ev.get("employee_code") or "",
        action="EVENT_REJECTED",
        from_status=None,
        to_status=None,
        event_type=ev.get("event_type"),
        event_id=event_id,
        actor_id=current_user["id"],
        notes=reason,
    )
    updated = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    return EmployeeLifecycleEventResponse(**updated)


@api_router.post("/admin/employee-lifecycle/escalate-approvals")
async def escalate_pending_lifecycle_approvals(current_user: dict = Depends(get_current_user)):
    """
    M2-2: mark stale approval requests as escalated and notify HR admins.
    Intended for cron / workflow_dispatch.
    """
    _require_admin(current_user)
    now = datetime.now(timezone.utc)
    pending = await db.employee_lifecycle_events.find(
        {"requires_approval": True, "approval_status": "PENDING", "escalated_at": None},
        {"_id": 0},
    ).to_list(500)
    escalated = 0
    for ev in pending:
        et = ev.get("event_type") or ""
        rule = approval_rule_for_event(et)
        if not rule:
            continue
        _, hours = rule
        try:
            raw = (ev.get("created_at") or "").replace("Z", "+00:00")
            created = datetime.fromisoformat(raw)
            if created.tzinfo is None:
                created = created.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        if now - created < timedelta(hours=hours):
            continue
        await db.employee_lifecycle_events.update_one(
            {"id": ev["id"]},
            {"$set": {"escalated_at": now.isoformat(), "updated_at": now.isoformat()}},
        )
        for rid in await _hr_escalation_recipient_ids():
            await create_notification(
                recipient_id=rid,
                notification_type="LIFECYCLE_APPROVAL_ESCALATION",
                title="Lifecycle approval escalated",
                message=f"Event {ev.get('event_type')} for {ev.get('employee_code')} pending since {ev.get('created_at')}",
                metadata={"event_id": ev.get("id"), "employee_code": ev.get("employee_code")},
            )
        escalated += 1
    return {"escalated": escalated}


@api_router.get("/employee-lifecycle/audit-log")
async def list_lifecycle_audit_log(
    employee_code: Optional[str] = None,
    page: int = 1,
    page_size: int = 50,
    current_user: dict = Depends(get_current_user),
):
    """M2-1: immutable-style audit trail for lifecycle actions."""
    _require_phase1_access(current_user, "employees_read")
    page = max(1, page)
    page_size = min(max(1, page_size), 200)
    q: Dict[str, Any] = {}
    if employee_code:
        q["employee_code"] = (employee_code or "").strip()
    total = await db[LIFECYCLE_AUDIT_COLLECTION].count_documents(q)
    rows = (
        await db[LIFECYCLE_AUDIT_COLLECTION].find(q, {"_id": 0})
        .sort("created_at", -1)
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )
    return {
        "items": rows,
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
    }


# --- M2-3 Compliance documents ---


def _compliance_doc_public(doc: Dict[str, Any]) -> Dict[str, Any]:
    out = {k: v for k, v in doc.items() if k != "content_base64"}
    return out


@api_router.post("/compliance/documents", response_model=ComplianceDocumentResponse)
async def create_compliance_document(
    payload: ComplianceDocumentCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    code = (payload.employee_code or "").strip()
    emp = await db.employees.find_one({"employee_code": code}, {"_id": 0, "id": 1})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")
    b64 = payload.content_base64
    if b64 and len(b64) > 1_500_000:
        raise HTTPException(status_code=400, detail="content_base64 too large (max ~1.1MB binary)")
    now = datetime.now(timezone.utc).isoformat()
    doc_id = str(uuid.uuid4())
    sla_days = int(os.environ.get("COMPLIANCE_VERIFY_SLA_DAYS", "14") or "14")
    doc = {
        "id": doc_id,
        "employee_code": code,
        "document_type": (payload.document_type or "").strip(),
        "title": (payload.title or "").strip(),
        "storage_uri": payload.storage_uri,
        "content_base64": b64,
        "status": "PENDING_VERIFY",
        "uploaded_by": current_user["id"],
        "uploaded_at": now,
        "verified_at": None,
        "verified_by": None,
        "expires_at": payload.expires_at,
        "sla_due_at": default_sla_due(now, days=sla_days),
        "reminder_sent_at": None,
        "sla_breached_at": None,
        "updated_at": now,
    }
    await db[COMPLIANCE_DOCS_COLLECTION].insert_one(doc)
    return ComplianceDocumentResponse(**_compliance_doc_public(doc))


@api_router.get("/compliance/documents", response_model=List[ComplianceDocumentResponse])
async def list_compliance_documents(
    employee_code: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_read")
    q: Dict[str, Any] = {}
    if employee_code:
        q["employee_code"] = (employee_code or "").strip()
    if status:
        q["status"] = status.upper()
    rows = await db[COMPLIANCE_DOCS_COLLECTION].find(q, {"_id": 0}).sort("uploaded_at", -1).to_list(500)
    return [ComplianceDocumentResponse(**_compliance_doc_public(r)) for r in rows]


@api_router.post("/compliance/documents/{doc_id}/verify", response_model=ComplianceDocumentResponse)
async def verify_compliance_document(
    doc_id: str,
    payload: ComplianceDocumentVerifyRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    row = await db[COMPLIANCE_DOCS_COLLECTION].find_one({"id": doc_id}, {"_id": 0})
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    now = datetime.now(timezone.utc).isoformat()
    patch = {
        "status": "VERIFIED",
        "verified_at": now,
        "verified_by": current_user["id"],
        "updated_at": now,
        "sla_breached_at": None,
    }
    if payload.notes:
        patch["verify_notes"] = payload.notes
    await db[COMPLIANCE_DOCS_COLLECTION].update_one({"id": doc_id}, {"$set": patch})
    updated = await db[COMPLIANCE_DOCS_COLLECTION].find_one({"id": doc_id}, {"_id": 0})
    return ComplianceDocumentResponse(**_compliance_doc_public(updated))


@api_router.post("/admin/compliance/scan-sla-breaches")
async def scan_compliance_sla_breaches(current_user: dict = Depends(get_current_user)):
    """Mark documents past verification SLA as breached (M2-3)."""
    _require_admin(current_user)
    now = datetime.now(timezone.utc).isoformat()
    rows = await db[COMPLIANCE_DOCS_COLLECTION].find(
        {"status": "PENDING_VERIFY", "verified_at": None},
        {"_id": 0},
    ).to_list(2000)
    marked = 0
    for r in rows:
        if r.get("sla_breached_at"):
            continue
        if is_past_iso(r.get("sla_due_at")):
            await db[COMPLIANCE_DOCS_COLLECTION].update_one(
                {"id": r["id"]},
                {"$set": {"sla_breached_at": now, "updated_at": now, "status": "SLA_BREACH"}},
            )
            marked += 1
    return {"sla_breaches_marked": marked}


@api_router.post("/admin/compliance/dispatch-document-reminders")
async def dispatch_compliance_document_reminders(current_user: dict = Depends(get_current_user)):
    """
    Remind owners about pending verification or upcoming expiry (M2-3).
    """
    _require_admin(current_user)
    now = datetime.now(timezone.utc)
    sent = 0

    pending = await db[COMPLIANCE_DOCS_COLLECTION].find(
        {"status": "PENDING_VERIFY", "reminder_sent_at": None},
        {"_id": 0},
    ).to_list(500)
    for r in pending:
        uid = r.get("uploaded_by")
        if not uid:
            continue
        await create_notification(
            recipient_id=uid,
            notification_type="COMPLIANCE_DOC_REMINDER",
            title="Compliance document pending verification",
            message=f"Document '{r.get('title')}' for {r.get('employee_code')} needs verification.",
            metadata={"document_id": r.get("id")},
        )
        await db[COMPLIANCE_DOCS_COLLECTION].update_one(
            {"id": r["id"]},
            {"$set": {"reminder_sent_at": now.isoformat(), "updated_at": now.isoformat()}},
        )
        sent += 1

    expiring = await db[COMPLIANCE_DOCS_COLLECTION].find(
        {"status": "VERIFIED", "expiry_reminder_sent_at": None, "expires_at": {"$ne": None}},
        {"_id": 0},
    ).to_list(500)
    for r in expiring:
        uid = r.get("uploaded_by")
        if not uid or not r.get("expires_at"):
            continue
        try:
            raw = str(r["expires_at"]).replace("Z", "+00:00")
            exp = datetime.fromisoformat(raw)
            if exp.tzinfo is None:
                exp = exp.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        if now <= exp <= now + timedelta(days=30):
            await create_notification(
                recipient_id=uid,
                notification_type="COMPLIANCE_DOC_EXPIRY",
                title="Compliance document expiring soon",
                message=f"'{r.get('title')}' for {r.get('employee_code')} expires at {r.get('expires_at')}",
                metadata={"document_id": r.get("id")},
            )
            await db[COMPLIANCE_DOCS_COLLECTION].update_one(
                {"id": r["id"]},
                {"$set": {"expiry_reminder_sent_at": now.isoformat(), "updated_at": now.isoformat()}},
            )
            sent += 1

    return {"notifications_sent": sent}


@api_router.get("/compliance/report/export")
async def export_compliance_report(
    export_format: Literal["csv"] = "csv",
    current_user: dict = Depends(get_current_user),
):
    """M2-3: compliance export for auditors."""
    _require_phase1_access(current_user, "employees_read")
    rows = await db[COMPLIANCE_DOCS_COLLECTION].find({}, {"_id": 0, "content_base64": 0}).sort("uploaded_at", -1).to_list(5000)
    if export_format != "csv":
        raise HTTPException(status_code=400, detail="Only csv supported")
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(
        [
            "id",
            "employee_code",
            "document_type",
            "title",
            "status",
            "uploaded_at",
            "verified_at",
            "verified_by",
            "expires_at",
            "sla_due_at",
            "sla_breached_at",
            "reminder_sent_at",
            "storage_uri",
        ]
    )
    for r in rows:
        w.writerow(
            [
                r.get("id"),
                r.get("employee_code"),
                r.get("document_type"),
                r.get("title"),
                r.get("status"),
                r.get("uploaded_at"),
                r.get("verified_at"),
                r.get("verified_by"),
                r.get("expires_at"),
                r.get("sla_due_at"),
                r.get("sla_breached_at"),
                r.get("reminder_sent_at"),
                r.get("storage_uri"),
            ]
        )
    buf.seek(0)
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": 'attachment; filename="compliance_documents.csv"'},
    )


@api_router.get("/employee-lifecycle/dashboard", response_model=EmployeeLifecycleDashboardResponse)
async def get_employee_lifecycle_dashboard(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_read")

    now = datetime.now(timezone.utc)
    start_30 = (now - timedelta(days=30)).isoformat()

    event_types = ["ONBOARDED", "ACTIVATED", "ROLE_CHANGED", "DOCUMENT_ADDED", "EXITED"]
    counts_by_event_type: Dict[str, int] = {}
    for t in event_types:
        counts_by_event_type[t] = await db.employee_lifecycle_events.count_documents({"event_type": t})

    last_30_days_events = await db.employee_lifecycle_events.count_documents({"created_at": {"$gte": start_30}})

    return EmployeeLifecycleDashboardResponse(
        total_events=sum(counts_by_event_type.values()),
        last_30_days_events=last_30_days_events,
        counts_by_event_type=counts_by_event_type,
        onboarded_total=counts_by_event_type.get("ONBOARDED", 0),
        activated_total=counts_by_event_type.get("ACTIVATED", 0),
        role_changed_total=counts_by_event_type.get("ROLE_CHANGED", 0),
        document_added_total=counts_by_event_type.get("DOCUMENT_ADDED", 0),
        exited_total=counts_by_event_type.get("EXITED", 0),
    )


async def process_employee_lifecycle_event(event_id: str) -> None:
    """
    Event-driven workflow handler:
    Apply lifecycle event to the canonical `employees` record.
    Any processing failure is captured in `employee_lifecycle_events.details`
    to keep the system fault-tolerant/observable.
    """
    event = await db.employee_lifecycle_events.find_one({"id": event_id}, {"_id": 0})
    if not event:
        return

    try:
        employee_code = (event.get("employee_code") or "").strip()
        event_type = event.get("event_type")
        details = event.get("details") or {}

        if not (employee_code and event_type):
            return

        if event.get("requires_approval") and event.get("approval_status") != "APPROVED":
            return

        # Idempotency: if already processed successfully, do nothing.
        if event.get("processing_status") == "PROCESSED":
            return

        emp = await db.employees.find_one({"employee_code": employee_code}, {"_id": 0})
        if not emp:
            raise ValueError("Employee not found for lifecycle processing")
        cur_status = (emp.get("status") or "ACTIVE").upper()
        ev_err = validate_lifecycle_event_for_status(str(event_type), cur_status)
        if ev_err:
            raise ValueError(ev_err)

        attempts = int(event.get("attempts") or 0) + 1
        await db.employee_lifecycle_events.update_one(
            {"id": event_id},
            {"$set": {"attempts": attempts, "processing_error": None}},
        )

        update_doc: Dict[str, Any] = {}
        if event_type == "ONBOARDED":
            update_doc["status"] = "ONBOARDING"
        elif event_type == "ACTIVATED":
            update_doc["status"] = "ACTIVE"
        elif event_type == "EXITED":
            update_doc["status"] = "EXITED"
        elif event_type == "ROLE_CHANGED":
            # Expected keys (best-effort):
            # - role_title: new title
            # - manager_id: new manager reference
            if isinstance(details, dict) and details.get("role_title"):
                update_doc["role_title"] = str(details.get("role_title")).strip()
            if isinstance(details, dict) and details.get("manager_id"):
                update_doc["manager_id"] = str(details.get("manager_id")).strip()

        # For DOCUMENT_ADDED: no direct canonical employee mutation in Phase-2 scope.

        new_status = cur_status
        if update_doc:
            if update_doc.get("status"):
                st_err = validate_direct_status_transition(cur_status, str(update_doc["status"]))
                if st_err:
                    raise ValueError(st_err)
            update_doc["updated_at"] = datetime.now(timezone.utc).isoformat()
            await db.employees.update_one(
                {"employee_code": employee_code},
                {"$set": update_doc},
            )
            if update_doc.get("status"):
                new_status = str(update_doc["status"]).upper()

        await _append_lifecycle_audit(
            employee_code=employee_code,
            action="EVENT_PROCESSED",
            from_status=cur_status,
            to_status=new_status,
            event_type=str(event_type),
            event_id=event_id,
            actor_id="system",
            notes=None,
        )

        await db.employee_lifecycle_events.update_one(
            {"id": event_id},
            {
                "$set": {
                    "processing_status": "PROCESSED",
                    "processed_at": datetime.now(timezone.utc).isoformat(),
                    "processing_error": None,
                }
            },
        )
    except Exception as e:
        # Record errors against the lifecycle event for troubleshooting.
        err = str(e)
        await db.employee_lifecycle_events.update_one(
            {"id": event_id},
            {
                "$set": {
                    "processing_status": "FAILED",
                    "processing_error": err,
                    "processed_at": datetime.now(timezone.utc).isoformat(),
                }
            },
        )


# ========================
# Phase-5/6 M7: Cost Optimization & Automation (MVP)
# ========================

async def _lifecycle_event_ids_for_reprocess(limit: int) -> List[str]:
    limit = max(1, min(int(limit or 50), 500))
    cursor = (
        db.employee_lifecycle_events.find(
            {
                "$or": [
                    {"processing_status": "FAILED"},
                    {
                        "processing_status": "PENDING",
                        "$or": [
                            {"requires_approval": False},
                            {"requires_approval": {"$exists": False}},
                            {"approval_status": "APPROVED"},
                        ],
                    },
                ]
            },
            {"_id": 0, "id": 1, "attempts": 1, "processing_status": 1},
        )
        .sort("attempts", 1)
        .limit(limit)
    )
    rows = await cursor.to_list(limit)
    return [str(r.get("id")) for r in rows if r.get("id")]


class AutomationLifecycleStatusResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    lifecycle_events_total: int
    lifecycle_events_pending: int
    lifecycle_events_processed: int
    lifecycle_events_failed: int
    workflow_rules_enabled: int = 0
    workflow_runs_succeeded_24h: int = 0


class AutomationReprocessLifecycleRequest(BaseModel):
    limit: int = 50


class AutomationReprocessLifecycleResponse(BaseModel):
    enqueued: int


@api_router.get("/automation/status", response_model=AutomationLifecycleStatusResponse)
async def get_automation_status(current_user: dict = Depends(get_current_user)):
    """
    MVP automation status:
    - shows lifecycle event processing health (PENDING/FAILED/PROCESSED)
    - can be extended later with job match orchestration and workflow engine
    """
    _require_phase1_access(current_user, "kpi_read")

    total = await db.employee_lifecycle_events.count_documents({})
    pending = await db.employee_lifecycle_events.count_documents({"processing_status": "PENDING"})
    failed = await db.employee_lifecycle_events.count_documents({"processing_status": "FAILED"})
    processed = await db.employee_lifecycle_events.count_documents({"processing_status": "PROCESSED"})

    since = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    wf_enabled = await db[COL_WORKFLOW_RULES].count_documents({"enabled": True})
    wf_ok = await db[COL_WORKFLOW_RUNS].count_documents({"status": "SUCCESS", "created_at": {"$gte": since}})

    return AutomationLifecycleStatusResponse(
        lifecycle_events_total=total,
        lifecycle_events_pending=pending,
        lifecycle_events_processed=processed,
        lifecycle_events_failed=failed,
        workflow_rules_enabled=wf_enabled,
        workflow_runs_succeeded_24h=wf_ok,
    )


@api_router.post("/automation/reprocess-lifecycle", response_model=AutomationReprocessLifecycleResponse)
async def reprocess_lifecycle_events(
    req: AutomationReprocessLifecycleRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    """
    MVP automation action:
    - Re-enqueue background processing for FAILED/PENDING lifecycle events.
    """
    _require_phase1_access(current_user, "lifecycle_write")

    limit = max(1, int(req.limit or 50))
    limit = min(limit, 200)

    ids = await _lifecycle_event_ids_for_reprocess(limit)
    for eid in ids:
        background_tasks.add_task(process_employee_lifecycle_event, eid)

    return AutomationReprocessLifecycleResponse(enqueued=len(ids))


WorkflowTriggerType = Literal[
    "MANUAL",
    "ON_LIFECYCLE_PENDING_THRESHOLD",
    "ON_SCHEDULE",
    "WEBHOOK_INBOUND",
]
WorkflowActionType = Literal["REPROCESS_LIFECYCLE", "NOTIFY_HR", "NOOP", "HTTP_WEBHOOK"]


class WorkflowAutomationRuleCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)
    enabled: bool = True
    trigger_type: WorkflowTriggerType = "MANUAL"
    trigger_config: Dict[str, Any] = Field(default_factory=dict)
    action_type: WorkflowActionType = "NOOP"
    action_config: Dict[str, Any] = Field(default_factory=dict)
    max_retries: int = Field(default=3, ge=1, le=8)
    retry_backoff_sec: float = Field(default=2.0, ge=0.5, le=60.0)
    # Visual / multi-step (React Flow graph); when set with nodes, overrides single action at run time.
    flow_graph: Optional[Dict[str, Any]] = None
    schedule_interval_minutes: Optional[int] = Field(default=None, ge=5, le=10080)
    inbound_webhook_secret: Optional[str] = Field(default=None, min_length=8, max_length=256)

    @model_validator(mode="after")
    def _workflow_create_consistency(self):
        if self.trigger_type == "ON_SCHEDULE" and self.schedule_interval_minutes is None:
            raise ValueError("ON_SCHEDULE rules require schedule_interval_minutes (5–10080).")
        if self.trigger_type == "WEBHOOK_INBOUND":
            if not self.inbound_webhook_secret or len(self.inbound_webhook_secret.strip()) < 8:
                raise ValueError("WEBHOOK_INBOUND rules require inbound_webhook_secret (min 8 chars) on create.")
        return self


class WorkflowAutomationRuleUpdate(BaseModel):
    name: Optional[str] = Field(default=None, min_length=1, max_length=200)
    enabled: Optional[bool] = None
    trigger_type: Optional[WorkflowTriggerType] = None
    trigger_config: Optional[Dict[str, Any]] = None
    action_type: Optional[WorkflowActionType] = None
    action_config: Optional[Dict[str, Any]] = None
    max_retries: Optional[int] = Field(default=None, ge=1, le=8)
    retry_backoff_sec: Optional[float] = Field(default=None, ge=0.5, le=60.0)
    flow_graph: Optional[Dict[str, Any]] = None
    schedule_interval_minutes: Optional[int] = Field(default=None, ge=5, le=10080)
    schedule_next_run_at: Optional[str] = None
    inbound_webhook_secret: Optional[str] = Field(default=None, min_length=8, max_length=256)


class WorkflowAutomationExecuteResponse(BaseModel):
    run_id: str
    status: str
    detail: Dict[str, Any] = Field(default_factory=dict)


class HrCopilotChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=4000)
    session_id: Optional[str] = None


class HrCopilotChatResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    reply: str
    intent: str
    session_id: str
    actions: List[Dict[str, Any]] = Field(default_factory=list)
    copilot_engine_version: str = COPILOT_ENGINE_VERSION
    audit_id: str
    intent_source: str = "rules"
    rule_intent: Optional[str] = None
    hf_model: Optional[str] = None
    hf_score: Optional[float] = None


class ManualWorkflowBaselineCreate(BaseModel):
    workflow_key: str = Field(..., min_length=1, max_length=64)
    label: str = Field(..., min_length=1, max_length=200)
    minutes_per_run: float = Field(..., ge=0.0, le=24 * 60.0)
    hourly_fully_loaded_cost_usd: float = Field(default=0.0, ge=0.0)


class ManualWorkflowBaselineUpdate(BaseModel):
    label: Optional[str] = Field(default=None, min_length=1, max_length=200)
    minutes_per_run: Optional[float] = Field(default=None, ge=0.0, le=24 * 60.0)
    hourly_fully_loaded_cost_usd: Optional[float] = Field(default=None, ge=0.0)


class CostOptimizationSummaryResponse(BaseModel):
    generated_at: str
    window_days: int
    baselines_count: int
    automation_runs_success: int
    automation_runs_failed: int
    estimated_minutes_saved: float
    estimated_cost_saved_usd: float


def _m7_savings_workflow_key(rule: Dict[str, Any]) -> str:
    fg = rule.get("flow_graph")
    if isinstance(fg, dict):
        for n in fg.get("nodes") or []:
            if not isinstance(n, dict):
                continue
            data = n.get("data") if isinstance(n.get("data"), dict) else {}
            at = str(data.get("action_type") or "").strip().upper()
            if at and at != "NOOP":
                return at
    return str(rule.get("action_type") or "NOOP").upper()


async def _m7_execute_single_action(
    rule: Dict[str, Any],
    background_tasks: BackgroundTasks,
    step_index: int = 0,
) -> Dict[str, Any]:
    at = (rule.get("action_type") or "NOOP").upper()
    cfg = rule.get("action_config") or {}
    if at == "NOOP":
        return {"detail": "noop", "step_index": step_index}
    if at == "REPROCESS_LIFECYCLE":
        limit = max(1, min(int(cfg.get("limit") or 50), 200))
        ids = await _lifecycle_event_ids_for_reprocess(limit)
        for eid in ids:
            background_tasks.add_task(process_employee_lifecycle_event, eid)
        return {"enqueued": len(ids), "limit": limit, "step_index": step_index}
    if at == "NOTIFY_HR":
        title = str(cfg.get("title") or "Workflow notification")
        message = str(cfg.get("message") or "An automation rule fired.")
        meta: Dict[str, Any] = {"rule_id": rule.get("id"), "workflow_engine": WORKFLOW_ENGINE_VERSION}
        n = 0
        for rid in await _hr_escalation_recipient_ids():
            await create_notification(rid, "WORKFLOW_AUTOMATION", title, message, meta)
            n += 1
        return {"notified_hr_users": n, "step_index": step_index}
    if at == "HTTP_WEBHOOK":
        out = await execute_http_webhook(cfg, rule=rule, extra_template_vars={"step_index": step_index})
        return {**out, "step_index": step_index}
    raise ValueError(f"Unsupported action_type: {at}")


async def _m7_execute_rule_action(rule: Dict[str, Any], background_tasks: BackgroundTasks) -> Dict[str, Any]:
    fg = rule.get("flow_graph")
    if isinstance(fg, dict) and isinstance(fg.get("nodes"), list) and len(fg["nodes"]) > 0:
        return await execute_flow_graph(
            fg,
            rule=rule,
            background_tasks=background_tasks,
            run_single_action=_m7_execute_single_action,
        )
    return await _m7_execute_single_action(rule, background_tasks, step_index=0)


async def _m7_bump_schedule_after_success(rule_id: str, rule: Dict[str, Any]) -> None:
    if str(rule.get("trigger_type") or "").upper() != "ON_SCHEDULE":
        return
    interval = int(rule.get("schedule_interval_minutes") or (rule.get("trigger_config") or {}).get("interval_minutes") or 60)
    interval = max(5, min(interval, 10080))
    nxt = (datetime.now(timezone.utc) + timedelta(minutes=interval)).isoformat()
    ts = datetime.now(timezone.utc).isoformat()
    await db[COL_WORKFLOW_RULES].update_one(
        {"id": rule_id},
        {"$set": {"schedule_next_run_at": nxt, "updated_at": ts}},
    )


async def _append_hr_copilot_audit(
    *,
    actor_id: Optional[str],
    session_id: str,
    intent: str,
    message_snippet: str,
    actions: List[Dict[str, Any]],
    reply_snippet: str,
    intent_resolution: Optional[Dict[str, Any]] = None,
) -> str:
    aid = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": aid,
        "actor_id": actor_id,
        "session_id": session_id,
        "intent": intent,
        "message_snippet": message_snippet[:500],
        "reply_snippet": reply_snippet[:500],
        "actions": actions,
        "intent_resolution": intent_resolution or {},
        "engine_version": COPILOT_ENGINE_VERSION,
        "created_at": now,
    }
    await db[COL_HR_COPILOT_AUDIT].insert_one(doc)
    return aid


@api_router.get("/admin/workflow-automation/rules", response_model=List[Dict[str, Any]])
async def admin_list_workflow_rules(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    rows = await db[COL_WORKFLOW_RULES].find({}, {"_id": 0}).sort("updated_at", -1).to_list(200)
    for r in rows:
        r.pop("inbound_webhook_secret_hash", None)
    return rows


@api_router.post("/admin/workflow-automation/rules", response_model=Dict[str, Any])
async def admin_create_workflow_rule(
    payload: WorkflowAutomationRuleCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    rid = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc: Dict[str, Any] = {
        "id": rid,
        "name": payload.name.strip(),
        "enabled": bool(payload.enabled),
        "trigger_type": payload.trigger_type,
        "trigger_config": dict(payload.trigger_config or {}),
        "action_type": payload.action_type,
        "action_config": dict(payload.action_config or {}),
        "max_retries": int(payload.max_retries),
        "retry_backoff_sec": float(payload.retry_backoff_sec),
        "engine_version": WORKFLOW_ENGINE_VERSION,
        "created_by": current_user.get("id"),
        "created_at": now,
        "updated_at": now,
        "flow_graph": payload.flow_graph if isinstance(payload.flow_graph, dict) else None,
    }
    if payload.trigger_type == "ON_SCHEDULE":
        doc["schedule_interval_minutes"] = int(payload.schedule_interval_minutes or 60)
        doc["schedule_next_run_at"] = now
    if payload.trigger_type == "WEBHOOK_INBOUND" and payload.inbound_webhook_secret:
        doc["inbound_webhook_secret_hash"] = hash_password(payload.inbound_webhook_secret.strip())
    await db[COL_WORKFLOW_RULES].insert_one(doc)
    out = {k: v for k, v in doc.items() if k != "inbound_webhook_secret_hash"}
    if payload.trigger_type == "WEBHOOK_INBOUND":
        out["inbound_webhook_path"] = f"/api/webhooks/workflow/inbound/{rid}"
        out["inbound_webhook_auth"] = "Header X-Workflow-Token or query ?token="
    return out


@api_router.put("/admin/workflow-automation/rules/{rule_id}", response_model=Dict[str, Any])
async def admin_update_workflow_rule(
    rule_id: str,
    payload: WorkflowAutomationRuleUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    ex = await db[COL_WORKFLOW_RULES].find_one({"id": rule_id}, {"_id": 0})
    if not ex:
        raise HTTPException(status_code=404, detail="Rule not found")
    patch = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    if "inbound_webhook_secret" in patch:
        sec = patch.pop("inbound_webhook_secret")
        if sec:
            patch["inbound_webhook_secret_hash"] = hash_password(str(sec).strip())
    if not patch:
        return ex
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db[COL_WORKFLOW_RULES].update_one({"id": rule_id}, {"$set": patch})
    upd = await db[COL_WORKFLOW_RULES].find_one({"id": rule_id}, {"_id": 0})
    if upd:
        upd.pop("inbound_webhook_secret_hash", None)
    return upd


@api_router.delete("/admin/workflow-automation/rules/{rule_id}")
async def admin_delete_workflow_rule(rule_id: str, current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    r = await db[COL_WORKFLOW_RULES].delete_one({"id": rule_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Rule not found")
    return {"ok": True, "id": rule_id}


async def _run_workflow_rule_once(
    rule_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict,
    *,
    manual_trigger: bool,
) -> Optional[WorkflowAutomationExecuteResponse]:
    rule = await db[COL_WORKFLOW_RULES].find_one({"id": rule_id}, {"_id": 0})
    if not rule:
        raise HTTPException(status_code=404, detail="Rule not found")
    if not rule.get("enabled", True):
        if manual_trigger:
            raise HTTPException(status_code=400, detail="Rule is disabled")
        return None

    pending = await db.employee_lifecycle_events.count_documents({"processing_status": "PENDING"})
    ok, skip_reason = should_execute_trigger(
        trigger_type=str(rule.get("trigger_type") or "MANUAL"),
        trigger_config=rule.get("trigger_config") or {},
        pending_lifecycle_count=pending,
        manual=manual_trigger,
        rule=rule,
    )
    if not ok:
        if manual_trigger:
            raise HTTPException(status_code=400, detail=skip_reason)
        return None

    sk = _m7_savings_workflow_key(rule)
    at_display = "FLOW_GRAPH" if isinstance(rule.get("flow_graph"), dict) and (rule.get("flow_graph") or {}).get("nodes") else rule.get("action_type")

    run_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    await db[COL_WORKFLOW_RUNS].insert_one(
        {
            "id": run_id,
            "rule_id": rule_id,
            "rule_name": rule.get("name"),
            "status": "RUNNING",
            "action_type": at_display,
            "savings_workflow_key": sk,
            "error": None,
            "attempts_used": 0,
            "triggered_by": current_user.get("id") or "inbound_webhook",
            "engine_version": WORKFLOW_ENGINE_VERSION,
            "created_at": now,
            "finished_at": None,
            "detail": {},
        }
    )

    max_att = max(1, min(int(rule.get("max_retries") or 3), 8))
    backoff = float(rule.get("retry_backoff_sec") or 2.0)

    async def _op():
        return await _m7_execute_rule_action(rule, background_tasks)

    try:
        detail = await run_with_retries(_op, max_attempts=max_att, backoff_sec=backoff, label=f"workflow_rule[{rule_id}]")
        fin = datetime.now(timezone.utc).isoformat()
        await db[COL_WORKFLOW_RUNS].update_one(
            {"id": run_id},
            {"$set": {"status": "SUCCESS", "finished_at": fin, "detail": detail, "attempts_used": max_att}},
        )
        await _m7_bump_schedule_after_success(rule_id, rule)
        await _m10_publish_safe(
            topic=TOPIC_WORKFLOW_RUN_COMPLETED,
            payload={
                "run_id": run_id,
                "rule_id": rule_id,
                "rule_name": rule.get("name"),
                "savings_workflow_key": sk,
            },
            idempotency_key=f"workflow_run:{run_id}:completed",
            correlation_id=run_id,
        )
        return WorkflowAutomationExecuteResponse(run_id=run_id, status="SUCCESS", detail=detail)
    except Exception as e:
        fin = datetime.now(timezone.utc).isoformat()
        await db[COL_WORKFLOW_RUNS].update_one(
            {"id": run_id},
            {"$set": {"status": "FAILED", "finished_at": fin, "error": str(e), "detail": {"error": str(e)}}},
        )
        await _m10_publish_safe(
            topic=TOPIC_WORKFLOW_RUN_FAILED,
            payload={"run_id": run_id, "rule_id": rule_id, "rule_name": rule.get("name"), "error": str(e)[:2000]},
            idempotency_key=f"workflow_run:{run_id}:failed",
            correlation_id=run_id,
        )
        if manual_trigger:
            raise HTTPException(status_code=500, detail=f"Workflow execution failed: {e}")
        raise


@api_router.post("/admin/workflow-automation/rules/{rule_id}/execute", response_model=WorkflowAutomationExecuteResponse)
async def admin_execute_workflow_rule(
    rule_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    out = await _run_workflow_rule_once(rule_id, background_tasks, current_user, manual_trigger=True)
    assert out is not None
    return out


@api_router.post("/webhooks/workflow/inbound/{rule_id}", response_model=WorkflowAutomationExecuteResponse)
async def workflow_inbound_webhook_trigger(
    rule_id: str,
    background_tasks: BackgroundTasks,
    x_workflow_token: Optional[str] = Header(None, alias="X-Workflow-Token"),
    token: Optional[str] = Query(None, description="Same secret as header; for simple integrations"),
):
    """
    Signed inbound trigger for rules with trigger_type WEBHOOK_INBOUND (no JWT).
    Authenticate with header X-Workflow-Token or query ?token= (plaintext matches stored bcrypt hash).
    """
    rule = await db[COL_WORKFLOW_RULES].find_one({"id": rule_id}, {"_id": 0})
    if not rule or not rule.get("enabled", True):
        raise HTTPException(status_code=404, detail="Rule not found")
    if str(rule.get("trigger_type") or "").upper() != "WEBHOOK_INBOUND":
        raise HTTPException(status_code=404, detail="Rule not found")
    raw_tok = x_workflow_token or token
    if not raw_tok:
        raise HTTPException(status_code=401, detail="Missing X-Workflow-Token or token query param")
    h = rule.get("inbound_webhook_secret_hash")
    if not h or not verify_password(raw_tok, h):
        raise HTTPException(status_code=401, detail="Invalid workflow token")
    actor = {"id": "inbound_webhook", "role": "admin"}
    out = await _run_workflow_rule_once(rule_id, background_tasks, actor, manual_trigger=True)
    if out is None:
        raise HTTPException(status_code=500, detail="Workflow did not run")
    return out


@api_router.get("/admin/workflow-automation/runs", response_model=List[Dict[str, Any]])
async def admin_list_workflow_runs(
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    limit = min(max(1, int(limit)), 200)
    return await db[COL_WORKFLOW_RUNS].find({}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)


@api_router.post("/admin/workflow-automation/dispatch-triggered", response_model=Dict[str, Any])
async def admin_dispatch_triggered_workflow_rules(
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    """
    Evaluates enabled rules: lifecycle backlog thresholds, due ON_SCHEDULE rules, etc.
    Skips MANUAL and WEBHOOK_INBOUND (those need explicit run or signed POST).
    """
    _require_admin(current_user)
    pending = await db.employee_lifecycle_events.count_documents({"processing_status": "PENDING"})
    rules = await db[COL_WORKFLOW_RULES].find({"enabled": True}, {"_id": 0}).to_list(100)
    executed = 0
    skipped = 0
    errors: List[str] = []
    for rule in rules:
        t = str(rule.get("trigger_type") or "MANUAL").upper()
        if t == "MANUAL":
            skipped += 1
            continue
        ok, reason = should_execute_trigger(
            trigger_type=t,
            trigger_config=rule.get("trigger_config") or {},
            pending_lifecycle_count=pending,
            manual=False,
            rule=rule,
        )
        if not ok:
            skipped += 1
            continue
        try:
            res = await _run_workflow_rule_once(rule["id"], background_tasks, current_user, manual_trigger=False)
            if res:
                executed += 1
            else:
                skipped += 1
        except HTTPException as he:
            errors.append(f"{rule.get('id')}: {he.detail}")
        except Exception as e:
            errors.append(f"{rule.get('id')}: {e}")
    return {"executed": executed, "skipped": skipped, "errors": errors[:20]}


@api_router.post("/hr-copilot/chat", response_model=HrCopilotChatResponse)
async def hr_copilot_chat(
    body: HrCopilotChatRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    intent, intent_meta = await resolve_copilot_intent_async(body.message)
    session_id = body.session_id or str(uuid.uuid4())
    actions: List[Dict[str, Any]] = []
    reply_parts: List[str] = []

    if intent == "empty":
        reply_parts.append("Please enter a question.")
    elif intent == "help":
        reply_parts.append(help_text())
    elif intent == "automation_status":
        total = await db.employee_lifecycle_events.count_documents({})
        pend = await db.employee_lifecycle_events.count_documents({"processing_status": "PENDING"})
        fail = await db.employee_lifecycle_events.count_documents({"processing_status": "FAILED"})
        proc = await db.employee_lifecycle_events.count_documents({"processing_status": "PROCESSED"})
        wf_en = await db[COL_WORKFLOW_RULES].count_documents({"enabled": True})
        reply_parts.append(
            f"Lifecycle events — total {total}, pending {pend}, processed {proc}, failed {fail}. "
            f"Enabled workflow rules: {wf_en}."
        )
        actions.append({"type": "automation_status", "status": "OK"})
    elif intent == "reprocess_lifecycle":
        if not _user_has_phase1_permission(current_user, "lifecycle_write"):
            reply_parts.append("Reprocessing requires lifecycle_write (admin, hr_admin, or recruiter).")
            actions.append({"type": "reprocess_lifecycle", "status": "DENIED_POLICY"})
        else:
            limit = 50
            ids = await _lifecycle_event_ids_for_reprocess(limit)
            for eid in ids:
                background_tasks.add_task(process_employee_lifecycle_event, eid)
            reply_parts.append(f"Queued {len(ids)} lifecycle event(s) for background reprocessing (limit {limit}).")
            actions.append({"type": "reprocess_lifecycle", "status": "EXECUTED", "enqueued": len(ids)})
    elif intent == "employee_lookup":
        if not _user_has_phase1_permission(current_user, "employees_read"):
            reply_parts.append("Employee lookup requires employees_read permission.")
            actions.append({"type": "employee_lookup", "status": "DENIED_POLICY"})
        else:
            code = extract_employee_code_hint(body.message)
            if not code:
                reply_parts.append("Include an employee code to look up (e.g. E1234).")
                actions.append({"type": "employee_lookup", "status": "NEED_CODE"})
            else:
                row = await db.employees.find_one(
                    {"employee_code": {"$regex": f"^{re.escape(code)}$", "$options": "i"}},
                    {"_id": 0, "employee_code": 1, "full_name": 1, "department": 1, "status": 1, "role_title": 1},
                )
                if not row:
                    reply_parts.append(f"No employee found for code `{code}`.")
                    actions.append({"type": "employee_lookup", "status": "NOT_FOUND"})
                else:
                    reply_parts.append(
                        f"{row.get('full_name') or '—'} ({row.get('employee_code')}) — "
                        f"{row.get('department') or '—'}, status {row.get('status') or '—'}, "
                        f"role {row.get('role_title') or '—'}."
                    )
                    actions.append({"type": "employee_lookup", "status": "OK", "employee_code": row.get("employee_code")})
    elif intent == "workflow_rules":
        if current_user.get("role") != "admin":
            reply_parts.append("Listing workflow rules requires admin access.")
            actions.append({"type": "workflow_rules", "status": "DENIED_POLICY"})
        else:
            rlist = await db[COL_WORKFLOW_RULES].find({}, {"_id": 0, "id": 1, "name": 1, "enabled": 1, "trigger_type": 1, "action_type": 1}).to_list(50)
            if not rlist:
                reply_parts.append("No workflow rules configured yet.")
            else:
                lines = [f"- {x.get('name')} ({x.get('id')}): {x.get('trigger_type')} → {x.get('action_type')} [{'on' if x.get('enabled') else 'off'}]" for x in rlist]
                reply_parts.append("Workflow rules:\n" + "\n".join(lines))
            actions.append({"type": "workflow_rules", "status": "OK", "count": len(rlist)})
    else:
        reply_parts.append(
            "I’m not sure how to help with that yet. Try “automation status”, “reprocess lifecycle events”, "
            "or “lookup employee E1234”."
        )
        actions.append({"type": "unknown", "status": "NO_MATCH"})

    reply = "\n".join(reply_parts)
    audit_id = await _append_hr_copilot_audit(
        actor_id=current_user.get("id"),
        session_id=session_id,
        intent=intent,
        message_snippet=body.message.strip(),
        actions=actions,
        reply_snippet=reply,
        intent_resolution=intent_meta,
    )
    return HrCopilotChatResponse(
        reply=reply,
        intent=intent,
        session_id=session_id,
        actions=actions,
        audit_id=audit_id,
        intent_source=str(intent_meta.get("intent_source") or "rules"),
        rule_intent=intent_meta.get("rule_intent"),
        hf_model=intent_meta.get("hf_model"),
        hf_score=float(intent_meta["hf_top_score"]) if intent_meta.get("hf_top_score") is not None else None,
    )


@api_router.get("/admin/hr-copilot/audit", response_model=List[Dict[str, Any]])
async def admin_list_hr_copilot_audit(
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    limit = min(max(1, int(limit)), 200)
    return await db[COL_HR_COPILOT_AUDIT].find({}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)


@api_router.get("/admin/cost-optimization/baselines", response_model=List[Dict[str, Any]])
async def admin_list_manual_baselines(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    return await db[COL_MANUAL_WORKFLOW_BASELINES].find({}, {"_id": 0}).sort("workflow_key", 1).to_list(200)


@api_router.post("/admin/cost-optimization/baselines", response_model=Dict[str, Any])
async def admin_create_manual_baseline(
    payload: ManualWorkflowBaselineCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    bid = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    wk = payload.workflow_key.strip().upper()
    doc = {
        "id": bid,
        "workflow_key": wk,
        "label": payload.label.strip(),
        "minutes_per_run": float(payload.minutes_per_run),
        "hourly_fully_loaded_cost_usd": float(payload.hourly_fully_loaded_cost_usd),
        "created_at": now,
        "updated_at": now,
    }
    await db[COL_MANUAL_WORKFLOW_BASELINES].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.put("/admin/cost-optimization/baselines/{baseline_id}", response_model=Dict[str, Any])
async def admin_update_manual_baseline(
    baseline_id: str,
    payload: ManualWorkflowBaselineUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    patch = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    if not patch:
        row = await db[COL_MANUAL_WORKFLOW_BASELINES].find_one({"id": baseline_id}, {"_id": 0})
        if not row:
            raise HTTPException(status_code=404, detail="Baseline not found")
        return row
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    r = await db[COL_MANUAL_WORKFLOW_BASELINES].update_one({"id": baseline_id}, {"$set": patch})
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Baseline not found")
    return await db[COL_MANUAL_WORKFLOW_BASELINES].find_one({"id": baseline_id}, {"_id": 0})


@api_router.delete("/admin/cost-optimization/baselines/{baseline_id}")
async def admin_delete_manual_baseline(baseline_id: str, current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    r = await db[COL_MANUAL_WORKFLOW_BASELINES].delete_one({"id": baseline_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Baseline not found")
    return {"ok": True, "id": baseline_id}


@api_router.get("/executive/cost-optimization-summary", response_model=CostOptimizationSummaryResponse)
async def get_cost_optimization_summary(
    window_days: int = 30,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    window_days = max(1, min(int(window_days or 30), 365))
    since = (datetime.now(timezone.utc) - timedelta(days=window_days)).isoformat()
    baselines = await db[COL_MANUAL_WORKFLOW_BASELINES].find({}, {"_id": 0}).to_list(500)
    bmap = baseline_map(baselines)
    runs = await db[COL_WORKFLOW_RUNS].find({"created_at": {"$gte": since}}, {"_id": 0}).to_list(5000)
    ok_runs = [r for r in runs if (r.get("status") or "").upper() == "SUCCESS"]
    fail_runs = [r for r in runs if (r.get("status") or "").upper() == "FAILED"]
    totals = compute_savings_totals(successful_runs=ok_runs, baselines=bmap)
    return CostOptimizationSummaryResponse(
        generated_at=datetime.now(timezone.utc).isoformat(),
        window_days=window_days,
        baselines_count=len(baselines),
        automation_runs_success=len(ok_runs),
        automation_runs_failed=len(fail_runs),
        estimated_minutes_saved=float(totals["estimated_minutes_saved"]),
        estimated_cost_saved_usd=float(totals["estimated_cost_saved_usd"]),
    )


# ========================
# Workforce Intelligence (Demand-Supply)
# ========================

def _priority_rank(priority: str) -> int:
    p = (priority or "").upper()
    if p == "HIGH":
        return 3
    if p == "MEDIUM":
        return 2
    return 1

async def _derived_supply_from_employees() -> Dict[str, int]:
    """
    Derive supply counts from employee profile skills (lower-cased skill keys).
    Used to support project-demand skills that are not present in workforce_skills.
    """
    employees = await db.employees.find({}, {"_id": 0, "skills": 1}).to_list(2000)
    derived: Dict[str, int] = {}
    for e in employees:
        uniq_skills = set(
            (s.strip().lower() for s in (e.get("skills") or []) if isinstance(s, str) and s.strip())
        )
        for sk in uniq_skills:
            derived[sk] = derived.get(sk, 0) + 1
    return derived

async def _get_project_demand_map() -> Dict[str, Dict[str, Any]]:
    """
    Returns demand aggregation by skill_name_lc:
      { "<skill_lc>": { "skill_name": <original>, "demand_count": <sum>, "priority": <best> } }
    """
    docs = await db.project_skill_demands.find({}, {"_id": 0, "skill_name": 1, "skill_name_lc": 1, "demand_count": 1, "priority": 1}).to_list(50000)
    out: Dict[str, Dict[str, Any]] = {}
    for d in docs:
        sk_lc = (d.get("skill_name_lc") or "").strip().lower()
        if not sk_lc:
            continue
        demand = max(0, int(d.get("demand_count") or 0))
        pri = (d.get("priority") or "MEDIUM").upper()
        if sk_lc not in out:
            out[sk_lc] = {"skill_name": d.get("skill_name") or d.get("skill_name_lc") or sk_lc, "demand_count": demand, "priority": pri}
        else:
            out[sk_lc]["demand_count"] += demand
            # Keep highest-priority override if present.
            if _priority_rank(pri) > _priority_rank(out[sk_lc].get("priority") or "MEDIUM"):
                out[sk_lc]["priority"] = pri
    return out


async def _get_project_allocations_supply_map() -> Dict[str, Dict[str, Any]]:
    """
    Returns allocated (supply) aggregation by skill_name_lc across all projects:
      { "<skill_lc>": { "skill_name": <original>, "allocated_count": <sum> } }
    """
    docs = await db.project_skill_allocations.find(
        {},
        {"_id": 0, "skill_name": 1, "skill_name_lc": 1, "allocated_count": 1},
    ).to_list(100000)
    out: Dict[str, Dict[str, Any]] = {}
    for d in docs:
        sk_lc = (d.get("skill_name_lc") or "").strip().lower()
        if not sk_lc:
            continue
        allocated = max(0, int(d.get("allocated_count") or 0))
        if sk_lc not in out:
            out[sk_lc] = {
                "skill_name": d.get("skill_name") or d.get("skill_name_lc") or sk_lc,
                "allocated_count": allocated,
            }
        else:
            out[sk_lc]["allocated_count"] += allocated
    return out

async def _get_skill_priority_map(skill_keys: List[str]) -> Dict[str, str]:
    if not skill_keys:
        return {}
    rows = await db.workforce_skills.find(
        {"skill_name_lc": {"$in": [k.strip().lower() for k in skill_keys if isinstance(k, str) and k.strip()]}},
        {"_id": 0, "skill_name_lc": 1, "priority": 1},
    ).to_list(50000)
    out: Dict[str, str] = {}
    for r in rows:
        sk_lc = (r.get("skill_name_lc") or "").strip().lower()
        if not sk_lc:
            continue
        out[sk_lc] = str(r.get("priority") or "MEDIUM").upper()
    return out

# ========================
# Workforce dashboard caching + audit (Phase-3 M3/M4)
# ========================

WORKFORCE_DASHBOARD_CACHE_TTL_SECONDS = int(os.environ.get("WORKFORCE_DASHBOARD_CACHE_TTL_SECONDS", "300"))


async def _get_dashboard_cache_ttl_hit(cache_doc: Optional[Dict[str, Any]], now_dt: datetime) -> bool:
    if not cache_doc:
        return False
    cached_at = cache_doc.get("cached_at")
    if isinstance(cached_at, datetime):
        # Mongo may store naive datetimes depending on driver/serializer settings.
        # Normalize to UTC so subtraction with `now_dt` (timezone-aware) is safe.
        if cached_at.tzinfo is None:
            cached_at = cached_at.replace(tzinfo=timezone.utc)
        return (now_dt - cached_at) <= timedelta(seconds=WORKFORCE_DASHBOARD_CACHE_TTL_SECONDS)
    # Fallback for older docs storing ISO strings.
    if isinstance(cached_at, str):
        try:
            parsed = datetime.fromisoformat(cached_at)
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
            return (now_dt - parsed) <= timedelta(seconds=WORKFORCE_DASHBOARD_CACHE_TTL_SECONDS)
        except Exception:
            return False
    return False


async def _write_workforce_dashboard_audit_log(
    *,
    dashboard: str,
    demand_source: str,
    supply_source: str,
    horizon_months: Optional[int],
    cache_hit: bool,
    refresh_requested: bool,
    generated_at: Optional[str],
    top_rows_count: Optional[int] = None,
    total_shortage: Optional[int] = None,
    total_bench: Optional[int] = None,
):
    now_dt = datetime.now(timezone.utc)
    doc = {
        "id": str(uuid.uuid4()),
        "dashboard": dashboard,
        "demand_source": demand_source,
        "supply_source": supply_source,
        "horizon_months": horizon_months,
        "cache_hit": bool(cache_hit),
        "refresh_requested": bool(refresh_requested),
        "generated_at": generated_at,
        "top_rows_count": top_rows_count,
        "total_shortage": total_shortage,
        "total_bench": total_bench,
        "created_at": now_dt,
    }
    await db.workforce_dashboard_audit_logs.insert_one(doc)

@api_router.get("/workforce/intelligence", response_model=WorkforceIntelligenceResponse)
async def get_workforce_intelligence(
    horizon_months: int = 1,
    refresh: bool = False,
    current_user: dict = Depends(get_current_user),
):
    """
    Phase-3 M3:
    - demand forecasting (best-effort heuristic based on current demand + supply gap)
    - supply mapping from `workforce_skills`
    - forecast gap analytics for proactive hiring planning

    This MVP runs near real-time using current inventory data; it can be extended later
    with historical signals when available.
    """
    _require_phase1_access(current_user, "kpi_read")

    horizon_months = max(1, int(horizon_months or 1))

    now_dt = datetime.now(timezone.utc)
    project_demand_count = await db.project_skill_demands.count_documents({})
    allocation_count = await db.project_skill_allocations.count_documents({})
    demand_source = "projects" if project_demand_count > 0 else "workforce_skills"
    supply_source = "allocations" if allocation_count > 0 else "employees"
    if not refresh:
        cached_doc = await db.workforce_intelligence_forecast_cache.find_one(
            {"horizon_months": horizon_months, "demand_source": demand_source},
            {"_id": 0},
        )
        if await _get_dashboard_cache_ttl_hit(cached_doc, now_dt):
            # If supply mode changed, don't reuse cached data.
            if cached_doc.get("supply_source") != supply_source:
                cached_doc = None
            else:
                cached_data = cached_doc.get("data")
                if not cached_data:
                    pass
                else:
                    cached_resp = WorkforceIntelligenceResponse(**cached_data)
                    await _write_workforce_dashboard_audit_log(
                        dashboard="workforce_intelligence",
                        demand_source=demand_source,
                        supply_source=supply_source,
                        horizon_months=horizon_months,
                        cache_hit=True,
                        refresh_requested=False,
                        generated_at=cached_resp.generated_at,
                        top_rows_count=len(cached_resp.top_forecast_gaps or []),
                    )
                    return cached_resp

    priority_mult = {"HIGH": 1.15, "MEDIUM": 1.08, "LOW": 1.03}

    forecast_rows: List[WorkforceIntelligenceSkillForecast] = []
    demand_current_total = 0
    workforce_supply_total = 0
    demand_forecast_total = 0

    # Supply mapping is derived via `list_skill_inventory` (includes derived supply from employee profiles).
    skill_rows: List[SkillInventoryResponse] = await list_skill_inventory(current_user)

    # If project skill demands are present, use them as demand_source (project-driven demand).
    demand_map: Optional[Dict[str, Dict[str, Any]]] = None
    if project_demand_count > 0:
        demand_map = await _get_project_demand_map()

    # Pre-build inventory lookup.
    inv_by_skill_lc: Dict[str, SkillInventoryResponse] = {}
    for r in skill_rows:
        key = (r.skill_name or "").strip().lower()
        if key:
            inv_by_skill_lc[key] = r

    derived_supply: Dict[str, int] = {}
    if demand_map:
        derived_supply = await _derived_supply_from_employees()

    allocation_supply: Dict[str, Dict[str, Any]] = {}
    if allocation_count > 0:
        allocation_supply = await _get_project_allocations_supply_map()

    horizon_growth = 1 + min(0.35, 0.06 * max(0, horizon_months - 1))

    if demand_map:
        # Forecast only for skills defined by projects (demand_source=projects).
        for sk_lc, d in demand_map.items():
            inv_row = inv_by_skill_lc.get(sk_lc)
            if allocation_count > 0:
                supply_count = int(allocation_supply.get(sk_lc, {}).get("allocated_count", 0))
            else:
                supply_count = int(inv_row.supply_count or 0) if inv_row else int(derived_supply.get(sk_lc, 0))
            demand_current = max(0, int(d.get("demand_count") or 0))
            priority = str((inv_row.priority if inv_row else d.get("priority")) or "MEDIUM").upper()
            pm = priority_mult.get(priority, 1.08)
            name = inv_row.skill_name if inv_row else (d.get("skill_name") or sk_lc)

            # Gap ratio: 0 when supply >= demand, approaches 1 when demand >> supply.
            gap_ratio = (max(0, demand_current - supply_count) / (demand_current + 1)) if demand_current > 0 else 0.0

            demand_forecast = int(round(demand_current * pm * (1 + 0.2 * gap_ratio) * horizon_growth))
            forecast_gap = max(0, demand_forecast - supply_count)
            forecast_gap_pct = round((forecast_gap / demand_forecast) * 100.0, 2) if demand_forecast > 0 else 0.0

            forecast_rows.append(
                WorkforceIntelligenceSkillForecast(
                    skill_name=name,
                    priority=priority,
                    demand_current=demand_current,
                    supply_count=supply_count,
                    demand_forecast=demand_forecast,
                    forecast_gap=forecast_gap,
                    forecast_gap_pct=forecast_gap_pct,
                )
            )
            demand_current_total += demand_current
            workforce_supply_total += supply_count
            demand_forecast_total += demand_forecast
    else:
        # Default: workforce_skills demand_source (near real-time heuristic).
        for s in skill_rows:
            name = s.skill_name or ""
            if not name:
                continue

            demand_current = max(0, int(s.demand_count or 0))
            supply_count = max(0, int(s.supply_count or 0))
            if allocation_count > 0:
                sk_lc = (name or "").strip().lower()
                supply_count = int(allocation_supply.get(sk_lc, {}).get("allocated_count", 0))
            priority = str(s.priority or "MEDIUM").upper()
            pm = priority_mult.get(priority, 1.08)

            gap_ratio = (max(0, demand_current - supply_count) / (demand_current + 1)) if demand_current > 0 else 0.0
            demand_forecast = int(round(demand_current * pm * (1 + 0.2 * gap_ratio) * horizon_growth))
            forecast_gap = max(0, demand_forecast - supply_count)
            forecast_gap_pct = round((forecast_gap / demand_forecast) * 100.0, 2) if demand_forecast > 0 else 0.0

            forecast_rows.append(
                WorkforceIntelligenceSkillForecast(
                    skill_name=name,
                    priority=priority,
                    demand_current=demand_current,
                    supply_count=supply_count,
                    demand_forecast=demand_forecast,
                    forecast_gap=forecast_gap,
                    forecast_gap_pct=forecast_gap_pct,
                )
            )
            demand_current_total += demand_current
            workforce_supply_total += supply_count
            demand_forecast_total += demand_forecast

    forecast_rows.sort(key=lambda x: x.forecast_gap, reverse=True)
    forecast_gap_total = sum(r.forecast_gap for r in forecast_rows)
    generated_at = datetime.now(timezone.utc).isoformat()

    resp = WorkforceIntelligenceResponse(
        horizon_months=horizon_months,
        generated_at=generated_at,
        skills_total=len(forecast_rows),
        demand_current_total=demand_current_total,
        workforce_supply_total=workforce_supply_total,
        demand_forecast_total=demand_forecast_total,
        forecast_gap_total=forecast_gap_total,
        top_forecast_gaps=forecast_rows[:10],
    )

    await db.workforce_intelligence_forecast_cache.update_one(
        {"horizon_months": horizon_months, "demand_source": demand_source},
        {"$set": {"cached_at": now_dt, "data": resp.model_dump(), "supply_source": supply_source}},
        upsert=True,
    )
    await _write_workforce_dashboard_audit_log(
        dashboard="workforce_intelligence",
        demand_source=demand_source,
        supply_source=supply_source,
        horizon_months=horizon_months,
        cache_hit=False,
        refresh_requested=refresh,
        generated_at=resp.generated_at,
        top_rows_count=len(resp.top_forecast_gaps or []),
    )

    return resp


# --- M3: historical pipeline, baseline model API, monitoring (Workforce Intelligence) ---


@api_router.get("/workforce/intelligence/model-forecast", response_model=WorkforceIntelModelForecastResponse)
async def get_workforce_intel_model_forecast(
    horizon_months: int = 1,
    current_user: dict = Depends(get_current_user),
):
    """
    M3-2: serve demand forecast from the **active** registered baseline model (per-skill linear trend on snapshots).
    `horizon_months` is treated as discrete snapshot-step lookahead (same cadence as ETL).
    """
    _require_phase1_access(current_user, "kpi_read")
    horizon_months = max(1, int(horizon_months or 1))
    st = await db[COL_MODEL_STATE].find_one({"_id": MODEL_STATE_DOC_ID}, {"_id": 0, "active_version_id": 1})
    version_id = (st or {}).get("active_version_id")
    if not version_id:
        raise HTTPException(status_code=404, detail="No active workforce intel model; train and activate (admin).")

    mdoc = await db[COL_MODELS].find_one({"version_id": version_id}, {"_id": 0})
    if not mdoc:
        raise HTTPException(status_code=404, detail="Active model version not found in registry.")

    per_skill = ((mdoc.get("params") or {}).get("per_skill")) or {}
    live_rows, demand_source, supply_source = await extract_workforce_intel_feature_rows(db)

    out_rows: List[WorkforceIntelModelForecastRow] = []
    demand_forecast_total = 0
    for r in live_rows:
        sk = (r.get("skill_name_lc") or "").strip().lower()
        if not sk or sk not in per_skill:
            continue
        p = BaselineParams.from_dict(per_skill[sk])
        d_hat = predict_demand(p, steps_ahead=horizon_months)
        supply_count = max(0, int(r.get("supply_count") or 0))
        demand_current = max(0, int(r.get("demand_current") or 0))
        fg = max(0, d_hat - supply_count)
        fg_pct = round((fg / d_hat) * 100.0, 2) if d_hat > 0 else 0.0
        out_rows.append(
            WorkforceIntelModelForecastRow(
                skill_name=str(r.get("skill_name") or sk),
                skill_name_lc=sk,
                priority=str(r.get("priority") or "MEDIUM"),
                demand_current=demand_current,
                supply_count=supply_count,
                demand_forecast_model=d_hat,
                forecast_gap=fg,
                forecast_gap_pct=fg_pct,
                horizon_steps=horizon_months,
            )
        )
        demand_forecast_total += d_hat

    out_rows.sort(key=lambda x: x.forecast_gap, reverse=True)
    generated_at = datetime.now(timezone.utc).isoformat()
    forecast_gap_total = sum(r.forecast_gap for r in out_rows)
    return WorkforceIntelModelForecastResponse(
        version_id=version_id,
        horizon_months=horizon_months,
        generated_at=generated_at,
        demand_source=demand_source,
        supply_source=supply_source,
        skills_total=len(out_rows),
        demand_forecast_total=demand_forecast_total,
        forecast_gap_total=forecast_gap_total,
        top_forecast_gaps=out_rows[:50],
    )


@api_router.post("/admin/workforce-intel/etl/snapshot")
async def admin_workforce_intel_etl_snapshot(
    current_user: dict = Depends(get_current_user),
):
    """M3-1: extract features + DQ + persist snapshot rows."""
    _require_admin(current_user)
    return await etl_snapshot(db, enforce_dq=True, actor_id=current_user.get("id"))


@api_router.post("/admin/workforce-intel/etl/backfill")
async def admin_workforce_intel_etl_backfill(
    payload: M3EtlBackfillRequest,
    current_user: dict = Depends(get_current_user),
):
    """M3-1: deterministic demo backfill for labs without historical snapshots."""
    _require_admin(current_user)
    return await etl_backfill_demo(
        db,
        days=payload.days,
        seed=payload.seed,
        actor_id=current_user.get("id"),
    )


@api_router.get("/admin/workforce-intel/etl/last-run")
async def admin_workforce_intel_etl_last_run(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    rows = await db[COL_ETL_RUNS].find({}, {"_id": 0}).sort("ended_at", -1).limit(1).to_list(1)
    return rows[0] if rows else {}


@api_router.post("/admin/workforce-intel/models/train")
async def admin_workforce_intel_train_model(
    payload: M3TrainModelRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    M3-2: fit baseline v1 on `workforce_intel_hist_features` for the current demand_source slice.
    """
    _require_admin(current_user)
    _rows, demand_source, _ss = await extract_workforce_intel_feature_rows(db)
    series, times = await load_demand_series_by_skill(db, demand_source, max_snapshots=payload.max_snapshots)
    if len(times) < 2:
        raise HTTPException(
            status_code=400,
            detail="Not enough history (need >= 2 snapshots). Run ETL snapshot repeatedly or use demo backfill.",
        )

    trained = fit_per_skill_baseline(series, min_points=2)
    mape, mae, n_ev, _per = evaluate_on_history(series)
    version_id = datetime.now(timezone.utc).strftime("v1-%Y%m%dT%H%M%SZ-") + str(uuid.uuid4())[:8]
    params = {sk: p.to_dict() for sk, p in trained.items()}
    now = datetime.now(timezone.utc)
    doc = {
        "version_id": version_id,
        "created_at": now,
        "active": bool(payload.activate),
        "algorithm": "baseline_linear_per_skill",
        "demand_source": demand_source,
        "params": {"per_skill": params},
        "train_metrics": {
            "mape_pct_holdout": round(mape, 4),
            "mae_holdout": round(mae, 4),
            "holdout_points": n_ev,
            "snapshot_count": len(times),
            "skills_trained": len(trained),
        },
        "created_by": current_user.get("id"),
        "rollback_notes": "Activate prior version via POST /admin/workforce-intel/models/{version_id}/activate",
    }
    await db[COL_MODELS].insert_one(doc)

    if payload.activate:
        await db[COL_MODEL_STATE].update_one(
            {"_id": MODEL_STATE_DOC_ID},
            {"$set": {"active_version_id": version_id, "updated_at": now, "updated_by": current_user.get("id")}},
            upsert=True,
        )

    return {"version_id": version_id, "train_metrics": doc["train_metrics"], "activated": bool(payload.activate)}


@api_router.post("/admin/workforce-intel/models/{version_id}/activate")
async def admin_workforce_intel_activate_model(version_id: str, current_user: dict = Depends(get_current_user)):
    """M3-2: rollback / promote — point active model to an existing registry version."""
    _require_admin(current_user)
    exists = await db[COL_MODELS].find_one({"version_id": version_id}, {"_id": 1})
    if not exists:
        raise HTTPException(status_code=404, detail="Unknown version_id")
    now = datetime.now(timezone.utc)
    await db[COL_MODEL_STATE].update_one(
        {"_id": MODEL_STATE_DOC_ID},
        {"$set": {"active_version_id": version_id, "updated_at": now, "updated_by": current_user.get("id")}},
        upsert=True,
    )
    await db[COL_MODELS].update_many({}, {"$set": {"active": False}})
    await db[COL_MODELS].update_one({"version_id": version_id}, {"$set": {"active": True}})
    return {"active_version_id": version_id}


@api_router.get("/admin/workforce-intel/models")
async def admin_workforce_intel_list_models(
    limit: int = 30,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    limit = min(max(1, int(limit)), 100)
    st = await db[COL_MODEL_STATE].find_one({"_id": MODEL_STATE_DOC_ID}, {"_id": 0})
    active = (st or {}).get("active_version_id")
    cur = db[COL_MODELS].find({}, {"_id": 0, "params": 0}).sort("created_at", -1).limit(limit)
    rows = await cur.to_list(limit)
    for r in rows:
        r["is_active"] = r.get("version_id") == active
    return {"active_version_id": active, "items": rows}


@api_router.post("/admin/workforce-intel/monitoring/evaluate")
async def admin_workforce_intel_monitoring_evaluate(current_user: dict = Depends(get_current_user)):
    """M3-3: MAPE/MAE vs current live features + drift alerts + retrain flag."""
    _require_admin(current_user)
    return await evaluate_active_model_vs_current(db, actor_id=current_user.get("id"))


@api_router.get("/admin/workforce-intel/monitoring/summary")
async def admin_workforce_intel_monitoring_summary(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    st = await db[COL_MONITORING_STATE].find_one({"_id": MONITORING_STATE_DOC_ID}, {"_id": 0})
    evals = await db[COL_EVAL_RUNS].find({}, {"_id": 0}).sort("evaluated_at", -1).limit(10).to_list(10)
    drift = await db[COL_DRIFT_EVENTS].find({}, {"_id": 0}).sort("created_at", -1).limit(10).to_list(10)
    policy = await retrain_trigger_evaluation(db)
    return {"monitoring_state": st, "recent_evaluations": evals, "recent_drift_events": drift, "retrain_policy": policy}


@api_router.get("/admin/workforce-intel/monitoring/retrain-policy")
async def admin_workforce_intel_retrain_policy(current_user: dict = Depends(get_current_user)):
    """M3-3: expose retraining trigger status (driven by last evaluation thresholds)."""
    _require_admin(current_user)
    return await retrain_trigger_evaluation(db)


@api_router.get("/workforce/resource-optimization", response_model=ResourceOptimizationResponse)
async def get_resource_optimization(
    refresh: bool = False,
    current_user: dict = Depends(get_current_user),
):
    """
    Phase-3 M4 MVP:
    - Uses workforce skill supply/demand (derived supply from employee profiles where possible)
    - Produces utilization, bench, shortage, and allocation status per skill

    Full project-resource mapping is not present in this repo; this endpoint delivers the
    closest utilization/bench/alert governance dashboard using available Phase-1 data.
    """
    _require_phase1_access(current_user, "kpi_read")

    now_dt = datetime.now(timezone.utc)
    project_demand_count = await db.project_skill_demands.count_documents({})
    allocation_count = await db.project_skill_allocations.count_documents({})
    demand_source = "projects" if project_demand_count > 0 else "workforce_skills"
    supply_source = "allocations" if allocation_count > 0 else "employees"
    if not refresh:
        cached_doc = await db.workforce_resource_optimization_cache.find_one(
            {"demand_source": demand_source},
            {"_id": 0},
        )
        if await _get_dashboard_cache_ttl_hit(cached_doc, now_dt):
            # If supply mode changed, don't reuse cached data.
            if cached_doc.get("supply_source") != supply_source:
                cached_doc = None
            else:
                cached_data = cached_doc.get("data")
                if not cached_data:
                    pass
                else:
                    cached_resp = ResourceOptimizationResponse(**cached_data)
                    await _write_workforce_dashboard_audit_log(
                        dashboard="resource_optimization",
                        demand_source=demand_source,
                        supply_source=supply_source,
                        horizon_months=None,
                        cache_hit=True,
                        refresh_requested=False,
                        generated_at=cached_resp.generated_at,
                        top_rows_count=None,
                        total_shortage=cached_resp.total_shortage,
                        total_bench=cached_resp.total_bench,
                    )
                    return cached_resp

    skill_rows: List[SkillInventoryResponse] = await list_skill_inventory(current_user)

    generated_at = datetime.now(timezone.utc).isoformat()

    under: List[ResourceOptimizationSkillMetrics] = []
    over: List[ResourceOptimizationSkillMetrics] = []
    total_bench = 0
    total_shortage = 0
    total_demand = 0
    total_supply = 0

    demand_map: Optional[Dict[str, Dict[str, Any]]] = None
    if project_demand_count > 0:
        demand_map = await _get_project_demand_map()

    inv_by_skill_lc: Dict[str, SkillInventoryResponse] = {}
    for r in skill_rows:
        key = (r.skill_name or "").strip().lower()
        if key:
            inv_by_skill_lc[key] = r

    derived_supply: Dict[str, int] = {}
    allocation_supply: Dict[str, Dict[str, Any]] = {}
    all_keys: set[str] = set(inv_by_skill_lc.keys())
    if demand_map:
        derived_supply = await _derived_supply_from_employees()
        all_keys |= set(demand_map.keys())
    if allocation_count > 0:
        allocation_supply = await _get_project_allocations_supply_map()
        all_keys |= set(allocation_supply.keys())

    for sk_lc in all_keys:
        inv_row = inv_by_skill_lc.get(sk_lc)
        if allocation_count > 0:
            supply = int(allocation_supply.get(sk_lc, {}).get("allocated_count", 0))
        else:
            supply = int(inv_row.supply_count or 0) if inv_row else int(derived_supply.get(sk_lc, 0))

        if demand_map:
            d = demand_map.get(sk_lc) or {}
            demand = int(d.get("demand_count") or 0)
            priority = str(d.get("priority") or (inv_row.priority if inv_row else "MEDIUM")).upper()
            skill_name = inv_row.skill_name if inv_row else (d.get("skill_name") or sk_lc)
        else:
            demand = int(inv_row.demand_count or 0) if inv_row else 0
            priority = str(inv_row.priority or "MEDIUM").upper() if inv_row else "MEDIUM"
            skill_name = inv_row.skill_name if inv_row else sk_lc

        total_demand += demand
        total_supply += supply

        shortage = max(0, demand - supply)
        bench = max(0, supply - demand)
        utilization_pct = round((supply / demand) * 100.0, 2) if demand > 0 else 0.0

        if shortage > 0:
            status = "UNDER_ALLOCATED"
        elif bench > 0:
            status = "OVER_ALLOCATED"
        else:
            status = "BALANCED"

        metrics = ResourceOptimizationSkillMetrics(
            skill_name=skill_name,
            priority=priority,
            demand_count=demand,
            supply_count=supply,
            utilization_pct=utilization_pct,
            allocation_status=status,
            bench_count=bench,
            shortage_count=shortage,
        )

        if status == "UNDER_ALLOCATED":
            under.append(metrics)
            total_shortage += shortage
        elif status == "OVER_ALLOCATED":
            over.append(metrics)
            total_bench += bench

    # Sort alerts for quick visibility.
    under.sort(key=lambda x: x.shortage_count, reverse=True)
    over.sort(key=lambda x: x.bench_count, reverse=True)

    skills_total = len(all_keys)
    resp = ResourceOptimizationResponse(
        generated_at=generated_at,
        skills_total=skills_total,
        total_bench=total_bench,
        total_shortage=total_shortage,
        total_demand=total_demand,
        total_supply=total_supply,
        under_allocated=under,
        over_allocated=over,
    )

    await db.workforce_resource_optimization_cache.update_one(
        {"demand_source": demand_source},
        {"$set": {"cached_at": now_dt, "data": resp.model_dump(), "supply_source": supply_source}},
        upsert=True,
    )
    await _write_workforce_dashboard_audit_log(
        dashboard="resource_optimization",
        demand_source=demand_source,
        supply_source=supply_source,
        horizon_months=None,
        cache_hit=False,
        refresh_requested=refresh,
        generated_at=resp.generated_at,
        top_rows_count=None,
        total_shortage=resp.total_shortage,
        total_bench=resp.total_bench,
    )

    return resp


# --- M4: capacity constraints, solver, what-if scenarios, approvals ---


@api_router.get("/workforce/resource-optimization/settings", response_model=AllocationOptimizationSettingsResponse)
async def get_allocation_optimization_settings(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    await ensure_default_settings(db)
    s = await get_merged_settings(db)
    return AllocationOptimizationSettingsResponse(
        max_projects_per_employee=int(s.get("max_projects_per_employee") or 3),
        max_seats_per_employee_per_project=int(s.get("max_seats_per_employee_per_project") or 1),
        shortage_penalty_hard=float(s.get("shortage_penalty_hard") or 10.0),
        shortage_penalty_soft=float(s.get("shortage_penalty_soft") or 3.0),
        utilization_weight=float(s.get("utilization_weight") or 4.0),
        target_utilization_pct=float(s.get("target_utilization_pct") or 85.0),
    )


@api_router.put("/workforce/resource-optimization/settings", response_model=AllocationOptimizationSettingsResponse)
async def put_allocation_optimization_settings(
    payload: AllocationOptimizationSettingsUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    await ensure_default_settings(db)
    patch = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    if patch:
        await db[COL_ALLOCATION_SETTINGS].update_one(
            {"_id": SETTINGS_DOC_ID},
            {"$set": patch},
            upsert=True,
        )
    s = await get_merged_settings(db)
    return AllocationOptimizationSettingsResponse(
        max_projects_per_employee=int(s.get("max_projects_per_employee") or 3),
        max_seats_per_employee_per_project=int(s.get("max_seats_per_employee_per_project") or 1),
        shortage_penalty_hard=float(s.get("shortage_penalty_hard") or 10.0),
        shortage_penalty_soft=float(s.get("shortage_penalty_soft") or 3.0),
        utilization_weight=float(s.get("utilization_weight") or 4.0),
        target_utilization_pct=float(s.get("target_utilization_pct") or 85.0),
    )


@api_router.post("/workforce/resource-optimization/solve", response_model=Dict[str, Any])
async def post_resource_optimization_solve(current_user: dict = Depends(get_current_user)):
    """M4-2: deterministic allocation run against live DB demands (no persistence)."""
    _require_phase1_access(current_user, "kpi_read")
    return await run_allocation_solve(db)


@api_router.post("/workforce/resource-optimization/simulate", response_model=Dict[str, Any])
async def post_resource_optimization_simulate(
    payload: AllocationSimulateRequest,
    current_user: dict = Depends(get_current_user),
):
    """M4-3: what-if run with demand/constraint overrides (no persistence)."""
    _require_phase1_access(current_user, "kpi_read")
    return await run_allocation_solve(
        db,
        demand_overrides=payload.demand_overrides,
        constraint_overrides=payload.constraint_overrides,
    )


@api_router.post("/workforce/resource-optimization/scenarios", response_model=Dict[str, Any])
async def create_allocation_scenario(
    payload: AllocationScenarioCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    result = payload.result
    if not result:
        result = await run_allocation_solve(
            db,
            demand_overrides=payload.demand_overrides,
            constraint_overrides=payload.constraint_overrides,
        )
    sid = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": sid,
        "name": (payload.name or "").strip() or "Scenario",
        "description": (payload.description or "").strip() or None,
        "status": "DRAFT",
        "payload": {
            "demand_overrides": payload.demand_overrides,
            "constraint_overrides": payload.constraint_overrides,
        },
        "result": result,
        "created_at": now,
        "created_by": current_user.get("id"),
        "submitted_at": None,
        "approved_by": None,
        "approved_at": None,
        "rejection_reason": None,
    }
    await db[COL_ALLOCATION_SCENARIOS].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.get("/workforce/resource-optimization/scenarios", response_model=List[Dict[str, Any]])
async def list_allocation_scenarios(
    limit: int = 30,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    limit = min(max(1, int(limit)), 100)
    cur = (
        db[COL_ALLOCATION_SCENARIOS]
        .find({}, {"_id": 0, "result": 0})
        .sort("created_at", -1)
        .limit(limit)
    )
    return await cur.to_list(limit)


@api_router.get("/workforce/resource-optimization/scenarios/compare", response_model=Dict[str, Any])
async def compare_allocation_scenarios(
    scenario_a_id: str,
    scenario_b_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    a = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_a_id}, {"_id": 0, "result": 1, "name": 1})
    b = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_b_id}, {"_id": 0, "result": 1, "name": 1})
    if not a or not b:
        raise HTTPException(status_code=404, detail="Scenario not found")
    ra = a.get("result") or {}
    rb = b.get("result") or {}
    return {
        "scenario_a": {"id": scenario_a_id, "name": a.get("name")},
        "scenario_b": {"id": scenario_b_id, "name": b.get("name")},
        "comparison": compare_solve_results(ra, rb),
    }


@api_router.get("/workforce/resource-optimization/scenarios/{scenario_id}", response_model=Dict[str, Any])
async def get_allocation_scenario(scenario_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    doc = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    return doc


@api_router.post("/workforce/resource-optimization/scenarios/{scenario_id}/submit", response_model=Dict[str, Any])
async def submit_allocation_scenario(scenario_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "skills_write")
    doc = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    if doc.get("status") not in ("DRAFT", "REJECTED"):
        raise HTTPException(status_code=400, detail="Scenario cannot be submitted from this state")
    now = datetime.now(timezone.utc).isoformat()
    await db[COL_ALLOCATION_SCENARIOS].update_one(
        {"id": scenario_id},
        {"$set": {"status": "PENDING_APPROVAL", "submitted_at": now}},
    )
    for rid in await _hr_escalation_recipient_ids():
        await create_notification(
            recipient_id=rid,
            notification_type="ALLOCATION_SCENARIO_PENDING_APPROVAL",
            title="Allocation scenario pending approval",
            message=f"Scenario '{doc.get('name')}' submitted for approval.",
            metadata={"scenario_id": scenario_id},
        )
    updated = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    return updated


@api_router.post("/workforce/resource-optimization/scenarios/{scenario_id}/approve", response_model=Dict[str, Any])
async def approve_allocation_scenario(scenario_id: str, current_user: dict = Depends(get_current_user)):
    _require_allocation_approver(current_user)
    doc = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    if doc.get("status") != "PENDING_APPROVAL":
        raise HTTPException(status_code=400, detail="Scenario is not pending approval")
    now = datetime.now(timezone.utc).isoformat()
    await db[COL_ALLOCATION_SCENARIOS].update_one(
        {"id": scenario_id},
        {"$set": {"status": "APPROVED", "approved_by": current_user.get("id"), "approved_at": now, "rejection_reason": None}},
    )
    updated = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    return updated


@api_router.post("/workforce/resource-optimization/scenarios/{scenario_id}/reject", response_model=Dict[str, Any])
async def reject_allocation_scenario(
    scenario_id: str,
    payload: AllocationScenarioRejectRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_allocation_approver(current_user)
    doc = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    if doc.get("status") != "PENDING_APPROVAL":
        raise HTTPException(status_code=400, detail="Scenario is not pending approval")
    now = datetime.now(timezone.utc).isoformat()
    reason = (payload.reason or "").strip() or "Rejected"
    await db[COL_ALLOCATION_SCENARIOS].update_one(
        {"id": scenario_id},
        {
            "$set": {
                "status": "REJECTED",
                "rejection_reason": reason,
                "approved_by": None,
                "approved_at": None,
                "updated_at": now,
            }
        },
    )
    updated = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    return updated


@api_router.post("/workforce/resource-optimization/scenarios/{scenario_id}/apply", response_model=Dict[str, Any])
async def apply_allocation_scenario(
    scenario_id: str,
    payload: AllocationScenarioApplyRequest,
    current_user: dict = Depends(get_current_user),
):
    """Persist solver assignments into `project_skill_allocations` (requires APPROVED)."""
    _require_allocation_approver(current_user)
    doc = await db[COL_ALLOCATION_SCENARIOS].find_one({"id": scenario_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    if doc.get("status") != "APPROVED":
        raise HTTPException(status_code=400, detail="Scenario must be APPROVED before apply")
    result = doc.get("result") or {}
    return await apply_assignments_to_project_allocations(
        db,
        result,
        actor_id=current_user.get("id") or "",
        dry_run=payload.dry_run,
    )


#
# ========================
# Phase-3 M4: Project CRUD + Skill Demand MVP
# ========================
#

@api_router.post("/projects", response_model=ProjectResponse)
async def create_project(
    payload: ProjectCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    now = datetime.now(timezone.utc).isoformat()
    project_id = str(uuid.uuid4())
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Project name is required")
    doc = {
        "id": project_id,
        "name": name,
        "description": payload.description,
        "client_name": (payload.client_name or "").strip() or None,
        "business_unit": (payload.business_unit or "").strip() or None,
        "project_type": (str(payload.project_type).upper() if payload.project_type else None),
        "status": (payload.status or "ACTIVE").upper(),
        "start_date": payload.start_date,
        "end_date": payload.end_date,
        "project_manager_id": (payload.project_manager_id or "").strip() or None,
        "budget": payload.budget,
        "billing_type": (str(payload.billing_type).upper() if payload.billing_type else None),
        "required_skills": [r.model_dump() for r in (payload.required_skills or [])],
        "milestones": [m.model_dump() for m in (payload.milestones or [])],
        "documents": [d.model_dump() for d in (payload.documents or [])],
        "created_at": now,
        "updated_at": None,
    }
    await db.projects.insert_one(doc)
    return ProjectResponse(**doc)


@api_router.get("/projects", response_model=List[ProjectResponse])
async def list_projects(
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    query: Dict[str, Any] = {}
    if status:
        query["status"] = status.upper()
    rows = await db.projects.find(query, {"_id": 0}).sort("created_at", -1).to_list(200)
    return [ProjectResponse(**r) for r in rows]


@api_router.put("/projects/{project_id}", response_model=ProjectResponse)
async def update_project(
    project_id: str,
    payload: ProjectUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    existing = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Project not found")

    update_doc = payload.model_dump(exclude_none=True)
    if "name" in update_doc and isinstance(update_doc["name"], str):
        update_doc["name"] = update_doc["name"].strip()
    if "status" in update_doc and isinstance(update_doc["status"], str):
        update_doc["status"] = update_doc["status"].upper()
    update_doc["updated_at"] = datetime.now(timezone.utc).isoformat()

    if update_doc:
        await db.projects.update_one({"id": project_id}, {"$set": update_doc})

    updated = await db.projects.find_one({"id": project_id}, {"_id": 0})
    return ProjectResponse(**updated)


@api_router.delete("/projects/{project_id}")
async def delete_project(
    project_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    res = await db.projects.delete_one({"id": project_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")
    await db.project_skill_demands.delete_many({"project_id": project_id})
    await db.project_skill_allocations.delete_many({"project_id": project_id})
    await db.allocations.delete_many({"project_id": project_id})
    return {"message": "Project deleted"}


# ========================
# Project Section APIs (Enterprise Project Management)
# ========================

def _today_key() -> str:
    return datetime.now(timezone.utc).date().isoformat()


def _norm_lc(s: Optional[str]) -> str:
    return str(s or "").strip().lower()


def _safe_days_between(a: Optional[str], b: Optional[str]) -> Optional[int]:
    try:
        if not a or not b:
            return None
        da = date.fromisoformat(str(a)[:10])
        db = date.fromisoformat(str(b)[:10])
        return int((db - da).days)
    except Exception:
        return None


def _project_master_to_response(doc: Dict[str, Any]) -> ProjectMasterResponse:
    start = doc.get("start_date")
    end = doc.get("end_date")
    planned_duration_days = _safe_days_between(start, end)
    return ProjectMasterResponse(
        **{k: v for k, v in doc.items() if k != "_id"},
        planned_duration_days=planned_duration_days,
    )


async def _next_project_code(db) -> str:
    # Simple sequential code: PRJ-00001 (atomic-ish using a counters collection)
    now = datetime.now(timezone.utc).isoformat()
    res = await db.counters.find_one_and_update(
        {"_id": "project_code"},
        {"$inc": {"seq": 1}, "$setOnInsert": {"created_at": now}},
        upsert=True,
        return_document=True,
    )
    seq = int((res or {}).get("seq") or 1)
    return f"PRJ-{seq:05d}"


def _validate_project_master(payload: Dict[str, Any]):
    name = str(payload.get("project_name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="project_name is required")
    st = payload.get("start_date")
    en = payload.get("end_date")
    if st and en and str(en)[:10] < str(st)[:10]:
        raise HTTPException(status_code=400, detail="end_date cannot be before start_date")
    pt = _norm_lc(payload.get("project_type") or "external")
    if pt not in PROJECT_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid project_type. Allowed: {', '.join(PROJECT_TYPES)}")
    pr = _norm_lc(payload.get("project_priority") or "medium")
    if pr not in PROJECT_PRIORITIES:
        raise HTTPException(status_code=400, detail=f"Invalid project_priority. Allowed: {', '.join(PROJECT_PRIORITIES)}")
    ph = _norm_lc(payload.get("project_health") or "green")
    if ph not in PROJECT_HEALTH:
        raise HTTPException(status_code=400, detail=f"Invalid project_health. Allowed: {', '.join(PROJECT_HEALTH)}")
    ps = _norm_lc(payload.get("project_status") or "draft")
    if ps not in PROJECT_LIFECYCLE_STATES:
        raise HTTPException(status_code=400, detail=f"Invalid project_status. Allowed: {', '.join(PROJECT_LIFECYCLE_STATES)}")
    if payload.get("project_budget") is not None and float(payload.get("project_budget") or 0) < 0:
        raise HTTPException(status_code=400, detail="project_budget cannot be negative")
    if payload.get("expected_revenue") is not None and float(payload.get("expected_revenue") or 0) < 0:
        raise HTTPException(status_code=400, detail="expected_revenue cannot be negative")


@api_router.get("/project-section/dashboard/summary")
async def project_section_dashboard_summary(
    q: Optional[str] = None,
    status: Optional[str] = None,
    business_unit: Optional[str] = None,
    manager_id: Optional[str] = None,
    client_name: Optional[str] = None,
    from_: Optional[str] = Query(default=None, alias="from"),
    to: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_project_view(current_user)
    query: Dict[str, Any] = {"is_deleted": {"$ne": True}}
    if q:
        query["$or"] = [
            {"project_name": {"$regex": q, "$options": "i"}},
            {"project_code": {"$regex": q, "$options": "i"}},
            {"client_name": {"$regex": q, "$options": "i"}},
        ]
    if status:
        query["project_status"] = _norm_lc(status)
    if business_unit:
        query["business_unit"] = {"$regex": f"^{re.escape(business_unit)}$", "$options": "i"}
    if manager_id:
        query["project_manager_id"] = manager_id
    if client_name:
        query["client_name"] = {"$regex": f"^{re.escape(client_name)}$", "$options": "i"}
    if from_ or to:
        rng: Dict[str, Any] = {}
        if from_:
            rng["$gte"] = str(from_)[:10]
        if to:
            rng["$lte"] = str(to)[:10]
        query["start_date"] = rng

    total = await db.project_masters.count_documents(query)
    active = await db.project_masters.count_documents({**query, "project_status": "active"})
    completed = await db.project_masters.count_documents({**query, "project_status": {"$in": ["completed", "closed"]}})
    on_hold = await db.project_masters.count_documents({**query, "project_status": "on_hold"})

    today = _today_key()
    delayed = await db.project_masters.count_documents(
        {**query, "end_date": {"$lt": today}, "project_status": {"$in": ["active", "on_hold", "approved"]}}
    )

    pipeline = [
        {"$match": query},
        {"$group": {"_id": "$project_status", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ]
    by_status = [{"key": r["_id"], "count": r["count"]} for r in (await db.project_masters.aggregate(pipeline).to_list(50))]

    quick = (
        await db.project_masters.find(query, {"_id": 0})
        .sort("updated_at", -1)
        .limit(10)
        .to_list(10)
    )
    quick_list = [
        {
            "id": r.get("id"),
            "project_code": r.get("project_code"),
            "project_name": r.get("project_name"),
            "project_status": r.get("project_status"),
            "end_date": r.get("end_date"),
        }
        for r in quick
    ]

    return {
        "kpi": {
            "total_projects": total,
            "active_projects": active,
            "completed_projects": completed,
            "on_hold_projects": on_hold,
            "delayed_projects": delayed,
        },
        "distributions": {"by_status": by_status},
        "quick_list": quick_list,
    }


@api_router.get("/project-section/projects")
async def project_section_list_projects(
    page: int = 1,
    page_size: int = 25,
    q: Optional[str] = None,
    status: Optional[str] = None,
    priority: Optional[str] = None,
    business_unit: Optional[str] = None,
    client_name: Optional[str] = None,
    include_archived: bool = False,
    current_user: dict = Depends(get_current_user),
):
    _require_project_view(current_user)
    page = max(1, page)
    page_size = min(max(1, page_size), 200)
    query: Dict[str, Any] = {"is_deleted": {"$ne": True}}
    if not include_archived:
        query["is_archived"] = {"$ne": True}
    if q:
        query["$or"] = [
            {"project_name": {"$regex": q, "$options": "i"}},
            {"project_code": {"$regex": q, "$options": "i"}},
            {"client_name": {"$regex": q, "$options": "i"}},
        ]
    if status:
        query["project_status"] = _norm_lc(status)
    if priority:
        query["project_priority"] = _norm_lc(priority)
    if business_unit:
        query["business_unit"] = {"$regex": f"^{re.escape(business_unit)}$", "$options": "i"}
    if client_name:
        query["client_name"] = {"$regex": f"^{re.escape(client_name)}$", "$options": "i"}

    total = await db.project_masters.count_documents(query)
    rows = (
        await db.project_masters.find(query, {"_id": 0})
        .sort("updated_at", -1)
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )
    return {
        "items": [_project_master_to_response(r).model_dump() for r in rows],
        "page": page,
        "page_size": page_size,
        "total": total,
        "total_pages": max(1, (total + page_size - 1) // page_size),
    }


@api_router.post("/project-section/projects", response_model=ProjectMasterResponse)
async def project_section_create_project(
    payload: ProjectMasterCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["project_name"] = str(doc.get("project_name") or "").strip()
    doc["project_type"] = _norm_lc(doc.get("project_type") or "external")
    doc["project_priority"] = _norm_lc(doc.get("project_priority") or "medium")
    doc["project_health"] = _norm_lc(doc.get("project_health") or "green")
    doc["project_status"] = _norm_lc(doc.get("project_status") or "draft")
    if not doc.get("project_code"):
        doc["project_code"] = await _next_project_code(db)
    else:
        doc["project_code"] = str(doc["project_code"]).strip().upper()
    _validate_project_master(doc)

    dup = await db.project_masters.find_one(
        {"project_code": doc["project_code"], "is_deleted": {"$ne": True}},
        {"_id": 0, "id": 1},
    )
    if dup:
        raise HTTPException(status_code=409, detail="project_code already exists")

    project_uuid = str(uuid.uuid4())
    doc2 = {
        "id": project_uuid,
        "project_id": project_uuid,
        **doc,
        "tags": [str(t).strip() for t in (doc.get("tags") or []) if str(t).strip()],
        "is_archived": False,
        "is_deleted": False,
        "created_at": now,
        "updated_at": now,
        "created_by": current_user.get("id"),
        "updated_by": current_user.get("id"),
    }
    await db.project_masters.insert_one(doc2)
    await db.project_lifecycle_history.insert_one(
        {
            "id": str(uuid.uuid4()),
            "project_id": project_uuid,
            "from_state": None,
            "to_state": doc2["project_status"],
            "reason": "Created",
            "changed_by": current_user.get("id"),
            "changed_at": now,
        }
    )
    return _project_master_to_response(doc2)


@api_router.get("/project-section/projects/{project_id}", response_model=ProjectMasterResponse)
async def project_section_get_project(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    doc = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Project not found")
    return _project_master_to_response(doc)


@api_router.put("/project-section/projects/{project_id}", response_model=ProjectMasterResponse)
async def project_section_update_project(
    project_id: str,
    payload: ProjectMasterUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    existing = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Project not found")
    patch = payload.model_dump(exclude_none=True)
    if "project_name" in patch:
        patch["project_name"] = str(patch["project_name"]).strip()
    if "project_code" in patch and patch.get("project_code"):
        patch["project_code"] = str(patch["project_code"]).strip().upper()
        dup = await db.project_masters.find_one(
            {"project_code": patch["project_code"], "id": {"$ne": project_id}, "is_deleted": {"$ne": True}},
            {"_id": 0, "id": 1},
        )
        if dup:
            raise HTTPException(status_code=409, detail="project_code already exists")
    for k in ("project_type", "project_priority", "project_health", "project_status"):
        if k in patch and patch.get(k) is not None:
            patch[k] = _norm_lc(patch.get(k))
    merged = {**existing, **patch}
    _validate_project_master(merged)

    now = datetime.now(timezone.utc).isoformat()
    patch["updated_at"] = now
    patch["updated_by"] = current_user.get("id")
    await db.project_masters.update_one({"id": project_id}, {"$set": patch})
    doc = await db.project_masters.find_one({"id": project_id}, {"_id": 0})
    return _project_master_to_response(doc)


@api_router.post("/project-section/projects/{project_id}/archive", response_model=ProjectMasterResponse)
async def project_section_archive_project(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    now = datetime.now(timezone.utc).isoformat()
    res = await db.project_masters.update_one(
        {"id": project_id, "is_deleted": {"$ne": True}},
        {"$set": {"is_archived": True, "updated_at": now, "updated_by": current_user.get("id")}},
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")
    doc = await db.project_masters.find_one({"id": project_id}, {"_id": 0})
    return _project_master_to_response(doc)


@api_router.post("/project-section/projects/{project_id}/clone", response_model=ProjectMasterResponse)
async def project_section_clone_project(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    existing = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Project not found")
    now = datetime.now(timezone.utc).isoformat()
    new_id = str(uuid.uuid4())
    new_code = await _next_project_code(db)
    clone = {k: v for k, v in existing.items() if k not in ("_id", "id", "project_id")}
    clone.update(
        {
            "id": new_id,
            "project_id": new_id,
            "project_code": new_code,
            "project_name": f'{existing.get("project_name")} (Clone)',
            "project_status": "draft",
            "is_archived": False,
            "is_deleted": False,
            "created_at": now,
            "updated_at": now,
            "created_by": current_user.get("id"),
            "updated_by": current_user.get("id"),
        }
    )
    await db.project_masters.insert_one(clone)
    await db.project_lifecycle_history.insert_one(
        {
            "id": str(uuid.uuid4()),
            "project_id": new_id,
            "from_state": None,
            "to_state": "draft",
            "reason": "Cloned",
            "changed_by": current_user.get("id"),
            "changed_at": now,
        }
    )
    return _project_master_to_response(clone)


def _allowed_lifecycle_transition(frm: str, to: str) -> bool:
    frm = _norm_lc(frm)
    to = _norm_lc(to)
    allowed = {
        "draft": {"proposed", "cancelled"},
        "proposed": {"under_review", "cancelled"},
        "under_review": {"approved", "cancelled"},
        "approved": {"active", "on_hold", "cancelled"},
        "active": {"on_hold", "completed", "cancelled"},
        "on_hold": {"active", "cancelled"},
        "completed": {"closed"},
        "closed": set(),
        "cancelled": set(),
    }
    return to in allowed.get(frm, set())


@api_router.post("/project-section/projects/{project_id}/lifecycle/transition")
async def project_section_lifecycle_transition(
    project_id: str,
    payload: ProjectLifecycleTransitionRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    doc = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Project not found")
    frm = _norm_lc(doc.get("project_status"))
    to_state = _norm_lc(payload.to_state)
    if to_state not in PROJECT_LIFECYCLE_STATES:
        raise HTTPException(status_code=400, detail="Invalid lifecycle state")
    if not _allowed_lifecycle_transition(frm, to_state):
        raise HTTPException(status_code=400, detail=f"Invalid transition: {frm} -> {to_state}")

    # Require approval before activation (approved -> active)
    if frm == "approved" and to_state == "active":
        # Create approval request if not exists approved
        pending = await db.project_approvals.find_one(
            {
                "project_id": project_id,
                "type": "lifecycle_transition",
                "status": {"$in": ["pending", "submitted"]},
            },
            {"_id": 0},
        )
        if pending:
            raise HTTPException(status_code=400, detail="Activation requires approval (pending already created)")
        now = datetime.now(timezone.utc).isoformat()
        appr_id = str(uuid.uuid4())
        await db.project_approvals.insert_one(
            {
                "id": appr_id,
                "project_id": project_id,
                "type": "lifecycle_transition",
                "payload": {"from_state": frm, "to_state": to_state, "reason": payload.reason},
                "status": "pending",
                "requested_by": current_user.get("id"),
                "requested_at": now,
                "decided_by": None,
                "decided_at": None,
                "decision_reason": None,
            }
        )
        return {"message": "Approval required", "approval_id": appr_id}

    now = datetime.now(timezone.utc).isoformat()
    await db.project_masters.update_one(
        {"id": project_id},
        {"$set": {"project_status": to_state, "updated_at": now, "updated_by": current_user.get("id")}},
    )
    await db.project_lifecycle_history.insert_one(
        {
            "id": str(uuid.uuid4()),
            "project_id": project_id,
            "from_state": frm,
            "to_state": to_state,
            "reason": (payload.reason or "").strip() or None,
            "changed_by": current_user.get("id"),
            "changed_at": now,
        }
    )
    return {"message": "Transition applied", "from_state": frm, "to_state": to_state}


# ---- Project Section submodules: Demand, Risks, Issues, Docs, Notes, Approvals, Closure, AI ----

MANDATORY_FLAG_VALUES = {"mandatory", "optional"}


class ProjectDemandRow(BaseModel):
    role_name: str
    skill_name: str
    skill_category: Optional[str] = None
    competency_level: Optional[str] = None
    minimum_experience: Optional[float] = None
    preferred_experience: Optional[float] = None
    certification_required: Optional[str] = None
    mandatory_or_optional: Optional[str] = "mandatory"
    demand_count: int = 0
    fulfilled_count: int = 0
    billable_or_non_billable: Optional[str] = None
    planned_start_date: Optional[str] = None
    planned_end_date: Optional[str] = None
    allocation_percentage: Optional[int] = None
    demand_priority: Optional[str] = None
    hiring_required_flag: Optional[bool] = None
    remarks: Optional[str] = None


class ProjectDemandBulkUpsertRequest(BaseModel):
    rows: List[ProjectDemandRow]
    mode: Literal["skip", "upsert"] = "upsert"


class ProjectRiskCreate(BaseModel):
    title: str
    category: Optional[str] = None
    description: Optional[str] = None
    probability: Optional[int] = None  # 1..5
    impact: Optional[int] = None  # 1..5
    severity: Optional[str] = None
    mitigation_plan: Optional[str] = None
    owner_employee_id: Optional[str] = None
    target_date: Optional[str] = None
    status: Optional[str] = "open"


class ProjectIssueCreate(BaseModel):
    title: str
    category: Optional[str] = None
    description: Optional[str] = None
    severity: Optional[str] = None
    owner_employee_id: Optional[str] = None
    raised_date: Optional[str] = None
    due_date: Optional[str] = None
    resolution: Optional[str] = None
    status: Optional[str] = "open"
    escalated_flag: Optional[bool] = False


class ProjectDocumentCreate(BaseModel):
    doc_name: str
    category: str
    url: Optional[str] = None
    version: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    remarks: Optional[str] = None


class ProjectNoteCreate(BaseModel):
    type: str = "note"  # meeting/mom/comment/announcement/reminder
    title: Optional[str] = None
    body: str
    pinned: bool = False


class ProjectClosureUpsert(BaseModel):
    closure_date: Optional[str] = None
    closure_reason: Optional[str] = None
    client_signoff_status: Optional[str] = None
    financial_closure_status: Optional[str] = None
    document_closure_status: Optional[str] = None
    resource_release_status: Optional[str] = None
    success_rating: Optional[int] = None  # 1..5
    lessons_learned: Optional[str] = None
    closure_summary: Optional[str] = None
    checklist: Dict[str, bool] = Field(default_factory=dict)


@api_router.get("/project-section/projects/{project_id}/demands")
async def project_section_list_demands(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    proj = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0, "id": 1})
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    rows = await db.project_demands.find({"project_id": project_id}, {"_id": 0}).sort("updated_at", -1).to_list(5000)
    for r in rows:
        r["open_count"] = max(0, int(r.get("demand_count") or 0) - int(r.get("fulfilled_count") or 0))
    return rows


@api_router.post("/project-section/projects/{project_id}/demands/bulk")
async def project_section_bulk_upsert_demands(
    project_id: str,
    payload: ProjectDemandBulkUpsertRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    proj = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0, "id": 1})
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    now = datetime.now(timezone.utc).isoformat()
    mode = payload.mode or "upsert"
    created = updated = skipped = 0
    for row in payload.rows or []:
        role_name = str(row.role_name or "").strip()
        skill_name = str(row.skill_name or "").strip()
        if not role_name or not skill_name:
            continue
        demand = max(0, int(row.demand_count or 0))
        fulfilled = max(0, int(row.fulfilled_count or 0))
        if fulfilled > demand:
            raise HTTPException(status_code=400, detail="fulfilled_count cannot exceed demand_count")
        mand = _norm_lc(row.mandatory_or_optional or "mandatory")
        if mand not in MANDATORY_FLAG_VALUES:
            mand = "mandatory"
        key = {"project_id": project_id, "role_name_lc": role_name.lower(), "skill_name_lc": skill_name.lower()}
        existing = await db.project_demands.find_one(key, {"_id": 0, "id": 1})
        doc = {
            **key,
            "role_name": role_name,
            "skill_name": skill_name,
            "skill_category": row.skill_category,
            "competency_level": row.competency_level,
            "minimum_experience": row.minimum_experience,
            "preferred_experience": row.preferred_experience,
            "certification_required": row.certification_required,
            "mandatory_or_optional": mand,
            "demand_count": demand,
            "fulfilled_count": fulfilled,
            "billable_or_non_billable": row.billable_or_non_billable,
            "planned_start_date": row.planned_start_date,
            "planned_end_date": row.planned_end_date,
            "allocation_percentage": row.allocation_percentage,
            "demand_priority": row.demand_priority,
            "hiring_required_flag": row.hiring_required_flag,
            "remarks": row.remarks,
            "updated_at": now,
            "updated_by": current_user.get("id"),
        }
        if existing:
            if mode == "skip":
                skipped += 1
                continue
            await db.project_demands.update_one(key, {"$set": doc})
            updated += 1
        else:
            await db.project_demands.insert_one(
                {
                    "id": str(uuid.uuid4()),
                    "created_at": now,
                    "created_by": current_user.get("id"),
                    **doc,
                }
            )
            created += 1
    return {"created": created, "updated": updated, "skipped": skipped}


@api_router.get("/project-section/projects/{project_id}/risks")
async def project_section_list_risks(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_risks.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return rows


@api_router.post("/project-section/projects/{project_id}/risks")
async def project_section_create_risk(project_id: str, payload: ProjectRiskCreate, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    if not str(payload.title or "").strip():
        raise HTTPException(status_code=400, detail="title is required")
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["title"] = str(doc["title"]).strip()
    doc["project_id"] = project_id
    doc["risk_id"] = f"RSK-{uuid.uuid4().hex[:8].upper()}"
    doc["created_at"] = now
    doc["updated_at"] = now
    doc["created_by"] = current_user.get("id")
    await db.project_risks.insert_one({"id": str(uuid.uuid4()), **doc})
    return doc


@api_router.get("/project-section/projects/{project_id}/issues")
async def project_section_list_issues(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_issues.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return rows


@api_router.post("/project-section/projects/{project_id}/issues")
async def project_section_create_issue(project_id: str, payload: ProjectIssueCreate, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    if not str(payload.title or "").strip():
        raise HTTPException(status_code=400, detail="title is required")
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["title"] = str(doc["title"]).strip()
    doc["project_id"] = project_id
    doc["issue_id"] = f"ISS-{uuid.uuid4().hex[:8].upper()}"
    doc["created_at"] = now
    doc["updated_at"] = now
    doc["created_by"] = current_user.get("id")
    await db.project_issues.insert_one({"id": str(uuid.uuid4()), **doc})
    return doc


@api_router.get("/project-section/projects/{project_id}/documents")
async def project_section_list_documents(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_documents.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return rows


@api_router.post("/project-section/projects/{project_id}/documents")
async def project_section_add_document(project_id: str, payload: ProjectDocumentCreate, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    if not str(payload.doc_name or "").strip():
        raise HTTPException(status_code=400, detail="doc_name is required")
    if not str(payload.category or "").strip():
        raise HTTPException(status_code=400, detail="category is required")
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["doc_name"] = str(doc["doc_name"]).strip()
    doc["category"] = str(doc["category"]).strip()
    doc["project_id"] = project_id
    doc["created_at"] = now
    doc["updated_at"] = now
    doc["created_by"] = current_user.get("id")
    await db.project_documents.insert_one({"id": str(uuid.uuid4()), **doc})
    return doc


@api_router.get("/project-section/projects/{project_id}/notes")
async def project_section_list_notes(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_notes.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return rows


@api_router.post("/project-section/projects/{project_id}/notes")
async def project_section_add_note(project_id: str, payload: ProjectNoteCreate, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    if not str(payload.body or "").strip():
        raise HTTPException(status_code=400, detail="body is required")
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["project_id"] = project_id
    doc["created_at"] = now
    doc["created_by"] = current_user.get("id")
    await db.project_notes.insert_one({"id": str(uuid.uuid4()), **doc})
    return doc


@api_router.get("/project-section/projects/{project_id}/allocations/summary")
async def project_section_allocation_summary(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    allocs = await db.allocations.find({"project_id": project_id}, {"_id": 0}).to_list(5000)
    total = len(allocs)
    pending = len([a for a in allocs if (a.get("approval_status") or "").upper() == "PENDING"])
    active = len([a for a in allocs if (a.get("status") or "").upper() == "ACTIVE"])
    return {"project_id": project_id, "total_allocations": total, "active_allocations": active, "pending_allocations": pending, "items": allocs}


@api_router.get("/project-section/approvals")
async def project_section_list_approvals(
    status: Optional[str] = None,
    project_id: Optional[str] = None,
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
):
    _require_project_view(current_user)
    q: Dict[str, Any] = {}
    if status:
        q["status"] = _norm_lc(status)
    if project_id:
        q["project_id"] = project_id
    rows = await db.project_approvals.find(q, {"_id": 0}).sort("requested_at", -1).limit(min(max(1, limit), 200)).to_list(min(max(1, limit), 200))
    return rows


@api_router.post("/project-section/approvals/{approval_id}/action")
async def project_section_approval_action(
    approval_id: str,
    payload: ProjectApprovalActionRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_project_approve(current_user)
    doc = await db.project_approvals.find_one({"id": approval_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Approval not found")
    if doc.get("status") not in ("pending", "submitted"):
        raise HTTPException(status_code=400, detail="Approval is not actionable")
    now = datetime.now(timezone.utc).isoformat()
    action = payload.action
    if action == "approve":
        await db.project_approvals.update_one(
            {"id": approval_id},
            {"$set": {"status": "approved", "decided_by": current_user.get("id"), "decided_at": now, "decision_reason": None}},
        )
        # Apply lifecycle transition approval if present
        if doc.get("type") == "lifecycle_transition":
            pl = doc.get("payload") or {}
            to_state = _norm_lc(pl.get("to_state"))
            frm = _norm_lc(pl.get("from_state"))
            await db.project_masters.update_one(
                {"id": doc.get("project_id")},
                {"$set": {"project_status": to_state, "updated_at": now, "updated_by": current_user.get("id")}},
            )
            await db.project_lifecycle_history.insert_one(
                {
                    "id": str(uuid.uuid4()),
                    "project_id": doc.get("project_id"),
                    "from_state": frm,
                    "to_state": to_state,
                    "reason": "Approved transition",
                    "changed_by": current_user.get("id"),
                    "changed_at": now,
                }
            )
    else:
        await db.project_approvals.update_one(
            {"id": approval_id},
            {"$set": {"status": "rejected", "decided_by": current_user.get("id"), "decided_at": now, "decision_reason": (payload.reason or "").strip() or "Rejected"}},
        )
    updated = await db.project_approvals.find_one({"id": approval_id}, {"_id": 0})
    return updated


@api_router.get("/project-section/projects/{project_id}/closure")
async def project_section_get_closure(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    doc = await db.project_closure.find_one({"project_id": project_id}, {"_id": 0})
    return doc or {"project_id": project_id, "checklist": {}}


@api_router.put("/project-section/projects/{project_id}/closure")
async def project_section_upsert_closure(
    project_id: str,
    payload: ProjectClosureUpsert,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    proj = await db.project_masters.find_one({"id": project_id, "is_deleted": {"$ne": True}}, {"_id": 0, "project_status": 1})
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")
    # If closing, enforce checklist
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["project_id"] = project_id
    doc["updated_at"] = now
    doc["updated_by"] = current_user.get("id")
    existing = await db.project_closure.find_one({"project_id": project_id}, {"_id": 0})
    if existing:
        await db.project_closure.update_one({"project_id": project_id}, {"$set": doc})
    else:
        await db.project_closure.insert_one({"id": str(uuid.uuid4()), "created_at": now, "created_by": current_user.get("id"), **doc})
    return await db.project_closure.find_one({"project_id": project_id}, {"_id": 0})


@api_router.get("/project-section/projects/{project_id}/ai-recommendations")
async def project_section_ai_recommendations(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    # Mock AI layer (future-ready contract)
    demands = await db.project_demands.find({"project_id": project_id}, {"_id": 0}).to_list(200)
    open_demands = [
        {
            "role_name": d.get("role_name"),
            "skill_name": d.get("skill_name"),
            "open_count": max(0, int(d.get("demand_count") or 0) - int(d.get("fulfilled_count") or 0)),
        }
        for d in demands
        if int(d.get("demand_count") or 0) > int(d.get("fulfilled_count") or 0)
    ][:10]
    return {
        "project_id": project_id,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "recommended_resources": [],
        "predicted_delay_risk": {"risk_level": "medium", "reasons": ["Limited data (mock model)"]},
        "skill_gap_insights": open_demands,
        "budget_overrun_warning": {"risk_level": "low", "reasons": []},
        "upcoming_staffing_gaps": open_demands,
    }


@api_router.get("/project-section/analytics")
async def project_section_analytics(current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    base = {"is_deleted": {"$ne": True}, "is_archived": {"$ne": True}}
    by_status = await db.project_masters.aggregate([{"$match": base}, {"$group": {"_id": "$project_status", "count": {"$sum": 1}}}, {"$sort": {"count": -1}}]).to_list(50)
    by_priority = await db.project_masters.aggregate([{"$match": base}, {"$group": {"_id": "$project_priority", "count": {"$sum": 1}}}, {"$sort": {"count": -1}}]).to_list(50)
    by_bu = await db.project_masters.aggregate([{"$match": base}, {"$group": {"_id": "$business_unit", "count": {"$sum": 1}}}, {"$sort": {"count": -1}}]).to_list(50)
    return {
        "by_status": [{"key": r["_id"] or "unknown", "count": r["count"]} for r in by_status],
        "by_priority": [{"key": r["_id"] or "unknown", "count": r["count"]} for r in by_priority],
        "by_business_unit": [{"key": r["_id"] or "unknown", "count": r["count"]} for r in by_bu],
    }


# ---- Planning & Structuring (WBS) ----


def _wbs_dependency_adjacency(items: List[Dict[str, Any]]) -> Dict[str, List[str]]:
    """Directed edges: predecessor -> successor (dep must finish before dependent)."""
    ids = {str(it["id"]) for it in items}
    adj: Dict[str, List[str]] = {i: [] for i in ids}
    for it in items:
        nid = str(it["id"])
        for dep in it.get("depends_on_ids") or []:
            d = str(dep)
            if d in ids and d != nid:
                adj[d].append(nid)
    return adj


def _wbs_graph_has_cycle(adj: Dict[str, List[str]]) -> bool:
    UNSEEN, VISITING, DONE = 0, 1, 2
    state = {n: UNSEEN for n in adj}

    def dfs(u: str) -> bool:
        state[u] = VISITING
        for v in adj.get(u, []):
            if state.get(v) == VISITING:
                return True
            if state.get(v) == UNSEEN and dfs(v):
                return True
        state[u] = DONE
        return False

    for n in adj:
        if state[n] == UNSEEN and dfs(n):
            return True
    return False


async def _project_wbs_cycles_ok(db, project_id: str, pending_item: Optional[Dict[str, Any]] = None):
    rows = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0, "id": 1, "depends_on_ids": 1}).to_list(5000)
    items: List[Dict[str, Any]] = [{"id": r["id"], "depends_on_ids": r.get("depends_on_ids") or []} for r in rows]
    if pending_item:
        pid = str(pending_item["id"])
        replaced = False
        for i, it in enumerate(items):
            if str(it["id"]) == pid:
                items[i] = {"id": pid, "depends_on_ids": pending_item.get("depends_on_ids") or []}
                replaced = True
                break
        if not replaced:
            items.append({"id": pid, "depends_on_ids": pending_item.get("depends_on_ids") or []})
    adj = _wbs_dependency_adjacency(items)
    if _wbs_graph_has_cycle(adj):
        raise HTTPException(status_code=400, detail="WBS dependencies contain a cycle")


async def _wbs_parent_is_valid(db, project_id: str, node_id: str, parent_id: Optional[str]):
    if not parent_id:
        return
    if parent_id == node_id:
        raise HTTPException(status_code=400, detail="parent_id cannot be self")
    desc: set = set()
    q = [node_id]
    while q:
        cur = q.pop()
        kids = await db.project_wbs_items.find({"project_id": project_id, "parent_id": cur}, {"_id": 0, "id": 1}).to_list(500)
        for k in kids:
            cid = k["id"]
            if cid not in desc:
                desc.add(cid)
                q.append(cid)
    if parent_id in desc:
        raise HTTPException(status_code=400, detail="parent_id cannot be a descendant of this item")


class ProjectWbsItemCreate(BaseModel):
    type: Literal["phase", "milestone", "deliverable", "task", "subtask"] = "task"
    parent_id: Optional[str] = None
    name: str
    owner_employee_id: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    status: Optional[str] = "pending"
    percent_complete: Optional[int] = 0
    estimated_effort_hours: Optional[float] = None
    depends_on_ids: List[str] = Field(default_factory=list)
    order: Optional[int] = None


class ProjectWbsItemUpdate(BaseModel):
    type: Optional[Literal["phase", "milestone", "deliverable", "task", "subtask"]] = None
    parent_id: Optional[str] = None
    name: Optional[str] = None
    owner_employee_id: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    status: Optional[str] = None
    percent_complete: Optional[int] = None
    estimated_effort_hours: Optional[float] = None
    depends_on_ids: Optional[List[str]] = None
    order: Optional[int] = None


@api_router.get("/project-section/projects/{project_id}/planning/wbs")
async def project_section_wbs_list(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0}).sort("order", 1).to_list(5000)
    return rows


@api_router.post("/project-section/projects/{project_id}/planning/wbs")
async def project_section_wbs_create(
    project_id: str,
    payload: ProjectWbsItemCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    name = str(payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name is required")
    now = datetime.now(timezone.utc).isoformat()
    # Dependency validation: no self-dependency, and dependencies must exist in same project
    deps = [d for d in (payload.depends_on_ids or []) if d]
    if len(set(deps)) != len(deps):
        deps = list(dict.fromkeys(deps))
    for dep_id in deps:
        if dep_id == "NEW":
            continue
        ex = await db.project_wbs_items.find_one({"project_id": project_id, "id": dep_id}, {"_id": 0, "id": 1})
        if not ex:
            raise HTTPException(status_code=400, detail=f"Invalid depends_on_id: {dep_id}")
    # Choose an order if missing
    order = payload.order
    if order is None:
        last = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0, "order": 1}).sort("order", -1).limit(1).to_list(1)
        order = int((last[0].get("order") or 0) + 1) if last else 1

    doc = {
        "id": str(uuid.uuid4()),
        "project_id": project_id,
        "type": payload.type,
        "parent_id": payload.parent_id,
        "name": name,
        "owner_employee_id": payload.owner_employee_id,
        "start_date": payload.start_date,
        "end_date": payload.end_date,
        "status": payload.status or "pending",
        "percent_complete": max(0, min(int(payload.percent_complete or 0), 100)),
        "estimated_effort_hours": payload.estimated_effort_hours,
        "depends_on_ids": deps,
        "order": int(order),
        "created_at": now,
        "updated_at": now,
        "created_by": current_user.get("id"),
        "updated_by": current_user.get("id"),
    }
    await _wbs_parent_is_valid(db, project_id, doc["id"], payload.parent_id)
    await _project_wbs_cycles_ok(db, project_id, pending_item={"id": doc["id"], "depends_on_ids": deps})
    await db.project_wbs_items.insert_one(doc)
    return doc


@api_router.put("/project-section/projects/{project_id}/planning/wbs/{wbs_id}")
async def project_section_wbs_update(
    project_id: str,
    wbs_id: str,
    payload: ProjectWbsItemUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    existing = await db.project_wbs_items.find_one({"project_id": project_id, "id": wbs_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="WBS item not found")
    patch = payload.model_dump(exclude_none=True)
    if "name" in patch:
        patch["name"] = str(patch["name"]).strip()
        if not patch["name"]:
            raise HTTPException(status_code=400, detail="name cannot be empty")
    if "percent_complete" in patch:
        patch["percent_complete"] = max(0, min(int(patch["percent_complete"] or 0), 100))
    if "depends_on_ids" in patch:
        deps = [d for d in (patch.get("depends_on_ids") or []) if d]
        if wbs_id in deps:
            raise HTTPException(status_code=400, detail="Item cannot depend on itself")
        for dep_id in deps:
            ex = await db.project_wbs_items.find_one({"project_id": project_id, "id": dep_id}, {"_id": 0, "id": 1})
            if not ex:
                raise HTTPException(status_code=400, detail=f"Invalid depends_on_id: {dep_id}")
        patch["depends_on_ids"] = deps
        await _project_wbs_cycles_ok(db, project_id, pending_item={"id": wbs_id, "depends_on_ids": deps})
    # Parent validation: parent must exist and not be self
    if "parent_id" in patch and patch["parent_id"]:
        if patch["parent_id"] == wbs_id:
            raise HTTPException(status_code=400, detail="parent_id cannot be self")
        ex = await db.project_wbs_items.find_one({"project_id": project_id, "id": patch["parent_id"]}, {"_id": 0, "id": 1})
        if not ex:
            raise HTTPException(status_code=400, detail="Invalid parent_id")
        await _wbs_parent_is_valid(db, project_id, wbs_id, patch["parent_id"])
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    patch["updated_by"] = current_user.get("id")
    await db.project_wbs_items.update_one({"project_id": project_id, "id": wbs_id}, {"$set": patch})
    return await db.project_wbs_items.find_one({"project_id": project_id, "id": wbs_id}, {"_id": 0})


@api_router.delete("/project-section/projects/{project_id}/planning/wbs/{wbs_id}")
async def project_section_wbs_delete(project_id: str, wbs_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    # Prevent delete if has children
    child = await db.project_wbs_items.find_one({"project_id": project_id, "parent_id": wbs_id}, {"_id": 0, "id": 1})
    if child:
        raise HTTPException(status_code=400, detail="Cannot delete item with children")
    res = await db.project_wbs_items.delete_one({"project_id": project_id, "id": wbs_id})
    if res.deleted_count == 0:
        raise HTTPException(status_code=404, detail="WBS item not found")
    return {"message": "Deleted"}


class ProjectWbsReorderRequest(BaseModel):
    ordered_ids: List[str]


@api_router.post("/project-section/projects/{project_id}/planning/wbs/reorder")
async def project_section_wbs_reorder(
    project_id: str,
    payload: ProjectWbsReorderRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_project_edit(current_user)
    ordered = [str(i).strip() for i in (payload.ordered_ids or []) if str(i).strip()]
    existing = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0, "id": 1}).to_list(5000)
    existing_ids = {str(r["id"]) for r in existing}
    if len(ordered) != len(existing_ids) or set(ordered) != existing_ids:
        raise HTTPException(status_code=400, detail="ordered_ids must list every WBS item id exactly once")
    now = datetime.now(timezone.utc).isoformat()
    uid = current_user.get("id")
    for idx, wid in enumerate(ordered, start=1):
        await db.project_wbs_items.update_one(
            {"project_id": project_id, "id": wid},
            {"$set": {"order": idx, "updated_at": now, "updated_by": uid}},
        )
    return {"message": "ok", "count": len(ordered)}


@api_router.get("/project-section/projects/{project_id}/planning/wbs/validate-graph")
async def project_section_wbs_validate_graph(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0, "id": 1, "depends_on_ids": 1}).to_list(5000)
    items = [{"id": r["id"], "depends_on_ids": r.get("depends_on_ids") or []} for r in rows]
    adj = _wbs_dependency_adjacency(items)
    cyclic = _wbs_graph_has_cycle(adj)
    return {"project_id": project_id, "acyclic": not cyclic, "node_count": len(items)}


# ---- Lifecycle history ----

@api_router.get("/project-section/projects/{project_id}/lifecycle/history")
async def project_section_lifecycle_history(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_lifecycle_history.find({"project_id": project_id}, {"_id": 0}).sort("changed_at", -1).to_list(2000)
    return rows


# ---- Financial Management ----

class ProjectFinancialUpsert(BaseModel):
    planned_budget: Optional[float] = None
    approved_budget: Optional[float] = None
    revised_budget: Optional[float] = None
    planned_cost: Optional[float] = None
    actual_cost: Optional[float] = None
    committed_cost: Optional[float] = None
    planned_revenue: Optional[float] = None
    actual_revenue: Optional[float] = None
    billing_rate: Optional[float] = None
    cost_rate: Optional[float] = None
    capex_or_opex: Optional[str] = None
    invoice_status: Optional[str] = None
    currency: Optional[str] = None
    period_month: Optional[str] = None  # YYYY-MM for snapshotting


@api_router.get("/project-section/projects/{project_id}/finance")
async def project_section_finance_get(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    doc = await db.project_financials.find_one({"project_id": project_id}, {"_id": 0})
    return doc or {"project_id": project_id}


@api_router.put("/project-section/projects/{project_id}/finance")
async def project_section_finance_upsert(project_id: str, payload: ProjectFinancialUpsert, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    now = datetime.now(timezone.utc).isoformat()
    patch = payload.model_dump(exclude_none=True)
    for k in [
        "planned_budget",
        "approved_budget",
        "revised_budget",
        "planned_cost",
        "actual_cost",
        "committed_cost",
        "planned_revenue",
        "actual_revenue",
        "billing_rate",
        "cost_rate",
    ]:
        if k in patch and patch[k] is not None and float(patch[k]) < 0:
            raise HTTPException(status_code=400, detail=f"{k} cannot be negative")

    # Compute derived metrics
    pb = float(patch.get("planned_budget") or 0)
    ab = float(patch.get("approved_budget") or 0)
    rb = float(patch.get("revised_budget") or 0)
    ac = float(patch.get("actual_cost") or 0)
    ar = float(patch.get("actual_revenue") or 0)
    cost = ac
    rev = ar
    gross_margin = (rev - cost) if rev or cost else None
    margin_pct = (gross_margin / rev * 100.0) if gross_margin is not None and rev > 0 else None

    patch.update(
        {
            "project_id": project_id,
            "gross_margin": gross_margin,
            "margin_percentage": margin_pct,
            "budget_variance": (rb or ab or pb) - cost if (rb or ab or pb) else None,
            "cost_variance": (float(patch.get("planned_cost") or 0) - cost) if patch.get("planned_cost") is not None else None,
            "updated_at": now,
            "updated_by": current_user.get("id"),
        }
    )
    existing = await db.project_financials.find_one({"project_id": project_id}, {"_id": 0})
    if existing:
        await db.project_financials.update_one({"project_id": project_id}, {"$set": patch})
    else:
        await db.project_financials.insert_one({"id": str(uuid.uuid4()), "created_at": now, "created_by": current_user.get("id"), **patch})

    # Optional snapshot
    period = patch.get("period_month")
    if period:
        await db.project_financial_snapshots.insert_one(
            {
                "id": str(uuid.uuid4()),
                "project_id": project_id,
                "period_month": str(period),
                "snapshot_at": now,
                "payload": {k: v for k, v in patch.items() if k not in ("updated_at", "updated_by", "created_at", "created_by")},
            }
        )
    return await db.project_financials.find_one({"project_id": project_id}, {"_id": 0})


@api_router.get("/project-section/projects/{project_id}/finance/snapshots")
async def project_section_finance_snapshots(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_financial_snapshots.find({"project_id": project_id}, {"_id": 0}).sort("snapshot_at", -1).to_list(200)
    return rows


# ---- Execution & Tracking: status reports ----

class ProjectStatusReportCreate(BaseModel):
    period: Optional[str] = None  # e.g. 2026-W16
    summary: str
    health: Optional[str] = None
    blockers: Optional[str] = None
    next_steps: Optional[str] = None


@api_router.get("/project-section/projects/{project_id}/execution/status-reports")
async def project_section_status_reports_list(project_id: str, current_user: dict = Depends(get_current_user)):
    _require_project_view(current_user)
    rows = await db.project_status_reports.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(200)
    return rows


@api_router.post("/project-section/projects/{project_id}/execution/status-reports")
async def project_section_status_reports_create(project_id: str, payload: ProjectStatusReportCreate, current_user: dict = Depends(get_current_user)):
    _require_project_edit(current_user)
    if not str(payload.summary or "").strip():
        raise HTTPException(status_code=400, detail="summary is required")
    now = datetime.now(timezone.utc).isoformat()
    doc = payload.model_dump()
    doc["project_id"] = project_id
    doc["created_at"] = now
    doc["created_by"] = current_user.get("id")
    await db.project_status_reports.insert_one({"id": str(uuid.uuid4()), **doc})
    return doc


def _date_key(d: Optional[str]) -> Optional[str]:
    """Normalize date-like strings for lexical comparisons (YYYY-MM-DD or ISO)."""
    if not d or not isinstance(d, str):
        return None
    s = d.strip()
    if not s:
        return None
    return s[:10]


@api_router.get("/project-section/projects/{project_id}/execution/alerts")
async def project_section_execution_alerts(project_id: str, current_user: dict = Depends(get_current_user)):
    """Overdue WBS items, breached milestones, and dependency-blocked work."""
    _require_project_view(current_user)
    from datetime import date

    today = date.today().isoformat()
    rows = await db.project_wbs_items.find({"project_id": project_id}, {"_id": 0}).to_list(5000)
    by_id = {str(r["id"]): r for r in rows}

    def is_open(r: Dict[str, Any]) -> bool:
        st = _norm_lc(str(r.get("status") or ""))
        if st in ("completed", "done", "closed", "cancelled"):
            return False
        try:
            pc = int(r.get("percent_complete") or 0)
        except (TypeError, ValueError):
            pc = 0
        if pc >= 100:
            return False
        return True

    overdue_items: List[Dict[str, Any]] = []
    milestone_breaches: List[Dict[str, Any]] = []
    dependency_blocked: List[Dict[str, Any]] = []

    for r in rows:
        end = _date_key(r.get("end_date"))
        if not end or end >= today:
            continue
        if not is_open(r):
            continue
        entry = {
            "id": r.get("id"),
            "name": r.get("name"),
            "type": r.get("type"),
            "end_date": r.get("end_date"),
            "status": r.get("status"),
            "percent_complete": r.get("percent_complete"),
            "alert": "overdue",
        }
        overdue_items.append(entry)
        if _norm_lc(str(r.get("type") or "")) == "milestone":
            milestone_breaches.append({**entry, "alert": "milestone_breach"})

    for r in rows:
        if not is_open(r):
            continue
        deps = r.get("depends_on_ids") or []
        for d in deps:
            pred = by_id.get(str(d))
            if pred and is_open(pred):
                dependency_blocked.append(
                    {
                        "id": r.get("id"),
                        "name": r.get("name"),
                        "type": r.get("type"),
                        "status": r.get("status"),
                        "alert": "dependency_blocked",
                        "blocked_by": {
                            "id": pred.get("id"),
                            "name": pred.get("name"),
                            "status": pred.get("status"),
                            "end_date": pred.get("end_date"),
                        },
                    }
                )
                break

    return {
        "project_id": project_id,
        "as_of": today,
        "overdue_items": overdue_items,
        "milestone_breaches": milestone_breaches,
        "dependency_blocked": dependency_blocked,
        "counts": {
            "overdue": len(overdue_items),
            "milestone_breaches": len(milestone_breaches),
            "dependency_blocked": len(dependency_blocked),
        },
    }


def _ranges_overlap(a_start: Optional[str], a_end: Optional[str], b_start: Optional[str], b_end: Optional[str]) -> bool:
    # Treat None end as open-ended; None start as "from -inf"
    as_ = _date_key(a_start)
    ae = _date_key(a_end)
    bs = _date_key(b_start)
    be = _date_key(b_end)
    left = as_ or "0000-01-01"
    right = ae or "9999-12-31"
    left2 = bs or "0000-01-01"
    right2 = be or "9999-12-31"
    return not (right < left2 or right2 < left)


async def _assert_no_overallocation(
    employee_id: str,
    start_date: Optional[str],
    end_date: Optional[str],
    new_pct: int,
    exclude_allocation_id: Optional[str] = None,
):
    """Reject if overlapping allocations exceed 100%."""
    emp_id = (employee_id or "").strip()
    if not emp_id:
        raise HTTPException(status_code=400, detail="employee_id is required")
    new_pct = int(new_pct or 0)
    if new_pct < 0 or new_pct > 100:
        raise HTTPException(status_code=400, detail="allocation_percentage must be between 0 and 100")

    query: Dict[str, Any] = {"employee_id": emp_id, "status": {"$in": ["ACTIVE", "PENDING"]}}
    if exclude_allocation_id:
        query["id"] = {"$ne": exclude_allocation_id}
    existing = await db.allocations.find(query, {"_id": 0}).to_list(5000)
    overlap_total = new_pct
    overlaps: List[Dict[str, Any]] = []
    for a in existing:
        if _ranges_overlap(start_date, end_date, a.get("start_date"), a.get("end_date")):
            overlap_total += int(a.get("allocation_percentage") or 0)
            overlaps.append(
                {
                    "allocation_id": a.get("id"),
                    "project_id": a.get("project_id"),
                    "allocation_percentage": int(a.get("allocation_percentage") or 0),
                    "start_date": a.get("start_date"),
                    "end_date": a.get("end_date"),
                    "status": a.get("status"),
                    "approval_status": a.get("approval_status"),
                }
            )
    if overlap_total > 100:
        raise HTTPException(
            status_code=409,
            detail={
                "message": "Overallocation detected",
                "employee_id": emp_id,
                "overlapping_total_pct": overlap_total,
                "overlaps": overlaps,
            },
        )


@api_router.post("/allocations", response_model=AllocationResponse)
async def create_allocation(
    payload: AllocationCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    project_id = (payload.project_id or "").strip()
    employee_id = (payload.employee_id or "").strip()
    if not project_id or not employee_id:
        raise HTTPException(status_code=400, detail="project_id and employee_id are required")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0, "id": 1})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    emp = await db.employees.find_one({"id": employee_id}, {"_id": 0, "id": 1})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")

    await _assert_no_overallocation(employee_id, payload.start_date, payload.end_date, payload.allocation_percentage)

    now = datetime.now(timezone.utc).isoformat()
    allocation_id = str(uuid.uuid4())
    doc = {
        "id": allocation_id,
        "project_id": project_id,
        "employee_id": employee_id,
        "role": (payload.role or "").strip() or None,
        "allocation_percentage": int(payload.allocation_percentage or 0),
        "start_date": payload.start_date,
        "end_date": payload.end_date,
        "billable": bool(payload.billable),
        "allocation_type": (str(payload.allocation_type).upper() if payload.allocation_type else None),
        "status": (payload.status or "PENDING").upper(),
        "approval_status": "PENDING",
        "cost_rate": payload.cost_rate,
        "billing_rate": payload.billing_rate,
        "created_at": now,
        "updated_at": None,
        "approved_by": None,
        "approved_at": None,
        "rejection_reason": None,
    }
    await db.allocations.insert_one(doc)
    return AllocationResponse(**doc)


@api_router.get("/allocations/project/{project_id}", response_model=List[AllocationResponse])
async def list_allocations_for_project(
    project_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    rows = await db.allocations.find({"project_id": project_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return [AllocationResponse(**r) for r in rows]


@api_router.get("/allocations/resource/{employee_id}", response_model=List[AllocationResponse])
async def list_allocations_for_employee(
    employee_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    rows = await db.allocations.find({"employee_id": employee_id}, {"_id": 0}).sort("created_at", -1).to_list(5000)
    return [AllocationResponse(**r) for r in rows]


@api_router.put("/allocations/{allocation_id}", response_model=AllocationResponse)
async def update_allocation(
    allocation_id: str,
    payload: AllocationUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    existing = await db.allocations.find_one({"id": allocation_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Allocation not found")

    update_doc = payload.model_dump(exclude_none=True)
    if "allocation_percentage" in update_doc:
        await _assert_no_overallocation(
            existing.get("employee_id") or "",
            update_doc.get("start_date", existing.get("start_date")),
            update_doc.get("end_date", existing.get("end_date")),
            int(update_doc.get("allocation_percentage") or 0),
            exclude_allocation_id=allocation_id,
        )
    if "status" in update_doc and isinstance(update_doc["status"], str):
        update_doc["status"] = update_doc["status"].upper()
    if "approval_status" in update_doc and isinstance(update_doc["approval_status"], str):
        update_doc["approval_status"] = update_doc["approval_status"].upper()
    if "allocation_type" in update_doc and update_doc.get("allocation_type"):
        update_doc["allocation_type"] = str(update_doc["allocation_type"]).upper()
    update_doc["updated_at"] = datetime.now(timezone.utc).isoformat()

    await db.allocations.update_one({"id": allocation_id}, {"$set": update_doc})
    updated = await db.allocations.find_one({"id": allocation_id}, {"_id": 0})
    return AllocationResponse(**updated)


@api_router.post("/allocations/{allocation_id}/approve", response_model=AllocationResponse)
async def approve_or_reject_allocation(
    allocation_id: str,
    payload: AllocationApproveRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    # Keep simple: same permission gate for now; can be tightened to Resource Manager later.
    existing = await db.allocations.find_one({"id": allocation_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Allocation not found")

    now = datetime.now(timezone.utc).isoformat()
    action = payload.action
    if action == "approve":
        patch = {
            "approval_status": "APPROVED",
            "status": "ACTIVE" if (existing.get("status") or "").upper() != "CLOSED" else "CLOSED",
            "approved_by": current_user.get("id"),
            "approved_at": now,
            "rejection_reason": None,
            "updated_at": now,
        }
    else:
        patch = {
            "approval_status": "REJECTED",
            "status": "PENDING",
            "approved_by": None,
            "approved_at": None,
            "rejection_reason": (payload.reason or "").strip() or "Rejected",
            "updated_at": now,
        }
    await db.allocations.update_one({"id": allocation_id}, {"$set": patch})
    updated = await db.allocations.find_one({"id": allocation_id}, {"_id": 0})
    return AllocationResponse(**updated)


@api_router.post("/projects/{project_id}/skill-demands/bulk", response_model=Dict[str, Any])
async def upsert_project_skill_demands(
    project_id: str,
    payload: ProjectSkillDemandBulkRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    mode = payload.mode or "upsert"
    if mode not in {"skip", "upsert"}:
        mode = "upsert"

    now = datetime.now(timezone.utc).isoformat()
    rows = payload.rows or []
    created = 0
    updated = 0
    skipped = 0

    for row in rows:
        sk_name = (row.skill_name or "").strip()
        if not sk_name:
            continue
        sk_lc = sk_name.lower()
        demand = max(0, int(row.demand_count or 0))
        pri = (row.priority or None)
        pri = str(pri).upper() if isinstance(pri, str) and pri.strip() else None
        dmin = demand if row.demand_min is None else max(0, int(row.demand_min))
        dmax = demand if row.demand_max is None else max(0, int(row.demand_max))
        dmax = max(dmin, dmax)
        ct = row.constraint_type
        ctype = str(ct).upper() if ct else "HARD"
        if ctype not in ("HARD", "SOFT"):
            ctype = "HARD"

        existing = await db.project_skill_demands.find_one(
            {"project_id": project_id, "skill_name_lc": sk_lc},
            {"_id": 0},
        )
        if existing:
            if mode == "skip":
                skipped += 1
                continue
            await db.project_skill_demands.update_one(
                {"id": existing["id"]},
                {
                    "$set": {
                        "skill_name": sk_name,
                        "demand_count": demand,
                        "priority": pri or existing.get("priority") or "MEDIUM",
                        "demand_min": dmin,
                        "demand_max": dmax,
                        "constraint_type": ctype,
                        "updated_at": now,
                    }
                },
            )
            updated += 1
        else:
            doc = {
                "id": str(uuid.uuid4()),
                "project_id": project_id,
                "skill_name": sk_name,
                "skill_name_lc": sk_lc,
                "demand_count": demand,
                "priority": pri or "MEDIUM",
                "demand_min": dmin,
                "demand_max": dmax,
                "constraint_type": ctype,
                "created_at": now,
                "updated_at": now,
            }
            await db.project_skill_demands.insert_one(doc)
            created += 1

    return {"project_id": project_id, "created": created, "updated": updated, "skipped": skipped}


@api_router.post("/projects/{project_id}/skill-demands/bulk-import", response_model=Dict[str, Any])
async def bulk_import_project_skill_demands(
    project_id: str,
    payload: ProjectSkillDemandBulkImportRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Bulk import (CSV-driven on frontend) with:
    - dry_run preview mode
    - per-row CREATE/UPDATE/FAILED results
    - audit log written into `import_audit_logs`
    """
    _require_phase1_access(current_user, "skills_write")

    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    results: List[Dict[str, Any]] = []
    created = 0
    updated = 0
    failed = 0

    existing_rows = await db.project_skill_demands.find(
        {"project_id": project_id},
        {"_id": 0, "id": 1, "skill_name_lc": 1, "priority": 1, "demand_count": 1, "skill_name": 1},
    ).to_list(100000)
    by_key = {str(r.get("skill_name_lc") or "").strip().lower(): r for r in existing_rows if r.get("skill_name_lc")}

    now_iso = datetime.now(timezone.utc).isoformat()
    for idx, row in enumerate(payload.rows, start=1):
        skill_name = (row.skill_name or "").strip()
        key = skill_name.lower()

        demand_count = max(0, int(row.demand_count or 0))
        priority = row.priority
        priority = str(priority).upper().strip() if isinstance(priority, str) and priority.strip() else "MEDIUM"
        dmin = demand_count if row.demand_min is None else max(0, int(row.demand_min))
        dmax = demand_count if row.demand_max is None else max(0, int(row.demand_max))
        dmax = max(dmin, dmax)
        ct = row.constraint_type
        ctype = str(ct).upper() if ct else "HARD"
        if ctype not in ("HARD", "SOFT"):
            ctype = "HARD"

        if not skill_name:
            failed += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "FAILED",
                    "reason": "skill_name is required",
                }
            )
            continue

        ex = by_key.get(key)
        if ex:
            if payload.mode == "skip":
                failed += 1
                results.append(
                    {
                        "row_number": idx,
                        "skill_name": skill_name,
                        "action": "FAILED",
                        "reason": "Duplicate skill demand (skipping)",
                    }
                )
                continue

            updated += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "UPDATE",
                    "reason": "",
                }
            )
            if not payload.dry_run:
                await db.project_skill_demands.update_one(
                    {"id": ex["id"]},
                    {
                        "$set": {
                            "skill_name": skill_name,
                            "demand_count": demand_count,
                            "priority": priority,
                            "demand_min": dmin,
                            "demand_max": dmax,
                            "constraint_type": ctype,
                            "updated_at": now_iso,
                        }
                    },
                )
        else:
            created += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "CREATE",
                    "reason": "",
                }
            )
            if not payload.dry_run:
                doc = {
                    "id": str(uuid.uuid4()),
                    "project_id": project_id,
                    "skill_name": skill_name,
                    "skill_name_lc": key,
                    "demand_count": demand_count,
                    "priority": priority,
                    "demand_min": dmin,
                    "demand_max": dmax,
                    "constraint_type": ctype,
                    "created_at": now_iso,
                    "updated_at": now_iso,
                }
                await db.project_skill_demands.insert_one(doc)
                by_key[key] = doc

    summary = {"created": created, "updated": updated, "failed": failed, "total": len(payload.rows)}
    audit = await _write_import_audit(
        "project_skill_demands",
        payload.mode,
        payload.dry_run,
        summary,
        results,
        current_user["id"],
    )
    return {"summary": summary, "rows": results, "audit_id": audit["id"]}


@api_router.post("/projects/{project_id}/skill-allocations/bulk-import", response_model=Dict[str, Any])
async def bulk_import_project_skill_allocations(
    project_id: str,
    payload: ProjectSkillAllocationBulkImportRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Bulk import (CSV-driven on frontend) for per-project skill supply allocations.
    Stores into `project_skill_allocations`.
    """
    _require_phase1_access(current_user, "skills_write")

    project = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    results: List[Dict[str, Any]] = []
    created = 0
    updated = 0
    failed = 0

    existing_rows = await db.project_skill_allocations.find(
        {"project_id": project_id},
        {"_id": 0, "id": 1, "skill_name_lc": 1},
    ).to_list(100000)
    by_key = {str(r.get("skill_name_lc") or "").strip().lower(): r for r in existing_rows if r.get("skill_name_lc")}

    now_iso = datetime.now(timezone.utc).isoformat()
    for idx, row in enumerate(payload.rows, start=1):
        skill_name = (row.skill_name or "").strip()
        key = skill_name.lower()
        allocated_count = max(0, int(row.allocated_count or 0))

        if not skill_name:
            failed += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "FAILED",
                    "reason": "skill_name is required",
                }
            )
            continue

        ex = by_key.get(key)
        if ex:
            if payload.mode == "skip":
                failed += 1
                results.append(
                    {
                        "row_number": idx,
                        "skill_name": skill_name,
                        "action": "FAILED",
                        "reason": "Duplicate skill allocation (skipping)",
                    }
                )
                continue

            updated += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "UPDATE",
                    "reason": "",
                }
            )
            if not payload.dry_run:
                await db.project_skill_allocations.update_one(
                    {"id": ex["id"]},
                    {
                        "$set": {
                            "skill_name": skill_name,
                            "allocated_count": allocated_count,
                            "updated_at": now_iso,
                        }
                    },
                )
        else:
            created += 1
            results.append(
                {
                    "row_number": idx,
                    "skill_name": skill_name,
                    "action": "CREATE",
                    "reason": "",
                }
            )
            if not payload.dry_run:
                doc = {
                    "id": str(uuid.uuid4()),
                    "project_id": project_id,
                    "skill_name": skill_name,
                    "skill_name_lc": key,
                    "allocated_count": allocated_count,
                    "created_at": now_iso,
                    "updated_at": now_iso,
                }
                await db.project_skill_allocations.insert_one(doc)
                by_key[key] = doc

    summary = {"created": created, "updated": updated, "failed": failed, "total": len(payload.rows)}
    audit = await _write_import_audit(
        "project_skill_allocations",
        payload.mode,
        payload.dry_run,
        summary,
        results,
        current_user["id"],
    )
    return {"summary": summary, "rows": results, "audit_id": audit["id"]}


@api_router.get("/projects/{project_id}/skill-allocations", response_model=List[ProjectSkillAllocationResponse])
async def list_project_skill_allocations(
    project_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    rows = (
        await db.project_skill_allocations.find({"project_id": project_id}, {"_id": 0})
        .sort("updated_at", -1)
        .to_list(500)
    )
    return [
        ProjectSkillAllocationResponse(
            project_id=r["project_id"],
            skill_name=r.get("skill_name") or r.get("skill_name_lc") or "",
            allocated_count=int(r.get("allocated_count") or 0),
            updated_at=r.get("updated_at") or r.get("created_at") or datetime.now(timezone.utc).isoformat(),
        )
        for r in rows
    ]


@api_router.get("/projects/{project_id}/skill-demands", response_model=List[ProjectSkillDemandResponse])
async def list_project_skill_demands(
    project_id: str,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    rows = (
        await db.project_skill_demands.find({"project_id": project_id}, {"_id": 0})
        .sort("updated_at", -1)
        .to_list(500)
    )
    out_demands: List[ProjectSkillDemandResponse] = []
    for r in rows:
        dc = int(r.get("demand_count") or 0)
        dmin = int(r["demand_min"]) if r.get("demand_min") is not None else dc
        dmax = int(r["demand_max"]) if r.get("demand_max") is not None else dmin
        dmax = max(dmin, dmax)
        out_demands.append(
            ProjectSkillDemandResponse(
                project_id=r["project_id"],
                skill_name=r.get("skill_name") or r.get("skill_name_lc") or "",
                demand_count=dc,
                priority=str(r.get("priority") or "MEDIUM").upper(),
                updated_at=r.get("updated_at") or r.get("created_at") or datetime.now(timezone.utc).isoformat(),
                demand_min=dmin,
                demand_max=dmax,
                constraint_type=str(r.get("constraint_type") or "HARD").upper(),
            )
        )
    return out_demands


@api_router.get("/workforce/training-recommendations", response_model=TrainingRecommendationsResponse)
async def get_training_recommendations(
    q: Optional[str] = None,
    page: int = 1,
    page_size: int = 10,
    max_skills_per_employee: int = 3,
    current_user: dict = Depends(get_current_user),
):
    """
    M5-1: Personalized learning paths from workforce skill gaps + optional DB path templates.
    Ranking: workforce skill priority (HIGH first) then gap size; excludes skills the employee already has.
    """
    _require_phase1_access(current_user, "kpi_read")

    page = max(1, int(page))
    page_size = min(max(1, int(page_size)), 50)
    max_skills_per_employee = min(max(1, int(max_skills_per_employee)), 10)

    skill_rows: List[SkillInventoryResponse] = await list_skill_inventory(current_user)
    gap_dicts = [s.model_dump() for s in skill_rows if int(s.gap) > 0][:80]

    emp_query: Dict[str, Any] = {}
    if q:
        emp_query["$or"] = [
            {"employee_code": {"$regex": q, "$options": "i"}},
            {"full_name": {"$regex": q, "$options": "i"}},
        ]

    total = await db.employees.count_documents(emp_query)
    total_pages = max(1, (total + page_size - 1) // page_size)

    employees = (
        await db.employees.find(
            emp_query,
            {"_id": 0, "employee_code": 1, "full_name": 1, "skills": 1},
        )
        .sort("created_at", -1)
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )

    templates = await load_path_templates_map(db)
    payloads = build_employee_recommendation_payloads(
        employees,
        gap_dicts,
        templates,
        max_skills_per_employee=max_skills_per_employee,
    )

    recommendations: List[EmployeeTrainingRecommendation] = []
    for p in payloads:
        rec_skills: List[TrainingSkillRecommendation] = []
        for s in p.get("recommended_skills") or []:
            steps = [
                TrainingPathStep(step_title=x.get("step_title") or "", description=x.get("description") or "")
                for x in (s.get("path_steps") or [])
            ]
            rec_skills.append(
                TrainingSkillRecommendation(
                    skill_name=s.get("skill_name") or "",
                    priority=str(s.get("priority") or "MEDIUM").upper(),
                    reason=str(s.get("reason") or ""),
                    path_steps=steps,
                )
            )
        recommendations.append(
            EmployeeTrainingRecommendation(
                employee_code=p.get("employee_code") or "",
                full_name=p.get("full_name") or "",
                recommended_skills=rec_skills,
            )
        )

    generated_at = datetime.now(timezone.utc).isoformat()
    return TrainingRecommendationsResponse(
        generated_at=generated_at,
        page=page,
        page_size=page_size,
        total=total,
        total_pages=total_pages,
        max_skills_per_employee=max_skills_per_employee,
        recommendations=recommendations,
    )


# --- M5: learning paths, assignments, LMS catalog, certifications ---


@api_router.get("/workforce/training/learning-path-templates", response_model=List[Dict[str, Any]])
async def list_learning_path_templates(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    return await db[COL_LEARNING_PATH_TEMPLATES].find({}, {"_id": 0}).to_list(200)


@api_router.put("/workforce/training/learning-path-templates", response_model=Dict[str, Any])
async def upsert_learning_path_template(
    payload: LearningPathTemplateUpsert,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    sk = (payload.skill_name or "").strip()
    if not sk:
        raise HTTPException(status_code=400, detail="skill_name is required")
    sk_lc = sk.lower()
    now = datetime.now(timezone.utc).isoformat()
    ex = await db[COL_LEARNING_PATH_TEMPLATES].find_one({"skill_name_lc": sk_lc}, {"_id": 0, "id": 1})
    tid = ex["id"] if ex else str(uuid.uuid4())
    doc = {
        "id": tid,
        "skill_name": sk,
        "skill_name_lc": sk_lc,
        "steps": [s.model_dump() for s in payload.steps],
        "updated_at": now,
    }
    if not ex:
        doc["created_at"] = now
    await db[COL_LEARNING_PATH_TEMPLATES].update_one({"skill_name_lc": sk_lc}, {"$set": doc}, upsert=True)
    return await db[COL_LEARNING_PATH_TEMPLATES].find_one({"skill_name_lc": sk_lc}, {"_id": 0})


@api_router.post("/workforce/training/assignments", response_model=Dict[str, Any])
async def create_training_assignment(
    payload: TrainingAssignmentCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "skills_write")
    code = (payload.employee_code or "").strip()
    sk = (payload.skill_name or "").strip()
    if not code or not sk:
        raise HTTPException(status_code=400, detail="employee_code and skill_name required")
    emp = await db.employees.find_one({"employee_code": code}, {"_id": 0})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")
    sk_lc = sk.lower()
    templates = await load_path_templates_map(db)
    from m5_training.recommendation_rules import path_for_skill

    steps = path_for_skill(sk, sk_lc, templates)
    now = datetime.now(timezone.utc).isoformat()
    aid = str(uuid.uuid4())
    doc = {
        "id": aid,
        "employee_code": code,
        "skill_name": sk,
        "skill_name_lc": sk_lc,
        "path_snapshot": steps,
        "status": "ASSIGNED",
        "progress_pct": 0.0,
        "assigned_by": current_user.get("id"),
        "assigned_at": now,
        "updated_at": now,
        "completed_at": None,
    }
    await db[COL_ASSIGNMENTS].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.get("/workforce/training/assignments", response_model=List[Dict[str, Any]])
async def list_training_assignments(
    employee_code: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    q: Dict[str, Any] = {}
    if employee_code:
        q["employee_code"] = (employee_code or "").strip()
    if status:
        q["status"] = status.upper()
    return await db[COL_ASSIGNMENTS].find(q, {"_id": 0}).sort("updated_at", -1).to_list(500)


@api_router.patch("/workforce/training/assignments/{assignment_id}", response_model=Dict[str, Any])
async def update_training_assignment_progress(
    assignment_id: str,
    payload: TrainingProgressUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    ex = await db[COL_ASSIGNMENTS].find_one({"id": assignment_id}, {"_id": 0})
    if not ex:
        raise HTTPException(status_code=404, detail="Assignment not found")
    patch: Dict[str, Any] = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if payload.progress_pct is not None:
        patch["progress_pct"] = max(0.0, min(100.0, float(payload.progress_pct)))
    if payload.status:
        patch["status"] = payload.status.upper()
        if payload.status.upper() == "COMPLETED":
            patch["progress_pct"] = 100.0
            patch["completed_at"] = patch["updated_at"]
    await db[COL_ASSIGNMENTS].update_one({"id": assignment_id}, {"$set": patch})
    return await db[COL_ASSIGNMENTS].find_one({"id": assignment_id}, {"_id": 0})


@api_router.get("/workforce/training/catalog", response_model=List[Dict[str, Any]])
async def list_training_catalog(
    limit: int = 100,
    skill: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    limit = min(max(1, int(limit)), 500)
    q: Dict[str, Any] = {}
    if skill:
        q["skill_tags_lc"] = (skill or "").strip().lower()
    cur = db[COL_LMS_COURSES].find(q, {"_id": 0}).sort("synced_at", -1).limit(limit)
    return await cur.to_list(limit)


@api_router.post("/workforce/training/certifications", response_model=Dict[str, Any])
async def create_training_certification(
    payload: TrainingCertificationCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    code = (payload.employee_code or "").strip()
    if not code or not (payload.title or "").strip():
        raise HTTPException(status_code=400, detail="employee_code and title required")
    emp = await db.employees.find_one({"employee_code": code}, {"_id": 0, "employee_code": 1})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")
    now = datetime.now(timezone.utc).isoformat()
    cid = str(uuid.uuid4())
    doc = {
        "id": cid,
        "employee_code": code,
        "title": payload.title.strip(),
        "issued_at": payload.issued_at,
        "expires_at": payload.expires_at,
        "created_by": current_user.get("id"),
        "created_at": now,
        "expiry_reminder_sent_at": None,
    }
    await db[COL_CERTIFICATIONS].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.get("/workforce/training/certifications", response_model=List[Dict[str, Any]])
async def list_training_certifications(
    employee_code: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    q: Dict[str, Any] = {}
    if employee_code:
        q["employee_code"] = (employee_code or "").strip()
    return await db[COL_CERTIFICATIONS].find(q, {"_id": 0}).sort("expires_at", 1).to_list(500)


@api_router.get("/workforce/training/manager-summary", response_model=Dict[str, Any])
async def get_training_manager_summary(
    manager_employee_id: str,
    current_user: dict = Depends(get_current_user),
):
    """M5-3: dashboard cards for a manager's direct reports (manager_employee_id = employee UUID)."""
    _require_phase1_access(current_user, "kpi_read")
    mid = (manager_employee_id or "").strip()
    if not mid:
        raise HTTPException(status_code=400, detail="manager_employee_id is required")
    boss = await db.employees.find_one({"id": mid}, {"_id": 0, "id": 1})
    if not boss:
        raise HTTPException(status_code=404, detail="Manager employee not found")
    return await manager_team_training_summary(db, manager_employee_id=mid)


@api_router.post("/admin/training/lms/sync", response_model=Dict[str, Any])
async def admin_training_lms_sync(
    payload: LmsSyncRequest = Body(default_factory=lambda: LmsSyncRequest()),
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    return await run_lms_catalog_sync(
        db,
        provider=payload.provider or DEFAULT_LMS_PROVIDER,
        actor_id=current_user.get("id"),
    )


@api_router.get("/admin/training/lms/sync/last", response_model=Dict[str, Any])
async def admin_training_lms_sync_last(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    rows = await db[COL_LMS_SYNC_RUNS].find({}, {"_id": 0}).sort("ended_at", -1).limit(1).to_list(1)
    return rows[0] if rows else {}


@api_router.post("/admin/training/certifications/scan-expiry", response_model=Dict[str, Any])
async def admin_training_cert_scan_expiry(
    days_ahead: int = 30,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)

    async def notify(c: Dict[str, Any]) -> None:
        title = (c.get("title") or "Certification").strip()
        code = c.get("employee_code") or ""
        exp = c.get("expires_at") or ""
        for rid in await _hr_escalation_recipient_ids():
            await create_notification(
                recipient_id=rid,
                notification_type="TRAINING_CERT_EXPIRING",
                title="Certification expiring soon",
                message=f"{title} for employee {code} expires {exp}",
                metadata={"certification_id": c.get("id"), "employee_code": code},
            )

    return await scan_certification_expiry(db, days_ahead=days_ahead, notify_fn=notify)


# ========================
# Phase-4 M6: Employee Satisfaction & Engagement (Pulse MVP)
# ========================


def _compute_sentiment(rating: int, response_text: Optional[str]) -> Dict[str, Any]:
    """Delegates to versioned M6 sentiment pipeline (label + score only for legacy callers)."""
    s = m6_compute_sentiment(rating, response_text)
    return {"sentiment_label": s["sentiment_label"], "sentiment_score": s["sentiment_score"]}


@api_router.post("/employee-engagement/surveys", response_model=PulseSurveyResponse)
async def create_pulse_survey(
    payload: PulseSurveyCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")

    title = (payload.title or "").strip()
    question = (payload.question or "").strip()
    if not title or not question:
        raise HTTPException(status_code=400, detail="title and question are required")

    survey_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    doc = {
        "id": survey_id,
        "title": title,
        "question": question,
        "rating_min": int(payload.rating_min),
        "rating_max": int(payload.rating_max),
        "active": bool(payload.active),
        "target_all": bool(payload.target_all),
        "target_departments": [str(d).strip() for d in (payload.target_departments or []) if str(d).strip()],
        "created_at": now,
    }
    await db.employee_engagement_surveys.insert_one(doc)
    return PulseSurveyResponse(**doc)


@api_router.get("/employee-engagement/surveys", response_model=List[PulseSurveyResponse])
async def list_pulse_surveys(
    active_only: bool = True,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_read")

    query: Dict[str, Any] = {}
    if active_only:
        query["active"] = True

    surveys = await db.employee_engagement_surveys.find(query, {"_id": 0}).sort("created_at", -1).to_list(200)
    return [PulseSurveyResponse(**s) for s in surveys]


@api_router.post("/employee-engagement/responses", response_model=PulseSurveyResponseResponse)
async def submit_pulse_response(
    payload: PulseSurveyResponseCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")

    survey = await db.employee_engagement_surveys.find_one({"id": payload.survey_id}, {"_id": 0})
    if not survey:
        raise HTTPException(status_code=404, detail="Survey not found")
    if not survey.get("active", True):
        raise HTTPException(status_code=400, detail="Survey is not active")

    employee_code = (payload.employee_code or "").strip()
    if not employee_code:
        raise HTTPException(status_code=400, detail="employee_code is required")
    emp = await db.employees.find_one({"employee_code": employee_code}, {"_id": 0, "employee_code": 1})
    if not emp:
        raise HTTPException(status_code=404, detail="Employee not found")

    rating = int(payload.rating)
    rating_min = int(survey.get("rating_min") or 1)
    rating_max = int(survey.get("rating_max") or 5)
    if rating < rating_min or rating > rating_max:
        raise HTTPException(status_code=400, detail=f"rating must be between {rating_min} and {rating_max}")

    sent = m6_compute_sentiment(rating, payload.response_text)
    topic_primary = classify_topic(payload.response_text)

    now = datetime.now(timezone.utc).isoformat()
    resp_id = str(uuid.uuid4())
    doc = {
        "id": resp_id,
        "survey_id": payload.survey_id,
        "employee_code": employee_code,
        "rating": rating,
        "response_text": payload.response_text,
        "sentiment_label": sent["sentiment_label"],
        "sentiment_score": sent["sentiment_score"],
        "sentiment_pipeline_version": sent.get("sentiment_pipeline_version"),
        "topic_primary": topic_primary,
        "created_at": now,
        "updated_at": None,
    }
    await db.employee_engagement_responses.insert_one(doc)
    return PulseSurveyResponseResponse(**doc)


@api_router.get("/employee-engagement/responses", response_model=PulseSurveyResponsesPagedResponse)
async def list_pulse_responses(
    survey_id: Optional[str] = None,
    employee_code: Optional[str] = None,
    q: Optional[str] = None,
    page: int = 1,
    page_size: int = 20,
    sort_by: str = "created_at",
    sort_dir: Literal["asc", "desc"] = "desc",
    current_user: dict = Depends(get_current_user),
):
    _require_engagement_raw_privileged(current_user)
    await log_engagement_privacy_event(
        db,
        action="LIST_RAW_RESPONSES",
        actor_id=current_user.get("id"),
        survey_id=survey_id,
        detail={"page": page, "employee_code_filter": employee_code},
    )

    page = max(1, page)
    page_size = min(max(1, page_size), 100)
    allowed_sort = {"created_at", "rating", "sentiment_score", "sentiment_label"}
    if sort_by not in allowed_sort:
        sort_by = "created_at"
    direction = 1 if sort_dir == "asc" else -1

    query: Dict[str, Any] = {}
    if survey_id:
        query["survey_id"] = survey_id
    if employee_code:
        query["employee_code"] = employee_code
    if q:
        query["$or"] = [
            {"employee_code": {"$regex": q, "$options": "i"}},
            {"response_text": {"$regex": q, "$options": "i"}},
        ]

    total = await db.employee_engagement_responses.count_documents(query)
    rows = await (
        db.employee_engagement_responses.find(query, {"_id": 0})
        .sort(sort_by, direction)
        .skip((page - 1) * page_size)
        .limit(page_size)
        .to_list(page_size)
    )

    return PulseSurveyResponsesPagedResponse(
        items=[PulseSurveyResponseResponse(**r) for r in rows],
        page=page,
        page_size=page_size,
        total=total,
        total_pages=max(1, (total + page_size - 1) // page_size),
        sort_by=sort_by,
        sort_dir=sort_dir,
    )


@api_router.get("/employee-engagement/dashboard", response_model=PulseSurveyDashboardResponse)
async def get_employee_engagement_dashboard(
    survey_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    await log_engagement_privacy_event(
        db,
        action="VIEW_DASHBOARD_AGGREGATE",
        actor_id=current_user.get("id"),
        survey_id=survey_id,
        detail={"role": current_user.get("role")},
    )

    now = datetime.now(timezone.utc)
    start_30 = (now - timedelta(days=30)).isoformat()

    qbase: Dict[str, Any] = {}
    if survey_id:
        qbase["survey_id"] = survey_id.strip()

    total_responses = await db.employee_engagement_responses.count_documents(qbase)
    last_30_days_responses = await db.employee_engagement_responses.count_documents(
        {**qbase, "created_at": {"$gte": start_30}}
    )

    sample = (
        await db.employee_engagement_responses.find(
            qbase,
            {"_id": 0, "rating": 1, "response_text": 1, "sentiment_label": 1, "topic_primary": 1, "created_at": 1},
        )
        .sort("created_at", -1)
        .limit(5000)
        .to_list(5000)
    )

    if survey_id and total_responses < anonymity_min_threshold():
        base = {
            "last_30_days_responses": last_30_days_responses,
        }
        red = redacted_dashboard_payload(total_responses=total_responses, base=base)
        return PulseSurveyDashboardResponse(**red)

    ratings = [int(x.get("rating") or 0) for x in sample]
    avg_rating = round(sum(ratings) / len(ratings), 2) if ratings else 0.0

    sentiment_counts: Dict[str, int] = {}
    for label in ["POSITIVE", "NEUTRAL", "NEGATIVE"]:
        sentiment_counts[label] = sum(1 for x in sample if (x.get("sentiment_label") or "") == label)

    topic_counts = aggregate_topic_counts(sample)
    weekly_trend = weekly_rating_trends(sample, max_weeks=8)
    tier, rationale = confidence_tier(total_responses)

    return PulseSurveyDashboardResponse(
        total_responses=total_responses,
        avg_rating=avg_rating,
        last_30_days_responses=last_30_days_responses,
        sentiment_counts=sentiment_counts,
        topic_counts=topic_counts,
        weekly_trend=weekly_trend,
        display_confidence=tier,
        confidence_rationale=rationale,
        anonymity_note=None,
    )


@api_router.get("/employee-engagement/survey-templates", response_model=List[Dict[str, Any]])
async def list_engagement_survey_templates(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_read")
    return await db[COL_SURVEY_TEMPLATES].find({}, {"_id": 0}).sort("updated_at", -1).to_list(200)


@api_router.post("/employee-engagement/survey-templates", response_model=Dict[str, Any])
async def create_engagement_survey_template(
    payload: EngagementSurveyTemplateCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    tid = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": tid,
        "name": (payload.name or "").strip(),
        "description": (payload.description or "").strip() or None,
        "default_title": (payload.default_title or "").strip(),
        "default_question": (payload.default_question or "").strip(),
        "rating_min": int(payload.rating_min),
        "rating_max": int(payload.rating_max),
        "target_all": bool(payload.target_all),
        "target_departments": [str(d).strip() for d in (payload.target_departments or []) if str(d).strip()],
        "created_at": now,
        "updated_at": now,
    }
    if not doc["name"] or not doc["default_title"] or not doc["default_question"]:
        raise HTTPException(status_code=400, detail="name, default_title, and default_question are required")
    await db[COL_SURVEY_TEMPLATES].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.put("/employee-engagement/survey-templates/{template_id}", response_model=Dict[str, Any])
async def update_engagement_survey_template(
    template_id: str,
    payload: EngagementSurveyTemplateUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    ex = await db[COL_SURVEY_TEMPLATES].find_one({"id": template_id}, {"_id": 0})
    if not ex:
        raise HTTPException(status_code=404, detail="Template not found")
    patch = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    if "target_departments" in patch and patch["target_departments"] is not None:
        patch["target_departments"] = [str(d).strip() for d in patch["target_departments"] if str(d).strip()]
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db[COL_SURVEY_TEMPLATES].update_one({"id": template_id}, {"$set": patch})
    return await db[COL_SURVEY_TEMPLATES].find_one({"id": template_id}, {"_id": 0})


@api_router.delete("/employee-engagement/survey-templates/{template_id}")
async def delete_engagement_survey_template(template_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_write")
    r = await db[COL_SURVEY_TEMPLATES].delete_one({"id": template_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template not found")
    return {"message": "Template deleted", "id": template_id}


@api_router.post("/employee-engagement/surveys/from-template", response_model=PulseSurveyResponse)
async def create_survey_from_template(
    payload: EngagementSurveyFromTemplateCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    tpl = await db[COL_SURVEY_TEMPLATES].find_one({"id": payload.template_id}, {"_id": 0})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    title = (payload.title_override or tpl.get("default_title") or "").strip()
    question = (payload.question_override or tpl.get("default_question") or "").strip()
    if not title or not question:
        raise HTTPException(status_code=400, detail="title and question required")
    survey_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    doc = {
        "id": survey_id,
        "title": title,
        "question": question,
        "rating_min": int(tpl.get("rating_min") or 1),
        "rating_max": int(tpl.get("rating_max") or 5),
        "active": True,
        "target_all": bool(tpl.get("target_all", True)),
        "target_departments": list(tpl.get("target_departments") or []),
        "template_id": tpl.get("id"),
        "created_at": now,
    }
    await db.employee_engagement_surveys.insert_one(doc)
    return PulseSurveyResponse(**doc)


@api_router.post("/employee-engagement/survey-schedules", response_model=Dict[str, Any])
async def create_engagement_survey_schedule(
    payload: EngagementScheduleCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    tpl = await db[COL_SURVEY_TEMPLATES].find_one({"id": payload.template_id}, {"_id": 0})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    sid = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    now_iso = now.isoformat()
    nxt = payload.next_run_at or now_iso
    doc = {
        "id": sid,
        "template_id": payload.template_id,
        "cadence": str(payload.cadence).upper(),
        "enabled": bool(payload.enabled),
        "next_run_at": nxt,
        "last_run_at": None,
        "created_by": current_user.get("id"),
        "created_at": now_iso,
        "updated_at": now_iso,
    }
    await db[COL_SURVEY_SCHEDULES].insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.get("/employee-engagement/survey-schedules", response_model=List[Dict[str, Any]])
async def list_engagement_survey_schedules(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "employees_read")
    return await db[COL_SURVEY_SCHEDULES].find({}, {"_id": 0}).sort("next_run_at", 1).to_list(200)


@api_router.post("/admin/employee-engagement/schedules/dispatch-due", response_model=Dict[str, Any])
async def admin_dispatch_engagement_schedules(current_user: dict = Depends(get_current_user)):
    """M6-1: materialize surveys from templates when `next_run_at` is due."""
    _require_admin(current_user)
    now = datetime.now(timezone.utc)
    now_iso = now.isoformat()
    due = await db[COL_SURVEY_SCHEDULES].find({"enabled": True, "next_run_at": {"$lte": now_iso}}, {"_id": 0}).to_list(
        100
    )
    created_surveys = 0
    for sch in due:
        tpl = await db[COL_SURVEY_TEMPLATES].find_one({"id": sch.get("template_id")}, {"_id": 0})
        if not tpl:
            continue
        survey_id = str(uuid.uuid4())
        created = datetime.now(timezone.utc).isoformat()
        doc = {
            "id": survey_id,
            "title": str(tpl.get("default_title") or "Pulse check"),
            "question": str(tpl.get("default_question") or ""),
            "rating_min": int(tpl.get("rating_min") or 1),
            "rating_max": int(tpl.get("rating_max") or 5),
            "active": True,
            "target_all": bool(tpl.get("target_all", True)),
            "target_departments": list(tpl.get("target_departments") or []),
            "template_id": tpl.get("id"),
            "schedule_id": sch.get("id"),
            "created_at": created,
        }
        await db.employee_engagement_surveys.insert_one(doc)
        created_surveys += 1
        ref = parse_iso_dt(str(sch.get("next_run_at") or now_iso))
        nxt = next_run_after(str(sch.get("cadence") or "MONTHLY"), ref)
        await db[COL_SURVEY_SCHEDULES].update_one(
            {"id": sch.get("id")},
            {"$set": {"last_run_at": created, "next_run_at": nxt.isoformat(), "updated_at": created}},
        )
    return {"dispatched_schedules": len(due), "surveys_created": created_surveys}


@api_router.post("/admin/employee-engagement/surveys/{survey_id}/remind-participation", response_model=Dict[str, Any])
async def admin_remind_engagement_survey_participation(
    survey_id: str,
    current_user: dict = Depends(get_current_user),
):
    """M6-1: nudge likely participants (email-matched users) + HR summary."""
    _require_admin(current_user)
    survey = await db.employee_engagement_surveys.find_one({"id": survey_id}, {"_id": 0})
    if not survey:
        raise HTTPException(status_code=404, detail="Survey not found")

    emp_query: Dict[str, Any] = {}
    if not survey.get("target_all", True):
        depts = [d for d in (survey.get("target_departments") or []) if d]
        if depts:
            emp_query["department"] = {"$in": depts}

    employees = await db.employees.find(emp_query, {"_id": 0, "employee_code": 1, "email": 1, "email_lc": 1}).to_list(
        5000
    )
    responded = await db.employee_engagement_responses.distinct("employee_code", {"survey_id": survey_id})
    responded_set = {str(x).strip() for x in responded if x}
    pending_emps = [e for e in employees if (e.get("employee_code") or "").strip() not in responded_set]

    notified = 0
    for e in pending_emps:
        elc = e.get("email_lc") or _norm_email(e.get("email"))
        if not elc:
            continue
        user_doc = await db.users.find_one({"email": {"$regex": f"^{re.escape(elc)}$", "$options": "i"}}, {"_id": 0, "id": 1})
        if user_doc and user_doc.get("id"):
            await create_notification(
                recipient_id=user_doc["id"],
                notification_type="ENGAGEMENT_SURVEY_REMINDER",
                title="Pulse survey reminder",
                message=f"Please submit your response for: {survey.get('title')}",
                metadata={"survey_id": survey_id},
            )
            notified += 1

    for rid in await _hr_escalation_recipient_ids():
        await create_notification(
            recipient_id=rid,
            notification_type="ENGAGEMENT_PARTICIPATION_SUMMARY",
            title="Engagement participation snapshot",
            message=f"Survey '{survey.get('title')}': {len(responded_set)} responses, ~{len(pending_emps)} pending employees in scope.",
            metadata={"survey_id": survey_id, "notified_users": notified},
        )

    return {
        "survey_id": survey_id,
        "employees_in_scope": len(employees),
        "already_responded": len(responded_set),
        "pending_employees": len(pending_emps),
        "notifications_to_users": notified,
    }


@api_router.get("/employee-engagement/privacy-audit", response_model=List[Dict[str, Any]])
async def list_engagement_privacy_audit(
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    limit = min(max(1, int(limit)), 200)
    return await db[COL_PRIVACY_AUDIT].find({}, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)


@api_router.get("/workforce/retention", response_model=RetentionDashboardResponse)
async def get_high_skill_retention(current_user: dict = Depends(get_current_user)):
    """
    Phase-4 M8 MVP:
    - Identify critical skills from HIGH priority workforce inventory.
    - Compute a shortage-based risk score per critical skill.
    - For each employee who has one or more critical skills, compute average risk score.
    """
    _require_phase1_access(current_user, "kpi_read")

    generated_at = datetime.now(timezone.utc).isoformat()

    # Critical skills based on HIGH priority (fallback to top gaps if none).
    critical = await db.workforce_skills.find(
        {"priority": "HIGH"},
        {"_id": 0, "skill_name": 1, "priority": 1, "demand_count": 1, "supply_count": 1, "gap": 1, "notes": 1},
    ).to_list(200)

    if not critical:
        top = await db.workforce_skills.find(
            {},
            {"_id": 0, "skill_name": 1, "priority": 1, "demand_count": 1, "supply_count": 1, "gap": 1},
        ).sort("gap", -1).limit(5).to_list(50)
        critical = top

    critical_metrics: List[RetentionCriticalSkill] = []
    for c in critical:
        name = c.get("skill_name") or ""
        demand = max(0, int(c.get("demand_count") or 0))
        supply = max(0, int(c.get("supply_count") or 0))
        shortage = max(0, demand - supply)
        shortage_ratio = round(shortage / demand, 3) if demand > 0 else 0.0
        risk_score = round(shortage_ratio, 3)
        critical_metrics.append(
            RetentionCriticalSkill(
                skill_name=name,
                priority=str(c.get("priority") or "HIGH").upper(),
                demand_count=demand,
                supply_count=supply,
                shortage_count=shortage,
                shortage_ratio=shortage_ratio,
                risk_score=risk_score,
            )
        )

    # Employee risk based on critical skills present in their employee skills.
    employees = await db.employees.find({}, {"_id": 0, "employee_code": 1, "full_name": 1, "skills": 1}).to_list(5000)

    critical_by_name = {m.skill_name.strip().lower(): m for m in critical_metrics if m.skill_name}

    def risk_label(score: float) -> str:
        if score >= 0.7:
            return "HIGH"
        if score >= 0.4:
            return "MEDIUM"
        return "LOW"

    risk_employees: List[RetentionRiskEmployee] = []
    for e in employees:
        emp_code = e.get("employee_code") or ""
        emp_name = e.get("full_name") or ""
        emp_skills = set(
            (s.strip().lower() for s in (e.get("skills") or []) if isinstance(s, str) and s.strip())
        )
        matched = []
        for skill_key, m in critical_by_name.items():
            if skill_key in emp_skills:
                matched.append(RetentionRiskEmployeeSkill(skill_name=m.skill_name, risk_score=m.risk_score))

        if not matched:
            continue

        avg_risk = round(sum(x.risk_score for x in matched) / len(matched), 3) if matched else 0.0
        risk_employees.append(
            RetentionRiskEmployee(
                employee_code=emp_code,
                full_name=emp_name,
                critical_skills_matched=len(matched),
                risk_score=avg_risk,
                risk_label=risk_label(avg_risk),
                skills=matched,
            )
        )

    risk_employees.sort(key=lambda x: x.risk_score, reverse=True)
    top_risk_employees = risk_employees[:20]
    avg_risk_score = round(sum(x.risk_score for x in risk_employees) / len(risk_employees), 3) if risk_employees else 0.0

    attrition_v1: Optional[AttritionV1DashboardSummary] = None
    scored_rows = await db[COL_ATTRITION_SCORES_LATEST].find({}, {"_id": 0}).to_list(5000)
    if scored_rows:
        last_ts = max((str(r.get("computed_at") or "") for r in scored_rows), default="")
        avg_a = sum(float(r.get("attrition_risk") or 0) for r in scored_rows) / len(scored_rows)
        top_a = sorted(scored_rows, key=lambda r: float(r.get("attrition_risk") or 0), reverse=True)[:15]
        briefs: List[AttritionV1EmployeeBrief] = []
        for r in top_a:
            tf = r.get("top_factors") or []
            factors = [
                AttritionV1FactorBrief(
                    feature=str(x.get("feature") or ""),
                    contribution=float(x.get("contribution") or 0),
                    direction=str(x.get("direction") or ""),
                )
                for x in tf[:5]
                if isinstance(x, dict)
            ]
            briefs.append(
                AttritionV1EmployeeBrief(
                    employee_id=str(r.get("employee_id") or ""),
                    employee_code=str(r.get("employee_code") or ""),
                    full_name=str(r.get("full_name") or ""),
                    attrition_risk=float(r.get("attrition_risk") or 0),
                    confidence=float(r.get("confidence") or 0),
                    risk_band=str(r.get("risk_band") or "LOW"),
                    segments=[str(s) for s in (r.get("segments") or []) if s],
                    top_factors=factors,
                )
            )
        attrition_v1 = AttritionV1DashboardSummary(
            model_version=str(scored_rows[0].get("model_version") or ATTRITION_MODEL_VERSION),
            last_computed_at=last_ts or None,
            scored_employee_count=len(scored_rows),
            avg_attrition_risk=round(avg_a, 4),
            top_at_risk=briefs,
        )

    return RetentionDashboardResponse(
        generated_at=generated_at,
        critical_skills=critical_metrics,
        total_high_skill_employees=len(risk_employees),
        avg_risk_score=avg_risk_score,
        top_risk_employees=top_risk_employees,
        attrition_v1=attrition_v1,
    )


@api_router.post("/workforce/retention/v1/score-run")
async def m8_retention_score_run(current_user: dict = Depends(get_current_user)):
    """M8-1: batch score active employees; requires kpi_read."""
    _require_phase1_access(current_user, "kpi_read")
    out = await m8_retention_service.run_score_batch(db)
    return out


def _m8_model_public_view(doc: Dict[str, Any]) -> Dict[str, Any]:
    """Avoid shipping large base64 classifier blobs to the browser."""
    out = dict(doc)
    b64 = out.get("classifier_blob_b64")
    if isinstance(b64, str) and b64:
        out["classifier_blob_b64"] = f"<redacted len={len(b64)}>"
        out["classifier_present"] = True
    else:
        out["classifier_present"] = False
    return out


@api_router.post("/workforce/retention/v1/score-run-cron")
async def m8_retention_score_run_cron(request: Request):
    """
    Scheduled attrition score batch (no JWT). Requires env M8_SCORE_RUN_TOKEN and header
    X-M8-Score-Token matching it. Use GitHub Actions or internal cron.
    """
    expected = (os.environ.get("M8_SCORE_RUN_TOKEN") or "").strip()
    got = (request.headers.get("X-M8-Score-Token") or "").strip()
    if not expected or got != expected:
        raise HTTPException(status_code=401, detail="Invalid or missing X-M8-Score-Token")
    return await m8_retention_service.run_score_batch(db)


@api_router.post("/workforce/retention/v1/train")
async def m8_retention_train(payload: M8AttritionTrainRequest, current_user: dict = Depends(get_current_user)):
    """M8-1: refit default logistic weights from labeled employees (admin)."""
    _require_admin(current_user)
    labels = [{"employee_id": x.employee_id, "churned": x.churned} for x in (payload.labels or [])]
    return await m8_retention_service.train_and_store(
        db,
        labels,
        use_gradient_boosting=bool(payload.use_gradient_boosting),
        interaction_features=bool(payload.interaction_features),
    )


@api_router.get("/workforce/retention/v1/model")
async def m8_retention_get_model(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    doc = await m8_retention_service.get_model_state_doc(db)
    return _m8_model_public_view(doc)


@api_router.patch("/workforce/retention/v1/model")
async def m8_retention_patch_model(
    payload: M8ModelRuntimePatch,
    current_user: dict = Depends(get_current_user),
):
    """Switch ensemble_mode (linear / gb / avg) or interaction features without retraining."""
    _require_admin(current_user)
    p = payload.model_dump(exclude_none=True)
    updated = await m8_retention_service.patch_model_runtime_settings(
        db,
        ensemble_mode=p.get("ensemble_mode"),
        interaction_features_enabled=p.get("interaction_features_enabled"),
    )
    return _m8_model_public_view(updated)


@api_router.get("/workforce/retention/v1/scores")
async def m8_retention_list_scores(
    department: Optional[str] = None,
    segment: Optional[str] = None,
    band: Optional[str] = None,
    min_risk: Optional[float] = None,
    limit: int = 200,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    limit = min(max(1, int(limit)), 500)
    return await m8_retention_service.list_latest_scores(
        db,
        department=department,
        segment=segment,
        band=band,
        min_risk=min_risk,
        limit=limit,
    )


@api_router.get("/workforce/retention/v1/employees/{employee_id}/score")
async def m8_retention_employee_score(employee_id: str, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    row = await m8_retention_service.get_employee_latest_score(db, employee_id)
    if not row:
        raise HTTPException(status_code=404, detail="No score — run POST /workforce/retention/v1/score-run first")
    return row


@api_router.get("/workforce/retention/v1/segments/settings")
async def m8_retention_segment_settings_get(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    return await m8_retention_service.get_segment_settings(db)


@api_router.put("/workforce/retention/v1/segments/settings")
async def m8_retention_segment_settings_put(
    payload: M8SegmentSettingsUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    patch = payload.model_dump(exclude_none=True)
    return await m8_retention_service.save_segment_settings(db, patch)


@api_router.get("/workforce/retention/v1/playbooks")
async def m8_retention_playbooks_list(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    return await m8_retention_service.list_playbooks(db)


@api_router.post("/workforce/retention/v1/playbooks")
async def m8_retention_playbooks_create(
    payload: M8RetentionPlaybookCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    return await m8_retention_service.create_playbook(
        db,
        {
            "title": payload.title,
            "description": payload.description,
            "category": payload.category,
            "suggested_duration_days": payload.suggested_duration_days,
        },
    )


@api_router.get("/workforce/retention/v1/interventions")
async def m8_retention_interventions_list(
    employee_id: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 100,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    limit = min(max(1, int(limit)), 300)
    return await m8_retention_service.list_interventions(db, employee_id=employee_id, status=status, limit=limit)


@api_router.post("/workforce/retention/v1/interventions")
async def m8_retention_interventions_create(
    payload: M8RetentionInterventionCreate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    try:
        return await m8_retention_service.create_intervention(
            db,
            employee_id=payload.employee_id,
            playbook_id=payload.playbook_id,
            assigned_by=str(current_user.get("id") or ""),
            notes=payload.notes or "",
        )
    except ValueError as e:
        msg = str(e)
        if msg == "playbook_not_found":
            raise HTTPException(status_code=404, detail="Playbook not found")
        if msg == "employee_not_found":
            raise HTTPException(status_code=404, detail="Employee not found")
        raise HTTPException(status_code=400, detail=msg)


@api_router.patch("/workforce/retention/v1/interventions/{intervention_id}/timeline")
async def m8_retention_intervention_timeline(
    intervention_id: str,
    payload: M8RetentionTimelineEvent,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    try:
        return await m8_retention_service.append_timeline_event(
            db,
            intervention_id,
            event_type=payload.event_type,
            note=payload.note or "",
        )
    except ValueError:
        raise HTTPException(status_code=404, detail="Intervention not found")


@api_router.put("/workforce/retention/v1/interventions/{intervention_id}/outcome")
async def m8_retention_intervention_outcome(
    intervention_id: str,
    payload: M8RetentionOutcomeUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "employees_write")
    try:
        return await m8_retention_service.set_intervention_outcome(
            db,
            intervention_id,
            outcome=payload.outcome,
            note=payload.note or "",
        )
    except ValueError as e:
        if str(e) == "not_found":
            raise HTTPException(status_code=404, detail="Intervention not found")
        raise HTTPException(status_code=400, detail=str(e))


@api_router.get("/workforce/retention/v1/metrics")
async def m8_retention_metrics(current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    return await m8_retention_service.retention_metrics(db)

# ========================
# Phase-4/5 M9: Analytics & Executive Dashboard (MVP)
# ========================

class RetentionRiskEmployeeBrief(BaseModel):
    model_config = ConfigDict(extra="ignore")
    employee_code: str
    full_name: str
    risk_label: str
    risk_score: float
    critical_skills_matched: int


class StrategicExecutiveDashboardResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    generated_at: str
    # M9: analytics time window for engagement counts + M7 run window (default 30).
    analytics_window_days: int = 30

    # Existing operational KPIs (from Phase-1)
    employee_count: int
    active_employee_count: int
    attrition_count: int
    attrition_rate_pct: float
    avg_skills_per_employee: float
    top_skill_gaps: List[Dict[str, Any]]

    # Workforce forecast / planning KPIs
    workforce_horizon_months: int
    forecast_gap_total: int
    resource_total_shortage: int
    resource_total_bench: int

    # Engagement KPIs (Phase-4)
    engagement_total_responses: int
    engagement_avg_rating: float
    engagement_last_30_days_responses: int
    engagement_sentiment_counts: Dict[str, int]

    # Retention KPIs (Phase-4)
    retention_total_high_skill_employees: int
    retention_avg_risk_score: float
    retention_top_risk_employees: List[RetentionRiskEmployeeBrief]

    # M7: cost / automation (30d window, baseline-driven)
    automation_runs_succeeded_30d: int = 0
    automation_runs_failed_30d: int = 0
    cost_optimization_baselines_count: int = 0
    estimated_manual_minutes_saved_30d: float = 0.0
    estimated_cost_saved_usd_30d: float = 0.0


@api_router.get("/executive/strategic-dashboard", response_model=StrategicExecutiveDashboardResponse)
async def get_strategic_executive_dashboard(
    horizon_months: int = 3,
    window_days: int = 30,
    current_user: dict = Depends(get_current_user),
):
    """
    M9 MVP:
    - Combines operational KPIs + predictive workforce gap totals + engagement + retention.
    - `window_days` drives engagement/M7 run windows (default 30). Field names remain *_30d for compatibility.
    """
    _require_phase1_access(current_user, "kpi_read")
    horizon_months = max(1, int(horizon_months or 3))
    window_days = max(1, min(int(window_days or 30), 365))
    data = await build_strategic_dashboard_data(
        db,
        horizon_months=horizon_months,
        window_days=window_days,
        scope_employee_ids=None,
    )
    top = [
        RetentionRiskEmployeeBrief(
            employee_code=r["employee_code"],
            full_name=r["full_name"],
            risk_label=r["risk_label"],
            risk_score=r["risk_score"],
            critical_skills_matched=r["critical_skills_matched"],
        )
        for r in data.get("retention_top_risk_employees") or []
    ]
    return StrategicExecutiveDashboardResponse(
        generated_at=data["generated_at"],
        analytics_window_days=int(data.get("drill_window_days") or window_days),
        employee_count=data["employee_count"],
        active_employee_count=data["active_employee_count"],
        attrition_count=data["attrition_count"],
        attrition_rate_pct=data["attrition_rate_pct"],
        avg_skills_per_employee=data["avg_skills_per_employee"],
        top_skill_gaps=data["top_skill_gaps"],
        workforce_horizon_months=data["workforce_horizon_months"],
        forecast_gap_total=data["forecast_gap_total"],
        resource_total_shortage=data["resource_total_shortage"],
        resource_total_bench=data["resource_total_bench"],
        engagement_total_responses=data["engagement_total_responses"],
        engagement_avg_rating=data["engagement_avg_rating"],
        engagement_last_30_days_responses=data["engagement_last_30_days_responses"],
        engagement_sentiment_counts=data["engagement_sentiment_counts"],
        retention_total_high_skill_employees=data["retention_total_high_skill_employees"],
        retention_avg_risk_score=data["retention_avg_risk_score"],
        retention_top_risk_employees=top,
        automation_runs_succeeded_30d=data["automation_runs_succeeded_30d"],
        automation_runs_failed_30d=data["automation_runs_failed_30d"],
        cost_optimization_baselines_count=data["cost_optimization_baselines_count"],
        estimated_manual_minutes_saved_30d=data["estimated_manual_minutes_saved_30d"],
        estimated_cost_saved_usd_30d=data["estimated_cost_saved_usd_30d"],
    )


class M9MonthlySnapshotCreate(BaseModel):
    period: str = Field(..., description="YYYY-MM")
    horizon_months: int = 3
    window_days: int = 30
    department: Optional[str] = None
    manager_root_id: Optional[str] = None
    role_title_contains: Optional[str] = None


class M9ThresholdUpdate(BaseModel):
    warn: Optional[float] = None
    critical: Optional[float] = None
    higher_is_worse: Optional[bool] = None


class M9DefinitionUpdate(BaseModel):
    description: Optional[str] = None
    formula: Optional[str] = None
    owner_role: Optional[str] = None
    steward_team: Optional[str] = None
    source_system: Optional[str] = None


@api_router.get("/admin/m9/kpi-thresholds")
async def admin_m9_list_thresholds(current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    return {"items": await list_threshold_overrides(db)}


@api_router.put("/admin/m9/kpi-thresholds/{kpi_id}")
async def admin_m9_upsert_threshold(
    kpi_id: str,
    body: M9ThresholdUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    try:
        return await upsert_threshold_override(db, kpi_id, body.model_dump(exclude_unset=True))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api_router.delete("/admin/m9/kpi-thresholds/{kpi_id}")
async def admin_m9_delete_threshold(kpi_id: str, current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    ok = await delete_threshold_override(db, kpi_id)
    if not ok:
        raise HTTPException(status_code=404, detail="threshold override not found")
    return {"ok": True}


@api_router.put("/admin/m9/kpi-definitions/{kpi_id}")
async def admin_m9_upsert_definition(
    kpi_id: str,
    body: M9DefinitionUpdate,
    current_user: dict = Depends(get_current_user),
):
    _require_admin(current_user)
    kid = (kpi_id or "").strip()
    if not kid:
        raise HTTPException(status_code=400, detail="kpi_id required")
    patch = {k: v for k, v in body.model_dump(exclude_unset=True).items() if v is not None}
    if not patch:
        raise HTTPException(status_code=400, detail="no fields to update")
    patch["kpi_id"] = kid
    patch["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db[COL_M9_KPI_DEFINITIONS].update_one({"kpi_id": kid}, {"$set": patch}, upsert=True)
    items = await load_merged_kpi_definitions(db)
    row = next((d for d in items if d.get("kpi_id") == kid), None)
    if not row:
        raise HTTPException(status_code=404, detail="kpi not found")
    return row


@api_router.delete("/admin/m9/kpi-definitions/{kpi_id}")
async def admin_m9_delete_definition(kpi_id: str, current_user: dict = Depends(get_current_user)):
    _require_admin(current_user)
    ok = await delete_definition_override(db, kpi_id)
    if not ok:
        raise HTTPException(status_code=404, detail="definition override not found")
    items = await load_merged_kpi_definitions(db)
    row = next((d for d in items if d.get("kpi_id") == (kpi_id or "").strip()), None)
    if not row:
        raise HTTPException(status_code=404, detail="kpi not found")
    return row


@api_router.get("/executive/m9/kpi-definitions")
async def m9_list_kpi_definitions(current_user: dict = Depends(get_current_user)):
    """M9-1: KPI semantic layer — definitions + ownership (merged with Mongo overrides)."""
    _require_phase1_access(current_user, "kpi_read")
    return {"items": await load_merged_kpi_definitions(db)}


@api_router.get("/executive/m9/kpi-pack")
async def m9_kpi_pack(
    horizon_months: int = 3,
    window_days: int = 30,
    current_user: dict = Depends(get_current_user),
):
    """M9-1: Unified KPI retrieval contract (definitions + computed values + freshness)."""
    _require_phase1_access(current_user, "kpi_read")
    return await get_kpi_pack(db, horizon_months=horizon_months, window_days=window_days)


@api_router.get("/executive/m9/freshness")
async def m9_freshness(current_user: dict = Depends(get_current_user)):
    """M9-1: Data freshness vs SLA for executive sources."""
    _require_phase1_access(current_user, "kpi_read")
    return await compute_source_freshness(db)


@api_router.get("/executive/m9/drill-options")
async def m9_drill_options(current_user: dict = Depends(get_current_user)):
    """M9-2: Filter options for org / team / role drill."""
    _require_phase1_access(current_user, "kpi_read")
    return await drill_filter_options(db)


@api_router.get("/executive/m9/dashboard-bundle")
async def m9_executive_dashboard_bundle(
    horizon_months: int = 3,
    window_days: int = 30,
    department: Optional[str] = None,
    manager_root_id: Optional[str] = None,
    role_title_contains: Optional[str] = None,
    compare_period: Optional[str] = None,
    compare_against: Optional[str] = None,
    trends_months: int = 12,
    snapshot_limit: int = 24,
    current_user: dict = Depends(get_current_user),
):
    """M9: Pack, drill, definitions, trends, filter options, and snapshots in one response."""
    _require_phase1_access(current_user, "kpi_read")
    return await get_executive_dashboard_bundle(
        db,
        horizon_months=horizon_months,
        window_days=window_days,
        department=department,
        manager_root_id=manager_root_id,
        role_title_contains=role_title_contains,
        compare_period=compare_period,
        compare_against=compare_against,
        trends_months=trends_months,
        snapshot_limit=snapshot_limit,
    )


@api_router.get("/executive/m9/strategic-drill")
async def m9_strategic_drill(
    horizon_months: int = 3,
    window_days: int = 30,
    department: Optional[str] = None,
    manager_root_id: Optional[str] = None,
    role_title_contains: Optional[str] = None,
    compare_period: Optional[str] = None,
    compare_against: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    """M9-2: Strategic dashboard with org/team/role scope + linked time window (short-TTL cache)."""
    _require_phase1_access(current_user, "kpi_read")
    return await get_drill_dashboard_cached(
        db,
        horizon_months=horizon_months,
        window_days=window_days,
        department=department,
        manager_root_id=manager_root_id,
        role_title_contains=role_title_contains,
        compare_period=compare_period,
        compare_against=compare_against,
    )


@api_router.get("/executive/m9/kpi-compare")
async def m9_kpi_compare(
    period: str,
    against_period: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    """M9: Period-over-period deltas from leadership snapshots."""
    _require_phase1_access(current_user, "kpi_read")
    return await compare_snapshots(db, period=period.strip(), against_period=(against_period or "").strip() or None)


@api_router.get("/executive/m9/predictive-views")
async def m9_predictive_views(
    horizon_months: int = 3,
    window_days: int = 30,
    department: Optional[str] = None,
    manager_root_id: Optional[str] = None,
    role_title_contains: Optional[str] = None,
    trends_months: int = 12,
    current_user: dict = Depends(get_current_user),
):
    """M9: Attrition trend forecast + M8-based retention risk projections."""
    _require_phase1_access(current_user, "kpi_read")
    return await get_executive_predictive_views(
        db,
        horizon_months=horizon_months,
        window_days=window_days,
        department=department,
        manager_root_id=manager_root_id,
        role_title_contains=role_title_contains,
        trends_months=trends_months,
    )


@api_router.get("/executive/m9/trends")
async def m9_kpi_trends(
    months: int = 12,
    department: Optional[str] = None,
    manager_root_id: Optional[str] = None,
    role_title_contains: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    """M9: Historical KPI series from leadership snapshots."""
    _require_phase1_access(current_user, "kpi_read")
    return await get_kpi_trends(
        db,
        months=months,
        department=department,
        manager_root_id=manager_root_id,
        role_title_contains=role_title_contains,
    )


@api_router.post("/executive/m9/export-packs/monthly-snapshot")
async def m9_create_monthly_snapshot(
    payload: M9MonthlySnapshotCreate,
    current_user: dict = Depends(get_current_user),
):
    """M9-3: Persist a monthly leadership snapshot (JSON payload + downloadable CSV/PDF)."""
    _require_phase1_access(current_user, "kpi_read")
    try:
        return await create_monthly_snapshot_and_deliver(
            db,
            year_month=payload.period,
            horizon_months=payload.horizon_months,
            window_days=payload.window_days,
            department=payload.department,
            manager_root_id=payload.manager_root_id,
            role_title_contains=payload.role_title_contains,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api_router.post("/executive/m9/export-packs/full-leadership-pack")
async def m9_download_full_leadership_pack_zip(
    payload: M9MonthlySnapshotCreate,
    current_user: dict = Depends(get_current_user),
):
    """
    Week 11: One-click ZIP — persisted snapshot + JSON + CSV + PDF (PDF requires fpdf2).
    """
    _require_phase1_access(current_user, "kpi_read")
    try:
        out = await create_full_leadership_pack_zip(
            db,
            year_month=payload.period,
            horizon_months=payload.horizon_months,
            window_days=payload.window_days,
            department=payload.department,
            manager_root_id=payload.manager_root_id,
            role_title_contains=payload.role_title_contains,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    period = str(out.get("period") or "snapshot").replace("/", "-")
    fname = f"m9-full-leadership-pack-{period}.zip"
    return StreamingResponse(
        io.BytesIO(out["zip_bytes"]),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@api_router.post("/executive/m9/export-packs/monthly-snapshot-cron")
async def m9_monthly_snapshot_cron(request: Request):
    """
    M9-3: Scheduled monthly snapshot (no JWT). Requires env M9_SNAPSHOT_TOKEN and header
    X-M9-Snapshot-Token. Optional JSON body: { "period": "YYYY-MM", "horizon_months", "window_days" }.
    If period omitted, uses the previous calendar month (UTC).
    """
    expected = (os.environ.get("M9_SNAPSHOT_TOKEN") or "").strip()
    got = (request.headers.get("X-M9-Snapshot-Token") or "").strip()
    if not expected or got != expected:
        raise HTTPException(status_code=401, detail="Invalid or missing X-M9-Snapshot-Token")
    body: Dict[str, Any] = {}
    try:
        b = await request.json()
        if isinstance(b, dict):
            body = b
    except Exception:
        pass
    period_raw = body.get("period")
    period = str(period_raw).strip() if period_raw is not None else None
    if period == "":
        period = None
    try:
        hm = max(1, int(body.get("horizon_months") or 3))
        wd = max(1, min(int(body.get("window_days") or 30), 365))
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="horizon_months and window_days must be integers")
    scoped_limit = body.get("scoped_department_limit")
    try:
        return await create_monthly_cron_snapshots(
            db,
            year_month=period,
            horizon_months=hm,
            window_days=wd,
            scoped_department_limit=int(scoped_limit) if scoped_limit is not None else None,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@api_router.get("/executive/m9/export-packs")
async def m9_list_export_packs(limit: int = 24, current_user: dict = Depends(get_current_user)):
    _require_phase1_access(current_user, "kpi_read")
    return {"items": await list_snapshots(db, limit=limit)}


@api_router.get("/executive/m9/export-packs/{snapshot_id}/download")
async def m9_download_export_pack(
    snapshot_id: str,
    export_format: Literal["csv", "pdf", "json"] = Query("csv", alias="format"),
    current_user: dict = Depends(get_current_user),
):
    _require_phase1_access(current_user, "kpi_read")
    doc = await get_snapshot_doc(db, snapshot_id)
    if not doc:
        raise HTTPException(status_code=404, detail="snapshot not found")
    payload = doc.get("payload") or {}
    if export_format == "json":
        raw = json.dumps(payload, indent=2, default=str).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(raw),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="m9-snapshot-{snapshot_id}.json"'},
        )
    if export_format == "csv":
        raw = format_snapshot_csv(payload)
        return StreamingResponse(
            io.BytesIO(raw),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="m9-snapshot-{snapshot_id}.csv"'},
        )
    try:
        raw = format_snapshot_pdf(payload)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e))
    return StreamingResponse(
        io.BytesIO(raw),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="m9-snapshot-{snapshot_id}.pdf"'},
    )


@api_router.post("/executive/m9/export-packs/{snapshot_id}/deliver")
async def m9_deliver_export_pack(
    snapshot_id: str,
    current_user: dict = Depends(get_current_user),
    webhook_url: Optional[str] = None,
):
    """M9-3: Re-fire scheduled delivery hook (webhook)."""
    _require_phase1_access(current_user, "kpi_read")
    doc = await get_snapshot_doc(db, snapshot_id)
    if not doc:
        raise HTTPException(status_code=404, detail="snapshot not found")
    payload = doc.get("payload") or {}
    return await deliver_snapshot_webhook(payload, webhook_url=webhook_url)


# ========================
# M10: Event backbone (admin)
# ========================


@api_router.post("/admin/m10-events/replay")
async def admin_m10_events_replay(
    body: M10ReplayRequest,
    clear_idempotency: bool = Query(False),
    current_user: dict = Depends(get_current_user),
):
    """
    Reset matching events to PENDING for consumer replay.
    If clear_idempotency=true, removes consumer idempotency rows for those events (use after fixing handlers).
    """
    _require_admin(current_user)
    return await replay_events(db, body, clear_idempotency=clear_idempotency)


@api_router.get("/admin/m10-events/stats")
async def admin_m10_events_stats(current_user: dict = Depends(get_current_user)):
    """Outbox depth for ops dashboards."""
    _require_admin(current_user)
    pending = await db[COL_M10_EVENTS].count_documents({"status": "PENDING"})
    processing = await db[COL_M10_EVENTS].count_documents({"status": "PROCESSING"})
    done = await db[COL_M10_EVENTS].count_documents({"status": "DONE"})
    failed = await db[COL_M10_EVENTS].count_documents({"status": "FAILED"})
    idem = await db[COL_M10_IDEMPOTENCY].count_documents({})
    audit = await db[COL_M10_HANDLER_AUDIT].count_documents({})
    return {
        "outbox": {"PENDING": pending, "PROCESSING": processing, "DONE": done, "FAILED": failed},
        "idempotency_records": idem,
        "handler_audit_rows": audit,
    }


# ========================
# HEALTH CHECK
# ========================

# ========================
# Interview Proposal (HR Approval Scheduling) - M1
# ========================

def _next_business_day_utc(from_dt: Optional[datetime] = None) -> datetime:
    dt = from_dt or datetime.now(timezone.utc)
    dt_date = dt.date()
    while True:
        dt_date = dt_date + timedelta(days=1)
        if dt_date.weekday() < 5:  # Mon-Fri
            return datetime(dt_date.year, dt_date.month, dt_date.day, tzinfo=timezone.utc)

def _compute_default_proposed_slots_utc() -> List[InterviewProposedSlot]:
    base = _next_business_day_utc()
    # 2 slots: 11:00 and 15:00 UTC. Each slot is 60 minutes.
    slot_defs = [(11, 0), (15, 0)]
    slots: List[InterviewProposedSlot] = []
    for idx, (hh, mm) in enumerate(slot_defs):
        start = base.replace(hour=hh, minute=mm)
        end = start + timedelta(minutes=60)
        slots.append(
            InterviewProposedSlot(
                slot_index=idx,
                scheduled_start=start.isoformat(),
                scheduled_end=end.isoformat(),
            )
        )
    return slots

async def generate_interview_proposals_for_top_matches(
    job_id: str,
    top_matches: List[Dict[str, Any]],
    created_by: str,
    top_proposals: int = 20,
    mode: str = "VIRTUAL",
) -> int:
    if not top_matches:
        return 0
    proposed_slots = _compute_default_proposed_slots_utc()
    created = 0

    for match in top_matches[:top_proposals]:
        candidate = match.get("candidate") or {}
        candidate_id = candidate.get("id")
        if not candidate_id:
            continue

        existing = await db.interview_proposals.find_one(
            {"job_id": job_id, "candidate_id": candidate_id},
            {"_id": 0},
        )
        if existing:
            continue

        now = datetime.now(timezone.utc).isoformat()
        proposal_id = str(uuid.uuid4())
        doc = {
            "id": proposal_id,
            "job_id": job_id,
            "candidate_id": candidate_id,
            "round": 1,
            "mode": mode,
            "status": "PENDING",
            "proposed_slots": [s.model_dump() for s in proposed_slots],
            "application_id": None,
            "interview_id": None,
            "approved_by": None,
            "approved_at": None,
            "rejected_reason": None,
            "created_by": created_by,
            "created_at": now,
            "updated_at": None,
        }
        await db.interview_proposals.insert_one(doc)
        created += 1
    return created

def _require_hr_approver(current_user: dict):
    role = (current_user.get("role") or "").lower()
    if role not in {"admin", "recruiter", "hr_admin"}:
        raise HTTPException(status_code=403, detail="HR approval only")
    return current_user

@api_router.get("/jobs/{job_id}/interview-proposals", response_model=List[InterviewProposalResponse])
async def list_interview_proposals(job_id: str, current_user: dict = Depends(get_current_user)):
    proposals = (
        await db.interview_proposals.find({"job_id": job_id}, {"_id": 0})
        .sort("created_at", -1)
        .to_list(200)
    )
    out: List[InterviewProposalResponse] = []
    for p in proposals:
        candidate = await db.candidates.find_one({"id": p["candidate_id"]}, {"_id": 0})
        job = await db.jobs.find_one({"id": p["job_id"]}, {"_id": 0})
        out.append(
            InterviewProposalResponse(
                **p,
                candidate=candidate,
                job={"id": job["id"], "title": job.get("title")} if job else None,
            )
        )
    return out

@api_router.post("/interview-proposals/{proposal_id}/approve", response_model=InterviewProposalResponse)
async def approve_interview_proposal(
    proposal_id: str,
    payload: InterviewProposalApproveRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    _require_hr_approver(current_user)
    proposal = await db.interview_proposals.find_one({"id": proposal_id}, {"_id": 0})
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if proposal.get("status") != "PENDING":
        raise HTTPException(status_code=400, detail="Proposal is not pending")

    slot_index = int(payload.slot_index or 0)
    proposed_slots = proposal.get("proposed_slots") or []
    if slot_index < 0 or slot_index >= len(proposed_slots):
        raise HTTPException(status_code=400, detail="Invalid slot_index")

    chosen = proposed_slots[slot_index]
    scheduled_start = chosen["scheduled_start"]
    scheduled_end = chosen["scheduled_end"]

    try:
        raw_st = str(scheduled_start).replace("Z", "+00:00")
        st_dt = datetime.fromisoformat(raw_st)
        if st_dt.tzinfo is None:
            st_dt = st_dt.replace(tzinfo=timezone.utc)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid proposed slot start time")
    if st_dt < datetime.now(timezone.utc):
        raise HTTPException(status_code=400, detail="Cannot approve a slot in the past")

    job = await db.jobs.find_one({"id": proposal["job_id"]}, {"_id": 0})
    candidate = await db.candidates.find_one({"id": proposal["candidate_id"]}, {"_id": 0})
    if not job or not candidate:
        raise HTTPException(status_code=404, detail="Job/candidate not found")

    existing_app = await db.applications.find_one(
        {"job_id": proposal["job_id"], "candidate_id": proposal["candidate_id"]},
        {"_id": 0},
    )
    now = datetime.now(timezone.utc).isoformat()

    if existing_app:
        application_id = existing_app["id"]
        await db.applications.update_one(
            {"id": application_id},
            {"$set": {"stage": "INTERVIEW_1", "updated_at": now}},
        )
    else:
        fit_result = await compute_fit_score(job, candidate)
        application_id = str(uuid.uuid4())
        fit_score_doc = {
            "id": str(uuid.uuid4()),
            "job_id": proposal["job_id"],
            "candidate_id": proposal["candidate_id"],
            **fit_result,
            "computed_at": now,
        }
        await db.fit_scores.insert_one(fit_score_doc)
        app_doc = {
            "id": application_id,
            "job_id": proposal["job_id"],
            "candidate_id": proposal["candidate_id"],
            "stage": "INTERVIEW_1",
            "status": "ACTIVE",
            "fit_score_id": fit_score_doc["id"],
            "created_at": now,
            "updated_at": now,
        }
        await db.applications.insert_one(app_doc)
        await db.application_stage_history.insert_one(
            {
                "id": str(uuid.uuid4()),
                "application_id": application_id,
                "from_stage": None,
                "to_stage": "INTERVIEW_1",
                "changed_by": current_user["id"],
                "changed_at": now,
            }
        )

    invalidate_hiring_pack_cache(reason="interview_proposal_accepted")

    interview_id = str(uuid.uuid4())
    interview_doc = {
        "id": interview_id,
        "application_id": application_id,
        "candidate_id": proposal["candidate_id"],
        "job_id": proposal["job_id"],
        "round": proposal.get("round") or 1,
        "mode": proposal.get("mode") or "VIRTUAL",
        "scheduled_start": scheduled_start,
        "scheduled_end": scheduled_end,
        "meeting_link": payload.meeting_link,
        "status": "SCHEDULED",
        "interviewers": payload.interviewers or [],
        "notes": payload.notes,
        "feedback": [],
        "created_by": current_user["id"],
        "created_at": now,
        "calendar_provider": os.environ.get("CALENDAR_PROVIDER", "WEBHOOK"),
        "calendar_sync_status": "PENDING",
        "calendar_last_sync": None,
        "calendar_sync_detail": None,
        "reminder_sent_at": None,
    }
    await db.interviews.insert_one(interview_doc)

    async def _calendar_followup(iid: str):
        doc = await db.interviews.find_one({"id": iid}, {"_id": 0})
        if not doc:
            return
        res = await sync_interview_calendar_event(doc)
        await db.interviews.update_one(
            {"id": iid},
            {
                "$set": {
                    "calendar_sync_status": res.get("status"),
                    "calendar_last_sync": datetime.now(timezone.utc).isoformat(),
                    "calendar_sync_detail": res,
                }
            },
        )

    background_tasks.add_task(_calendar_followup, interview_id)

    if payload.interviewers and candidate and job:
        background_tasks.add_task(
            notify_interview_scheduled,
            interview_id,
            candidate.get("full_name"),
            job.get("title"),
            scheduled_start,
            payload.interviewers,
        )

    await db.interview_proposals.update_one(
        {"id": proposal_id},
        {
            "$set": {
                "status": "APPROVED",
                "approved_by": current_user["id"],
                "approved_at": now,
                "application_id": application_id,
                "interview_id": interview_id,
                "updated_at": now,
                "rejected_reason": None,
            }
        },
    )

    updated = await db.interview_proposals.find_one({"id": proposal_id}, {"_id": 0})
    candidate_full = await db.candidates.find_one({"id": updated["candidate_id"]}, {"_id": 0})
    job_full = await db.jobs.find_one({"id": updated["job_id"]}, {"_id": 0})
    return InterviewProposalResponse(
        **updated,
        candidate=candidate_full,
        job={"id": job_full["id"], "title": job_full.get("title")} if job_full else None,
    )

@api_router.post("/interview-proposals/{proposal_id}/reject", response_model=InterviewProposalResponse)
async def reject_interview_proposal(
    proposal_id: str,
    payload: InterviewProposalRejectRequest,
    current_user: dict = Depends(get_current_user),
):
    _require_hr_approver(current_user)
    proposal = await db.interview_proposals.find_one({"id": proposal_id}, {"_id": 0})
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    if proposal.get("status") != "PENDING":
        raise HTTPException(status_code=400, detail="Proposal is not pending")

    now = datetime.now(timezone.utc).isoformat()
    await db.interview_proposals.update_one(
        {"id": proposal_id},
        {
            "$set": {
                "status": "REJECTED",
                "approved_by": None,
                "approved_at": None,
                "rejected_reason": payload.reason or "Rejected by HR",
                "updated_at": now,
            }
        },
    )

    updated = await db.interview_proposals.find_one({"id": proposal_id}, {"_id": 0})
    candidate_full = await db.candidates.find_one({"id": updated["candidate_id"]}, {"_id": 0})
    job_full = await db.jobs.find_one({"id": updated["job_id"]}, {"_id": 0})
    return InterviewProposalResponse(
        **updated,
        candidate=candidate_full,
        job={"id": job_full["id"], "title": job_full.get("title")} if job_full else None,
    )


@api_router.post("/admin/interviews/dispatch-reminders")
async def admin_dispatch_interview_reminders(current_user: dict = Depends(get_current_user)):
    """
    M1-5: send 24h (or same-day window) reminders for scheduled interviews. Intended for cron / workflow_dispatch.
    """
    _require_admin(current_user)
    now = datetime.now(timezone.utc)
    window_end = now + timedelta(hours=24)
    docs = await db.interviews.find(
        {"status": "SCHEDULED", "reminder_sent_at": None},
        {"_id": 0},
    ).to_list(400)
    sent = 0
    for inv in docs:
        try:
            st_raw = str(inv.get("scheduled_start", "")).replace("Z", "+00:00")
            st = datetime.fromisoformat(st_raw)
            if st.tzinfo is None:
                st = st.replace(tzinfo=timezone.utc)
        except Exception:
            continue
        if not (now <= st <= window_end):
            continue
        iviewers = inv.get("interviewers") or []
        if not iviewers:
            continue
        cand = await db.candidates.find_one({"id": inv.get("candidate_id")}, {"_id": 0})
        job = await db.jobs.find_one({"id": inv.get("job_id")}, {"_id": 0})
        await notify_interview_reminder(
            inv["id"],
            (cand or {}).get("full_name") or "Candidate",
            (job or {}).get("title") or "Role",
            inv.get("scheduled_start") or "",
            iviewers,
        )
        await db.interviews.update_one(
            {"id": inv["id"]},
            {"$set": {"reminder_sent_at": now.isoformat()}},
        )
        sent += 1
    return {"dispatched": sent, "window_end": window_end.isoformat()}


# ========================
# INTERVIEW ROUTES
# ========================

@api_router.post("/interviews", response_model=InterviewResponse)
async def create_interview(
    interview_data: InterviewCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Schedule a new interview"""
    # Verify application exists
    application = await db.applications.find_one({"id": interview_data.application_id})
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    
    candidate = await db.candidates.find_one({"id": application["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": application["job_id"]}, {"_id": 0})
    
    interview_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    interview_doc = {
        "id": interview_id,
        "application_id": interview_data.application_id,
        "candidate_id": application["candidate_id"],
        "job_id": application["job_id"],
        "round": interview_data.round,
        "mode": interview_data.mode,
        "scheduled_start": interview_data.scheduled_start,
        "scheduled_end": interview_data.scheduled_end,
        "meeting_link": interview_data.meeting_link,
        "status": "SCHEDULED",
        "interviewers": interview_data.interviewers,
        "notes": interview_data.notes,
        "feedback": [],
        "created_by": current_user["id"],
        "created_at": now
    }
    await db.interviews.insert_one(interview_doc)
    
    # Send notifications to interviewers
    if candidate and job and interview_data.interviewers:
        background_tasks.add_task(
            notify_interview_scheduled,
            interview_id,
            candidate["full_name"],
            job["title"],
            interview_data.scheduled_start,
            interview_data.interviewers
        )
    
    return InterviewResponse(
        **interview_doc,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]} if job else None
    )

@api_router.get("/interviews", response_model=List[InterviewResponse])
async def list_interviews(
    application_id: Optional[str] = None,
    candidate_id: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """List interviews with optional filters"""
    query = {}
    if application_id:
        query["application_id"] = application_id
    if candidate_id:
        query["candidate_id"] = candidate_id
    if status:
        query["status"] = status
    
    interviews = await db.interviews.find(query, {"_id": 0}).sort("scheduled_start", -1).to_list(200)
    
    result = []
    for interview in interviews:
        candidate = await db.candidates.find_one({"id": interview["candidate_id"]}, {"_id": 0})
        job = await db.jobs.find_one({"id": interview["job_id"]}, {"_id": 0})
        result.append(InterviewResponse(
            **interview,
            candidate=candidate,
            job={"id": job["id"], "title": job["title"]} if job else None
        ))
    
    return result

@api_router.get("/interviews/{interview_id}", response_model=InterviewResponse)
async def get_interview(interview_id: str, current_user: dict = Depends(get_current_user)):
    """Get interview details"""
    interview = await db.interviews.find_one({"id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    candidate = await db.candidates.find_one({"id": interview["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": interview["job_id"]}, {"_id": 0})
    
    return InterviewResponse(
        **interview,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]} if job else None
    )

@api_router.put("/interviews/{interview_id}", response_model=InterviewResponse)
async def update_interview(
    interview_id: str,
    interview_data: InterviewUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update interview details"""
    interview = await db.interviews.find_one({"id": interview_id})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    update_doc = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if interview_data.scheduled_start:
        update_doc["scheduled_start"] = interview_data.scheduled_start
    if interview_data.scheduled_end:
        update_doc["scheduled_end"] = interview_data.scheduled_end
    if interview_data.meeting_link:
        update_doc["meeting_link"] = interview_data.meeting_link
    if interview_data.status:
        update_doc["status"] = interview_data.status
    if interview_data.notes:
        update_doc["notes"] = interview_data.notes
    
    await db.interviews.update_one({"id": interview_id}, {"$set": update_doc})
    
    updated = await db.interviews.find_one({"id": interview_id}, {"_id": 0})
    candidate = await db.candidates.find_one({"id": updated["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": updated["job_id"]}, {"_id": 0})
    
    return InterviewResponse(
        **updated,
        candidate=candidate,
        job={"id": job["id"], "title": job["title"]} if job else None
    )

@api_router.post("/interviews/{interview_id}/feedback")
async def add_interview_feedback(
    interview_id: str,
    feedback_data: InterviewFeedbackCreate,
    current_user: dict = Depends(get_current_user)
):
    """Add feedback for an interview"""
    interview = await db.interviews.find_one({"id": interview_id})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    feedback_doc = {
        "id": str(uuid.uuid4()),
        "reviewer_id": current_user["id"],
        "reviewer_name": current_user["full_name"],
        "decision": feedback_data.decision,
        "score": feedback_data.score,
        "strengths": feedback_data.strengths,
        "concerns": feedback_data.concerns,
        "notes": feedback_data.notes,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.interviews.update_one(
        {"id": interview_id},
        {"$push": {"feedback": feedback_doc}, "$set": {"status": "COMPLETED"}}
    )
    
    return {"message": "Feedback added successfully", "feedback": feedback_doc}

@api_router.delete("/interviews/{interview_id}")
async def cancel_interview(interview_id: str, current_user: dict = Depends(get_current_user)):
    """Cancel an interview"""
    interview = await db.interviews.find_one({"id": interview_id})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    await db.interviews.update_one(
        {"id": interview_id},
        {"$set": {"status": "CANCELLED", "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return {"message": "Interview cancelled"}

@api_router.get("/interviews/upcoming/me", response_model=List[InterviewResponse])
async def get_my_upcoming_interviews(current_user: dict = Depends(get_current_user)):
    """Get upcoming interviews for the current user"""
    now = datetime.now(timezone.utc).isoformat()
    interviews = await db.interviews.find(
        {
            "interviewers": current_user["id"],
            "status": "SCHEDULED",
            "scheduled_start": {"$gte": now}
        },
        {"_id": 0}
    ).sort("scheduled_start", 1).to_list(50)
    
    result = []
    for interview in interviews:
        candidate = await db.candidates.find_one({"id": interview["candidate_id"]}, {"_id": 0})
        job = await db.jobs.find_one({"id": interview["job_id"]}, {"_id": 0})
        result.append(InterviewResponse(
            **interview,
            candidate=candidate,
            job={"id": job["id"], "title": job["title"]} if job else None
        ))
    
    return result

# ========================
# NOTIFICATION ROUTES
# ========================

@api_router.get("/notifications", response_model=List[NotificationResponse])
async def list_notifications(
    unread_only: bool = False,
    current_user: dict = Depends(get_current_user)
):
    """Get notifications for the current user"""
    query = {"recipient_id": current_user["id"]}
    if unread_only:
        query["read"] = False
    
    notifications = await db.notifications.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    return [NotificationResponse(**n) for n in notifications]

@api_router.put("/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, current_user: dict = Depends(get_current_user)):
    """Mark a notification as read"""
    result = await db.notifications.update_one(
        {"id": notification_id, "recipient_id": current_user["id"]},
        {"$set": {"read": True}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Notification not found")
    return {"message": "Notification marked as read"}

@api_router.put("/notifications/read-all")
async def mark_all_notifications_read(current_user: dict = Depends(get_current_user)):
    """Mark all notifications as read"""
    await db.notifications.update_many(
        {"recipient_id": current_user["id"], "read": False},
        {"$set": {"read": True}}
    )
    return {"message": "All notifications marked as read"}

@api_router.get("/notifications/unread-count")
async def get_unread_notification_count(current_user: dict = Depends(get_current_user)):
    """Get count of unread notifications"""
    count = await db.notifications.count_documents(
        {"recipient_id": current_user["id"], "read": False}
    )
    return {"unread_count": count}

@api_router.get("/")
async def root():
    return {"message": "AAI-HRMS API is running", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc).isoformat()}

@api_router.get("/metrics")
async def metrics_snapshot(current_user: dict = Depends(get_current_user)):
    """Admin JSON snapshot of in-memory counters. For Prometheus use GET /metrics (no JWT)."""
    _require_admin(current_user)
    by_path = {}
    for path, stats in API_METRICS["by_path"].items():
        count = stats.get("count", 0)
        avg_ms = (stats.get("total_ms", 0.0) / count) if count else 0.0
        by_path[path] = {
            "count": count,
            "errors": stats.get("errors", 0),
            "avg_ms": round(avg_ms, 2),
        }

    return {
        "started_at": API_METRICS["started_at"],
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "total_requests": API_METRICS["total_requests"],
        "total_errors": API_METRICS["total_errors"],
        "by_path": by_path,
    }

@api_router.get("/transformation/modules")
async def get_transformation_modules(current_user: dict = Depends(get_current_user)):
    """
    BRD/SRS transformation roadmap for the UI (`/transformation`).
    Each module includes BRD/SRS bullets plus `achieved` (as-shipped) and remaining `gap`.
    Source of truth for narrative: `memory/TRANSFORMATION_BRD_SRS_PAGE.md` (kept in sync conceptually).
    """
    modules = [
        {
            "id": "M1",
            "name": "Talent Acquisition (Smart Hiring)",
            "current_state": "End-to-end hiring loop is live: jobs, candidates, pipeline, interviews, referrals, assessments, and admin connector configs.",
            "achieved": [
                "Multi-source ingestion (LinkedIn / Naukri / Monster) with OAuth refresh, paging, throttling, and connector health in Admin Integrations.",
                "Unified candidate store with dedup keys (email, phone, resume hash, name), merge API, and dedup audit trail.",
                "JD ↔ candidate matching with deterministic + optional LLM scoring and ranking explainability in UI.",
                "Hiring pipeline (kanban), interview proposals (HR approve/reject), optional calendar webhook sync, interview reminders dispatch.",
            ],
            "gap": "Production-grade external contracts (SLAs, rate limits per tenant), deeper Indeed/Glassdoor coverage, and full calendar provider packs beyond webhook bridge.",
            "target_state": "Unified talent intelligence platform with multi-source ranking and automation.",
            "brd": [
                "Centralized talent intelligence platform",
                "Seamless sourcing across portals and internal DB",
                "Reduced hiring time/cost with improved quality of hire",
            ],
            "srs_functional": [
                "Portal connectors (LinkedIn, Naukri, Monster, etc.)",
                "Central candidate repository",
                "AI ranking + duplicate profile detection",
                "Automated interview scheduling",
            ],
            "srs_non_functional": ["High API reliability", "Scalable ingestion", "Near real-time processing"],
            "status": "in_progress",
            "priority": "P0",
        },
        {
            "id": "M2",
            "name": "Employee Lifecycle Management",
            "current_state": "Employee master, org hierarchy, lifecycle events with approvals, compliance documents, and audit trail are implemented.",
            "achieved": [
                "Employee master (CRUD, import, paged APIs, LCD50 demo seed) with legal status / transition validation on status changes.",
                "Lifecycle events (create/update/delete), approval matrix (EXITED / ROLE_CHANGED), approve/reject, background processing, reprocess queue.",
                "Org APIs: direct reports, management chain, org hierarchy; immutable-style lifecycle audit log.",
                "Compliance documents: upload/verify, SLA due dates, breach + reminder admin scans, CSV export.",
            ],
            "gap": "Deep HRIS payroll/benefits integrations, bulk document workflows, and mobile-first employee self-service.",
            "target_state": "End-to-end employee lifecycle from onboarding to exit.",
            "brd": ["Central employee repository", "Lifecycle visibility and governance"],
            "srs_functional": ["Employee master", "Onboard-active-exit workflows", "Role/hierarchy mapping", "Document management"],
            "srs_non_functional": ["Secure data handling", "Role-based access control"],
            "status": "in_progress",
            "priority": "P0",
        },
        {
            "id": "M3",
            "name": "Workforce Intelligence (Demand-Supply)",
            "current_state": "Demand/supply views, historical feature store, baseline forecasting model, and monitoring hooks are available.",
            "achieved": [
                "Workforce Intelligence UI + APIs for demand vs supply and gap analytics tied to projects, skills, and allocations.",
                "ETL snapshot pipeline with DQ gates, synthetic backfill for demos, and stored feature history.",
                "Per-skill baseline (OLS) forecast train/serve, model registry with activate/rollback, drift evaluation and retrain policy signals.",
            ],
            "gap": "Richer ML beyond linear baseline, automated scheduled ETL in all deployments without admin setup, and executive narrative packs for workforce risk.",
            "target_state": "Predictive workforce planning with demand-supply balancing.",
            "brd": ["AI-driven workforce planning", "Skill-demand forecasting for proactive decisions"],
            "srs_functional": ["Skill inventory", "Demand forecasting", "Supply mapping", "Skill gap analytics"],
            "srs_non_functional": ["ML scalability", "Near real-time predictions"],
            "status": "in_progress",
            "priority": "P1",
        },
        {
            "id": "M4",
            "name": "Resource vs Project Optimization",
            "current_state": "Project skill demands, allocations, greedy solver, scenario lab, and approval-to-apply workflow are live.",
            "achieved": [
                "Project demands and allocations data model with min/max seats and HARD/SOFT constraints.",
                "Constraint-based solver with explain steps, tunable optimization settings API.",
                "Scenario simulate/save/compare, submit for approval, admin/hr approve or reject, apply-to-DB with notifications.",
            ],
            "gap": "Real-time utilization telemetry from timesheets, global bench dashboard, and MILP/CP-SAT solver options for larger portfolios.",
            "target_state": "Intelligent resource allocation and utilization governance.",
            "brd": ["Optimize workforce utilization and project outcomes"],
            "srs_functional": ["Project-resource mapping", "Utilization dashboards", "Over/under allocation alerts", "Bench tracking"],
            "srs_non_functional": ["Real-time updates", "High data accuracy"],
            "status": "in_progress",
            "priority": "P1",
        },
        {
            "id": "M5",
            "name": "Employee Training & Skill Development",
            "current_state": "Recommendations, assignments, LMS catalog sync (stub provider), certifications, and manager rollup exist.",
            "achieved": [
                "Skill-gap-driven training recommendations with templates and persisted assignments + progress PATCH.",
                "LMS adapter with stub provider, catalog sync runs, course catalog API filtered by skill.",
                "Certifications CRUD, expiry reminder scan, manager summary for direct reports (assignments + expiring certs).",
            ],
            "gap": "Live SCORM/xAPI LMS integrations, company-wide learning budgets, and automated enrollment into vendor systems.",
            "target_state": "AI-driven L&D platform for continuous upskilling.",
            "brd": ["Continuous skill enhancement with personalized development"],
            "srs_functional": ["Skill gap detection", "Personalized learning paths", "Training tracking", "Certification management"],
            "srs_non_functional": ["LMS integration", "Scalable recommendation engine"],
            "status": "in_progress",
            "priority": "P1",
        },
        {
            "id": "M6",
            "name": "Employee Satisfaction & Engagement",
            "current_state": "Pulse surveys, templates, schedules, sentiment/topics dashboard, and privacy-aware RBAC are shipped.",
            "achieved": [
                "Survey template CRUD with targeting; create pulse from template; schedules with admin dispatch of due runs.",
                "Deterministic sentiment + topic aggregation; engagement dashboard with confidence tiers and anonymity threshold.",
                "Raw response access limited to admin/hr_admin with privacy audit log; participation reminders.",
            ],
            "gap": "Multilingual NLP, continuous listening across channels (Slack/Teams), and benchmarked eNPS industry packs.",
            "target_state": "Employee experience platform with actionable engagement insights.",
            "brd": ["Improve engagement and retention through continuous listening"],
            "srs_functional": ["Pulse surveys", "Feedback management", "Sentiment engine", "Engagement dashboards"],
            "srs_non_functional": ["Privacy compliance", "Real-time analytics"],
            "status": "in_progress",
            "priority": "P2",
        },
        {
            "id": "M7",
            "name": "Cost Optimization & Automation",
            "current_state": "Workflow automation (multi-trigger, flow designer, webhooks), HR Copilot, and savings baselines are live.",
            "achieved": [
                "Workflow rules: lifecycle thresholds, schedules, inbound signed webhooks, HTTP outbound actions, React Flow designer UI.",
                "Execution history, retries, admin dispatch; lifecycle reprocess automation hook.",
                "HR Copilot chat with rule + optional HF NLI routing; conversation audit for admins.",
                "Cost baselines CRUD and executive savings estimates from successful automation runs.",
            ],
            "gap": "Enterprise-wide RPA/desktop automation, natural-language authoring of rules, and multi-tenant workflow templates marketplace.",
            "target_state": "Autonomous HR workflows with lower operating cost.",
            "brd": ["Reduce manual HR overhead with automation-first operations"],
            "srs_functional": ["HR chatbot", "Automated screening", "Auto scheduling", "Workflow automation engine"],
            "srs_non_functional": ["High availability", "Low-latency response"],
            "status": "in_progress",
            "priority": "P1",
        },
        {
            "id": "M8",
            "name": "High-Skill Talent Retention",
            "current_state": "Attrition risk scoring (linear + optional gradient boosting), explanations, and HRIS fields on employee records are implemented.",
            "achieved": [
                "Feature pipeline (tenure, engagement, comp pressure, training gaps) with train/score APIs and cron-friendly score-run endpoint.",
                "SHAP-style linear attributions; optional sklearn GBM with ensemble modes; latest scores persisted for UI/API.",
                "Employee Master captures comp band, promotion history, high-performer/critical-role flags, market percentile for model input.",
            ],
            "gap": "Calibrated production labels loop, manager nudges in flow of work, and integration to compensation planning tools.",
            "target_state": "Talent intelligence focused on critical talent retention.",
            "brd": ["Retain high-value workforce segments"],
            "srs_functional": ["High performer detection", "Attrition prediction", "Career path recommendations"],
            "srs_non_functional": ["Target ML accuracy threshold", "Secure insights access"],
            "status": "in_progress",
            "priority": "P2",
        },
        {
            "id": "M9",
            "name": "Analytics & Executive Dashboard",
            "current_state": "Executive KPIs, M9 semantic layer, drill-down scope, and leadership export packs (CSV/PDF/ZIP) are available.",
            "achieved": [
                "KPI catalog with ownership + formulas; merged definitions; KPI pack + freshness APIs; talent acquisition metrics slice.",
                "Strategic drill by department, manager subtree, or role title; cached drill dashboard API; Executive KPI page UX.",
                "Monthly leadership snapshots, full ZIP export, optional webhook delivery, cron snapshot route with shared-secret auth.",
            ],
            "gap": "Embedded analytics (Looker/PowerBI) and cross-tenant benchmarking datasets.",
            "target_state": "Executive decision cockpit with real-time and predictive insight.",
            "brd": ["Data-driven strategic HR decisions"],
            "srs_functional": ["KPI dashboards", "Drill-down analytics", "Predictive insight views"],
            "srs_non_functional": ["High performance dashboards", "Real-time refresh"],
            "status": "in_progress",
            "priority": "P0",
        },
        {
            "id": "M10",
            "name": "Architecture & Scalability",
            "current_state": "Monolith hosts modular packages; event outbox/consumer backbone and architecture blueprint docs exist.",
            "achieved": [
                "Documented bounded-context / gateway migration blueprint (M10_ARCHITECTURE_BLUEPRINT.md) and operator runbooks (DR, on-call, perf smoke).",
                "Versioned event envelope, Mongo outbox producer, in-process consumer with idempotency + admin replay and stats API.",
                "Lifecycle create and workflow run outcomes publish to M10 topics for downstream expansion.",
            ],
            "gap": "External message broker, dedicated consumer fleet, API gateway edge, and active-active multi-region data planes.",
            "target_state": "Enterprise-grade, cloud-native, fault-tolerant architecture.",
            "brd": ["Scale platform reliably for enterprise workload"],
            "srs_functional": ["Microservices", "API gateway", "Event-driven processing"],
            "srs_non_functional": ["Horizontal scalability", "Fault tolerance", "Cloud-native deployment"],
            "status": "in_progress",
            "priority": "P0",
        },
    ]

    return {
        "platform_vision": "AI-Powered Workforce Intelligence Platform",
        "business_impact": {
            "hiring": "Faster, cheaper, higher quality",
            "workforce": "Optimized utilization",
            "employees": "Higher satisfaction",
            "organization": "Data-driven decisions",
            "cost": "Reduced operational overhead",
        },
        "modules": modules,
    }

# Allocation Section (M10 staffing bridge) — modular router
api_router.include_router(
    create_allocation_section_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
        require_approve=_require_allocation_approver,
        assert_no_overallocation=_assert_no_overallocation,
    )
)
api_router.include_router(
    create_resource_section_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
        require_approve=_require_allocation_approver,
    )
)
api_router.include_router(
    create_training_development_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_high_skill_retention_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_career_trajectory_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_assessments_router(
        db=db,
        get_current_user=get_current_user,
        generate_with_ai=generate_assessment_with_ai,
        create_notification=create_notification,
        llm_chat=_llm_chat,
        require_admin=_require_admin,
    )
)
api_router.include_router(
    create_phase2_fit_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_employee_lifecycle_management_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_workforce_intelligence_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_cost_optimization_automation_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "kpi_read"),
        require_write=lambda u: _require_phase1_access(u, "skills_write"),
    )
)
api_router.include_router(
    create_employee_satisfaction_engagement_router(
        db=db,
        get_current_user=get_current_user,
        require_read=lambda u: _require_phase1_access(u, "engagement_read"),
        require_write=lambda u: _require_phase1_access(u, "engagement_write"),
        require_engagement_executive=_require_engagement_executive,
        require_engagement_ai=_require_engagement_ai,
    )
)

# Include router
app.include_router(api_router)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=[
        o.strip()
        for o in os.environ.get("CORS_ORIGINS", "*").split(",")
        if o.strip()
    ] or ["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

from pymongo.errors import OperationFailure


class _IndexSafeCollection:
    __slots__ = ("_coll",)

    def __init__(self, coll):
        self._coll = coll

    async def create_index(self, keys, **kwargs):
        try:
            return await self._coll.create_index(keys, **kwargs)
        except OperationFailure as exc:
            if exc.code in (85, 86):
                return None
            raise

    def __getattr__(self, name):
        return getattr(self._coll, name)


class _IndexSafeDb:
    __slots__ = ("_db",)

    def __init__(self, db_handle):
        self._db = db_handle

    def __getitem__(self, name):
        return _IndexSafeCollection(self._db[name])

    def __getattr__(self, name):
        val = getattr(self._db, name)
        if hasattr(val, "create_index"):
            return _IndexSafeCollection(val)
        return val


@app.on_event("startup")
async def ensure_phase1_indexes():
    # Data governance: enforce unique keys and query indexes for phase-1 modules.
    idx_db = _IndexSafeDb(db)
    await idx_db.employees.create_index("employee_code", unique=True, name="uq_employee_code")
    await idx_db.employees.create_index([("status", 1), ("department", 1)], name="ix_employee_status_dept")
    await idx_db.workforce_skills.create_index("skill_name_lc", unique=True, name="uq_skill_name_lc")
    await idx_db.workforce_skills.create_index([("gap", -1), ("priority", 1)], name="ix_skill_gap_priority")
    await idx_db.import_audit_logs.create_index([("module", 1), ("created_at", -1)], name="ix_import_audit_module_created")

    # M1 (Talent Acquisition) governance: candidate canonicalization + interview proposals.
    await idx_db.candidates.create_index("email_lc", name="ix_candidates_email_lc")
    await idx_db.candidates.create_index("full_name_lc", name="ix_candidates_full_name_lc")
    await idx_db.candidates.create_index("phone_lc", name="ix_candidates_phone_lc", sparse=True)
    await idx_db.candidates.create_index("resume_content_hash", name="ix_candidates_resume_hash", sparse=True)
    await idx_db.ingestion_jobs.create_index([("job_id", 1), ("created_at", -1)], name="ix_ingestion_jobs_job_created")
    await idx_db.candidate_dedup_audit.create_index(
        [("candidate_id", 1), ("created_at", -1)],
        name="ix_dedup_audit_candidate_created",
    )
    await idx_db.interview_proposals.create_index("id", unique=True, name="uq_interview_proposal_id")
    await idx_db.interview_proposals.create_index(
        [("job_id", 1), ("candidate_id", 1), ("status", 1)],
        name="ix_interview_proposals_job_candidate_status",
    )

    # M3/M4 (Project-demand MVP) governance: projects + project skill demands.
    await idx_db.projects.create_index("id", unique=True, name="uq_project_id")
    await idx_db.project_skill_demands.create_index(
        [("project_id", 1), ("skill_name_lc", 1)],
        unique=True,
        name="uq_project_skill_demand_project_skill",
    )
    await idx_db.project_skill_demands.create_index([("project_id", 1), ("updated_at", -1)], name="ix_project_skill_demands_project_updated")

    # M4 (Project allocations governance)
    await idx_db.project_skill_allocations.create_index("id", unique=True, name="uq_project_skill_allocation_id")
    await idx_db.project_skill_allocations.create_index(
        [("project_id", 1), ("skill_name_lc", 1)],
        unique=True,
        name="uq_project_skill_alloc_project_skill",
    )
    await idx_db.project_skill_allocations.create_index(
        [("project_id", 1), ("updated_at", -1)],
        name="ix_project_skill_allocations_project_updated",
    )

    # M6 (Project vs resource allocations)
    await idx_db.allocations.create_index("id", unique=True, name="uq_allocation_id")
    await idx_db.allocations.create_index([("project_id", 1), ("created_at", -1)], name="ix_allocations_project_created")
    await idx_db.allocations.create_index([("employee_id", 1), ("created_at", -1)], name="ix_allocations_employee_created")
    await idx_db.allocations.create_index(
        [("employee_id", 1), ("status", 1), ("approval_status", 1)],
        name="ix_allocations_employee_status_approval",
    )

    # M7 (Project Section) governance
    await idx_db.project_masters.create_index("id", unique=True, name="uq_project_master_id")
    await idx_db.project_masters.create_index("project_code", unique=True, name="uq_project_master_code")
    await idx_db.project_masters.create_index([("project_status", 1), ("project_priority", 1)], name="ix_project_master_status_priority")
    await idx_db.project_masters.create_index([("business_unit", 1)], name="ix_project_master_bu")
    await idx_db.project_masters.create_index([("client_name", 1)], name="ix_project_master_client")

    await idx_db.project_lifecycle_history.create_index([("project_id", 1), ("changed_at", -1)], name="ix_project_lifecycle_project_changed")
    await idx_db.project_demands.create_index(
        [("project_id", 1), ("role_name_lc", 1), ("skill_name_lc", 1)],
        unique=True,
        name="uq_project_demands_project_role_skill",
    )
    await idx_db.project_demands.create_index([("project_id", 1), ("updated_at", -1)], name="ix_project_demands_project_updated")
    await idx_db.project_risks.create_index([("project_id", 1), ("created_at", -1)], name="ix_project_risks_project_created")
    await idx_db.project_issues.create_index([("project_id", 1), ("created_at", -1)], name="ix_project_issues_project_created")
    await idx_db.project_documents.create_index([("project_id", 1), ("created_at", -1)], name="ix_project_docs_project_created")
    await idx_db.project_notes.create_index([("project_id", 1), ("created_at", -1)], name="ix_project_notes_project_created")
    await idx_db.project_approvals.create_index([("status", 1), ("requested_at", -1)], name="ix_project_approvals_status_requested")
    await idx_db.project_approvals.create_index([("project_id", 1), ("requested_at", -1)], name="ix_project_approvals_project_requested")
    await idx_db.project_closure.create_index("project_id", unique=True, name="uq_project_closure_project")
    await idx_db.project_wbs_items.create_index([("project_id", 1), ("order", 1)], name="ix_project_wbs_project_order")
    await idx_db.project_financials.create_index("project_id", unique=True, name="uq_project_financials_project")
    await idx_db.project_financial_snapshots.create_index([("project_id", 1), ("snapshot_at", -1)], name="ix_project_fin_snapshots_project_at")
    await idx_db.project_status_reports.create_index([("project_id", 1), ("created_at", -1)], name="ix_project_status_reports_project_at")

    # M3/M4 dashboard cache + audit telemetry.
    await idx_db.workforce_intelligence_forecast_cache.create_index(
        [("horizon_months", 1), ("demand_source", 1)],
        unique=True,
        name="uq_workforce_intelligence_cache_key",
    )
    await idx_db.workforce_resource_optimization_cache.create_index(
        [("demand_source", 1)],
        unique=True,
        name="uq_workforce_resource_optimization_cache_key",
    )
    await idx_db.workforce_dashboard_audit_logs.create_index(
        [("dashboard", 1), ("created_at", -1)],
        name="ix_workforce_dashboard_audit_dashboard_created",
    )
    await idx_db.workforce_dashboard_audit_logs.create_index([("cache_hit", 1)], name="ix_workforce_dashboard_audit_cache_hit")

    # M3 Workforce Intelligence: historical features + model registry + monitoring.
    await idx_db[COL_HIST_FEATURES].create_index("id", unique=True, name="uq_wf_intel_hist_row_id")
    await idx_db[COL_HIST_FEATURES].create_index(
        [("demand_source", 1), ("snapshot_at", -1), ("skill_name_lc", 1)],
        name="ix_wf_intel_hist_demand_snapshot_skill",
    )
    await idx_db[COL_HIST_FEATURES].create_index([("etl_run_id", 1)], name="ix_wf_intel_hist_etl_run")
    await idx_db[COL_ETL_RUNS].create_index([("ended_at", -1)], name="ix_wf_intel_etl_runs_ended")
    await idx_db[COL_MODELS].create_index([("version_id", 1)], unique=True, name="uq_wf_intel_model_version")
    await idx_db[COL_MODELS].create_index([("created_at", -1)], name="ix_wf_intel_models_created")
    await idx_db[COL_EVAL_RUNS].create_index([("evaluated_at", -1)], name="ix_wf_intel_eval_at")
    await idx_db[COL_DRIFT_EVENTS].create_index([("created_at", -1)], name="ix_wf_intel_drift_created")
    await idx_db[COL_DRIFT_EVENTS].create_index([("version_id", 1), ("created_at", -1)], name="ix_wf_intel_drift_version_created")

    # M4 Resource vs Project Optimization: scenarios + settings singleton.
    await idx_db[COL_ALLOCATION_SCENARIOS].create_index("id", unique=True, name="uq_m4_allocation_scenario_id")
    await idx_db[COL_ALLOCATION_SCENARIOS].create_index(
        [("status", 1), ("created_at", -1)],
        name="ix_m4_allocation_scenario_status_created",
    )

    # M5 Training & Skill Development
    await idx_db[COL_LEARNING_PATH_TEMPLATES].create_index(
        "skill_name_lc", unique=True, name="uq_training_path_template_skill_lc"
    )
    await idx_db[COL_ASSIGNMENTS].create_index("id", unique=True, name="uq_training_assignment_id")
    await idx_db[COL_ASSIGNMENTS].create_index([("employee_code", 1), ("status", 1)], name="ix_training_assign_emp_status")
    await idx_db[COL_LMS_COURSES].create_index(
        [("provider", 1), ("external_id", 1)],
        unique=True,
        name="uq_training_lms_course_provider_ext",
    )
    await idx_db[COL_LMS_COURSES].create_index([("skill_tags_lc", 1)], name="ix_training_lms_skill_tags")
    await idx_db[COL_LMS_SYNC_RUNS].create_index([("ended_at", -1)], name="ix_training_lms_sync_ended")
    await idx_db[COL_CERTIFICATIONS].create_index("id", unique=True, name="uq_training_cert_id")
    await idx_db[COL_CERTIFICATIONS].create_index([("employee_code", 1), ("expires_at", 1)], name="ix_training_cert_emp_exp")

    # Phase-2 (M2) governance: employee lifecycle events.
    await idx_db.employee_lifecycle_events.create_index("id", unique=True, name="uq_lifecycle_event_id")
    await idx_db.employee_lifecycle_events.create_index("employee_code", name="ix_lifecycle_employee_code")
    await idx_db.employee_lifecycle_events.create_index("event_type", name="ix_lifecycle_event_type")
    await idx_db.employee_lifecycle_events.create_index([("created_at", -1)], name="ix_lifecycle_created_at")
    await idx_db.employee_lifecycle_events.create_index(
        [("requires_approval", 1), ("approval_status", 1), ("created_at", -1)],
        name="ix_lifecycle_approval_pending",
    )

    await idx_db[LIFECYCLE_AUDIT_COLLECTION].create_index(
        [("employee_code", 1), ("created_at", -1)],
        name="ix_lifecycle_audit_emp_created",
    )
    await idx_db[COMPLIANCE_DOCS_COLLECTION].create_index("id", unique=True, name="uq_compliance_doc_id")
    await idx_db[COMPLIANCE_DOCS_COLLECTION].create_index(
        [("employee_code", 1), ("status", 1)],
        name="ix_compliance_emp_status",
    )
    await idx_db[COMPLIANCE_DOCS_COLLECTION].create_index([("status", 1), ("sla_due_at", 1)], name="ix_compliance_sla")

    # Phase-4 (M6) engagement governance: surveys + responses.
    await idx_db.employee_engagement_surveys.create_index("id", unique=True, name="uq_pulse_survey_id")
    await idx_db.employee_engagement_surveys.create_index([("active", 1), ("created_at", -1)], name="ix_pulse_survey_active_created")

    await idx_db.employee_engagement_responses.create_index("id", unique=True, name="uq_pulse_response_id")
    await idx_db.employee_engagement_responses.create_index([("survey_id", 1), ("created_at", -1)], name="ix_pulse_response_survey_created")
    await idx_db.employee_engagement_responses.create_index([("employee_code", 1), ("created_at", -1)], name="ix_pulse_response_employee_created")

    # M6: templates, schedules, privacy audit
    await idx_db[COL_SURVEY_TEMPLATES].create_index("id", unique=True, name="uq_engagement_survey_template_id")
    await idx_db[COL_SURVEY_TEMPLATES].create_index([("updated_at", -1)], name="ix_engagement_template_updated")
    await idx_db[COL_SURVEY_SCHEDULES].create_index("id", unique=True, name="uq_engagement_survey_schedule_id")
    await idx_db[COL_SURVEY_SCHEDULES].create_index([("enabled", 1), ("next_run_at", 1)], name="ix_engagement_schedule_due")
    await idx_db[COL_SURVEY_SCHEDULES].create_index("template_id", name="ix_engagement_schedule_template")
    await idx_db[COL_PRIVACY_AUDIT].create_index([("created_at", -1)], name="ix_engagement_privacy_audit_created")
    await idx_db[COL_PRIVACY_AUDIT].create_index([("survey_id", 1), ("created_at", -1)], name="ix_engagement_privacy_audit_survey")

    # M7 workflow automation, copilot audit, baselines
    await idx_db[COL_WORKFLOW_RULES].create_index("id", unique=True, name="uq_m7_workflow_rule_id")
    await idx_db[COL_WORKFLOW_RULES].create_index([("enabled", 1), ("updated_at", -1)], name="ix_m7_workflow_rule_enabled_updated")
    await idx_db[COL_WORKFLOW_RUNS].create_index("id", unique=True, name="uq_m7_workflow_run_id")
    await idx_db[COL_WORKFLOW_RUNS].create_index([("created_at", -1)], name="ix_m7_workflow_run_created")
    await idx_db[COL_WORKFLOW_RUNS].create_index([("rule_id", 1), ("created_at", -1)], name="ix_m7_workflow_run_rule_created")
    await idx_db[COL_WORKFLOW_RUNS].create_index([("status", 1), ("created_at", -1)], name="ix_m7_workflow_run_status_created")
    await idx_db[COL_HR_COPILOT_AUDIT].create_index([("created_at", -1)], name="ix_m7_copilot_audit_created")
    await idx_db[COL_HR_COPILOT_AUDIT].create_index([("session_id", 1), ("created_at", -1)], name="ix_m7_copilot_audit_session")
    await idx_db[COL_MANUAL_WORKFLOW_BASELINES].create_index("id", unique=True, name="uq_m7_manual_baseline_id")
    await idx_db[COL_MANUAL_WORKFLOW_BASELINES].create_index(
        "workflow_key", unique=True, name="uq_m7_manual_baseline_workflow_key"
    )

    # M8 retention: attrition v1 + interventions
    await idx_db[COL_ATTRITION_MODEL_STATE].create_index("id", unique=True, name="uq_m8_attrition_model_id")
    await idx_db[COL_ATTRITION_SCORES_LATEST].create_index("employee_id", unique=True, name="uq_m8_attrition_score_emp")
    await idx_db[COL_ATTRITION_SCORES_LATEST].create_index([("attrition_risk", -1)], name="ix_m8_attrition_risk")
    await idx_db[COL_ATTRITION_SCORES_LATEST].create_index([("risk_band", 1), ("department", 1)], name="ix_m8_attrition_band_dept")
    await idx_db[COL_ATTRITION_SCORES_LATEST].create_index("segments", name="ix_m8_attrition_segments")
    await idx_db[COL_RETENTION_SEGMENT_SETTINGS].create_index("id", unique=True, name="uq_m8_retention_segment_settings")
    await idx_db[COL_RETENTION_PLAYBOOKS].create_index("id", unique=True, name="uq_m8_retention_playbook_id")
    await idx_db[COL_RETENTION_INTERVENTIONS].create_index("id", unique=True, name="uq_m8_retention_intervention_id")
    await idx_db[COL_RETENTION_INTERVENTIONS].create_index(
        [("employee_id", 1), ("created_at", -1)],
        name="ix_m8_intervention_emp_created",
    )
    await idx_db[COL_RETENTION_INTERVENTIONS].create_index(
        [("status", 1), ("created_at", -1)],
        name="ix_m8_intervention_status_created",
    )

    # M9 analytics: KPI overrides + leadership snapshots
    await idx_db[COL_M9_KPI_DEFINITIONS].create_index("kpi_id", unique=True, name="uq_m9_kpi_definition_id")
    await idx_db[COL_M9_KPI_THRESHOLDS].create_index("kpi_id", unique=True, name="uq_m9_kpi_threshold_id")
    await idx_db[COL_M9_LEADERSHIP_SNAPSHOTS].create_index("id", unique=True, name="uq_m9_leadership_snapshot_id")
    await idx_db[COL_M9_LEADERSHIP_SNAPSHOTS].create_index(
        [("period", 1), ("created_at", -1)],
        name="ix_m9_snapshot_period_created",
    )
    await idx_db[COL_M9_LEADERSHIP_SNAPSHOTS].create_index(
        [("period", 1), ("snapshot_scope", 1)],
        unique=True,
        name="uq_m9_snapshot_period_scope",
        partialFilterExpression={"snapshot_scope": {"$exists": True, "$type": "string"}},
    )

    # M10 Allocation Section (staffing bridge — collections alongside M10 events module)
    await idx_db[COL_STAFFING_REQUESTS].create_index("id", unique=True, name="uq_alloc_sec_request_id")
    await idx_db[COL_STAFFING_REQUESTS].create_index([("project_id", 1), ("created_at", -1)], name="ix_alloc_sec_req_project")
    await idx_db[COL_STAFFING_REQUEST_HISTORY].create_index([("request_id", 1), ("at", -1)], name="ix_alloc_sec_req_hist_req")
    await idx_db[COL_CONFLICTS].create_index("id", unique=True, name="uq_alloc_sec_conflict_id")
    await idx_db[COL_CONFLICTS].create_index([("allocation_id", 1), ("resolution_status", 1)], name="ix_alloc_sec_conflict_alloc")
    await idx_db[COL_ROLL_EVENTS].create_index("id", unique=True, name="uq_alloc_sec_roll_id")
    await idx_db[COL_ROLL_EVENTS].create_index([("resource_id", 1), ("planned_rolloff_date", 1)], name="ix_alloc_sec_roll_emp")
    await idx_db[COL_CHANGES].create_index("id", unique=True, name="uq_alloc_sec_change_id")
    await idx_db[COL_CHANGES].create_index([("allocation_id", 1), ("changed_on", -1)], name="ix_alloc_sec_change_alloc")
    await idx_db[COL_RELEASES].create_index("id", unique=True, name="uq_alloc_sec_release_id")
    await idx_db[COL_WORKFLOW_APPROVALS].create_index("id", unique=True, name="uq_alloc_sec_wf_appr_id")
    await idx_db[COL_WORKFLOW_APPROVALS].create_index([("status", 1), ("submitted_at", -1)], name="ix_alloc_sec_wf_status")
    await idx_db[COL_BENCH_MATCHES].create_index("id", unique=True, name="uq_alloc_sec_bench_match_id")
    await idx_db[COL_NOTES].create_index("id", unique=True, name="uq_alloc_sec_note_id")
    await idx_db[COL_DOCUMENTS].create_index("id", unique=True, name="uq_alloc_sec_doc_id")
    await idx_db[COL_ALERTS].create_index("id", unique=True, name="uq_alloc_sec_alert_id")
    await idx_db[COL_ALERTS].create_index([("created_at", -1)], name="ix_alloc_sec_alert_created")
    await idx_db[COL_ACTIVITY_LOGS].create_index([("created_at", -1)], name="ix_alloc_sec_activity_created")
    await idx_db[COL_FORECAST_SNAPSHOTS].create_index("id", unique=True, name="uq_alloc_sec_forecast_id")
    await idx_db[COL_AI_INSIGHTS].create_index("id", unique=True, name="uq_alloc_sec_ai_insight_id")
    await idx_db[COL_AI_INSIGHTS].create_index([("generated_at", -1)], name="ix_alloc_sec_ai_gen")

    # M11 Resource Section (workforce deployability overlay)
    await idx_db[RS_COL_PROFILES].create_index("resource_id", unique=True, name="uq_res_sec_profile_resource")
    await idx_db[RS_COL_CLASSIFICATIONS].create_index("id", unique=True, name="uq_res_sec_class_id")
    await idx_db[RS_COL_CLASSIFICATIONS].create_index([("resource_id", 1), ("tag", 1)], name="ix_res_sec_class_resource_tag")
    await idx_db[RS_COL_SKILL_RECORDS].create_index("id", unique=True, name="uq_res_sec_skill_id")
    await idx_db[RS_COL_SKILL_RECORDS].create_index([("resource_id", 1), ("skill_name", 1)], name="ix_res_sec_skill_resource")
    await idx_db[RS_COL_AVAILABILITY].create_index("id", unique=True, name="uq_res_sec_avail_id")
    await idx_db[RS_COL_AVAILABILITY].create_index([("resource_id", 1), ("updated_on", -1)], name="ix_res_sec_avail_resource")
    await idx_db[RS_COL_UTIL_SNAPSHOTS].create_index("id", unique=True, name="uq_res_sec_util_id")
    await idx_db[RS_COL_UTIL_SNAPSHOTS].create_index([("resource_id", 1), ("snapshot_period", -1)], name="ix_res_sec_util_resource_period")
    await idx_db[RS_COL_BENCH_RECORDS].create_index("id", unique=True, name="uq_res_sec_bench_id")
    await idx_db[RS_COL_BENCH_RECORDS].create_index([("resource_id", 1), ("bench_start_date", -1)], name="ix_res_sec_bench_resource")
    await idx_db[RS_COL_READINESS].create_index("id", unique=True, name="uq_res_sec_readiness_id")
    await idx_db[RS_COL_READINESS].create_index([("resource_id", 1), ("calculated_on", -1)], name="ix_res_sec_readiness_resource")
    await idx_db[RS_COL_DEMAND_MATCHES].create_index("id", unique=True, name="uq_res_sec_match_id")
    await idx_db[RS_COL_DEMAND_MATCHES].create_index([("resource_id", 1), ("fit_score", -1)], name="ix_res_sec_match_resource_score")
    await idx_db[RS_COL_MOBILITY].create_index("id", unique=True, name="uq_res_sec_mobility_id")
    await idx_db[RS_COL_MOBILITY].create_index([("resource_id", 1), ("event_date", -1)], name="ix_res_sec_mobility_resource")
    await idx_db[RS_COL_CAREER].create_index("resource_id", unique=True, name="uq_res_sec_career_resource")
    await idx_db[RS_COL_LEARNING].create_index("id", unique=True, name="uq_res_sec_learning_id")
    await idx_db[RS_COL_CERTIFICATIONS].create_index("id", unique=True, name="uq_res_sec_cert_id")
    await idx_db[RS_COL_CERTIFICATIONS].create_index([("resource_id", 1), ("expiry_date", 1)], name="ix_res_sec_cert_resource_exp")
    await idx_db[RS_COL_COST_PROFILES].create_index("resource_id", unique=True, name="uq_res_sec_cost_resource")
    await idx_db[RS_COL_ATTENDANCE_IMPACT].create_index("id", unique=True, name="uq_res_sec_att_id")
    await idx_db[RS_COL_RESOURCE_DOCUMENTS].create_index("id", unique=True, name="uq_res_sec_doc_id")
    await idx_db[RS_COL_COMPLIANCE].create_index("id", unique=True, name="uq_res_sec_compliance_id")
    await idx_db[RS_COL_RESOURCE_NOTES].create_index("id", unique=True, name="uq_res_sec_note_id")
    await idx_db[RS_COL_RESOURCE_NOTES].create_index([("resource_id", 1), ("created_at", -1)], name="ix_res_sec_note_resource")
    await idx_db[RS_COL_ACTIVITY].create_index("id", unique=True, name="uq_res_sec_activity_id")
    await idx_db[RS_COL_ACTIVITY].create_index([("created_at", -1)], name="ix_res_sec_activity_created")
    await idx_db[RS_COL_APPROVALS].create_index("id", unique=True, name="uq_res_sec_appr_id")
    await idx_db[RS_COL_APPROVALS].create_index([("status", 1), ("submitted_on", -1)], name="ix_res_sec_appr_status")
    await idx_db[RS_COL_FORECASTS].create_index("id", unique=True, name="uq_res_sec_forecast_id")
    await idx_db[RS_COL_AI_INSIGHTS].create_index("id", unique=True, name="uq_res_sec_ai_id")
    await idx_db[RS_COL_AI_INSIGHTS].create_index([("resource_id", 1), ("generated_at", -1)], name="ix_res_sec_ai_resource_gen")

    # M10 event backbone (outbox + idempotency)
    await idx_db[COL_M10_EVENTS].create_index("event_id", unique=True, name="uq_m10_event_id")
    await idx_db[COL_M10_EVENTS].create_index([("status", 1), ("created_at", 1)], name="ix_m10_event_status_created")
    await idx_db[COL_M10_EVENTS].create_index([("topic", 1), ("created_at", -1)], name="ix_m10_event_topic_created")
    await idx_db[COL_M10_EVENTS].create_index(
        [("topic", 1), ("idempotency_key", 1)],
        unique=True,
        partialFilterExpression={"idempotency_key": {"$type": "string"}},
        name="uq_m10_event_topic_idempotency",
    )
    await idx_db[COL_M10_IDEMPOTENCY].create_index(
        [("consumer", 1), ("topic", 1), ("idempotency_key", 1)],
        unique=True,
        name="uq_m10_idempotency_consumer_topic_key",
    )
    await idx_db[COL_M10_HANDLER_AUDIT].create_index(
        [("event_id", 1), ("at", -1)], name="ix_m10_handler_audit_event_at"
    )
    await idx_db[COL_M10_HANDLER_AUDIT].create_index([("at", -1)], name="ix_m10_handler_audit_at")

    # Smart Hiring assessments
    await idx_db.assessments.create_index("id", unique=True, name="uq_hiring_assessment_id")
    await idx_db.assessments.create_index([("job_id", 1), ("created_at", -1)], name="ix_hiring_assessment_job_created")
    await idx_db.assessment_submissions.create_index("id", unique=True, name="uq_hiring_assess_sub_id")
    await idx_db.assessment_submissions.create_index("access_token", unique=True, name="uq_hiring_assess_sub_token")
    await idx_db.assessment_submissions.create_index(
        [("assessment_id", 1), ("candidate_id", 1)], name="ix_hiring_assess_sub_assess_cand"
    )
    await idx_db.assessment_submissions.create_index([("job_id", 1), ("status", 1)], name="ix_hiring_assess_sub_job_status")
    await idx_db.assessment_submissions.create_index([("invited_at", -1)], name="ix_hiring_assess_sub_invited")
    await idx_db.assessment_submissions.create_index(
        [("status", 1), ("reminder_sent_at", 1), ("invited_at", 1)],
        name="ix_hiring_assess_sub_reminder",
    )
    await idx_db.assessment_audit_log.create_index([("created_at", -1)], name="ix_hiring_assess_audit_created")
    await idx_db.assessment_audit_log.create_index([("assessment_id", 1), ("created_at", -1)], name="ix_hiring_assess_audit_assess")
    await idx_db.assessment_invite_emails.create_index([("status", 1), ("created_at", 1)], name="ix_hiring_assess_email_queue")

    # M12 Training & Development (enterprise LMS module)
    await idx_db[TD_COL_TRAINING_PROGRAMS].create_index("id", unique=True, name="uq_td_program_id")
    await idx_db[TD_COL_TRAINING_PROGRAMS].create_index("training_code", unique=True, name="uq_td_program_code", partialFilterExpression={"deleted_at": None})
    await idx_db[TD_COL_TRAINING_PROGRAMS].create_index([("status", 1), ("updated_at", -1)], name="ix_td_program_status_updated")
    await idx_db[TD_COL_TRAINING_BATCHES].create_index("id", unique=True, name="uq_td_batch_id")
    await idx_db[TD_COL_TRAINING_BATCHES].create_index([("training_id", 1), ("created_at", -1)], name="ix_td_batch_training_created")
    await idx_db[TD_COL_TRAINING_SESSIONS].create_index("id", unique=True, name="uq_td_session_id")
    await idx_db[TD_COL_TRAINING_SESSIONS].create_index([("training_id", 1), ("start_datetime", 1)], name="ix_td_session_training_start")
    await idx_db[TD_COL_TRAINING_ENROLLMENTS].create_index("id", unique=True, name="uq_td_enrollment_id")
    await idx_db[TD_COL_TRAINING_ENROLLMENTS].create_index([("training_id", 1), ("enrolled_on", -1)], name="ix_td_enr_training_enrolled")
    await idx_db[TD_COL_TRAINING_ENROLLMENTS].create_index([("employee_id", 1), ("enrollment_status", 1)], name="ix_td_enr_employee_status")
    await idx_db[TD_COL_TRAINING_ATTENDANCE].create_index("id", unique=True, name="uq_td_attendance_id")
    await idx_db[TD_COL_TRAINING_ATTENDANCE].create_index([("session_id", 1), ("employee_id", 1)], name="ix_td_att_session_employee")
    await idx_db[TD_COL_CATALOG_ITEMS].create_index("id", unique=True, name="uq_td_catalog_id")
    await idx_db[TD_COL_CATALOG_ITEMS].create_index([("catalog_type", 1), ("status", 1)], name="ix_td_catalog_type_status")
    await idx_db[TD_COL_EXTENDED_RECORDS].create_index("id", unique=True, name="uq_td_extended_id")
    await idx_db[TD_COL_EXTENDED_RECORDS].create_index([("record_type", 1), ("created_at", -1)], name="ix_td_extended_type_created")
    await idx_db[TD_COL_EXTENDED_RECORDS].create_index([("employee_id", 1), ("record_type", 1)], name="ix_td_extended_emp_type")
    await idx_db[TD_COL_APPROVAL_REQUESTS].create_index("id", unique=True, name="uq_td_appr_id")
    await idx_db[TD_COL_APPROVAL_REQUESTS].create_index([("status", 1), ("submitted_at", -1)], name="ix_td_appr_status_submitted")
    await idx_db[TD_COL_ASSESSMENTS].create_index("id", unique=True, name="uq_td_assessment_id")
    await idx_db[TD_COL_ASSESSMENTS].create_index([("training_id", 1)], name="ix_td_assessment_training")
    await idx_db[TD_COL_ASSESSMENT_RESULTS].create_index("id", unique=True, name="uq_td_assess_result_id")
    await idx_db[TD_COL_ASSESSMENT_RESULTS].create_index([("training_id", 1), ("employee_id", 1)], name="ix_td_assess_res_training_emp")

    # M13 High-Skill Talent Retention (strategic retention intelligence)
    await idx_db[HSR_COL_CRITICAL_TALENT_PROFILES].create_index("id", unique=True, name="uq_hsr_profile_id")
    await idx_db[HSR_COL_CRITICAL_TALENT_PROFILES].create_index(
        "talent_code",
        unique=True,
        name="uq_hsr_talent_code",
        partialFilterExpression={"deleted_at": None},
    )
    await idx_db[HSR_COL_CRITICAL_TALENT_PROFILES].create_index(
        [("current_risk_level", 1), ("updated_at", -1)],
        name="ix_hsr_profile_risk_updated",
    )
    await idx_db[HSR_COL_TALENT_CRITICALITY_TAGS].create_index("id", unique=True, name="uq_hsr_tag_id")
    await idx_db[HSR_COL_TALENT_CRITICALITY_TAGS].create_index(
        [("employee_id", 1), ("tag_type", 1), ("active_flag", 1)],
        name="ix_hsr_tags_emp_type_active",
    )
    await idx_db[HSR_COL_TALENT_SEGMENTS].create_index("id", unique=True, name="uq_hsr_segment_id")
    await idx_db[HSR_COL_TALENT_SEGMENTS].create_index(
        [("segment_type", 1), ("priority_score", -1)],
        name="ix_hsr_segments_type_priority",
    )
    await idx_db[HSR_COL_RISK_ASSESSMENTS].create_index("id", unique=True, name="uq_hsr_risk_id")
    await idx_db[HSR_COL_RISK_ASSESSMENTS].create_index([("employee_id", 1), ("assessed_on", -1)], name="ix_hsr_risk_emp_assessed")
    await idx_db[HSR_COL_ATTRITION_PREDICTIONS].create_index("id", unique=True, name="uq_hsr_pred_id")
    await idx_db[HSR_COL_ATTRITION_PREDICTIONS].create_index([("employee_id", 1), ("generated_at", -1)], name="ix_hsr_pred_emp_gen")
    await idx_db[HSR_COL_STAY_INTERVIEWS].create_index("id", unique=True, name="uq_hsr_stay_id")
    await idx_db[HSR_COL_STAY_INTERVIEWS].create_index([("employee_id", 1), ("scheduled_on", -1)], name="ix_hsr_stay_emp_sched")
    await idx_db[HSR_COL_RETENTION_CASES].create_index("id", unique=True, name="uq_hsr_case_id")
    await idx_db[HSR_COL_RETENTION_CASES].create_index([("status", 1), ("opened_on", -1)], name="ix_hsr_case_status_opened")
    await idx_db[HSR_COL_ENGAGEMENT_ACTION_PLANS].create_index("id", unique=True, name="uq_hsr_action_id")
    await idx_db[HSR_COL_ENGAGEMENT_ACTION_PLANS].create_index([("employee_id", 1), ("status", 1)], name="ix_hsr_action_emp_status")
    await idx_db[HSR_COL_EXIT_RISK_TRIGGERS].create_index("id", unique=True, name="uq_hsr_trigger_id")
    await idx_db[HSR_COL_EXIT_RISK_TRIGGERS].create_index([("severity", 1), ("detected_on", -1)], name="ix_hsr_trigger_sev_detected")
    await idx_db[HSR_COL_STABILITY_FORECASTS].create_index("id", unique=True, name="uq_hsr_forecast_id")
    await idx_db[HSR_COL_AI_RECOMMENDATIONS].create_index("id", unique=True, name="uq_hsr_ai_rec_id")
    await idx_db[HSR_COL_AI_FLIGHT_RISK].create_index("id", unique=True, name="uq_hsr_ai_risk_id")
    await idx_db[HSR_COL_SEARCH_LOGS].create_index("id", unique=True, name="uq_hsr_search_id")
    await idx_db[HSR_COL_SEARCH_LOGS].create_index([("ts", -1)], name="ix_hsr_search_ts")

    # M14 Employee Lifecycle Management (journey orchestration)
    await idx_db[ELM_COL_PREBOARDING].create_index("id", unique=True, name="uq_elm_preboarding_id")
    await idx_db[ELM_COL_PREBOARDING].create_index([("employee_id", 1), ("created_at", -1)], name="ix_elm_preboarding_emp_created")
    await idx_db[ELM_COL_ONBOARDING].create_index("id", unique=True, name="uq_elm_onboarding_id")
    await idx_db[ELM_COL_ONBOARDING].create_index([("employee_id", 1), ("created_at", -1)], name="ix_elm_onboarding_emp_created")
    await idx_db[ELM_COL_PROBATION].create_index("id", unique=True, name="uq_elm_probation_id")
    await idx_db[ELM_COL_PROBATION].create_index([("employee_id", 1), ("probation_end_date", 1)], name="ix_elm_probation_emp_end")
    await idx_db[ELM_COL_CONFIRMATION].create_index("id", unique=True, name="uq_elm_confirmation_id")
    await idx_db[ELM_COL_CONFIRMATION].create_index([("approval_status", 1), ("confirmation_due_date", 1)], name="ix_elm_confirmation_status_due")
    await idx_db[ELM_COL_EMPLOYEE_DOCUMENTS].create_index("id", unique=True, name="uq_elm_doc_id")
    await idx_db[ELM_COL_EMPLOYEE_DOCUMENTS].create_index([("employee_id", 1), ("uploaded_at", -1)], name="ix_elm_doc_emp_uploaded")
    await idx_db[ELM_COL_BGV].create_index("id", unique=True, name="uq_elm_bgv_id")
    await idx_db[ELM_COL_BGV].create_index([("employee_id", 1), ("bgv_overall_status", 1)], name="ix_elm_bgv_emp_status")
    await idx_db[ELM_COL_POLICY_CONSENTS].create_index("id", unique=True, name="uq_elm_consent_id")
    await idx_db[ELM_COL_POLICY_CONSENTS].create_index([("employee_id", 1), ("policy_type", 1)], name="ix_elm_consent_emp_policy")
    await idx_db[ELM_COL_ACCESS_PROVISIONING].create_index("id", unique=True, name="uq_elm_prov_id")
    await idx_db[ELM_COL_ACCESS_PROVISIONING].create_index([("employee_id", 1), ("provisioning_status", 1)], name="ix_elm_prov_emp_status")
    await idx_db[ELM_COL_PAYROLL_LINKAGE].create_index("id", unique=True, name="uq_elm_payroll_id")
    await idx_db[ELM_COL_PAYROLL_LINKAGE].create_index([("employee_id", 1), ("payroll_readiness_status", 1)], name="ix_elm_payroll_emp_status")
    await idx_db[ELM_COL_APPROVAL_REQUESTS].create_index("id", unique=True, name="uq_elm_appr_id")
    await idx_db[ELM_COL_APPROVAL_REQUESTS].create_index([("status", 1), ("submitted_at", -1)], name="ix_elm_appr_status_submitted")
    await idx_db[ELM_COL_RETENTION_SIGNALS].create_index("id", unique=True, name="uq_elm_ret_signal_id")
    await idx_db[ELM_COL_RETENTION_SIGNALS].create_index([("severity", 1), ("detected_on", -1)], name="ix_elm_ret_signal_sev_detected")
    await idx_db[ELM_COL_RESIGNATION].create_index("id", unique=True, name="uq_elm_resignation_id")
    await idx_db[ELM_COL_RESIGNATION].create_index([("approval_status", 1), ("resignation_submitted_on", -1)], name="ix_elm_resignation_status_submitted")
    await idx_db[ELM_COL_NOTICE].create_index("id", unique=True, name="uq_elm_notice_id")
    await idx_db[ELM_COL_EXIT_INTERVIEW].create_index("id", unique=True, name="uq_elm_exit_interview_id")
    await idx_db[ELM_COL_CLEARANCE].create_index("id", unique=True, name="uq_elm_clearance_id")
    await idx_db[ELM_COL_FORECASTS].create_index("id", unique=True, name="uq_elm_forecast_id")
    await idx_db[ELM_COL_AI_INSIGHTS].create_index("id", unique=True, name="uq_elm_ai_id")
    await idx_db[ELM_COL_AI_INSIGHTS].create_index([("generated_at", -1)], name="ix_elm_ai_generated")
    await idx_db[ELM_COL_LIFECYCLE_NOTES].create_index("id", unique=True, name="uq_elm_note_id")
    await idx_db[ELM_COL_LIFECYCLE_NOTES].create_index([("employee_id", 1), ("created_at", -1)], name="ix_elm_note_emp_created")
    await idx_db[ELM_COL_ACTIVITY_LOGS].create_index([("ts", -1)], name="ix_elm_activity_ts")

    # M15 Workforce Intelligence (strategic decision support)
    await idx_db[WFI_COL_SNAPSHOT_RECORDS].create_index("id", unique=True, name="uq_wfi_snapshot_id")
    await idx_db[WFI_COL_SNAPSHOT_RECORDS].create_index([("snapshot_date", -1)], name="ix_wfi_snapshot_date")
    await idx_db[WFI_COL_HEADCOUNT_RECORDS].create_index("id", unique=True, name="uq_wfi_headcount_id")
    await idx_db[WFI_COL_HEADCOUNT_RECORDS].create_index(
        [("snapshot_date", -1), ("business_unit", 1), ("department", 1), ("geography", 1)],
        name="ix_wfi_headcount_scope",
    )
    await idx_db[WFI_COL_DEMOGRAPHIC_SNAPSHOTS].create_index("id", unique=True, name="uq_wfi_demo_id")
    await idx_db[WFI_COL_DEMOGRAPHIC_SNAPSHOTS].create_index([("snapshot_date", -1), ("dimension_type", 1)], name="ix_wfi_demo_date_dim")
    await idx_db[WFI_COL_SKILL_VISIBILITY_RECORDS].create_index("id", unique=True, name="uq_wfi_skill_vis_id")
    await idx_db[WFI_COL_SKILL_VISIBILITY_RECORDS].create_index([("snapshot_date", -1), ("skill_name", 1)], name="ix_wfi_skill_vis_date_skill")
    await idx_db[WFI_COL_UTILIZATION_SNAPSHOTS].create_index("id", unique=True, name="uq_wfi_util_id")
    await idx_db[WFI_COL_UTILIZATION_SNAPSHOTS].create_index([("snapshot_date", -1), ("department", 1)], name="ix_wfi_util_date_dept")
    await idx_db[WFI_COL_ENGAGEMENT_VISIBILITY_RECORDS].create_index("id", unique=True, name="uq_wfi_eng_id")
    await idx_db[WFI_COL_PERFORMANCE_VISIBILITY_RECORDS].create_index("id", unique=True, name="uq_wfi_perf_id")
    await idx_db[WFI_COL_COMPLIANCE_VISIBILITY_RECORDS].create_index("id", unique=True, name="uq_wfi_comp_id")
    await idx_db[WFI_COL_COST_VISIBILITY_RECORDS].create_index("id", unique=True, name="uq_wfi_cost_id")
    await idx_db[WFI_COL_WORKFORCE_PLANS].create_index("id", unique=True, name="uq_wfi_plan_id")
    await idx_db[WFI_COL_DEMAND_SUPPLY_RECORDS].create_index("id", unique=True, name="uq_wfi_ds_id")
    await idx_db[WFI_COL_SCENARIO_MODELS].create_index("id", unique=True, name="uq_wfi_scenario_id")
    await idx_db[WFI_COL_MANAGER_EFFECTIVENESS_RECORDS].create_index("id", unique=True, name="uq_wfi_mgr_eff_id")
    await idx_db[WFI_COL_FORECASTS].create_index("id", unique=True, name="uq_wfi_forecast_id")
    await idx_db[WFI_COL_FORECASTS].create_index([("forecast_type", 1), ("generated_on", -1)], name="ix_wfi_forecast_type_gen")
    await idx_db[WFI_COL_ATTRITION_PREDICTIONS].create_index("id", unique=True, name="uq_wfi_attr_pred_id")
    await idx_db[WFI_COL_BURNOUT_PREDICTIONS].create_index("id", unique=True, name="uq_wfi_burn_pred_id")
    await idx_db[WFI_COL_SKILL_RISK_PREDICTIONS].create_index("id", unique=True, name="uq_wfi_skill_risk_id")
    await idx_db[WFI_COL_COST_RISK_PREDICTIONS].create_index("id", unique=True, name="uq_wfi_cost_risk_id")
    await idx_db[WFI_COL_COMPLIANCE_RISK_PREDICTIONS].create_index("id", unique=True, name="uq_wfi_comp_risk_id")
    await idx_db[WFI_COL_AI_RECOMMENDATIONS].create_index("id", unique=True, name="uq_wfi_ai_rec_id")
    await idx_db[WFI_COL_COPILOT_QUERIES].create_index("id", unique=True, name="uq_wfi_copilot_id")
    await idx_db[WFI_COL_COPILOT_QUERIES].create_index([("created_at", -1)], name="ix_wfi_copilot_created")
    await idx_db[WFI_COL_STRATEGIC_RISK_SNAPSHOTS].create_index("id", unique=True, name="uq_wfi_strat_risk_id")
    await idx_db[WFI_COL_STRATEGIC_OPPORTUNITY_SNAPSHOTS].create_index("id", unique=True, name="uq_wfi_strat_opp_id")
    await idx_db[WFI_COL_EXECUTIVE_SUMMARY_SNAPSHOTS].create_index("id", unique=True, name="uq_wfi_exec_id")
    await idx_db[WFI_COL_ACTIVITY_LOGS].create_index([("ts", -1)], name="ix_wfi_activity_ts")

    # M16 Cost Optimization & Automation
    for _coa_col in ALL_INDEXED_COLLECTIONS:
        await idx_db[_coa_col].create_index("id", unique=True, name=f"uq_{_coa_col}_id")
    await idx_db["coa_budget_spend_records"].create_index(
        [("fiscal_period", 1), ("business_unit", 1), ("department", 1)], name="ix_coa_budget_scope"
    )
    await idx_db["coa_cost_forecast_records"].create_index([("forecast_type", 1), ("generated_on", -1)], name="ix_coa_forecast_type_gen")
    await idx_db["coa_copilot_query_logs"].create_index([("created_at", -1)], name="ix_coa_copilot_created")

    # M17 Employee Satisfaction & Engagement
    await m17_ese_service.ensure_m17_indexes(db)

    try:
        recovered = await recover_stale_analyze_jobs(db)
        if recovered:
            logging.getLogger(__name__).info(
                "Career trajectory: recovered %s background analyze job(s)", recovered
            )
    except Exception:
        logging.getLogger(__name__).exception("Career trajectory job recovery failed on startup")

    if os.environ.get("M10_EVENT_CONSUMER_ENABLED", "1").strip().lower() not in ("0", "false", "no"):
        app.state.m10_consumer_task = spawn_consumer_task(db)


@app.on_event("shutdown")
async def shutdown_db_client():
    task = getattr(app.state, "m10_consumer_task", None)
    if task is not None and not task.done():
        task.cancel()
        try:
            await task
        except asyncio.CancelledError:
            pass
    client.close()
